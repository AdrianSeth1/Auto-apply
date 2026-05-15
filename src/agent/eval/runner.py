"""Eval runner: load fixtures, execute the runner callable, score outputs."""

from __future__ import annotations

import json
import time
from collections.abc import Callable
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from src.agent.eval.report import EvalCaseResult, EvalReport
from src.agent.eval.scorers import score_expectations
from src.core.config import PROJECT_ROOT

RunnerFn = Callable[[dict[str, Any]], "str | RunnerOutput"]
"""(case_input) -> output text or :class:`RunnerOutput`.

Plain-string return values are still supported for back-compat (the
agent_smoke suite uses them). Suites that want token / cost telemetry
in the report should return a :class:`RunnerOutput` instead.
"""


@dataclass
class RunnerOutput:
    """Rich runner return value: scored text plus per-case telemetry."""

    output: str
    prompt_tokens: int = 0
    output_tokens: int = 0
    cost_usd: float = 0.0


@dataclass
class EvalCase:
    id: str
    description: str
    input: dict[str, Any]
    expectations: list[dict[str, Any]]

    @classmethod
    def from_path(cls, path: Path) -> EvalCase:
        data = json.loads(path.read_text(encoding="utf-8"))
        return cls(
            id=str(data.get("id") or path.stem),
            description=str(data.get("description", "")),
            input=dict(data.get("input", {})),
            expectations=list(data.get("expectations", [])),
        )


def load_cases(suite_dir: Path) -> list[EvalCase]:
    suite_dir = Path(suite_dir)
    if not suite_dir.exists():
        raise FileNotFoundError(f"Suite directory not found: {suite_dir}")
    files = sorted(suite_dir.glob("*.json"))
    return [EvalCase.from_path(p) for p in files]


def run_eval(
    suite_name: str,
    cases: list[EvalCase],
    runner: RunnerFn,
) -> EvalReport:
    """Run every case through the runner and score the outputs."""
    report = EvalReport(suite=suite_name)
    for case in cases:
        t0 = time.monotonic()
        error: str | None = None
        output = ""
        prompt_tokens = output_tokens = 0
        cost_usd = 0.0
        try:
            raw = runner(case.input)
            if isinstance(raw, RunnerOutput):
                output = raw.output
                prompt_tokens = raw.prompt_tokens
                output_tokens = raw.output_tokens
                cost_usd = raw.cost_usd
            elif isinstance(raw, str):
                output = raw
            else:
                output = json.dumps(raw, ensure_ascii=False, default=str)
        except Exception as exc:  # noqa: BLE001 -- harness boundary
            error = f"{type(exc).__name__}: {exc}"
        elapsed = int((time.monotonic() - t0) * 1000)

        if error:
            report.cases.append(
                EvalCaseResult(
                    case_id=case.id,
                    passed=False,
                    output="",
                    expectations=[],
                    elapsed_ms=elapsed,
                    error=error,
                )
            )
            continue

        results = score_expectations(output, case.expectations)
        report.cases.append(
            EvalCaseResult(
                case_id=case.id,
                passed=all(r.passed for r in results) and bool(results),
                output=output,
                expectations=results,
                elapsed_ms=elapsed,
                prompt_tokens=prompt_tokens,
                output_tokens=output_tokens,
                cost_usd=cost_usd,
            )
        )
    return report


# ---------- built-in suites ----------


def _agent_smoke_runner(case_input: dict[str, Any]) -> str:
    """Drive the agent loop with a scripted LLM defined in the fixture.

    Each fixture supplies ``goal``, ``tools`` (allowlist), and a list of
    ``llm_responses`` to play back. This lets the suite exercise the
    full loop with zero non-determinism, suitable for CI.
    """
    from src.agent.core.loop import AgentSession, SessionLimits
    from src.agent.tools.base import get_default_registry

    goal = str(case_input.get("goal", ""))
    allowed = list(case_input.get("tools", ["finish"]))
    responses = list(case_input.get("llm_responses", []))
    limits = SessionLimits(
        max_steps=int(case_input.get("max_steps", 6)),
        step_timeout=int(case_input.get("step_timeout", 30)),
    )

    queue = list(responses)

    def scripted(_p: str, _s: str, _t: int) -> str:
        if not queue:
            raise RuntimeError("Scripted LLM ran out of responses.")
        return queue.pop(0)

    tools = get_default_registry().view(allowed)
    session = AgentSession(goal=goal, tools=tools, llm=scripted, limits=limits)
    result = session.run()
    if result.finished and result.answer is not None:
        return result.answer
    return f"[unfinished:{result.stop_reason}]"


def _form_filler_runner(case_input: dict[str, Any]) -> str:
    """Drive the agent loop end-to-end against a static HTML fixture.

    Each fixture supplies:
        html        - the form HTML
        profile     - applicant profile dict
        llm_responses - scripted LLM transcript (one element per turn)
        max_steps   - optional, defaults to 16
        url, title  - optional metadata pinned into the snapshot

    The runner emits a JSON envelope:
        {
          "stop_reason": <agent_result.stop_reason>,
          "finished":     <bool>,
          "answer":       <str|null>,
          "proposals":    [{field_id, field_label, field_type,
                            value, confidence, reasoning}, ...]
        }

    Scorers like ``field_mapping_match`` consume this directly.
    """
    import json as _json  # noqa: PLC0415

    from src.agent.core.loop import AgentSession, SessionLimits  # noqa: PLC0415
    from src.agent.tools.browser import (  # noqa: PLC0415
        build_browser_tools,
        build_snapshot_from_html,
    )
    from src.agent.tools.profile import ProfileLookupTool  # noqa: PLC0415
    from src.execution.agent_form_filler import build_goal  # noqa: PLC0415

    html = str(case_input.get("html", ""))
    profile = case_input.get("profile") or {}
    responses = list(case_input.get("llm_responses", []))
    if not html:
        raise ValueError("form_filler fixture missing 'html'")

    snapshot = build_snapshot_from_html(
        html,
        url=str(case_input.get("url", "")),
        title=str(case_input.get("title", "")),
    )
    bundle = build_browser_tools(snapshot)
    if isinstance(profile, dict):
        bundle.registry.register(ProfileLookupTool(profile))

    queue = list(responses)

    def scripted(_p: str, _s: str, _t: int) -> str:
        if not queue:
            raise RuntimeError("Scripted LLM ran out of responses.")
        return queue.pop(0)

    limits = SessionLimits(
        max_steps=int(case_input.get("max_steps", 16)),
        step_timeout=int(case_input.get("step_timeout", 30)),
    )
    goal = build_goal(
        profile_summary=", ".join(sorted(profile.keys())) if isinstance(profile, dict) else None,
        extra_context=case_input.get("extra_context"),
    )
    session = AgentSession(
        goal=goal, tools=bundle.registry, llm=scripted, limits=limits
    )
    result = session.run()

    proposals = []
    for prop in bundle.collector.latest():
        descriptor = snapshot.field_by_id(prop.field_id)
        proposals.append(
            {
                "field_id": prop.field_id,
                "field_label": descriptor.label if descriptor else "",
                "field_type": descriptor.field_type if descriptor else "",
                "value": prop.value,
                "confidence": prop.confidence,
                "reasoning": prop.reasoning,
            }
        )

    output = _json.dumps(
        {
            "stop_reason": result.stop_reason,
            "finished": result.finished,
            "answer": result.answer,
            "step_count": len(result.steps),
            "proposals": proposals,
        },
        ensure_ascii=False,
    )
    return RunnerOutput(
        output=output,
        prompt_tokens=result.total_prompt_tokens,
        output_tokens=result.total_output_tokens,
        cost_usd=result.total_cost_usd,
    )


def _materials_docx_patch_runner(case_input: dict[str, Any]) -> str:
    """Phase 15.9 runner for the docx patch suite.

    Each fixture supplies an IR payload + an inline DOCX description
    (``source_paragraphs``: list of {style, text}); the runner builds
    the DOCX in a temp dir, runs the Phase 15.2 patcher, and returns a
    JSON envelope summarising the result paragraphs so scorers can
    assert against named-style preservation and bullet swap.
    """
    import json as _json
    import tempfile
    from pathlib import Path as _Path

    from docx import Document

    from src.generation.docx_patch import PatchFallback, patch_resume_docx
    from src.generation.ir import ResumeBullet, ResumeDocument, ResumeItem

    source_paragraphs = list(case_input.get("source_paragraphs") or [])
    ir_payload = dict(case_input.get("resume_ir") or {})

    # Build the source DOCX in a temp dir.
    with tempfile.TemporaryDirectory(prefix="eval_docx_patch_") as tmp:
        tmp_path = _Path(tmp)
        src = tmp_path / "source.docx"
        out = tmp_path / "out.docx"
        doc = Document()
        for spec in source_paragraphs:
            style = (spec or {}).get("style") or "Normal"
            text = (spec or {}).get("text") or ""
            kind = (spec or {}).get("kind") or "paragraph"
            if kind == "heading":
                doc.add_heading(text, level=int(spec.get("level") or 1))
            else:
                doc.add_paragraph(text, style=style)
        doc.save(str(src))

        # Build IR (use defaults for required fields the case omits).
        experiences = [
            ResumeItem(
                source_id=f"exp-{i}",
                source_type="experience",
                name=item.get("name", f"Role {i}"),
                bullets=[
                    ResumeBullet(
                        text=b.get("text", ""),
                        source_id=f"exp-{i}",
                        source_type="experience",
                        source_entity=item.get("name", "Org"),
                    )
                    for b in (item.get("bullets") or [])
                ],
            )
            for i, item in enumerate(ir_payload.get("experiences") or [])
        ]
        document = ResumeDocument(
            target_role=ir_payload.get("target_role", "SWE"),
            company=ir_payload.get("company", "Co"),
            summary=list(ir_payload.get("summary") or []),
            skills=dict(ir_payload.get("skills") or {}),
            experiences=experiences,
        )

        envelope: dict[str, Any] = {"patched": False, "fallback": False}
        try:
            report = patch_resume_docx(src, document, output_path=out)
            envelope["patched"] = True
            envelope["operations"] = [op.kind for op in report.operations]
            patched = Document(str(out))
            envelope["paragraphs"] = [
                {"style": p.style.name, "text": p.text}
                for p in patched.paragraphs
            ]
        except PatchFallback as exc:
            envelope["fallback"] = True
            envelope["failure"] = str(exc)

    return _json.dumps(envelope, default=str)


def _materials_latex_template_runner(case_input: dict[str, Any]) -> str:
    """Phase 15.9 runner for the LaTeX manifest-adapter suite.

    Each fixture supplies a tex template body + a list of field
    mappings + an IR payload. The runner renders the .tex (no
    compile -- compile is a separate manual-smoke item) and returns
    the rendered body so the scorer can assert commands present /
    absent / escaped.
    """
    import json as _json
    import tempfile
    from pathlib import Path as _Path

    from src.documents.latex_renderer import (
        ManifestRenderError,
        render_resume_tex,
    )
    from src.documents.templates import (
        LatexConfig,
        LatexFieldMapping,
        TemplateManifest,
    )
    from src.generation.ir import ResumeBullet, ResumeDocument, ResumeItem

    template_text = str(case_input.get("template") or "{{resume.commands}}")
    mappings = [
        LatexFieldMapping(
            ir_field=str(m.get("ir_field") or ""),
            command=str(m.get("command") or ""),
            arity=int(m.get("arity") or 1),  # type: ignore[arg-type]
            wrap_with_braces=bool(m.get("wrap_with_braces", True)),
            second_ir_field=m.get("second_ir_field"),
        )
        for m in case_input.get("field_mappings") or []
    ]
    ir_payload = dict(case_input.get("resume_ir") or {})
    document = ResumeDocument(
        target_role=ir_payload.get("target_role", "SWE"),
        company=ir_payload.get("company", "Co"),
        header=dict(ir_payload.get("header") or {}),
        summary=list(ir_payload.get("summary") or []),
        skills=dict(ir_payload.get("skills") or {}),
        experiences=[
            ResumeItem(
                source_id=f"exp-{i}",
                source_type="experience",
                name=e.get("name", f"R{i}"),
                bullets=[
                    ResumeBullet(
                        text=b.get("text", ""),
                        source_id=f"exp-{i}",
                        source_type="experience",
                        source_entity=e.get("name", "Org"),
                    )
                    for b in (e.get("bullets") or [])
                ],
            )
            for i, e in enumerate(ir_payload.get("experiences") or [])
        ],
    )
    manifest = TemplateManifest(
        template_id="eval-latex",
        document_type="resume",
        template_format="latex",
        renderer="latex",
        latex=LatexConfig(
            compile_engine="pdflatex",
            field_mappings=mappings,
            strict_field_coverage=bool(case_input.get("strict_field_coverage", False)),
        ),
    )

    with tempfile.TemporaryDirectory(prefix="eval_latex_tpl_") as tmp:
        tmp_path = _Path(tmp)
        tex_in = tmp_path / "template.tex"
        tex_in.write_text(template_text, encoding="utf-8")
        tex_out = tmp_path / "out.tex"
        envelope: dict[str, Any] = {"rendered": False}
        try:
            render_resume_tex(tex_in, document, tex_out, manifest)
            envelope["rendered"] = True
            envelope["body"] = tex_out.read_text(encoding="utf-8")
        except ManifestRenderError as exc:
            envelope["rendered"] = False
            envelope["error"] = str(exc)

    return _json.dumps(envelope, default=str)


def _cover_letter_runner(case_input: dict[str, Any]) -> str:
    """Phase 15.9 runner for the cover-letter suite.

    The fixture supplies a snapshot stub, evidence bullets, a fake
    LLM output, and an expected decision. The runner instantiates
    :class:`AgentCoverLetter` with a stub ``llm_fn`` and returns a
    JSON envelope describing the dispatch decision + drift report.
    """
    import json as _json
    import uuid as _uuid
    from dataclasses import dataclass as _dc

    from src.generation.agent_cover_letter import AgentCoverLetter

    @_dc
    class _Snap:
        id: _uuid.UUID = _uuid.uuid4()
        title: str = case_input.get("snapshot_title", "Backend Intern")
        description: str = case_input.get("snapshot_description", "")
        location: str = case_input.get("snapshot_location", "Remote")
        employment_type: str = "intern"
        requirements: dict[str, Any] = None  # type: ignore[assignment]
        raw_data: dict[str, Any] = None  # type: ignore[assignment]

        def __post_init__(self) -> None:
            if self.requirements is None:
                self.requirements = dict(case_input.get("requirements") or {})
            if self.raw_data is None:
                self.raw_data = dict(case_input.get("raw_data") or {})

    profile = dict(case_input.get("profile") or {})
    evidence = list(case_input.get("evidence") or [])
    llm_output = case_input.get("llm_output")

    if llm_output is None:
        llm_fn = None
    elif llm_output == "__raise__":
        def llm_fn(prompt: str, system: str = "") -> str:
            raise RuntimeError("eval-injected llm failure")
    else:
        text = str(llm_output)

        def llm_fn(prompt: str, system: str = "") -> str:  # type: ignore[misc]
            return text

    orchestrator = AgentCoverLetter(
        job_snapshot=_Snap(),
        profile_data=profile,
        llm_fn=llm_fn,
    )
    result = orchestrator.run(
        evidence_bullets=evidence or None,
        use_agent=bool(case_input.get("use_agent", True)),
    )
    envelope: dict[str, Any] = {
        "decision": result.decision,
        "has_document": result.document is not None,
        "fact_drift": (
            {
                "blocking": result.fact_drift.has_blocking_drift,
                "number_drift": result.fact_drift.number_drift,
                "entity_drift": result.fact_drift.entity_drift,
            }
            if result.fact_drift is not None
            else None
        ),
        "agent_error": result.agent_error,
    }
    return _json.dumps(envelope, default=str)


def _filter_borderline_runner(case_input: dict[str, Any]) -> str:
    """Phase 16.4 runner for the borderline edge-case agent suite.

    Fixture shape::

        {
          "input": {
            "breakdown": {              # ScoreBreakdown-like dict, see below
              "final_score": 0.45,
              "skill_overlap": 0.7,
              "keyword_similarity": 0.2,
              "rule_bonus": 1.0,
              "quality_multiplier": 1.0,
              "disqualified": false,
              "rules": [                # optional; each entry becomes a RuleResult
                {"rule_id": "...", "passed": true,  "reason": "OK"},
                {"rule_id": "...", "passed": false, "reason": "...",
                 "evidence_excerpt": "..."}
              ]
            },
            "llm_output": "..." | null  # raw agent string; null => no llm_fn
                                        # "__raise__" => llm_fn raises
            "use_agent": true           # optional, defaults true
          }
        }

    Runner instantiates :class:`EdgeCaseAgent` with a stub ``llm_fn``
    that returns the fixture's ``llm_output`` verbatim, runs the
    agent, and emits the decision as a JSON envelope so the existing
    ``json_field_equals`` / ``json_field_contains`` scorers (Phase
    15.9) can assert against it.

    The envelope shape::

        {
          "kind": "agent_ok" | "agent_malformed" | "agent_error" | "not_invoked",
          "verdict": "surface" | "reject" | "abstain",
          "confidence": <float>,
          "rationale": "...",
          "final_score": <float>,
          "is_borderline": <bool>
        }
    """
    import json as _json

    from src.matching.edge_case_agent import EdgeCaseAgent, is_borderline
    from src.matching.rules import RuleResult, RuleVerdict
    from src.matching.scorer import ScoreBreakdown

    bd_in = dict(case_input.get("breakdown") or {})
    rules_in = list(bd_in.get("rules") or [])
    rules = [
        RuleResult(
            rule_id=str(r.get("rule_id") or r.get("rule_name") or "rule"),
            rule_name=str(r.get("rule_name") or r.get("rule_id") or "rule"),
            passed=bool(r.get("passed", True)),
            reason=str(r.get("reason") or ""),
            verdict=("pass" if r.get("passed", True) else "fail"),
            evidence_excerpt=r.get("evidence_excerpt"),
        )
        for r in rules_in
    ]
    disqualified = bool(bd_in.get("disqualified", False))
    final_score = float(bd_in.get("final_score", 0.5))

    breakdown = ScoreBreakdown(
        job_id=str(bd_in.get("job_id") or "fixture-job"),
        company=str(bd_in.get("company") or "FixtureCo"),
        title=str(bd_in.get("title") or "Borderline Role"),
        final_score=final_score,
        skill_overlap=float(bd_in.get("skill_overlap", 0.0)),
        keyword_similarity=float(bd_in.get("keyword_similarity", 0.0)),
        rule_bonus=float(bd_in.get("rule_bonus", 0.0 if disqualified else 1.0)),
        quality_multiplier=float(bd_in.get("quality_multiplier", 1.0)),
        rule_verdict=RuleVerdict(
            job_id=str(bd_in.get("job_id") or "fixture-job"),
            passed=not disqualified,
            results=rules,
        ),
        disqualified=disqualified,
        disqualify_reasons=[r.reason for r in rules if not r.passed],
        disqualify_results=[r for r in rules if not r.passed],
        job_snapshot_id=bd_in.get("job_snapshot_id"),
    )

    llm_output = case_input.get("llm_output")
    if llm_output is None:
        llm_fn = None
    elif llm_output == "__raise__":
        def llm_fn(prompt: str, tools: dict[str, Any]) -> str:
            raise RuntimeError("eval-injected llm failure")
    else:
        text = str(llm_output)

        def llm_fn(prompt: str, tools: dict[str, Any]) -> str:  # type: ignore[misc]
            return text

    agent = EdgeCaseAgent(breakdown, llm_fn=llm_fn)
    decision = agent.run(use_agent=bool(case_input.get("use_agent", True)))

    envelope = {
        "kind": decision.kind,
        "verdict": decision.verdict,
        "confidence": decision.confidence,
        "rationale": decision.rationale,
        "agent_error": decision.agent_error,
        "final_score": breakdown.final_score,
        "is_borderline": is_borderline(breakdown.final_score),
    }
    return _json.dumps(envelope, default=str)


_BUILTIN_SUITES: dict[str, tuple[Path, RunnerFn]] = {
    "agent_smoke": (
        PROJECT_ROOT / "tests" / "agent_evals" / "fixtures" / "agent_smoke",
        _agent_smoke_runner,
    ),
    "form_filler": (
        PROJECT_ROOT / "tests" / "agent_evals" / "fixtures" / "form_filler",
        _form_filler_runner,
    ),
    # Phase 15.9 suites.
    "materials_docx_patch": (
        PROJECT_ROOT / "tests" / "agent_evals" / "fixtures" / "materials_docx_patch",
        _materials_docx_patch_runner,
    ),
    "materials_latex_template": (
        PROJECT_ROOT / "tests" / "agent_evals" / "fixtures" / "materials_latex_template",
        _materials_latex_template_runner,
    ),
    "cover_letter": (
        PROJECT_ROOT / "tests" / "agent_evals" / "fixtures" / "cover_letter",
        _cover_letter_runner,
    ),
    # Phase 16.4 suite -- borderline filter edge-case agent.
    "filter_borderline": (
        PROJECT_ROOT / "tests" / "agent_evals" / "fixtures" / "filter_borderline",
        _filter_borderline_runner,
    ),
}


def list_suites() -> list[str]:
    return sorted(_BUILTIN_SUITES)


def run_suite(name: str) -> EvalReport:
    if name not in _BUILTIN_SUITES:
        raise KeyError(f"Unknown suite '{name}'. Available: {list_suites()}.")
    fixtures_dir, runner = _BUILTIN_SUITES[name]
    cases = load_cases(fixtures_dir)
    return run_eval(name, cases, runner)
