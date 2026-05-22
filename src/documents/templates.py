"""Document template package management.

Templates are first-class assets: a DOCX file owns visual style definitions,
while manifest.json describes named styles and content capacity constraints.
"""

from __future__ import annotations

import json
import logging
import re
import shutil
from pathlib import Path
from typing import Literal

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Inches, Pt
from pydantic import BaseModel, Field

from src.core.config import PROJECT_ROOT

logger = logging.getLogger("autoapply.documents.templates")

TEMPLATE_REGISTRY: dict[str, Path] = {}
TEMPLATE_ROOT = PROJECT_ROOT / "data" / "templates"
DEFAULT_TEMPLATE_IDS = {
    "resume": "ats_single_column_v1",
    "cover_letter": "classic_v1",
}
_TEMPLATE_ID_RE = re.compile(r"^[A-Za-z0-9][A-Za-z0-9_-]{0,99}$")

DocumentType = Literal["resume", "cover_letter"]
TemplateFormat = Literal["docx", "latex"]
TemplateRenderer = Literal["docx", "latex"]
TemplateOutput = Literal["docx", "pdf", "tex"]


class TemplatePage(BaseModel):
    size: str = "letter"
    max_pages: int = 1


class TemplateSection(BaseModel):
    enabled: bool = True
    max_items: int | None = None
    max_bullets_per_item: int | None = None
    max_words_per_bullet: int | None = None
    max_lines: int | None = None


class TemplateCapacity(BaseModel):
    max_pages: int = 1
    max_sections: int | None = None
    max_experience_items: int | None = None
    max_project_items: int | None = None
    max_bullets_total: int | None = None
    max_words_per_bullet: int | None = None
    max_skill_lines: int | None = None


class TemplateStyleOverride(BaseModel):
    """User-editable font/spacing tweaks for one named DOCX style.

    Stored on the manifest so the Template Editor can let users adjust
    font family, size, bold, line spacing, and paragraph spacing
    without leaving the browser. Re-applied to ``template.docx`` on
    every manifest save so the next render picks up the new look.
    """

    font: str | None = None
    size: int | None = None
    bold: bool | None = None
    italic: bool | None = None
    line_spacing: float | None = None
    # Paragraph spacing in points -- "space before this paragraph" and
    # "space after this paragraph". Pre-render defaults vary per style
    # (e.g. Resume.SectionHeading has space_before=6pt to give air
    # between sections); overrides let users tighten or relax that.
    space_before_pt: float | None = None
    space_after_pt: float | None = None


LatexEngine = Literal["pdflatex", "xelatex", "lualatex"]


class LatexFieldMapping(BaseModel):
    """Phase 15.3: maps a resume / cover-letter IR field to a LaTeX
    command in a template. Consumed by the 15.4 manifest-adapter
    dispatcher so the renderer never has to bake template-specific
    commands into Python.

    Example::

        LatexFieldMapping(
            ir_field="header.name",
            command="resumeheadername",
            arity=1,
            wrap_with_braces=True,
        )

    Renders to ``\\resumeheadername{Alice Smith}``.
    """

    ir_field: str  # dotted path into the IR (header.name, skills.must_have, ...)
    command: str  # bare LaTeX command name (no leading backslash)
    arity: Literal[0, 1, 2] = 1
    wrap_with_braces: bool = True
    # For arity-2 commands like ``\experienceitem{Title}{Dates}`` --
    # the second slot reads from this IR field.
    second_ir_field: str | None = None


class LatexConfig(BaseModel):
    """Phase 15.3: LaTeX-specific template package configuration.

    Optional on the manifest -- DOCX-only packages do not populate
    this. The 15.4 renderer reads ``compile_engine`` to choose between
    ``pdflatex``, ``xelatex``, and ``lualatex``; the
    ``escape_allowlist`` lets a template opt out of escaping for
    characters it wants to pass through (e.g. a template that prints
    raw URLs with ``\\url{}`` does not need ``%`` escaped inside).
    ``required_packages`` is purely advisory -- listed at validation
    time but never auto-installed.

    Per D024: arbitrary LaTeX may be imported, but it is not active
    until a manifest exists and a sample compile passes."""

    compile_engine: LatexEngine = "pdflatex"
    # Files (relative to the template package dir) that should be
    # copied next to the rendered ``.tex`` so ``\\input`` /
    # ``\\includegraphics`` succeed. Validated to stay inside the
    # package dir at template-install time (D013 mirror).
    assets: list[str] = Field(default_factory=list)
    # Characters whose default LaTeX escaping is suppressed for this
    # template (e.g. a template that escapes ``&`` itself via a custom
    # column macro can put ``"&"`` here).
    escape_allowlist: list[str] = Field(default_factory=list)
    # Required LaTeX packages, listed for the operator to install --
    # not auto-installed. Maps package name -> minimum version, or
    # empty string for "any".
    required_packages: dict[str, str] = Field(default_factory=dict)
    # IR-field -> LaTeX-command mapping for templates that use custom
    # commands (the default ``latex_engine.py`` resume_template uses
    # placeholder substitution; custom templates use the mapping
    # approach so the agent / renderer never hard-codes commands).
    field_mappings: list[LatexFieldMapping] = Field(default_factory=list)
    # Sample IR file relative to the package dir (used by 15.8
    # adapter assistant to validate a manifest via a real compile).
    sample_ir: str | None = None
    # Strict mode: if True, the renderer refuses to render an IR that
    # has fields not declared in ``field_mappings``. Defaults False
    # so existing placeholder-style templates keep working.
    strict_field_coverage: bool = False


class TemplateManifest(BaseModel):
    template_id: str
    document_type: DocumentType
    template_format: TemplateFormat = "docx"
    renderer: TemplateRenderer = "docx"
    supported_outputs: list[TemplateOutput] = Field(default_factory=lambda: ["docx", "pdf"])
    name: str = ""
    description: str = ""
    page: TemplatePage = Field(default_factory=TemplatePage)
    styles: dict[str, str] = Field(default_factory=dict)
    sections: dict[str, TemplateSection] = Field(default_factory=dict)
    section_order: list[str] = Field(default_factory=list)
    capacity: TemplateCapacity = Field(default_factory=TemplateCapacity)
    blocks: dict[str, str] = Field(default_factory=dict)
    # Phase 15.3: LaTeX-specific options live in their own sub-model
    # so DOCX-only packages do not need to know about it.
    latex: LatexConfig | None = None
    # User-editable per-style font/size/spacing overrides for DOCX
    # templates. Keys are the same logical style keys used in
    # ``styles`` (e.g. "name", "normal", "body"). Empty for LaTeX
    # templates -- the .tex source owns the look there.
    style_overrides: dict[str, TemplateStyleOverride] = Field(default_factory=dict)
    # Target page count: fitting expands content to roughly fill this,
    # and the post-render validator errors when the produced PDF
    # differs. Kept in sync with ``capacity.max_pages`` so older code
    # paths reading the legacy field still get the right value.
    target_pages: int = 1
    # Filename pattern preset used when materializing artifacts. See
    # :data:`src.documents.file_manager.FILENAME_PATTERNS` for the
    # supported values.
    filename_pattern: str = "company_role_date"
    # Label substituted into ``type_custom_seq`` pattern; ignored by
    # other patterns.
    filename_custom_label: str = ""
    # Optional "second font" used for bold runs the LLM emits via
    # ``**inline**`` markup. Empty string / None means "use the same
    # font as the surrounding body" -- which is what most resumes
    # want. Pairing Arial body with a Georgia / Cambria emphasis font
    # is the classic typographic trick this enables.
    emphasis_font: str = ""


class TemplatePackage(BaseModel):
    template_id: str
    document_type: DocumentType
    directory: Path
    template_path: Path
    manifest_path: Path
    manifest: TemplateManifest


def register_template(name: str, path: Path) -> None:
    if not path.exists():
        raise FileNotFoundError(f"Template not found: {path}")
    TEMPLATE_REGISTRY[name] = path
    logger.debug("Registered template '%s' at %s", name, path)


def get_template_path(name: str) -> Path:
    if name not in TEMPLATE_REGISTRY:
        raise KeyError(f"Template '{name}' not registered. Available: {list(TEMPLATE_REGISTRY)}")
    return TEMPLATE_REGISTRY[name]


def discover_templates(template_dir: Path) -> None:
    """Auto-register all .docx files found in template_dir."""
    if not template_dir.exists():
        logger.warning("Template directory not found: %s", template_dir)
        return
    for path in template_dir.glob("*.docx"):
        register_template(path.stem, path)
    logger.info("Discovered %d templates in %s", len(TEMPLATE_REGISTRY), template_dir)


def ensure_template_package(
    document_type: DocumentType,
    template_id: str | None = None,
    *,
    template_root: Path = TEMPLATE_ROOT,
) -> TemplatePackage:
    """Create the default template package if needed and return it."""
    template_id = template_id or DEFAULT_TEMPLATE_IDS[document_type]
    package_dir = _template_package_dir(document_type, template_id, template_root)
    package_dir.mkdir(parents=True, exist_ok=True)

    manifest_path = package_dir / "manifest.json"

    if manifest_path.exists():
        package = load_template_package(document_type, template_id, template_root=template_root)
        _ensure_required_markers(package)
        _write_sample_assets(package)
        return load_template_package(document_type, template_id, template_root=template_root)

    template_path = package_dir / "template.docx"

    manifest_path.write_text(
        json.dumps(_default_manifest(document_type, template_id), indent=2) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    if not template_path.exists():
        if document_type == "resume":
            _create_default_resume_template(template_path)
        else:
            _create_default_cover_letter_template(template_path)

    package = load_template_package(document_type, template_id, template_root=template_root)
    _ensure_required_markers(package)
    _write_sample_assets(package)
    return load_template_package(document_type, template_id, template_root=template_root)


def list_template_packages(
    document_type: DocumentType | None = None,
    *,
    template_root: Path = TEMPLATE_ROOT,
) -> dict[str, list[dict]]:
    """Return discovered template packages grouped by document type."""
    document_types = [document_type] if document_type else ["resume", "cover_letter"]
    grouped: dict[str, list[dict]] = {kind: [] for kind in document_types}
    for kind in document_types:
        ensure_template_package(kind, template_root=template_root)
        kind_dir = template_root / kind
        for manifest_path in sorted(kind_dir.glob("*/manifest.json")):
            template_id = manifest_path.parent.name
            try:
                package = load_template_package(kind, template_id, template_root=template_root)
            except Exception as exc:
                logger.warning(
                    "Skipping invalid template package %s: %s",
                    manifest_path.parent,
                    exc,
                )
                continue
            grouped[kind].append(serialize_template_package(package))
    return grouped


def save_uploaded_template_package(
    *,
    document_type: DocumentType,
    filename: str,
    content: bytes,
    template_name: str | None = None,
    template_root: Path = TEMPLATE_ROOT,
) -> dict:
    """Persist an uploaded DOCX or single-file LaTeX template package."""
    if document_type not in DEFAULT_TEMPLATE_IDS:
        raise ValueError("Unsupported template document type.")
    suffix = Path(filename).suffix.lower()
    if suffix not in {".docx", ".tex"}:
        raise ValueError("Only .docx and .tex templates are supported.")
    if not content:
        raise ValueError("Template upload is empty.")

    if suffix == ".tex":
        return _save_uploaded_latex_template_package(
            document_type=document_type,
            filename=filename,
            content=content,
            template_name=template_name,
            template_root=template_root,
        )

    return _save_uploaded_docx_template_package(
        document_type=document_type,
        filename=filename,
        content=content,
        template_name=template_name,
        template_root=template_root,
    )


def create_latex_template_package(
    *,
    document_type: DocumentType,
    template_name: str | None = None,
    description: str | None = None,
    template_root: Path = TEMPLATE_ROOT,
) -> dict:
    """Create a blank single-file LaTeX template package and return metadata."""
    if document_type not in DEFAULT_TEMPLATE_IDS:
        raise ValueError("Unsupported template document type.")

    default_name = "LaTeX Resume Template" if document_type == "resume" else "LaTeX Cover Letter"
    display_name = (template_name or default_name).strip() or default_name
    template_id = _unique_template_id(template_root / document_type, _slugify(display_name))
    package_dir = template_root / document_type / template_id
    package_dir.mkdir(parents=True, exist_ok=False)

    template_path = package_dir / "template.tex"
    manifest_path = package_dir / "manifest.json"
    try:
        manifest_payload = _default_latex_manifest(document_type, template_id)
        manifest_payload["name"] = display_name
        if description is not None:
            manifest_payload["description"] = description.strip()
        template_path.write_text(
            _default_latex_template(document_type),
            encoding="utf-8",
            newline="\n",
        )
        manifest_path.write_text(
            json.dumps(manifest_payload, indent=2) + "\n",
            encoding="utf-8",
            newline="\n",
        )
        package = load_template_package(document_type, template_id, template_root=template_root)
        _write_sample_assets(package)
        return serialize_template_package(package)
    except Exception:
        shutil.rmtree(package_dir, ignore_errors=True)
        raise


def get_template_package_detail(
    document_type: DocumentType,
    template_id: str,
    *,
    template_root: Path = TEMPLATE_ROOT,
) -> dict:
    """Return template metadata plus editable content for text-based templates."""
    package = load_template_package(document_type, template_id, template_root=template_root)
    serialized = serialize_template_package(package)
    is_latex = package.manifest.renderer == "latex"
    serialized["content"] = (
        package.template_path.read_text(encoding="utf-8") if is_latex else None
    )
    if not is_latex:
        serialized["editable_styles"] = editable_style_options(document_type)
        serialized["style_overrides"] = {
            key: override.model_dump(mode="json")
            for key, override in package.manifest.style_overrides.items()
        }
    else:
        serialized["editable_styles"] = []
        serialized["style_overrides"] = {}
    return serialized


def delete_template_package(
    document_type: DocumentType,
    template_id: str,
    *,
    template_root: Path = TEMPLATE_ROOT,
) -> None:
    """Delete a template package directory.

    Refuses to delete the seeded defaults (DEFAULT_TEMPLATE_IDS[document_type])
    so the app always has a renderable fallback. Validates the template id
    via _template_package_dir() so callers cannot escape the document-type
    root via path traversal. No-op when the package directory does not exist.
    """

    if document_type not in {"resume", "cover_letter"}:
        raise ValueError("Unsupported template document type.")

    if template_id == DEFAULT_TEMPLATE_IDS.get(document_type):
        raise ValueError("Built-in default templates cannot be deleted.")

    package_dir = _template_package_dir(document_type, template_id, template_root)
    if not package_dir.exists():
        raise FileNotFoundError(f"Template '{template_id}' not found.")
    if not package_dir.is_dir():
        raise ValueError("Template path is not a directory.")

    shutil.rmtree(package_dir)


# Default font/size/bold/line-spacing values applied by
# ``_ensure_default_styles`` -- mirrored here so the UI can show what a
# style looks like out of the box (and so an empty override falls back
# to a predictable baseline rather than whatever happens to be in the
# DOCX). Keep in sync with ``_ensure_default_styles``.
# Keep this table in lockstep with ``_ensure_default_styles`` so what the
# UI shows as the baseline is exactly what the renderer applies when no
# override is present. Adding a property here also surfaces it in the
# Template Editor's per-style row -- the frontend reads
# ``supports_line_spacing`` etc. to decide which inputs to render.
_EDITABLE_RESUME_STYLES: tuple[tuple[str, str, dict], ...] = (
    (
        "name",
        "Resume.Name",
        {"font": "Arial", "size": 16, "bold": True, "line_spacing": 1.0,
         "space_before_pt": 0, "space_after_pt": 0},
    ),
    (
        "contact",
        "Resume.Contact",
        {"font": "Arial", "size": 9, "bold": False, "line_spacing": 1.0,
         "space_before_pt": 0, "space_after_pt": 0},
    ),
    (
        "section_heading",
        "Resume.SectionHeading",
        {"font": "Arial", "size": 11, "bold": True, "line_spacing": 1.0,
         "space_before_pt": 6, "space_after_pt": 0},
    ),
    (
        "item_title",
        "Resume.ItemTitle",
        {"font": "Arial", "size": 10, "bold": True, "line_spacing": 1.0,
         "space_before_pt": 0, "space_after_pt": 0},
    ),
    (
        "item_subtitle",
        "Resume.ItemSubtitle",
        {"font": "Arial", "size": 9, "bold": False, "line_spacing": 1.0,
         "space_before_pt": 0, "space_after_pt": 0},
    ),
    (
        "item_meta",
        "Resume.ItemMeta",
        {"font": "Arial", "size": 9, "bold": False, "line_spacing": 1.0,
         "space_before_pt": 0, "space_after_pt": 0},
    ),
    (
        "normal",
        "Resume.Normal",
        {"font": "Arial", "size": 9, "bold": False, "line_spacing": 1.0,
         "space_before_pt": 0, "space_after_pt": 0},
    ),
    (
        "bullet",
        "Resume.Bullet",
        {"font": "Arial", "size": 9, "bold": False, "line_spacing": 1.0,
         "space_before_pt": 0, "space_after_pt": 0},
    ),
    (
        "skill_category",
        "Resume.SkillCategory",
        {"font": "Arial", "size": 9, "bold": True, "line_spacing": 1.0,
         "space_before_pt": 0, "space_after_pt": 0},
    ),
    (
        "skill_line",
        "Resume.SkillLine",
        {"font": "Arial", "size": 9, "bold": False, "line_spacing": 1.0,
         "space_before_pt": 0, "space_after_pt": 0},
    ),
)

_EDITABLE_COVER_LETTER_STYLES: tuple[tuple[str, str, dict], ...] = (
    (
        "header",
        "CoverLetter.Header",
        {"font": "Times New Roman", "size": 11, "bold": True, "line_spacing": 1.0,
         "space_before_pt": 0, "space_after_pt": 0},
    ),
    (
        "date",
        "CoverLetter.Date",
        {"font": "Times New Roman", "size": 11, "bold": False, "line_spacing": 1.0,
         "space_before_pt": 10, "space_after_pt": 0},
    ),
    (
        "recipient",
        "CoverLetter.Recipient",
        {"font": "Times New Roman", "size": 11, "bold": False, "line_spacing": 1.0,
         "space_before_pt": 0, "space_after_pt": 10},
    ),
    (
        "body",
        "CoverLetter.Body",
        {"font": "Times New Roman", "size": 11, "bold": False, "line_spacing": 1.05,
         "space_before_pt": 0, "space_after_pt": 7},
    ),
    (
        "signature",
        "CoverLetter.Signature",
        {"font": "Times New Roman", "size": 11, "bold": False, "line_spacing": 1.0,
         "space_before_pt": 7, "space_after_pt": 0},
    ),
)


def editable_style_options(document_type: DocumentType) -> list[dict]:
    """Return the UI-friendly editable-style table for a document type.

    Used by the Template Editor to render a row per named DOCX style.
    Each entry carries the manifest ``styles`` key, the Word style name
    it maps to, a human label, and the baseline defaults the editor
    uses when ``style_overrides`` is empty. Every style row supports
    every property in the override schema -- the renderer simply leaves
    a property at the default when the user did not touch it.
    """
    if document_type not in DEFAULT_TEMPLATE_IDS:
        raise ValueError("Unsupported template document type.")
    table = (
        _EDITABLE_RESUME_STYLES
        if document_type == "resume"
        else _EDITABLE_COVER_LETTER_STYLES
    )
    return [
        {
            "key": key,
            "style_name": style_name,
            "label": key.replace("_", " ").title(),
            "defaults": dict(defaults),
            # Kept for back-compat with the older Vue template, but all
            # editable styles now support line_spacing + paragraph spacing.
            "supports_line_spacing": True,
            "supports_paragraph_spacing": True,
            "supports_italic": True,
        }
        for key, style_name, defaults in table
    ]


def update_docx_template_styles(
    *,
    document_type: DocumentType,
    template_id: str,
    overrides: dict[str, dict],
    template_name: str | None = None,
    description: str | None = None,
    target_pages: int | None = None,
    filename_pattern: str | None = None,
    filename_custom_label: str | None = None,
    emphasis_font: str | None = None,
    template_root: Path = TEMPLATE_ROOT,
) -> dict:
    """Apply user-edited font/size/spacing overrides to a DOCX template.

    Persists the overrides on ``manifest.json`` and re-applies them to
    ``template.docx`` so the next render uses the new look. Refuses
    LaTeX templates -- their look lives in the .tex source which the
    text editor already exposes.
    """
    package = load_template_package(document_type, template_id, template_root=template_root)
    if package.manifest.renderer != "docx":
        raise ValueError("Only DOCX templates accept style overrides.")

    editable_keys = {entry["key"] for entry in editable_style_options(document_type)}
    parsed_overrides: dict[str, TemplateStyleOverride] = {}
    for key, payload in (overrides or {}).items():
        if key not in editable_keys:
            raise ValueError(f"Unknown style key: {key}")
        if not isinstance(payload, dict):
            raise ValueError(f"Invalid override payload for {key}.")
        parsed_overrides[key] = TemplateStyleOverride.model_validate(payload)

    manifest_updates: dict = {
        "style_overrides": parsed_overrides,
        "name": (
            template_name.strip() if template_name is not None else package.manifest.name
        ),
        "description": (
            description.strip() if description is not None else package.manifest.description
        ),
    }
    _apply_template_settings_updates(
        manifest_updates,
        package.manifest,
        target_pages=target_pages,
        filename_pattern=filename_pattern,
        filename_custom_label=filename_custom_label,
        emphasis_font=emphasis_font,
    )
    manifest = package.manifest.model_copy(update=manifest_updates)
    package.manifest_path.write_text(
        json.dumps(manifest.model_dump(mode="json"), indent=2) + "\n",
        encoding="utf-8",
        newline="\n",
    )

    doc = Document(str(package.template_path))
    _ensure_default_styles(doc, document_type)
    _apply_style_overrides(doc, document_type, parsed_overrides)
    doc.save(str(package.template_path))

    package = load_template_package(document_type, template_id, template_root=template_root)
    _write_sample_assets(package)
    return get_template_package_detail(
        document_type, template_id, template_root=template_root
    )


def _apply_style_overrides(
    doc: Document,
    document_type: DocumentType,
    overrides: dict[str, TemplateStyleOverride],
) -> None:
    """Layer user font/size/bold/italic/line-spacing/paragraph-spacing
    overrides on top of the existing style definitions left by
    ``_ensure_default_styles``.

    Touches only the fields the user actually set so layout-affecting
    properties not exposed in the UI (tab stops, base style) survive.
    """
    table = (
        _EDITABLE_RESUME_STYLES
        if document_type == "resume"
        else _EDITABLE_COVER_LETTER_STYLES
    )
    style_name_by_key = {key: style_name for key, style_name, _ in table}
    for key, override in overrides.items():
        style_name = style_name_by_key.get(key)
        if style_name is None:
            continue
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        if override.font:
            style.font.name = override.font
        if override.size:
            style.font.size = Pt(override.size)
        if override.bold is not None:
            style.font.bold = override.bold
        if override.italic is not None:
            style.font.italic = override.italic
        if override.line_spacing is not None:
            style.paragraph_format.line_spacing = override.line_spacing
        if override.space_before_pt is not None:
            style.paragraph_format.space_before = Pt(override.space_before_pt)
        if override.space_after_pt is not None:
            style.paragraph_format.space_after = Pt(override.space_after_pt)


def update_latex_template_package(
    *,
    document_type: DocumentType,
    template_id: str,
    content: str,
    template_name: str | None = None,
    description: str | None = None,
    target_pages: int | None = None,
    filename_pattern: str | None = None,
    filename_custom_label: str | None = None,
    emphasis_font: str | None = None,
    template_root: Path = TEMPLATE_ROOT,
) -> dict:
    """Update editable LaTeX template content and metadata."""
    package = load_template_package(document_type, template_id, template_root=template_root)
    if package.manifest.renderer != "latex":
        raise ValueError("Only LaTeX templates can be edited as text.")
    if not content.strip():
        raise ValueError("Template content is required.")

    package.template_path.write_text(content, encoding="utf-8", newline="\n")
    manifest_updates: dict = {
        "name": template_name.strip() if template_name is not None else package.manifest.name,
        "description": (
            description.strip() if description is not None else package.manifest.description
        ),
    }
    _apply_template_settings_updates(
        manifest_updates,
        package.manifest,
        target_pages=target_pages,
        filename_pattern=filename_pattern,
        filename_custom_label=filename_custom_label,
        emphasis_font=emphasis_font,
    )
    manifest = package.manifest.model_copy(update=manifest_updates)
    package.manifest_path.write_text(
        json.dumps(manifest.model_dump(mode="json"), indent=2) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    package = load_template_package(document_type, template_id, template_root=template_root)
    _write_sample_assets(package)
    return get_template_package_detail(
        document_type, template_id, template_root=template_root
    )


def _apply_template_settings_updates(
    updates: dict,
    manifest: TemplateManifest,
    *,
    target_pages: int | None,
    filename_pattern: str | None,
    filename_custom_label: str | None,
    emphasis_font: str | None = None,
) -> None:
    """Apply optional template-level settings into a model_copy(update=...) dict.

    Keeps capacity.max_pages in lockstep with target_pages so legacy
    callers (e.g. validators using capacity.max_pages directly) keep
    working without a migration.
    """
    from src.documents.file_manager import FILENAME_PATTERNS  # noqa: PLC0415

    if target_pages is not None:
        if target_pages < 1 or target_pages > 5:
            raise ValueError("target_pages must be between 1 and 5.")
        updates["target_pages"] = target_pages
        new_capacity = manifest.capacity.model_copy(update={"max_pages": target_pages})
        updates["capacity"] = new_capacity
    if filename_pattern is not None:
        if filename_pattern not in FILENAME_PATTERNS:
            raise ValueError(
                f"filename_pattern must be one of {FILENAME_PATTERNS}."
            )
        updates["filename_pattern"] = filename_pattern
    if filename_custom_label is not None:
        updates["filename_custom_label"] = filename_custom_label.strip()
    if emphasis_font is not None:
        updates["emphasis_font"] = emphasis_font.strip()


def _save_uploaded_docx_template_package(
    *,
    document_type: DocumentType,
    filename: str,
    content: bytes,
    template_name: str | None,
    template_root: Path,
) -> dict:
    """Persist an uploaded DOCX as a template package and return metadata."""

    display_name = (template_name or Path(filename).stem or "Uploaded Template").strip()
    template_id = _unique_template_id(template_root / document_type, _slugify(display_name))
    package_dir = template_root / document_type / template_id
    package_dir.mkdir(parents=True, exist_ok=False)

    template_path = package_dir / "template.docx"
    manifest_path = package_dir / "manifest.json"
    try:
        template_path.write_bytes(content)
        doc = Document(str(template_path))
        manifest_payload = _default_manifest(document_type, template_id)
        manifest_payload["name"] = display_name
        manifest_payload["description"] = f"Uploaded {document_type.replace('_', ' ')} template."
        manifest_path.write_text(
            json.dumps(manifest_payload, indent=2) + "\n",
            encoding="utf-8",
            newline="\n",
        )
        manifest = TemplateManifest.model_validate(manifest_payload)
        _ensure_default_styles(doc, document_type)
        doc.save(str(template_path))
        package = TemplatePackage(
            template_id=template_id,
            document_type=document_type,
            directory=package_dir,
            template_path=template_path,
            manifest_path=manifest_path,
            manifest=manifest,
        )
        _ensure_required_markers(package)
        package = load_template_package(document_type, template_id, template_root=template_root)
        _write_sample_assets(package)
        return serialize_template_package(package)
    except Exception:
        shutil.rmtree(package_dir, ignore_errors=True)
        raise


def _save_uploaded_latex_template_package(
    *,
    document_type: DocumentType,
    filename: str,
    content: bytes,
    template_name: str | None,
    template_root: Path,
) -> dict:
    """Persist an uploaded single-file LaTeX template package and return metadata."""
    try:
        template_text = content.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise ValueError("LaTeX templates must be UTF-8 text files.") from exc
    if "\x00" in template_text:
        raise ValueError("LaTeX template contains invalid null bytes.")

    display_name = (template_name or Path(filename).stem or "Uploaded LaTeX Template").strip()
    template_id = _unique_template_id(template_root / document_type, _slugify(display_name))
    package_dir = template_root / document_type / template_id
    package_dir.mkdir(parents=True, exist_ok=False)

    template_path = package_dir / "template.tex"
    manifest_path = package_dir / "manifest.json"
    try:
        manifest_payload = _default_latex_manifest(document_type, template_id)
        manifest_payload["name"] = display_name
        manifest_payload["description"] = (
            f"Uploaded {document_type.replace('_', ' ')} LaTeX template."
        )
        template_path.write_text(template_text, encoding="utf-8", newline="\n")
        manifest_path.write_text(
            json.dumps(manifest_payload, indent=2) + "\n",
            encoding="utf-8",
            newline="\n",
        )
        package = load_template_package(document_type, template_id, template_root=template_root)
        _write_sample_assets(package)
        return serialize_template_package(package)
    except Exception:
        shutil.rmtree(package_dir, ignore_errors=True)
        raise


def serialize_template_package(package: TemplatePackage) -> dict:
    """Serialize template metadata for APIs/UI without leaking system internals."""
    preview_pdf = package.directory / "preview.pdf"
    preview_png = package.directory / "preview.png"
    is_default = package.template_id == DEFAULT_TEMPLATE_IDS.get(package.document_type)
    return {
        "template_id": package.template_id,
        "document_type": package.document_type,
        "template_format": package.manifest.template_format,
        "renderer": package.manifest.renderer,
        "supported_outputs": package.manifest.supported_outputs,
        "name": package.manifest.name or package.template_id.replace("_", " ").title(),
        "description": package.manifest.description,
        "manifest": package.manifest.model_dump(mode="json"),
        "preview_pdf": _public_asset_path(preview_pdf),
        "preview_png": _public_asset_path(preview_png),
        "validation": validate_template_package(package),
        "is_default": is_default,
    }


def validate_template_package(package: TemplatePackage) -> dict:
    """Check the template has the styles and block markers declared by manifest."""
    issues = []
    try:
        if package.manifest.renderer == "latex":
            text = package.template_path.read_text(encoding="utf-8")
            if not text.strip():
                issues.append(
                    {
                        "type": "empty_template",
                        "severity": "error",
                        "message": "Template content is empty.",
                    }
                )
        else:
            doc = Document(str(package.template_path))
            style_names = {style.name for style in doc.styles}
            for style in package.manifest.styles.values():
                if style not in style_names:
                    issues.append({"type": "missing_style", "message": f"Missing style: {style}"})
            text = _document_text(doc)

        for marker in package.manifest.blocks.values():
            if marker and marker not in text:
                issues.append(
                    {
                        "type": "missing_block",
                        "severity": "error",
                        "message": (
                            f"Missing block marker: {marker}. Add this marker exactly where "
                            "AutoApply should insert the generated content."
                        ),
                    }
                )
    except Exception as exc:
        issues.append({"type": "template_unreadable", "severity": "error", "message": str(exc)})
    return {"ok": not issues, "issues": issues}


def load_template_package(
    document_type: DocumentType,
    template_id: str | None = None,
    *,
    template_root: Path = TEMPLATE_ROOT,
) -> TemplatePackage:
    """Load a template package from templates/<document_type>/<template_id>."""
    template_id = template_id or DEFAULT_TEMPLATE_IDS[document_type]
    package_dir = _template_package_dir(document_type, template_id, template_root)
    manifest_path = package_dir / "manifest.json"

    if not manifest_path.exists():
        raise FileNotFoundError(f"Template manifest not found: {manifest_path}")

    manifest = TemplateManifest.model_validate_json(manifest_path.read_text(encoding="utf-8"))
    template_name = "template.tex" if manifest.renderer == "latex" else "template.docx"
    template_path = package_dir / template_name
    if not template_path.exists():
        raise FileNotFoundError(f"Template file not found: {template_path}")
    if manifest.document_type != document_type:
        raise ValueError(
            f"Template {template_id} is {manifest.document_type}, not {document_type}"
        )
    return TemplatePackage(
        template_id=template_id,
        document_type=document_type,
        directory=package_dir,
        template_path=template_path,
        manifest_path=manifest_path,
        manifest=manifest,
    )


def default_manifest(document_type: Literal["resume", "cover_letter"]) -> TemplateManifest:
    template_id = DEFAULT_TEMPLATE_IDS[document_type]
    return TemplateManifest.model_validate(_default_manifest(document_type, template_id))


def _default_manifest(document_type: str, template_id: str) -> dict:
    if document_type == "cover_letter":
        return {
            "template_id": template_id,
            "document_type": "cover_letter",
            "template_format": "docx",
            "renderer": "docx",
            "supported_outputs": ["docx", "pdf"],
            "name": "Classic Cover Letter",
            "description": "Simple one-page cover letter with editable Word styles.",
            "page": {"size": "letter", "max_pages": 1},
            "styles": {
                "header": "CoverLetter.Header",
                "date": "CoverLetter.Date",
                "recipient": "CoverLetter.Recipient",
                "body": "CoverLetter.Body",
                "signature": "CoverLetter.Signature",
            },
            "sections": {
                "header": {"enabled": True},
                "body": {"enabled": True, "max_items": 5},
            },
            "section_order": ["header", "recipient", "body", "signature"],
            "capacity": {"max_pages": 1, "max_sections": 4},
            "blocks": {"body": "{{cover_letter.body}}"},
        }

    return {
        "template_id": template_id,
        "document_type": "resume",
        "template_format": "docx",
        "renderer": "docx",
        "supported_outputs": ["docx", "pdf"],
        "name": "ATS Single Column",
        "description": "One-page ATS-friendly resume with named Word styles and tab-stop dates.",
        "page": {"size": "letter", "max_pages": 1},
        "styles": {
            "name": "Resume.Name",
            "contact": "Resume.Contact",
            "section_heading": "Resume.SectionHeading",
            "item_title": "Resume.ItemTitle",
            "item_subtitle": "Resume.ItemSubtitle",
            "item_meta": "Resume.ItemMeta",
            "normal": "Resume.Normal",
            "bullet": "Resume.Bullet",
            "skill_category": "Resume.SkillCategory",
            "skill_line": "Resume.SkillLine",
        },
        "sections": {
            "education": {"enabled": True, "max_items": 2},
            "experience": {
                "enabled": True,
                "max_items": 3,
                "max_bullets_per_item": 4,
                "max_words_per_bullet": 24,
            },
            "projects": {
                "enabled": True,
                "max_items": 3,
                "max_bullets_per_item": 3,
                "max_words_per_bullet": 22,
            },
            "skills": {"enabled": True, "max_lines": 4},
        },
        "section_order": ["header", "education", "skills", "projects", "experience"],
        "capacity": {
            "max_pages": 1,
            "max_sections": 5,
            "max_experience_items": 3,
            "max_project_items": 3,
            "max_bullets_total": 13,
            "max_words_per_bullet": 24,
            "max_skill_lines": 4,
        },
        "blocks": {"sections": "{{resume.sections}}"},
    }


def _default_latex_manifest(document_type: str, template_id: str) -> dict:
    manifest = _default_manifest(document_type, template_id)
    manifest["template_format"] = "latex"
    manifest["renderer"] = "latex"
    manifest["supported_outputs"] = ["tex", "pdf"]
    manifest["styles"] = {}
    if document_type == "cover_letter":
        manifest["name"] = "LaTeX Cover Letter"
        manifest["description"] = "Single-file LaTeX cover letter template."
    else:
        manifest["name"] = "LaTeX Resume"
        manifest["description"] = "Single-file LaTeX resume template."
    return manifest


def _default_latex_template(document_type: str) -> str:
    if document_type == "cover_letter":
        return r"""\documentclass[11pt,letterpaper]{article}
\usepackage[margin=0.8in]{geometry}
\usepackage[hidelinks]{hyperref}
\setlength{\parindent}{0pt}
\setlength{\parskip}{8pt}

\begin{document}

\begin{flushright}
{\large\textbf{ {{applicant.name}} }}\\
{{applicant.contact}}
\end{flushright}

{{date}}

{{recipient.block}}

Dear Hiring Manager,

{{cover_letter.body}}

Sincerely,\\[1.5em]
{{signature}}

Enclosure

\end{document}
"""

    return r"""\documentclass[10pt,letterpaper]{article}
\usepackage[margin=0.65in]{geometry}
\usepackage[hidelinks]{hyperref}
\setlength{\parindent}{0pt}
\setlength{\parskip}{4pt}

\begin{document}

{\LARGE\textbf{ {{full_name}} }}\\
{{contact}}\\
{{links}}

{{resume.sections}}

\end{document}
"""


def _create_default_resume_template(path: Path) -> None:
    doc = Document()
    _set_page(doc, "resume")
    _ensure_default_styles(doc, "resume")
    doc.add_paragraph("{{full_name}}", style="Resume.Name")
    doc.add_paragraph("{{contact}}", style="Resume.Contact")
    doc.add_paragraph("{{links}}", style="Resume.Contact")
    doc.add_paragraph("{{resume.sections}}", style="Resume.Normal")
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def _create_default_cover_letter_template(path: Path) -> None:
    doc = Document()
    _set_page(doc, "cover_letter")
    _ensure_default_styles(doc, "cover_letter")
    doc.add_paragraph("{{applicant.name}}", style="CoverLetter.Header")
    doc.add_paragraph("{{applicant.contact}}", style="CoverLetter.Header")
    doc.add_paragraph("{{date}}", style="CoverLetter.Date")
    doc.add_paragraph("{{recipient.block}}", style="CoverLetter.Recipient")
    doc.add_paragraph("Dear Hiring Manager,", style="CoverLetter.Body")
    doc.add_paragraph("{{cover_letter.body}}", style="CoverLetter.Body")
    doc.add_paragraph("Sincerely,", style="CoverLetter.Signature")
    doc.add_paragraph("{{signature}}", style="CoverLetter.Signature")
    doc.add_paragraph("Enclosure", style="CoverLetter.Signature")
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def _ensure_style(
    doc: Document,
    name: str,
    *,
    size: int,
    bold: bool = False,
    base: str | None = None,
    space_before: int = 0,
    space_after: int = 0,
    font: str = "Arial",
    line_spacing: float | None = None,
    right_tab: bool = False,
    alignment: int | None = None,
) -> None:
    try:
        style = doc.styles[name]
    except KeyError:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    if base:
        try:
            style.base_style = doc.styles[base]
        except KeyError:
            pass
    style.font.name = font
    style.font.size = Pt(size)
    style.font.bold = bold
    style.paragraph_format.space_before = Pt(space_before)
    style.paragraph_format.space_after = Pt(space_after)
    if line_spacing is not None:
        style.paragraph_format.line_spacing = line_spacing
    if right_tab:
        style.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), WD_TAB_ALIGNMENT.RIGHT)
    if alignment is not None:
        style.paragraph_format.alignment = alignment


def _ensure_default_styles(doc: Document, document_type: str) -> None:
    if document_type == "resume":
        _ensure_style(doc, "Resume.Name", size=16, bold=True)
        _ensure_style(doc, "Resume.Contact", size=9)
        _ensure_style(doc, "Resume.SectionHeading", size=11, bold=True, space_before=6)
        _ensure_style(doc, "Resume.ItemTitle", size=10, bold=True, right_tab=True)
        _ensure_style(doc, "Resume.ItemSubtitle", size=9)
        _ensure_style(doc, "Resume.ItemMeta", size=9)
        _ensure_style(doc, "Resume.Normal", size=9)
        _ensure_style(doc, "Resume.SkillCategory", size=9, bold=True)
        _ensure_style(doc, "Resume.SkillLine", size=9)
        _ensure_style(doc, "Resume.Bullet", size=9, base="List Bullet")
        return
    cover_font = "Times New Roman"
    _ensure_style(
        doc,
        "CoverLetter.Header",
        size=11,
        bold=True,
        font=cover_font,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
    )
    _ensure_style(doc, "CoverLetter.Date", size=11, font=cover_font, space_before=10)
    _ensure_style(doc, "CoverLetter.Recipient", size=11, font=cover_font, space_after=10)
    _ensure_style(
        doc, "CoverLetter.Body", size=11, font=cover_font, space_after=7, line_spacing=1.05
    )
    _ensure_style(doc, "CoverLetter.Signature", size=11, font=cover_font, space_before=7)


def _set_page(doc: Document, document_type: str = "resume") -> None:
    section = doc.sections[0]
    if document_type == "cover_letter":
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        return
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)


def _write_sample_assets(package: TemplatePackage) -> None:
    package_dir = package.directory
    sample_name = (
        "sample_resume.json" if package.document_type == "resume" else "sample_cover_letter.json"
    )
    sample_path = package_dir / sample_name
    if not sample_path.exists():
        sample_path.write_text("{}\n", encoding="utf-8", newline="\n")
    style_lock = package_dir / "style.lock.json"
    style_lock.write_text(
        json.dumps(
            {
                "template_id": package.template_id,
                "document_type": package.document_type,
                "styles": package.manifest.styles,
                "blocks": package.manifest.blocks,
            },
            indent=2,
        )
        + "\n",
        encoding="utf-8",
        newline="\n",
    )


def _ensure_required_markers(package: TemplatePackage) -> None:
    if not package.manifest.blocks:
        return
    if package.manifest.renderer == "latex":
        return
    try:
        doc = Document(str(package.template_path))
    except Exception:
        return
    if (
        package.document_type == "cover_letter"
        and package.template_id == DEFAULT_TEMPLATE_IDS["cover_letter"]
    ):
        _set_page(doc, "cover_letter")
        _ensure_default_styles(doc, "cover_letter")
    # Re-apply persisted user style edits after the idempotent default
    # pass above. Without this, ``ensure_template_package`` (called from
    # every regenerate path) would silently undo the Template Library's
    # Edit Styles changes on the next render.
    if package.manifest.style_overrides:
        _apply_style_overrides(
            doc, package.document_type, package.manifest.style_overrides
        )
    changed = False
    is_default_cover_letter = (
        package.document_type == "cover_letter"
        and package.template_id == DEFAULT_TEMPLATE_IDS["cover_letter"]
    )
    if is_default_cover_letter:
        changed = _ensure_default_cover_letter_layout(doc, package.manifest.styles) or changed
        changed = True
    text = _document_text(doc)
    if package.document_type == "resume":
        style = package.manifest.styles.get("normal")
    else:
        style = package.manifest.styles.get("body")
    for marker in package.manifest.blocks.values():
        if marker and marker not in text:
            _add_marker_paragraph(doc, marker, style)
            changed = True
    if package.document_type == "cover_letter":
        changed = _ensure_cover_letter_frame(
            doc,
            package.manifest.styles,
            repair_legacy_placeholders=is_default_cover_letter,
        ) or changed
    if changed:
        doc.save(str(package.template_path))


def _ensure_default_cover_letter_layout(doc: Document, styles: dict[str, str]) -> bool:
    expected = [
        "{{applicant.name}}",
        "{{applicant.contact}}",
        "{{date}}",
        "{{recipient.block}}",
        "Dear Hiring Manager,",
        "{{cover_letter.body}}",
        "Sincerely,",
        "{{signature}}",
        "Enclosure",
    ]
    texts = [(paragraph.text or "").strip() for paragraph in doc.paragraphs]
    if texts == expected:
        return False

    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)

    frame = [
        ("{{applicant.name}}", styles.get("header")),
        ("{{applicant.contact}}", styles.get("header")),
        ("{{date}}", styles.get("date")),
        ("{{recipient.block}}", styles.get("recipient")),
        ("Dear Hiring Manager,", styles.get("body")),
        ("{{cover_letter.body}}", styles.get("body")),
        ("Sincerely,", styles.get("signature")),
        ("{{signature}}", styles.get("signature")),
        ("Enclosure", styles.get("signature")),
    ]
    for text, style in frame:
        _add_marker_paragraph(doc, text, style)
    return True


def _ensure_cover_letter_frame(
    doc: Document,
    styles: dict[str, str],
    *,
    repair_legacy_placeholders: bool = False,
) -> bool:
    """Repair old default cover-letter templates that only had body text.

    Existing ``classic_v1`` packages were created before the renderer
    guaranteed a salutation and sign-off. Add those lines around the
    body marker so generated DOCX/PDF artifacts are complete letters.
    """
    paragraphs = list(doc.paragraphs)
    texts = [(paragraph.text or "").strip() for paragraph in paragraphs]
    marker_idx = next(
        (idx for idx, text in enumerate(texts) if "{{cover_letter.body}}" in text),
        None,
    )
    if marker_idx is None:
        return False

    changed = False
    if repair_legacy_placeholders:
        for paragraph in paragraphs:
            if (paragraph.text or "").strip() == "{{recipient.company}}":
                _replace_template_paragraph_text(paragraph, "{{recipient.block}}")
                changed = True

    has_date = any("{{date}}" in text for text in texts)
    if not has_date:
        insert_before = paragraphs[marker_idx]
        company_idx = next(
            (idx for idx, text in enumerate(texts) if "{{recipient.company}}" in text),
            None,
        )
        if company_idx is not None:
            insert_before = paragraphs[company_idx]
        _insert_template_paragraph_before(
            insert_before,
            "{{date}}",
            styles.get("date"),
        )
        changed = True

    has_salutation = any(text.lower().startswith("dear ") for text in texts)
    if not has_salutation:
        _insert_template_paragraph_before(
            paragraphs[marker_idx],
            "Dear Hiring Manager,",
            styles.get("body"),
        )
        changed = True

    closing_para = _find_cover_letter_closing_paragraph(doc)
    if closing_para is None:
        marker_para = next(
            paragraph
            for paragraph in doc.paragraphs
            if "{{cover_letter.body}}" in (paragraph.text or "")
        )
        closing_para = _insert_template_paragraph_after(
            marker_para,
            "Sincerely,",
            styles.get("signature"),
        )
        changed = True
    signature_para = next(
        (paragraph for paragraph in doc.paragraphs if "{{signature}}" in (paragraph.text or "")),
        None,
    )
    if signature_para is None:
        signature_para = _insert_template_paragraph_after(
            closing_para,
            "{{signature}}",
            styles.get("signature"),
        )
        changed = True
    has_enclosure = any(
        (paragraph.text or "").strip().lower() == "enclosure"
        for paragraph in doc.paragraphs
    )
    if not has_enclosure:
        _insert_template_paragraph_after(
            signature_para,
            "Enclosure",
            styles.get("signature"),
        )
        changed = True
    return changed


def _find_cover_letter_closing_paragraph(doc: Document):
    closings = {"sincerely", "best regards", "kind regards"}
    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip().lower().rstrip(",")
        if text in closings:
            return paragraph
    return None


def _replace_template_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
        return
    paragraph.add_run(text)


def _insert_template_paragraph_before(paragraph, text: str, style: str | None) -> None:
    from docx.oxml import OxmlElement  # noqa: PLC0415
    from docx.text.paragraph import Paragraph  # noqa: PLC0415

    new_element = OxmlElement("w:p")
    paragraph._p.addprevious(new_element)
    new_paragraph = Paragraph(new_element, paragraph._parent)
    _style_template_paragraph(new_paragraph, style)
    new_paragraph.add_run(text)


def _insert_template_paragraph_after(paragraph, text: str, style: str | None):
    from docx.oxml import OxmlElement  # noqa: PLC0415
    from docx.text.paragraph import Paragraph  # noqa: PLC0415

    new_element = OxmlElement("w:p")
    paragraph._p.addnext(new_element)
    new_paragraph = Paragraph(new_element, paragraph._parent)
    _style_template_paragraph(new_paragraph, style)
    new_paragraph.add_run(text)
    return new_paragraph


def _style_template_paragraph(paragraph, style: str | None) -> None:
    if not style:
        return
    try:
        paragraph.style = style
    except KeyError:
        pass


def _add_marker_paragraph(doc: Document, marker: str, style: str | None) -> None:
    if style:
        try:
            doc.add_paragraph(marker, style=style)
            return
        except KeyError:
            pass
    doc.add_paragraph(marker)


def _document_text(doc: Document) -> str:
    parts = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(parts)


def _slugify(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", "_", value.lower()).strip("_")[:60] or "template"


def _template_package_dir(document_type: str, template_id: str, template_root: Path) -> Path:
    if not _TEMPLATE_ID_RE.match(template_id):
        raise ValueError("Invalid template id.")

    type_root = (template_root / document_type).resolve()
    package_dir = (type_root / template_id).resolve()
    try:
        package_dir.relative_to(type_root)
    except ValueError as exc:
        raise ValueError("Invalid template id.") from exc
    return package_dir


def _public_asset_path(path: Path) -> str | None:
    if not path.exists():
        return None
    try:
        return path.resolve().relative_to(PROJECT_ROOT).as_posix()
    except ValueError:
        return path.name


def _unique_template_id(root: Path, template_id: str) -> str:
    candidate = template_id
    index = 2
    while (root / candidate).exists():
        candidate = f"{template_id}_{index}"
        index += 1
    return candidate
