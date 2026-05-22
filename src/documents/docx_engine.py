"""Word document creation and editing engine.

Builds tailored resumes from a base template using block-based assembly.
The template uses {{PLACEHOLDER}} syntax for variable substitution.
Section blocks (experience, education, skills, projects) are rebuilt from
the applicant's bullet pool selection.
"""

from __future__ import annotations

import logging
import re
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

from src.documents.templates import TemplateManifest, default_manifest

logger = logging.getLogger("autoapply.documents.docx_engine")

PLACEHOLDER_RE = re.compile(r"\{\{([\w.]+)\}\}")


# ---------------------------------------------------------------------------
# Core document operations
# ---------------------------------------------------------------------------


def substitute_placeholders(doc: Document, variables: dict[str, str]) -> None:
    """Replace {{KEY}} placeholders in all paragraphs and table cells."""
    for para in _iter_paragraphs(doc):
        _substitute_para(para, variables)


def _iter_paragraphs(doc: Document):
    """Yield all paragraphs including those inside tables."""
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def _substitute_para(para, variables: dict[str, str]) -> None:
    """Substitute placeholders within a paragraph, preserving run formatting."""
    # Consolidate paragraph text first for matching
    full_text = "".join(run.text for run in para.runs)
    if "{{" not in full_text:
        return

    # Replace in the full text
    def replacer(m: re.Match) -> str:
        key = m.group(1)
        return variables.get(key, m.group(0))

    new_text = PLACEHOLDER_RE.sub(replacer, full_text)

    if new_text == full_text:
        return

    # Write back to first run, clear the rest
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ""


# ---------------------------------------------------------------------------
# Block-based resume builder
# ---------------------------------------------------------------------------


def build_resume(
    template_path: Path,
    identity: dict[str, Any],
    education: list[dict],
    experiences: list[dict],
    projects: list[dict],
    skills: dict[str, list[str]],
    selected_bullets: dict[str, list[str]],  # entity_name -> [bullet texts]
    output_path: Path,
) -> Path:
    """Build a tailored resume from a template and structured data.

    Args:
        template_path: Base .docx template.
        identity: Identity section from profile.
        education: List of education records.
        experiences: List of work experience records.
        projects: List of project records.
        skills: Skills dict by category.
        selected_bullets: Maps entity (company/project name) to selected bullet texts.
        output_path: Where to save the generated .docx.

    Returns:
        Path to the generated .docx file.
    """
    doc = Document(str(template_path))

    # Step 1: Substitute header placeholders
    variables = _build_header_variables(identity)
    substitute_placeholders(doc, variables)

    # Step 2: Rebuild section blocks
    _rebuild_education_block(doc, education)
    _rebuild_experience_block(doc, experiences, selected_bullets)
    _rebuild_projects_block(doc, projects, selected_bullets)
    _rebuild_skills_block(doc, skills)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info("Saved resume to %s", output_path)
    return output_path


def build_resume_from_ir(
    template_path: Path,
    document,
    output_path: Path,
    manifest: TemplateManifest | None = None,
) -> Path:
    """Build a DOCX resume from a validated ResumeDocument IR.

    This keeps the current DOCX-first renderer while decoupling content planning
    from template rendering.
    """
    manifest = manifest or default_manifest("resume")
    styles = manifest.styles
    # Activate the manifest's optional "emphasis font" (the second font
    # the user picked in Template Library) for the duration of this
    # render. ``_add_styled_paragraph`` reads it when emitting bold
    # runs so a Markdown ``**FastAPI**`` highlight can use a different
    # face from the body text. Cleared in ``finally`` so a render fault
    # never leaks a font setting into the next call.
    _set_active_emphasis_font(getattr(manifest, "emphasis_font", None))
    try:
        doc = Document(str(template_path))
        substitute_placeholders(doc, _resume_template_variables(document))

        if not _render_resume_markers(doc, document, styles, manifest):
            _clear_document_body(doc)
            _render_resume_sections(_DocxSink(doc), document, styles, include_header=True)

        _remove_empty_placeholder_paragraphs(doc)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        logger.info("Saved IR resume to %s", output_path)
        return output_path
    finally:
        _set_active_emphasis_font(None)


def _render_resume_markers(
    doc: Document,
    document,
    styles: dict[str, str],
    manifest: TemplateManifest,
) -> bool:
    rendered = False
    sections_marker = manifest.blocks.get("sections", "{{resume.sections}}")
    marker_para = _find_marker_paragraph(doc, sections_marker)
    if marker_para is not None:
        sink = _DocxSink(doc, marker_para)
        _render_resume_sections(sink, document, styles, include_header=False)
        _remove_paragraph(marker_para)
        rendered = True

    section_markers = {
        "header": manifest.blocks.get("header", "{{resume.header}}"),
        "education": manifest.blocks.get("education", "{{resume.education}}"),
        "skills": manifest.blocks.get("skills", "{{resume.skills}}"),
        "experience": manifest.blocks.get("experience", "{{resume.experience}}"),
        "projects": manifest.blocks.get("projects", "{{resume.projects}}"),
    }
    # Any legacy ``{{resume.summary}}`` marker in a user template is
    # stripped so it cannot leak through as visible placeholder text.
    legacy_summary_marker = manifest.blocks.get("summary", "{{resume.summary}}")
    legacy_summary_para = _find_marker_paragraph(doc, legacy_summary_marker)
    if legacy_summary_para is not None:
        _remove_paragraph(legacy_summary_para)
        rendered = True
    for section, marker in section_markers.items():
        marker_para = _find_marker_paragraph(doc, marker)
        if marker_para is None:
            continue
        sink = _DocxSink(doc, marker_para)
        _render_resume_section(section, sink, document, styles)
        _remove_paragraph(marker_para)
        rendered = True
    return rendered


def _render_resume_sections(
    sink,
    document,
    styles: dict[str, str],
    *,
    include_header: bool,
) -> None:
    # Build the divider lookup once -- match tokens are normalised so
    # both ``"experience"`` and ``"custom:VOLUNTEER EXPERIENCE"`` work.
    dividers_after = _normalise_divider_set(getattr(document, "dividers_after", None))
    rendered = set()
    rendered_custom_titles: set[str] = set()
    for section in _resolved_section_order(document):
        if section == "header" and not include_header:
            continue
        if section in rendered:
            continue
        rendered.add(section)
        if section.startswith("custom:"):
            rendered_custom_titles.add(section.split(":", 1)[1].strip().lower())
        elif section in ("custom", "custom_sections"):
            for custom in getattr(document, "custom_sections", []) or []:
                rendered_custom_titles.add(custom.title.strip().lower())
        _render_resume_section(section, sink, document, styles)
        if _section_wants_divider(section, dividers_after):
            _add_horizontal_rule(sink)

    # Catch-all: any CustomSection that wasn't explicitly placed by
    # ``section_order`` gets rendered at the end. This is what lets the
    # candidate's VOLUNTEER / AWARDS / etc. survive without forcing the
    # user to hand-tune section ordering -- the canonical sections still
    # come first, the rest land below them in profile order.
    for custom in getattr(document, "custom_sections", []) or []:
        if custom.title.strip().lower() in rendered_custom_titles:
            continue
        _render_ir_custom_section(sink, custom, styles)
        rendered_custom_titles.add(custom.title.strip().lower())
        token = f"custom:{custom.title}"
        if _section_wants_divider(token, dividers_after):
            _add_horizontal_rule(sink)


def _normalise_divider_set(raw) -> set[str]:
    """Normalise ``dividers_after`` into a comparable lowercase set."""
    if not raw:
        return set()
    out: set[str] = set()
    for value in raw:
        token = str(value or "").strip().lower()
        if not token:
            continue
        out.add(token)
        # Accept the "custom:" prefix form as well as the bare title.
        if token.startswith("custom:"):
            out.add(token.split(":", 1)[1].strip())
    return out


def _section_wants_divider(section: str, dividers_after: set[str]) -> bool:
    """True if a horizontal rule should follow ``section``."""
    if not dividers_after:
        return False
    section_lower = (section or "").strip().lower()
    if section_lower in dividers_after:
        return True
    if section_lower.startswith("custom:"):
        title = section_lower.split(":", 1)[1].strip()
        if title in dividers_after:
            return True
    return False


def _render_resume_section(section: str, sink, document, styles: dict[str, str]) -> None:
    if section == "header":
        _render_ir_header(sink, document.header, styles)
    elif section == "education":
        _render_ir_education(sink, document.education, styles)
    elif section == "skills":
        _render_ir_skills(sink, document.skills, styles)
    elif section == "experience":
        _render_ir_experience(sink, document.experiences, styles)
    elif section == "projects":
        _render_ir_projects(sink, document.projects, styles)
    elif section == "custom" or section == "custom_sections":
        # Render every custom section the IR carries. Use this catch-all
        # token in ``section_order`` to control where the block lands;
        # otherwise the caller in ``_render_resume_sections`` will
        # append them after the named sections.
        for custom in getattr(document, "custom_sections", []) or []:
            _render_ir_custom_section(sink, custom, styles)
    elif section.startswith("custom:"):
        # ``custom:VOLUNTEER EXPERIENCE`` in section_order pins a single
        # named custom section at a specific position.
        title = section.split(":", 1)[1].strip()
        for custom in getattr(document, "custom_sections", []) or []:
            if custom.title.strip().lower() == title.lower():
                _render_ir_custom_section(sink, custom, styles)
                break


def _render_ir_custom_section(sink, custom, styles: dict[str, str]) -> None:
    """Render one CustomSection block onto the sink.

    Uses the same Resume.SectionHeading + Resume.ItemTitle / Subtitle
    / Bullet styles as the canonical sections so a free-form section
    visually matches the rest of the resume. Empty fields are skipped
    (no "Organization: None" leaks).
    """
    entries = [
        entry
        for entry in getattr(custom, "entries", []) or []
        if (entry.title or entry.organization or entry.details or entry.bullets)
    ]
    if not entries:
        return
    _add_section_heading(sink, custom.title, styles)
    for entry in entries:
        primary = _clean_field(entry.title) or _clean_field(entry.organization)
        dates = _format_date_range(entry.start_date, entry.end_date)
        if primary or dates:
            _add_styled_paragraph(
                sink, _join_tab(primary, dates), styles.get("item_title")
            )
        subtitle_parts = []
        if entry.title and entry.organization and primary != entry.organization:
            subtitle_parts.append(entry.organization)
        if entry.location:
            subtitle_parts.append(entry.location)
        subtitle = _join_nonempty(subtitle_parts)
        if subtitle:
            _add_styled_paragraph(sink, subtitle, styles.get("item_subtitle"))
        if entry.details:
            _add_styled_paragraph(sink, entry.details, styles.get("item_meta"))
        for bullet in entry.bullets:
            cleaned = _clean_field(bullet)
            if cleaned:
                _add_resume_bullet(sink, cleaned, styles)


def build_cover_letter_from_ir(
    document,
    output_path: Path,
    *,
    template_path: Path | None = None,
    manifest: TemplateManifest | None = None,
) -> Path:
    """Build a DOCX cover letter from CoverLetterDocument IR."""
    manifest = manifest or default_manifest("cover_letter")
    styles = manifest.styles
    _set_active_emphasis_font(getattr(manifest, "emphasis_font", None))
    try:
        doc = (
            Document(str(template_path))
            if template_path and template_path.exists()
            else Document()
        )
        substitute_placeholders(doc, _cover_letter_template_variables(document))

        marker = manifest.blocks.get("body", "{{cover_letter.body}}")
        marker_para = _find_marker_paragraph(doc, marker)
        if marker_para is not None:
            sink = _DocxSink(doc, marker_para)
            _render_cover_letter_body(sink, document, styles)
            _remove_paragraph(marker_para)
        else:
            _clear_document_body(doc)
            _render_cover_letter_fallback(_DocxSink(doc), document, styles)

        _remove_empty_placeholder_paragraphs(doc)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        logger.info("Saved cover letter to %s", output_path)
        return output_path
    finally:
        _set_active_emphasis_font(None)


def _render_cover_letter_fallback(sink, document, styles: dict[str, str]) -> None:
    applicant = document.applicant or {}
    recipient = document.recipient or {}

    name = applicant.get("name") or applicant.get("full_name") or ""
    if name:
        paragraph = _add_styled_paragraph(sink, str(name), styles.get("header"))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    contact_lines = _cover_letter_contact_lines(applicant)
    if contact_lines:
        paragraph = _add_styled_paragraph(
            sink, "\n".join(contact_lines), styles.get("header")
        )
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    _add_styled_paragraph(sink, _cover_letter_date(), styles.get("date"))
    recipient_lines = _cover_letter_recipient_lines(recipient)
    if recipient_lines:
        _add_styled_paragraph(sink, "\n".join(recipient_lines), styles.get("recipient"))
    _add_styled_paragraph(sink, "Dear Hiring Manager,", styles.get("body"))
    _render_cover_letter_body(sink, document, styles)
    _add_styled_paragraph(sink, "Sincerely,", styles.get("signature"))
    if name:
        _add_styled_paragraph(sink, str(name), styles.get("signature"))
    _add_styled_paragraph(sink, "Enclosure", styles.get("signature"))


def _render_cover_letter_body(sink, document, styles: dict[str, str]) -> None:
    for paragraph in document.paragraphs:
        if paragraph.text:
            _add_styled_paragraph(sink, paragraph.text, styles.get("body"))


def _clear_document_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


class _DocxSink:
    """Append paragraphs to a document or after a template block marker."""

    def __init__(self, doc: Document, anchor: Paragraph | None = None):
        self.doc = doc
        self.anchor = anchor

    def add_paragraph(self, text: str = "", style: str | None = None):
        if self.anchor is None:
            return self.doc.add_paragraph(text, style=style)
        paragraph = _insert_paragraph_after(self.anchor, text, style)
        self.anchor = paragraph
        return paragraph


def _insert_paragraph_after(
    paragraph: Paragraph,
    text: str = "",
    style: str | None = None,
) -> Paragraph:
    new_element = OxmlElement("w:p")
    paragraph._p.addnext(new_element)
    new_paragraph = Paragraph(new_element, paragraph._parent)
    if style:
        try:
            new_paragraph.style = style
        except KeyError:
            pass
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def _find_marker_paragraph(doc: Document, marker: str) -> Paragraph | None:
    for paragraph in _iter_paragraphs(doc):
        if marker in paragraph.text:
            return paragraph
    return None


def _remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def _remove_empty_placeholder_paragraphs(doc: Document) -> None:
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if text in {"{{links}}", "{{applicant.contact}}", "{{signature}}"}:
            _remove_paragraph(paragraph)


def _resume_template_variables(document) -> dict[str, str]:
    identity = document.header or {}
    contact = _join_nonempty(
        [identity.get("email"), identity.get("phone"), identity.get("location")]
    )
    links = _join_nonempty(
        [identity.get("linkedin_url"), identity.get("github_url"), identity.get("portfolio_url")]
    )
    return {
        "full_name": str(identity.get("full_name") or identity.get("name") or ""),
        "contact": contact,
        "links": links,
    }


def _cover_letter_template_variables(document) -> dict[str, str]:
    applicant = document.applicant or {}
    recipient = document.recipient or {}
    name = applicant.get("name") or applicant.get("full_name") or ""
    contact = "\n".join(_cover_letter_contact_lines(applicant))
    recipient_block = "\n".join(_cover_letter_recipient_lines(recipient))
    return {
        "applicant.name": str(name),
        "applicant.contact": contact,
        "date": _cover_letter_date(),
        "recipient.block": recipient_block,
        "recipient.company": str(recipient.get("company") or ""),
        "signature": str(name),
    }


def _cover_letter_date() -> str:
    today = datetime.now()
    return f"{today:%B} {today.day}, {today:%Y}"


def _cover_letter_contact_lines(applicant: dict[str, Any]) -> list[str]:
    return [
        line
        for line in (
            _clean_cover_letter_location(applicant.get("location")),
            _clean_field(applicant.get("phone")),
            _clean_field(applicant.get("email")),
        )
        if line
    ]


def _cover_letter_recipient_lines(recipient: dict[str, Any]) -> list[str]:
    hiring_manager = _clean_field(recipient.get("hiring_manager"))
    company = _clean_field(recipient.get("company"))
    location = _clean_cover_letter_location(recipient.get("location"))
    lines = []
    if hiring_manager:
        lines.append(hiring_manager)
    elif company:
        lines.append("Hiring Team")
    if company:
        lines.append(company)
    if location:
        lines.append(location)
    return lines


def _clean_cover_letter_location(value) -> str:
    text = _clean_field(value)
    if not text:
        return ""
    text = re.sub(r"\s*\([^)]*\)", "", text)
    text = re.sub(
        r"\s*(?:[-–—|,]|\b)\s*(?:hybrid|remote|on-?site|in-person)\b.*$",
        "",
        text,
        flags=re.IGNORECASE,
    )
    return text.strip(" ,-|–—")


def _resolved_section_order(document) -> list[str]:
    # The caller's ``section_order`` is the source of truth. Only fall
    # back to the default when the caller did not provide an order.
    # ``summary`` is filtered out unconditionally -- this system never
    # renders a Summary section regardless of what callers / legacy
    # manifests request.
    default_order = ["header", "education", "skills", "experience", "projects"]
    explicit = [
        section
        for section in document.section_order
        if section in default_order and section != "summary"
    ]
    return explicit or default_order


def _render_ir_header(doc: Document, identity: dict[str, Any], styles: dict[str, str]) -> None:
    name = identity.get("full_name") or identity.get("name") or ""
    if name:
        _add_styled_paragraph(doc, str(name), styles.get("name"), fallback="Title")

    contact = _join_nonempty(
        [identity.get("email"), identity.get("phone"), identity.get("location")]
    )
    if contact:
        _add_styled_paragraph(doc, contact, styles.get("contact"))

    links = _join_nonempty(
        [identity.get("linkedin_url"), identity.get("github_url"), identity.get("portfolio_url")]
    )
    if links:
        _add_styled_paragraph(doc, links, styles.get("contact"))


def _render_ir_education(doc: Document, education: list[dict], styles: dict[str, str]) -> None:
    if not education:
        return
    _add_section_heading(doc, "Education", styles)
    for edu in education:
        institution = edu.get("institution", "")
        dates = _format_date_range(edu.get("start_date", ""), edu.get("end_date", ""))
        _add_styled_paragraph(doc, _join_tab(institution, dates), styles.get("item_title"))
        degree = " ".join(part for part in [edu.get("degree", ""), edu.get("field", "")] if part)
        gpa = _gpa_label(edu.get("gpa"))
        details = _join_nonempty([degree, edu.get("location"), gpa])
        if details:
            _add_styled_paragraph(doc, details, styles.get("item_meta"))
        courses = edu.get("relevant_courses", [])
        if courses:
            course_names = ", ".join(c.get("name", "") for c in courses if isinstance(c, dict))
            if course_names:
                _add_styled_paragraph(
                    doc,
                    f"Relevant coursework: {course_names}",
                    styles.get("normal"),
                )


def _render_ir_skills(doc: Document, skills: dict[str, list[str]], styles: dict[str, str]) -> None:
    rows = [(label, skills.get(key, [])) for key, label in _skill_label_map().items()]
    rows.extend(
        (key.replace("_", " ").title(), values)
        for key, values in skills.items()
        if key not in _skill_label_map()
    )
    rows = [(label, values) for label, values in rows if values]
    if not rows:
        return
    _add_section_heading(doc, "Skills", styles)
    for label, values in rows:
        _add_styled_paragraph(doc, f"{label}: {', '.join(values)}", styles.get("skill_line"))


def _render_ir_experience(doc: Document, experiences: list, styles: dict[str, str]) -> None:
    if not experiences:
        return
    _add_section_heading(doc, "Experience", styles)
    for item in experiences:
        org = item.organization or item.name
        dates = _format_date_range(item.start_date, item.end_date)
        _add_styled_paragraph(doc, _join_tab(item.title or org, dates), styles.get("item_title"))
        subtitle = _join_nonempty([org, item.location])
        if subtitle:
            _add_styled_paragraph(doc, subtitle, styles.get("item_subtitle"))
        for bullet in item.bullets:
            _add_resume_bullet(doc, bullet.text, styles)


def _render_ir_projects(doc: Document, projects: list, styles: dict[str, str]) -> None:
    if not projects:
        return
    _add_section_heading(doc, "Projects", styles)
    for item in projects:
        dates = _format_date_range(item.start_date, item.end_date)
        tech = ", ".join(item.tech_stack)
        _add_styled_paragraph(doc, _join_tab(item.name, dates), styles.get("item_title"))
        if tech:
            _add_styled_paragraph(doc, tech, styles.get("item_subtitle"))
        if item.meta:
            _add_styled_paragraph(doc, item.meta, styles.get("item_meta"))
        for bullet in item.bullets:
            _add_resume_bullet(doc, bullet.text, styles)


def _add_section_heading(doc: Document, title: str, styles: dict[str, str]) -> None:
    _add_styled_paragraph(doc, title, styles.get("section_heading"), fallback="Heading 2")


# Module-level "current emphasis font" -- set by ``build_resume_from_ir`` /
# ``build_cover_letter_from_ir`` when the active manifest declares one,
# read inside ``_add_styled_paragraph`` to swap the font on bold/italic
# runs. We keep it as module state rather than threading a parameter
# through every sub-renderer because the call graph is wide
# (header / education / skills / experience / projects / custom) and
# every call site would otherwise need a new kwarg. The renderer is
# always invoked from a single Python thread per render, so module
# state is safe here. Always cleared in a ``finally`` so a render that
# raises does not leak the font into the next call.
_ACTIVE_EMPHASIS_FONT: str | None = None


def _set_active_emphasis_font(font: str | None) -> None:
    global _ACTIVE_EMPHASIS_FONT
    _ACTIVE_EMPHASIS_FONT = (font or "").strip() or None


def _add_styled_paragraph(
    doc: Document,
    text: str,
    style_name: str | None,
    *,
    fallback: str | None = None,
):
    """Append a paragraph honouring inline ``**bold**`` / ``*italic*``
    markers from the LLM.

    Backwards compatible: ``text`` with no markers renders identically
    to the previous single-run implementation. When markers are present
    the paragraph is built run by run so Word shows the formatting and
    the style's named font is preserved on every run. A configured
    ``manifest.emphasis_font`` (the "second font" the user picked in
    Template Library) overrides the font for bold runs only -- italics
    keep the body font so the contrast is purely structural rather than
    typographic noise.
    """
    from src.generation.inline_format import (  # noqa: PLC0415
        is_divider_paragraph,
        parse_inline_markup,
    )

    if is_divider_paragraph(text):
        return _add_horizontal_rule(doc)

    style_candidates = [style for style in (style_name, fallback) if style]
    paragraph = None
    for style in style_candidates:
        try:
            paragraph = doc.add_paragraph("", style=style)
            break
        except KeyError:
            continue
    if paragraph is None:
        paragraph = doc.add_paragraph("")

    runs = parse_inline_markup(text)
    if not runs or (len(runs) == 1 and not runs[0].bold and not runs[0].italic):
        # Fast path -- a single plain run is the common case; assign
        # directly so existing behaviour (and downstream
        # ``paragraph.runs[0].text`` access) keeps working.
        plain = runs[0].text if runs else ""
        if plain:
            paragraph.add_run(plain)
        return paragraph

    emphasis_font = _ACTIVE_EMPHASIS_FONT
    for run in runs:
        if not run.text:
            continue
        word_run = paragraph.add_run(run.text)
        if run.bold:
            word_run.bold = True
            if emphasis_font:
                word_run.font.name = emphasis_font
        if run.italic:
            word_run.italic = True
    return paragraph


def _add_horizontal_rule(doc) -> None:
    """Insert a horizontal rule paragraph.

    Implemented as an empty paragraph with a bottom border so it shows
    up the same in Word, LibreOffice, and the PDF converter. We avoid
    actual ``<w:hr/>`` because that element is not part of the OOXML
    standard -- the bottom-border trick is the canonical workaround.
    """
    from docx.oxml import OxmlElement  # noqa: PLC0415
    from docx.oxml.ns import qn  # noqa: PLC0415

    paragraph = doc.add_paragraph("")
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")  # 1/8 pt units -> ~0.75pt rule
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "808080")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    return paragraph


def _add_resume_bullet(doc: Document, text: str, styles: dict[str, str]) -> None:
    _add_styled_paragraph(doc, text, styles.get("bullet"), fallback="List Bullet")


def _add_bold_paragraph(doc: Document, text: str) -> None:
    if not text:
        return
    paragraph = doc.add_paragraph()
    paragraph.add_run(text).bold = True


def _add_bullet(doc: Document, text: str) -> None:
    try:
        doc.add_paragraph(text, style="List Bullet")
    except KeyError:
        doc.add_paragraph(f"• {text}")


def _format_date_range(start: str, end: str) -> str:
    if start and end:
        return f"{start} – {end}"
    return start or end or ""


def _gpa_label(value) -> str:
    """Return ``"GPA: 3.8"``-style label, or empty string when unset.

    A blank profile field comes through as ``None`` / ``""`` / ``0``,
    none of which belong on a resume. Centralising the formatting here
    means downstream renderers cannot accidentally print "GPA: None"
    by interpolating directly into an f-string.
    """
    text = _clean_field(value)
    if not text:
        return ""
    if text.lower() in {"none", "n/a", "na", "null", "0", "0.0", "0.00"}:
        return ""
    return f"GPA: {text}"


def _clean_field(value) -> str:
    """Return a trimmed string, treating None / "None" / "null" as empty.

    Profile fields imported from heterogeneous sources sometimes round-trip
    through ``str(None)`` and arrive at the renderer as the literal text
    "None". This helper keeps every renderer call site one line of code
    away from doing the right thing.
    """
    if value is None:
        return ""
    text = str(value).strip()
    if not text:
        return ""
    if text.lower() in {"none", "null", "n/a", "na", "nan"}:
        return ""
    return text


def _join_nonempty(values: list) -> str:
    return " | ".join(
        cleaned for cleaned in (_clean_field(value) for value in values) if cleaned
    )


def _join_tab(left: str, right: str) -> str:
    if left and right:
        return f"{left}\t{right}"
    return left or right or ""


def _skill_label_map() -> dict[str, str]:
    return {
        "languages": "Languages",
        "frameworks": "Frameworks",
        "databases": "Databases",
        "tools": "Tools & DevOps",
        "domains": "Domains",
        "soft_skills": "Soft Skills",
        "certifications": "Certifications",
    }


def _build_header_variables(identity: dict) -> dict[str, str]:
    return {
        "FULL_NAME": identity.get("full_name", ""),
        "EMAIL": identity.get("email", ""),
        "PHONE": identity.get("phone", ""),
        "LOCATION": identity.get("location", ""),
        "LINKEDIN": identity.get("linkedin_url", ""),
        "GITHUB": identity.get("github_url", ""),
        "PORTFOLIO": identity.get("portfolio_url", ""),
    }


def _find_section_paragraph(doc: Document, marker: str) -> int | None:
    """Find the index of the paragraph containing the section marker."""
    for i, para in enumerate(doc.paragraphs):
        if marker in para.text:
            return i
    return None


def _clear_section_content(doc: Document, start_idx: int, next_markers: list[str]) -> int:
    """Remove paragraphs from start_idx until a next section marker is hit.

    Returns the index where the next section begins.
    """
    i = start_idx + 1
    while i < len(doc.paragraphs):
        text = doc.paragraphs[i].text.strip()
        if any(m in text for m in next_markers):
            break
        p = doc.paragraphs[i]._element
        p.getparent().remove(p)
        # Note: after removal, len(doc.paragraphs) shrinks, but i stays same
    return i


def _add_paragraph_after(
    doc: Document, ref_idx: int, text: str, style: str = "Normal", bold: bool = False
) -> None:
    """Insert a new paragraph after the paragraph at ref_idx."""
    ref_para = doc.paragraphs[ref_idx]
    new_para = OxmlElement("w:p")
    ref_para._element.addnext(new_para)

    # Reload reference after DOM change
    target_para = doc.paragraphs[ref_idx + 1]
    run = target_para.add_run(text)
    run.bold = bold
    try:
        target_para.style = doc.styles[style]
    except KeyError:
        pass


# ---------------------------------------------------------------------------
# Section rebuilders — these operate on known template markers
# ---------------------------------------------------------------------------


def _rebuild_education_block(doc: Document, education: list[dict]) -> None:
    """Replace {{EDUCATION_BLOCK}} section with formatted education entries."""
    idx = _find_section_paragraph(doc, "{{EDUCATION_BLOCK}}")
    if idx is None:
        logger.debug("No {{EDUCATION_BLOCK}} marker found in template")
        return

    # Clear the marker text and any existing sample content below it (P1 fix)
    if doc.paragraphs[idx].runs:
        doc.paragraphs[idx].runs[0].text = ""
    _clear_section_content(
        doc, idx, ["{{EXPERIENCE_BLOCK}}", "{{PROJECTS_BLOCK}}", "{{SKILLS_BLOCK}}"]
    )

    insert_idx = idx
    for edu in education:
        institution = edu.get("institution", "")
        degree = f"{edu.get('degree', '')} in {edu.get('field', '')}"
        dates = f"{edu.get('start_date', '')} – {edu.get('end_date', '')}"
        gpa = f"GPA: {edu['gpa']}" if edu.get("gpa") else ""

        _add_paragraph_after(doc, insert_idx, f"{institution} | {dates}", bold=True)
        insert_idx += 1
        _add_paragraph_after(doc, insert_idx, degree)
        insert_idx += 1
        if gpa:
            _add_paragraph_after(doc, insert_idx, gpa)
            insert_idx += 1

        courses = edu.get("relevant_courses", [])
        if courses:
            course_names = ", ".join(c.get("name", "") for c in courses if isinstance(c, dict))
            _add_paragraph_after(doc, insert_idx, f"Relevant coursework: {course_names}")
            insert_idx += 1


def _rebuild_experience_block(
    doc: Document,
    experiences: list[dict],
    selected_bullets: dict[str, list[str]],
) -> None:
    """Replace {{EXPERIENCE_BLOCK}} with tailored experience entries."""
    idx = _find_section_paragraph(doc, "{{EXPERIENCE_BLOCK}}")
    if idx is None:
        logger.debug("No {{EXPERIENCE_BLOCK}} marker found in template")
        return

    if doc.paragraphs[idx].runs:
        doc.paragraphs[idx].runs[0].text = ""
    _clear_section_content(doc, idx, ["{{PROJECTS_BLOCK}}", "{{SKILLS_BLOCK}}"])

    insert_idx = idx
    for exp in experiences:
        company = exp.get("company", "")
        title = exp.get("title", "")
        dates = f"{exp.get('start_date', '')} – {exp.get('end_date', '')}"
        location = exp.get("location", "")

        _add_paragraph_after(doc, insert_idx, f"{company} | {location}", bold=True)
        insert_idx += 1
        _add_paragraph_after(doc, insert_idx, f"{title} | {dates}")
        insert_idx += 1

        entity_key = f"{company} - {title}"
        bullets = selected_bullets.get(entity_key, [])
        if not bullets:
            # Fall back to all bullets in the experience
            bullets = [b["text"] for b in exp.get("bullets", []) if isinstance(b, dict)]

        for bullet_text in bullets:
            _add_paragraph_after(doc, insert_idx, f"• {bullet_text}")
            insert_idx += 1


def _rebuild_projects_block(
    doc: Document,
    projects: list[dict],
    selected_bullets: dict[str, list[str]],
) -> None:
    """Replace {{PROJECTS_BLOCK}} with tailored project entries."""
    idx = _find_section_paragraph(doc, "{{PROJECTS_BLOCK}}")
    if idx is None:
        logger.debug("No {{PROJECTS_BLOCK}} marker found in template")
        return

    if doc.paragraphs[idx].runs:
        doc.paragraphs[idx].runs[0].text = ""
    _clear_section_content(doc, idx, ["{{SKILLS_BLOCK}}"])

    insert_idx = idx
    for proj in projects:
        name = proj.get("name", "")
        tech = ", ".join(proj.get("tech_stack", []))
        dates = f"{proj.get('start_date', '')} – {proj.get('end_date', '')}".strip(" –")

        header = f"{name} | {tech}"
        if dates:
            header += f" | {dates}"

        _add_paragraph_after(doc, insert_idx, header, bold=True)
        insert_idx += 1

        bullets = selected_bullets.get(name, [])
        if not bullets:
            bullets = [b["text"] for b in proj.get("bullets", []) if isinstance(b, dict)]

        for bullet_text in bullets:
            _add_paragraph_after(doc, insert_idx, f"• {bullet_text}")
            insert_idx += 1


def _rebuild_skills_block(doc: Document, skills: dict[str, list[str]]) -> None:
    """Replace {{SKILLS_BLOCK}} with a formatted skills summary."""
    idx = _find_section_paragraph(doc, "{{SKILLS_BLOCK}}")
    if idx is None:
        logger.debug("No {{SKILLS_BLOCK}} marker found in template")
        return

    if doc.paragraphs[idx].runs:
        doc.paragraphs[idx].runs[0].text = ""
    _clear_section_content(doc, idx, [])  # last section — clear to end of doc

    insert_idx = idx
    label_map = {
        "languages": "Languages",
        "frameworks": "Frameworks",
        "databases": "Databases",
        "tools": "Tools & DevOps",
        "domains": "Domains",
    }
    for key, label in label_map.items():
        items = skills.get(key, [])
        if items:
            _add_paragraph_after(doc, insert_idx, f"{label}: {', '.join(items)}")
            insert_idx += 1


# ---------------------------------------------------------------------------
# Minimal template creator (used when no template .docx exists yet)
# ---------------------------------------------------------------------------


def create_default_template(output_path: Path) -> Path:
    """Create a basic resume template .docx with all expected markers."""
    doc = Document()

    # Header
    doc.add_heading("{{FULL_NAME}}", level=1)
    doc.add_paragraph("{{EMAIL}} | {{PHONE}} | {{LOCATION}}")
    doc.add_paragraph("{{LINKEDIN}} | {{GITHUB}}")

    doc.add_heading("Education", level=2)
    doc.add_paragraph("{{EDUCATION_BLOCK}}")

    doc.add_heading("Experience", level=2)
    doc.add_paragraph("{{EXPERIENCE_BLOCK}}")

    doc.add_heading("Projects", level=2)
    doc.add_paragraph("{{PROJECTS_BLOCK}}")

    doc.add_heading("Skills", level=2)
    doc.add_paragraph("{{SKILLS_BLOCK}}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info("Created default template at %s", output_path)
    return output_path
