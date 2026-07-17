"""Validation checks for generated document IR."""

from __future__ import annotations

import re
from collections import Counter
from math import ceil
from pathlib import Path

from src.documents.page_count import get_docx_page_count, get_pdf_page_count
from src.generation.ir import CoverLetterDocument, ResumeDocument, ValidationIssue, ValidationResult


def validate_resume_document(
    document: ResumeDocument,
    *,
    jd_tags: list[str] | None = None,
    max_bullet_words: int = 32,
    max_estimated_pages: int = 1,
) -> ValidationResult:
    """Run deterministic safety and fit checks before rendering."""
    issues: list[ValidationIssue] = []
    bullets = [
        bullet
        for item in [*document.experiences, *document.projects]
        for bullet in item.bullets
    ]

    if not document.header.get("full_name"):
        issues.append(
            ValidationIssue(
                type="missing_header_name",
                severity="error",
                section="header",
                message="Resume header is missing the applicant name.",
            )
        )

    if not bullets:
        issues.append(
            ValidationIssue(
                type="no_evidence_bullets",
                severity="warning",
                message="No evidence bullets were selected for this resume.",
            )
        )

    for section, items in (("experience", document.experiences), ("projects", document.projects)):
        for item in items:
            if not item.bullets:
                issues.append(
                    ValidationIssue(
                        type="empty_item",
                        severity="info",
                        section=section,
                        item=item.name,
                        source_id=item.source_id,
                        message=f"{item.name} has no selected bullets.",
                    )
                )
            for index, bullet in enumerate(item.bullets):
                word_count = _word_count(bullet.text)
                if word_count > max_bullet_words:
                    issues.append(
                        ValidationIssue(
                            type="bullet_too_long",
                            severity="warning",
                            section=section,
                            item=item.name,
                            source_id=bullet.source_id,
                            message="Bullet exceeds the target word budget.",
                            details={
                                "bullet_index": index,
                                "current_words": word_count,
                                "max_words": max_bullet_words,
                            },
                        )
                    )

                if bullet.original_text:
                    added_numbers = _added_numbers(bullet.original_text, bullet.text)
                    if added_numbers:
                        issues.append(
                            ValidationIssue(
                                type="added_unverified_number",
                                severity="error",
                                section=section,
                                item=item.name,
                                source_id=bullet.source_id,
                                message="Rewritten bullet appears to introduce new numbers.",
                                details={"numbers": added_numbers},
                            )
                        )

    experience_bullet_count = sum(len(item.bullets) for item in document.experiences)
    if document.experiences and experience_bullet_count < 5:
        issues.append(
            ValidationIssue(
                type="thin_professional_experience",
                severity="warning",
                section="experience",
                message=(
                    "Professional experience carries fewer than five evidence bullets; "
                    "do not trade work history for project or skills density."
                ),
                details={"experience_bullet_count": experience_bullet_count},
            )
        )

    repeated_verbs = _repeated_action_verbs([bullet.text for bullet in bullets])
    if repeated_verbs:
        issues.append(
            ValidationIssue(
                type="repeated_action_verbs",
                severity="info",
                message="Several bullets start with the same action verb.",
                details={"verbs": repeated_verbs},
            )
        )

    coverage = _keyword_coverage(document, jd_tags or [])
    if jd_tags and coverage["coverage_ratio"] < 0.35:
        issues.append(
            ValidationIssue(
                type="low_keyword_coverage",
                severity="warning",
                message="Resume covers few target JD keywords.",
                details=coverage,
            )
        )

    estimated_pages = _estimate_pages(document)
    estimated_page_fill = _estimate_resume_page_fill(document)
    if estimated_pages > max_estimated_pages:
        issues.append(
            ValidationIssue(
                type="estimated_page_overflow",
                severity="warning",
                message="Resume may exceed the configured page target.",
                details={"estimated_pages": estimated_pages, "max_pages": max_estimated_pages},
            )
        )

    metrics = {
        "bullet_count": len(bullets),
        "experience_count": len(document.experiences),
        "experience_bullet_count": experience_bullet_count,
        "project_count": len(document.projects),
        "estimated_pages": estimated_pages,
        "estimated_resume_page_fill_ratio": estimated_page_fill,
        "font_family": "Arial",
        "body_font_size_pt": 9,
        "page_margin_inches": {"top": 0.55, "bottom": 0.55, "left": 0.6, "right": 0.6},
        **coverage,
    }
    return ValidationResult(
        ok=not any(issue.severity == "error" for issue in issues),
        issues=issues,
        metrics=metrics,
    )


def validate_resume_artifacts(
    validation: ValidationResult,
    *,
    docx_path: Path | None,
    pdf_path: Path | None = None,
    pdf_attempted: bool = False,
    max_pages: int = 1,
    target_pages: int | None = None,
    document_type: str = "resume",
) -> ValidationResult:
    """Add deterministic renderer/file-system checks to an existing validation result.

    ``target_pages`` (when supplied) drives the strict page-match check
    used by the Template Library's "Expected pages" setting: a render
    that produces more *or* fewer pages than configured raises an error
    so the review queue can demand a regenerate. ``max_pages`` is kept
    around for callers that only care about an upper bound (legacy
    behaviour) and defaults to ``target_pages`` when both are passed.
    """
    target = target_pages if target_pages is not None else max_pages
    issues = list(validation.issues)
    metrics = dict(validation.metrics)

    docx_ok = bool(docx_path and docx_path.exists() and docx_path.stat().st_size > 0)
    pdf_ok = bool(pdf_path and pdf_path.exists() and pdf_path.stat().st_size > 0)
    pdf_pages = get_pdf_page_count(pdf_path)
    docx_pages = get_docx_page_count(docx_path, rendered_pdf_path=pdf_path)
    metrics.update(
        {
            "docx_generated": docx_ok,
            "docx_path": str(docx_path) if docx_path else None,
            "docx_page_count": docx_pages,
            "pdf_generated": pdf_ok,
            "pdf_path": str(pdf_path) if pdf_path else None,
            "pdf_page_count": pdf_pages,
            "target_pages": target,
            **_docx_layout_metrics(docx_path, document_type=document_type),
        }
    )

    if not docx_ok:
        issues.append(
            ValidationIssue(
                type="docx_generation_failed",
                severity="error",
                message="DOCX renderer did not produce a valid file.",
                details={"path": str(docx_path) if docx_path else None},
            )
        )

    if pdf_attempted and not pdf_ok:
        issues.append(
            ValidationIssue(
                type="pdf_generation_failed",
                severity="warning",
                message="PDF conversion did not produce a valid file.",
                details={"path": str(pdf_path) if pdf_path else None},
            )
        )

    rendered_pages = pdf_pages or docx_pages
    if rendered_pages is not None and target_pages is not None:
        if rendered_pages > target:
            issues.append(
                ValidationIssue(
                    type="rendered_page_overflow",
                    severity="error",
                    message=(
                        f"Rendered document is {rendered_pages} pages but template "
                        f"targets exactly {target}. Regenerate with less content."
                    ),
                    details={"page_count": rendered_pages, "target_pages": target},
                )
            )
        elif rendered_pages < target:
            issues.append(
                ValidationIssue(
                    type="rendered_page_underflow",
                    severity="error",
                    message=(
                        f"Rendered document is {rendered_pages} pages but template "
                        f"targets exactly {target}. Regenerate with more content."
                    ),
                    details={"page_count": rendered_pages, "target_pages": target},
                )
            )
    elif rendered_pages is not None and rendered_pages > max_pages:
        # Legacy callers that did not pass target_pages keep the
        # softer warning so we do not change behaviour for them.
        issues.append(
            ValidationIssue(
                type="rendered_page_overflow",
                severity="warning",
                message="Rendered document exceeds the configured page target.",
                details={"page_count": rendered_pages, "max_pages": max_pages},
            )
        )

    return ValidationResult(
        ok=not any(issue.severity == "error" for issue in issues),
        issues=issues,
        metrics=metrics,
    )


def validate_cover_letter_artifacts(
    validation: ValidationResult | None = None,
    *,
    docx_path: Path | None,
    pdf_path: Path | None = None,
    pdf_attempted: bool = False,
    max_pages: int = 1,
    target_pages: int | None = None,
) -> ValidationResult:
    """Validate rendered cover letter files."""
    base = validation or ValidationResult(ok=True, issues=[], metrics={})
    return validate_resume_artifacts(
        base,
        docx_path=docx_path,
        pdf_path=pdf_path,
        pdf_attempted=pdf_attempted,
        max_pages=max_pages,
        target_pages=target_pages,
        document_type="cover_letter",
    )


def validate_cover_letter_document(
    document: CoverLetterDocument,
    *,
    # 2026-07-16: aligned with cover_letter._length_window_for (min 180 /
    # target 240). The old 260-word minimum flagged 20/20 letters in the
    # first real batch — including the good ones — making the warning
    # meaningless. Keep these two thresholds in sync.
    min_words: int = 180,
    max_words: int = 430,
    min_paragraphs: int = 4,
    min_estimated_page_fill: float = 0.45,
    target_pages: int = 1,
) -> ValidationResult:
    """Validate cover-letter content and rough page-fit before rendering."""
    # Single-page cover letter is the canonical case. For 2-page
    # templates expand the word/paragraph budget proportionally so the
    # validator does not flag a deliberately long letter as too long.
    scale = max(1, int(target_pages))
    min_words = min_words * scale
    max_words = max_words * scale
    min_paragraphs = min_paragraphs if scale == 1 else max(min_paragraphs, 4 + (scale - 1) * 2)
    issues: list[ValidationIssue] = []
    paragraphs = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]
    body_text = "\n\n".join(paragraphs)
    word_count = _word_count(body_text)
    paragraph_count = len(paragraphs)
    estimated_page_fill = _estimate_cover_letter_page_fill(document)

    if not document.applicant.get("name"):
        issues.append(
            ValidationIssue(
                type="missing_applicant_name",
                severity="warning",
                section="header",
                message="Cover letter header is missing the applicant name.",
            )
        )
    if paragraph_count < min_paragraphs:
        issues.append(
            ValidationIssue(
                type="cover_letter_too_few_paragraphs",
                severity="warning",
                section="body",
                message="Cover letter body has too few paragraphs.",
                details={"paragraph_count": paragraph_count, "min_paragraphs": min_paragraphs},
            )
        )
    if word_count < min_words:
        issues.append(
            ValidationIssue(
                type="cover_letter_too_short",
                severity="warning",
                section="body",
                message="Cover letter body is below the target word budget.",
                details={"word_count": word_count, "min_words": min_words},
            )
        )
    if word_count > max_words:
        issues.append(
            ValidationIssue(
                type="cover_letter_too_long",
                severity="warning",
                section="body",
                message="Cover letter body exceeds the target word budget.",
                details={"word_count": word_count, "max_words": max_words},
            )
        )
    if estimated_page_fill < min_estimated_page_fill:
        issues.append(
            ValidationIssue(
                type="cover_letter_underfilled_page",
                severity="warning",
                section="layout",
                message="Cover letter is likely to underfill a one-page document.",
                details={
                    "estimated_page_fill_ratio": estimated_page_fill,
                    "min_estimated_page_fill_ratio": min_estimated_page_fill,
                },
            )
        )
    # 2026-07-16: the old heuristic (regex for any "At <Company>, <X>"
    # sentence) flagged 20/20 letters because the generation prompt
    # *instructs* that style ("At SDS, I redesigned onboarding…").
    # Only flag genuine dumps: a resume bullet copied near-verbatim.
    evidence_bullets = list((document.metadata or {}).get("evidence_bullets") or [])
    if _contains_verbatim_bullet(body_text, evidence_bullets):
        issues.append(
            ValidationIssue(
                type="raw_evidence_dump",
                severity="warning",
                section="body",
                message=(
                    "Cover letter copies a resume bullet near-verbatim "
                    "instead of retelling it."
                ),
            )
        )
    if re.search(r"</?[a-z][^>]*>", body_text, flags=re.IGNORECASE):
        issues.append(
            ValidationIssue(
                type="raw_html_markup",
                severity="error",
                section="body",
                message="Cover letter contains raw HTML markup and must be regenerated.",
            )
        )
    quality_issues = list((document.metadata or {}).get("quality_issues") or [])
    for issue in quality_issues:
        issues.append(
            ValidationIssue(
                type="cover_letter_quality_check",
                severity="warning",
                section="body",
                message=f"Cover letter quality check flagged: {issue}",
            )
        )

    return ValidationResult(
        ok=not any(issue.severity == "error" for issue in issues),
        issues=issues,
        metrics={
            "cover_letter_word_count": word_count,
            "cover_letter_paragraph_count": paragraph_count,
            "estimated_page_fill_ratio": estimated_page_fill,
            "font_family": "Times New Roman",
            "font_size_pt": 11,
            "page_margin_inches": 0.85,
            "cover_letter_quality_issues": quality_issues,
        },
    )


def _contains_verbatim_bullet(body_text: str, evidence_bullets: list[str]) -> bool:
    """True when a substantial evidence bullet appears near-verbatim.

    Retelling evidence is the whole point of the letter; copying a long
    bullet word-for-word is lazy dumping. Normalize whitespace/case and
    look for the first ~12 words of each bullet as a contiguous run.
    """
    normalized_body = " ".join(re.findall(r"[a-z0-9]+", body_text.lower()))
    for bullet in evidence_bullets:
        tokens = re.findall(r"[a-z0-9]+", str(bullet).lower())
        if len(tokens) < 8:
            continue
        probe = " ".join(tokens[:12])
        if probe and probe in normalized_body:
            return True
    return False


def _estimate_cover_letter_page_fill(document: CoverLetterDocument) -> float:
    paragraphs = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]
    # Times New Roman 11pt, 0.85" margins: roughly 85-95 chars per line and
    # 44-48 usable lines on a Letter page. Include fixed frame lines
    # for name/contact/date/company/salutation/signature plus paragraph gaps.
    body_lines = sum(max(1, ceil(len(text) / 90)) for text in paragraphs)
    frame_lines = 7
    paragraph_gap_lines = max(0, len(paragraphs) - 1)
    return round(min(1.0, (frame_lines + paragraph_gap_lines + body_lines) / 46), 2)


def _docx_layout_metrics(docx_path: Path | None, *, document_type: str) -> dict:
    if not docx_path or not docx_path.exists():
        return {}
    try:
        from docx import Document  # noqa: PLC0415
    except ImportError:
        return {}
    try:
        doc = Document(str(docx_path))
    except Exception:
        return {}

    style_name = "CoverLetter.Body" if document_type == "cover_letter" else "Resume.Normal"
    metrics: dict = {}
    try:
        style = doc.styles[style_name]
        metrics["docx_body_style"] = style_name
        metrics["docx_body_font_family"] = style.font.name
        metrics["docx_body_font_size_pt"] = (
            round(style.font.size.pt, 1) if style.font.size is not None else None
        )
    except KeyError:
        metrics["docx_body_style"] = None

    section = doc.sections[0]
    metrics["docx_margins_inches"] = {
        "top": round(section.top_margin.inches, 2),
        "bottom": round(section.bottom_margin.inches, 2),
        "left": round(section.left_margin.inches, 2),
        "right": round(section.right_margin.inches, 2),
    }
    return metrics


def validate_latex_artifacts(
    validation: ValidationResult | None = None,
    *,
    tex_path: Path | None,
    pdf_path: Path | None = None,
    pdf_attempted: bool = False,
    max_pages: int = 1,
    target_pages: int | None = None,
) -> ValidationResult:
    """Validate rendered LaTeX files without requiring a DOCX artifact."""
    validation = validation or ValidationResult(ok=True, issues=[], metrics={})
    issues = list(validation.issues)
    metrics = dict(validation.metrics)

    tex_ok = bool(tex_path and tex_path.exists() and tex_path.stat().st_size > 0)
    pdf_ok = bool(pdf_path and pdf_path.exists() and pdf_path.stat().st_size > 0)
    pdf_pages = get_pdf_page_count(pdf_path)
    metrics.update(
        {
            "tex_generated": tex_ok,
            "tex_path": str(tex_path) if tex_path else None,
            "pdf_generated": pdf_ok,
            "pdf_path": str(pdf_path) if pdf_path else None,
            "pdf_page_count": pdf_pages,
        }
    )

    if not tex_ok:
        issues.append(
            ValidationIssue(
                type="tex_generation_failed",
                severity="error",
                message="LaTeX renderer did not produce a valid .tex file.",
                details={"path": str(tex_path) if tex_path else None},
            )
        )

    if pdf_attempted and not pdf_ok:
        issues.append(
            ValidationIssue(
                type="pdf_generation_failed",
                severity="warning",
                message="LaTeX PDF compilation did not produce a valid file.",
                details={"path": str(pdf_path) if pdf_path else None},
            )
        )

    target = target_pages if target_pages is not None else max_pages
    metrics["target_pages"] = target
    if pdf_pages is not None and target_pages is not None:
        if pdf_pages > target:
            issues.append(
                ValidationIssue(
                    type="rendered_page_overflow",
                    severity="error",
                    message=(
                        f"Rendered document is {pdf_pages} pages but template "
                        f"targets exactly {target}. Regenerate with less content."
                    ),
                    details={"page_count": pdf_pages, "target_pages": target},
                )
            )
        elif pdf_pages < target:
            issues.append(
                ValidationIssue(
                    type="rendered_page_underflow",
                    severity="error",
                    message=(
                        f"Rendered document is {pdf_pages} pages but template "
                        f"targets exactly {target}. Regenerate with more content."
                    ),
                    details={"page_count": pdf_pages, "target_pages": target},
                )
            )
    elif pdf_pages is not None and pdf_pages > max_pages:
        issues.append(
            ValidationIssue(
                type="rendered_page_overflow",
                severity="warning",
                message="Rendered document exceeds the configured page target.",
                details={"page_count": pdf_pages, "max_pages": max_pages},
            )
        )

    return ValidationResult(
        ok=not any(issue.severity == "error" for issue in issues),
        issues=issues,
        metrics=metrics,
    )


def _word_count(value: str) -> int:
    return len(re.findall(r"\b[\w+#.]+\b", value))


def _numbers(value: str) -> set[str]:
    return set(re.findall(r"\b\d+(?:\.\d+)?%?\+?\b", value))


def _added_numbers(original: str, rewritten: str) -> list[str]:
    return sorted(_numbers(rewritten) - _numbers(original))


def _repeated_action_verbs(bullets: list[str]) -> dict[str, int]:
    verbs = []
    for bullet in bullets:
        match = re.search(r"[A-Za-z]+", bullet)
        if match:
            verbs.append(match.group(0).lower())
    counts = Counter(verbs)
    return {verb: count for verb, count in counts.items() if count >= 3}


def _keyword_coverage(document: ResumeDocument, jd_tags: list[str]) -> dict:
    normalized_tags = {_normalize(tag) for tag in jd_tags if _normalize(tag)}
    if not normalized_tags:
        return {"covered_keywords": [], "missing_keywords": [], "coverage_ratio": 0.0}

    text_parts = []
    for values in document.skills.values():
        text_parts.extend(values)
    for item in [*document.experiences, *document.projects]:
        text_parts.extend([item.name, item.title, item.organization, *item.tech_stack])
        text_parts.extend(bullet.text for bullet in item.bullets)

    haystack = {
        _normalize(token)
        for token in re.findall(r"[A-Za-z][A-Za-z0-9+#.]+", " ".join(text_parts))
    }
    covered = sorted(tag for tag in normalized_tags if tag in haystack)
    missing = sorted(normalized_tags - set(covered))
    return {
        "covered_keywords": covered,
        "missing_keywords": missing,
        "coverage_ratio": round(len(covered) / len(normalized_tags), 3),
    }


def _estimate_pages(document: ResumeDocument) -> int:
    bullet_count = sum(len(item.bullets) for item in [*document.experiences, *document.projects])
    item_count = len(document.education) + len(document.experiences) + len(document.projects)
    skill_count = sum(len(values) for values in document.skills.values())
    estimated_lines = 5 + item_count * 2 + bullet_count * 2 + max(1, skill_count // 6)
    return max(1, (estimated_lines + 42) // 43)


def _estimate_resume_page_fill(document: ResumeDocument) -> float:
    bullet_count = sum(len(item.bullets) for item in [*document.experiences, *document.projects])
    item_count = len(document.education) + len(document.experiences) + len(document.projects)
    skill_count = sum(len(values) for values in document.skills.values())
    estimated_lines = 5 + item_count * 2 + bullet_count * 2 + max(1, skill_count // 6)
    return round(min(1.0, estimated_lines / 43), 2)


def _normalize(value: str) -> str:
    return re.sub(r"[^a-z0-9+#.]+", "_", value.lower().strip()).strip("_")
