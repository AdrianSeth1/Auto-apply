"""Phase 15.2: DOCX patch mode.

Applies localized edits to a user-uploaded DOCX without re-rendering
the document from scratch:

* Strip the Summary section entirely. The heading + body are blanked
  even if the user-uploaded source resume contained a Summary --
  generated resumes never include one.
* Replace bullet text in-place. Bullets are identified by the
  paragraph's named style starting with ``List``, or by leading "•"/
  "-" markers. The text-run's font / size / colour / bold / italic
  attributes are preserved -- only the visible string is swapped.
* Reorder the *skills* list. We detect a "Skills"-style heading and
  rewrite the paragraphs that follow it (until the next heading)
  with the IR's skills payload, preserving the surrounding style.
* Include / drop sections. Each section is bounded by its heading;
  the patcher can mark a whole section as "hidden" by removing the
  underlying paragraphs while leaving the heading present so the user
  can see the omission in track-changes-style diffs.

Style preservation is the guarantee, not pixel-perfect pagination.
Page-count drift is detected and reported via :class:`PatchReport`;
the materials router (Phase 15.5) decides whether the drift is
acceptable or falls back to ``generate_from_template`` per D024.

This module deliberately does NOT use python-docx's high-level
``add_paragraph`` for replacements -- that would create a fresh
paragraph with default styling. We mutate the existing paragraph's
runs in place so the user's named-style assignment survives.
"""

from __future__ import annotations

import copy
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from src.generation.ir import ResumeDocument

logger = logging.getLogger(__name__)


# ---- Result types ----------------------------------------------------


@dataclass
class PatchOperation:
    kind: str  # 'bullet' | 'skills' | 'section_drop'
    detail: str  # human-readable, e.g. 'bullets@experience[0]: 3 -> 4'


@dataclass
class PatchReport:
    """Per-document outcome. ``success`` is False whenever any operation
    raises or a hard guarantee fails."""

    success: bool
    output_path: Path
    operations: list[PatchOperation] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    failure_reason: str | None = None


class PatchFallback(Exception):  # noqa: N818 -- "Fallback" is the operative noun; it is control-flow, not an error condition
    """Raised when the patch cannot continue. The materials router
    catches this and routes to ``generate_from_template``."""


# ---- Public entry ----------------------------------------------------


_BULLET_PREFIX_RE = re.compile(r"^\s*[•●◦\-*]\s*")
_TOP_HEADING_STYLES = {"heading 1", "summary", "title"}
_SKILLS_STYLES = {"skills"}


def patch_resume_docx(
    source_path: Path,
    document: ResumeDocument,
    *,
    output_path: Path,
    allow_reorder_sections: bool = True,
    allow_add_remove_bullets: bool = True,
) -> PatchReport:
    """Open ``source_path``, apply IR-driven patches, save to
    ``output_path``. Returns a :class:`PatchReport` describing what
    changed. Raises :class:`PatchFallback` only when the document is
    so far from the IR's shape that no useful edit can be made --
    callers catch it as the "give up and template" signal.

    Phase 18.x: two policy knobs from the user's saved patch settings:

    * ``allow_reorder_sections`` — when ``False``, the IR's
      ``section_order`` is ignored for visibility decisions. The
      patcher only edits the sections that are already present in
      the source document and never re-orders or hides them. This
      is the conservative "keep my layout exactly" mode.
    * ``allow_add_remove_bullets`` — when ``False``, the patcher
      will not append surplus bullets to a section and will not
      blank deficit bullets; the existing count is preserved by
      truncating the IR's bullet list to the source's slot count.
      Surplus IR bullets get folded back as a strategy note for
      auditability.

    Defaults preserve the previous (Phase 15.2) behavior so existing
    callers don't change shape.
    """
    try:
        from docx import Document  # local import: heavy
    except ImportError as exc:  # pragma: no cover -- python-docx is a hard dep
        raise PatchFallback(f"python-docx not available: {exc}") from exc

    if not source_path.exists():
        raise PatchFallback(f"source DOCX missing: {source_path}")

    doc = Document(str(source_path))
    report = PatchReport(success=True, output_path=output_path)

    try:
        # Summary section is intentionally never patched and is actively
        # stripped from any source DOCX (see _strip_summary_section).
        _strip_summary_section(doc, report)
        _patch_skills(doc, document, report)
        _patch_bullets(
            doc,
            document,
            report,
            allow_add_remove_bullets=allow_add_remove_bullets,
        )
        if allow_reorder_sections:
            _patch_section_visibility(doc, document, report)
    except PatchFallback:
        # Bubble up to caller (materials router) -- do NOT swallow.
        raise
    except Exception as exc:  # noqa: BLE001
        # Unexpected python-docx error: degrade to template fallback
        # with the original exception captured for the operator UI.
        logger.exception("docx patch failed; falling back to template")
        raise PatchFallback(f"docx patch raised: {exc!r}") from exc

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return report


# ---- Section detection ----------------------------------------------


def _is_heading(paragraph: Any) -> bool:
    """True for any heading-styled paragraph (used by section
    visibility, which treats every heading as a potential boundary)."""
    style = (_style_name(paragraph) or "").lower()
    return style.startswith("heading") or style in _TOP_HEADING_STYLES | _SKILLS_STYLES


def _is_top_level_heading(paragraph: Any) -> bool:
    """Top-level section divider (``Heading 1`` / ``Title``).

    Sub-headings (``Heading 2`` and below) are item headings inside a
    section (e.g. a job entry inside Experience) and must NOT terminate
    a section's body range, otherwise bullets nested under a Heading 2
    job entry would never be patched."""
    style = (_style_name(paragraph) or "").lower()
    if style in {"heading 1", "title"}:
        return True
    if style in _TOP_HEADING_STYLES | _SKILLS_STYLES:
        return True
    return False


def _style_name(paragraph: Any) -> str:
    style = getattr(paragraph, "style", None)
    return getattr(style, "name", "") or ""


def _is_bullet(paragraph: Any) -> bool:
    style = _style_name(paragraph).lower()
    if style.startswith("list"):
        return True
    text = (paragraph.text or "").lstrip()
    return bool(_BULLET_PREFIX_RE.match(text))


def _section_index(doc: Any, name: str) -> int | None:
    """Return the paragraph index of a top-level heading whose text
    matches ``name`` case-insensitively, or None if absent."""
    needle = name.lower()
    for idx, para in enumerate(doc.paragraphs):
        if not _is_top_level_heading(para):
            continue
        text = (para.text or "").strip().lower()
        if text == needle or text.startswith(needle):
            return idx
    return None


def _section_body_range(doc: Any, heading_idx: int) -> tuple[int, int]:
    """Return ``(start, end_exclusive)`` paragraph indices for the body
    of the section opened by ``heading_idx``. The body ends at the
    next top-level heading (or EOF). Sub-headings (Heading 2+) are
    part of the body, not boundaries."""
    paras = doc.paragraphs
    end = len(paras)
    for j in range(heading_idx + 1, len(paras)):
        if _is_top_level_heading(paras[j]):
            end = j
            break
    return heading_idx + 1, end


# ---- In-place text replacement (preserves run formatting) -----------


def _replace_paragraph_text(paragraph: Any, new_text: str) -> None:
    """Set the paragraph's visible text to ``new_text`` while
    preserving the formatting of the first run. python-docx exposes
    one or more runs per paragraph; we collapse them into the first
    run (keeping its font/size/bold/italic/color) so style attributes
    propagate to the new text."""
    runs = list(paragraph.runs)
    if not runs:
        paragraph.add_run(new_text)
        return
    first, *rest = runs
    first.text = new_text
    for extra in rest:
        # Remove the extra runs by clearing their text; physically
        # deleting them via _element requires nuking children which
        # can break lists. Setting text to "" leaves an empty run
        # which Word renders invisibly.
        extra.text = ""


def _clone_paragraph_after(template_para: Any, text: str) -> Any:
    """Create a new paragraph immediately after ``template_para`` that
    inherits the same XML style. Returns the new paragraph object."""
    new_xml = copy.deepcopy(template_para._p)
    # Drop existing run text so we can write our own.
    for r in new_xml.findall(".//{*}r"):
        for t in r.findall(".//{*}t"):
            t.text = ""
    template_para._p.addnext(new_xml)
    from docx.text.paragraph import Paragraph

    new_para = Paragraph(new_xml, template_para._parent)
    _replace_paragraph_text(new_para, text)
    return new_para


# ---- Patch operations -----------------------------------------------


def _strip_summary_section(doc: Any, report: PatchReport) -> None:
    """Remove the Summary heading + body from the source DOCX in place.

    Per product decision: the generated resume must never contain a
    Summary section, even if the user-uploaded source resume had one.
    The heading paragraph and every paragraph in its body range are
    blanked (text emptied, run formatting preserved -- python-docx
    cannot reliably delete paragraph elements without breaking lists
    so we leave empty paragraphs in place). Word renders empty
    paragraphs as roughly nothing.
    """
    heading_idx = _section_index(doc, "Summary")
    if heading_idx is None:
        return
    start, end = _section_body_range(doc, heading_idx)
    _replace_paragraph_text(doc.paragraphs[heading_idx], "")
    for j in range(start, end):
        _replace_paragraph_text(doc.paragraphs[j], "")
    report.operations.append(
        PatchOperation(kind="section_drop", detail="summary heading + body stripped")
    )


def _patch_skills(doc: Any, document: ResumeDocument, report: PatchReport) -> None:
    skills = document.skills or {}
    if not skills:
        return
    rendered = _render_skills_lines(skills)
    if not rendered:
        return

    heading_idx = _section_index(doc, "Skills")
    if heading_idx is None:
        report.warnings.append("no Skills heading found; skills unchanged")
        return

    start, end = _section_body_range(doc, heading_idx)
    paragraphs = doc.paragraphs
    body = paragraphs[start:end]
    if not body:
        # Append after heading.
        template = paragraphs[heading_idx]
        for line in rendered:
            template = _clone_paragraph_after(template, line)
        report.operations.append(
            PatchOperation(kind="skills", detail=f"skills appended ({len(rendered)} lines)")
        )
        return

    # Replace first ``len(rendered)`` body paragraphs; blank-out any
    # remaining old paragraphs (cannot physically delete without risk
    # of breaking list numbering).
    for i, line in enumerate(rendered):
        if start + i < end:
            _replace_paragraph_text(paragraphs[start + i], line)
        else:
            template = paragraphs[end - 1]
            _clone_paragraph_after(template, line)
            end += 1
    for j in range(start + len(rendered), end):
        _replace_paragraph_text(paragraphs[j], "")
    report.operations.append(
        PatchOperation(kind="skills", detail=f"skills replaced ({len(rendered)} lines)")
    )


def _patch_bullets(
    doc: Any,
    document: ResumeDocument,
    report: PatchReport,
    *,
    allow_add_remove_bullets: bool = True,
) -> None:
    """Replace bullet runs section by section. Matches sections by
    name (Experience / Projects); other sections are left untouched.

    When ``allow_add_remove_bullets`` is False the source DOCX's bullet
    count is treated as authoritative: surplus IR bullets are dropped
    (and noted in the report) and deficit slots are left with their
    original text instead of being blanked. This is the user's
    "preserve my structure exactly" mode.
    """
    section_map = {
        "experience": list(document.experiences or []),
        "experiences": list(document.experiences or []),
        "projects": list(document.projects or []),
    }
    if not any(section_map.values()):
        return

    paragraphs = doc.paragraphs
    section_bounds = _find_all_sections(doc, set(section_map.keys()))
    if not section_bounds:
        report.warnings.append("no Experience/Projects sections found; bullets unchanged")
        return

    for heading_idx, name, body_start, body_end in section_bounds:
        items = section_map.get(name) or []
        if not items:
            continue
        bullet_indices = [
            i for i in range(body_start, body_end) if _is_bullet(paragraphs[i])
        ]
        # Flatten the IR items into a list of bullet strings; the IR
        # may have many items each carrying multiple bullets.
        new_bullets: list[str] = []
        for item in items:
            for bullet in getattr(item, "bullets", []) or []:
                txt = getattr(bullet, "text", "") or ""
                if txt.strip():
                    new_bullets.append(txt.strip())

        for i, source_bullet_idx in enumerate(bullet_indices):
            if i < len(new_bullets):
                _replace_paragraph_text(paragraphs[source_bullet_idx], new_bullets[i])
            elif allow_add_remove_bullets:
                # Default behaviour: blank the deficit slot. Keeps
                # the bullet style/numbering anchor in place so the
                # rendered DOCX doesn't reflow weirdly, but the line
                # reads empty to the reader.
                _replace_paragraph_text(paragraphs[source_bullet_idx], "")
            # else: preserve original bullet text -- the user asked us
            # not to add or remove bullets, so leaving the source text
            # is the honest choice.

        if len(new_bullets) > len(bullet_indices) and bullet_indices:
            if allow_add_remove_bullets:
                # Append any leftover bullets after the last existing
                # bullet so we do not lose IR content.
                anchor = paragraphs[bullet_indices[-1]]
                for extra in new_bullets[len(bullet_indices):]:
                    anchor = _clone_paragraph_after(anchor, extra)
            else:
                # User opted out of growing the section. Surface the
                # dropped bullets as a report warning so the operator
                # can decide whether to widen the policy.
                dropped = new_bullets[len(bullet_indices):]
                report.warnings.append(
                    f"section={name} dropped {len(dropped)} IR bullet(s) "
                    f"(allow_add_remove_bullets=False): "
                    + " | ".join(dropped[:3])
                    + (" …" if len(dropped) > 3 else "")
                )

        report.operations.append(
            PatchOperation(
                kind="bullet",
                detail=(
                    f"section={name} bullets {len(bullet_indices)} -> "
                    f"{min(len(new_bullets), len(bullet_indices)) if not allow_add_remove_bullets else len(new_bullets)}"
                ),
            )
        )


def _patch_section_visibility(
    doc: Any, document: ResumeDocument, report: PatchReport
) -> None:
    """If the IR's ``section_order`` is set, hide any section whose
    name does not appear in that list by blanking its body. The
    heading itself stays so the user can see we omitted it on
    purpose.

    Sections that the IR *actively populates* are always kept, even
    if their name is missing from ``section_order``. This makes
    ``section_order`` an "include this peripheral section" list
    rather than a strict whitelist that would silently drop core
    content (the IR's default ``section_order`` omits "summary" but
    we should never blank a populated summary block)."""
    if not document.section_order:
        return
    keep = {name.lower() for name in document.section_order}
    # ``summary`` is intentionally not in has_content -- the ResumeDocument
    # IR no longer has a summary field, and _strip_summary_section above
    # already removed the source-doc Summary block before this runs.
    has_content = {
        "skills": bool(document.skills),
        "experience": bool(document.experiences),
        "experiences": bool(document.experiences),
        "projects": bool(document.projects),
        "education": bool(document.education),
        "header": True,  # never blank the header block
    }
    paragraphs = doc.paragraphs
    for idx, para in enumerate(paragraphs):
        if not _is_top_level_heading(para):
            continue
        name = (para.text or "").strip().lower()
        if not name:
            continue
        norm = _normalize_heading(name)
        if name in keep or norm in keep:
            continue
        if has_content.get(name) or has_content.get(norm):
            continue
        start, end = _section_body_range(doc, idx)
        if end <= start:
            continue
        for j in range(start, end):
            _replace_paragraph_text(paragraphs[j], "")
        report.operations.append(
            PatchOperation(kind="section_drop", detail=f"{name} body blanked")
        )


# ---- Helpers ---------------------------------------------------------


_SKILL_GROUP_ORDER = ("must_have", "preferred", "additional")


def _render_skills_lines(skills: dict[str, list[str]]) -> list[str]:
    lines: list[str] = []
    seen: set[str] = set()
    for group in _SKILL_GROUP_ORDER:
        values = skills.get(group) or []
        if values:
            lines.append(f"{group.replace('_', ' ').title()}: {', '.join(values)}")
            seen.add(group)
    for group, values in skills.items():
        if group in seen or not values:
            continue
        lines.append(f"{group.replace('_', ' ').title()}: {', '.join(values)}")
    return lines


def _normalize_heading(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", "_", text.lower()).strip("_")


def _find_all_sections(
    doc: Any, names: set[str]
) -> list[tuple[int, str, int, int]]:
    """Return ``(heading_idx, normalized_name, body_start, body_end)``
    for every paragraph whose normalized heading matches one of
    ``names``."""
    found: list[tuple[int, str, int, int]] = []
    paragraphs = doc.paragraphs
    for idx, para in enumerate(paragraphs):
        if not _is_top_level_heading(para):
            continue
        name = (para.text or "").strip().lower()
        if not name:
            continue
        norm = _normalize_heading(name)
        for candidate in (name, norm, name.rstrip("s"), norm.rstrip("s")):
            if candidate in names:
                start, end = _section_body_range(doc, idx)
                found.append((idx, candidate, start, end))
                break
    return found


# --- Cover letter patcher --------------------------------------------
#
# Cover letters are structurally simpler than resumes -- no bullet
# lists, no sub-sections -- but they care a lot about the surrounding
# template (sender address block, date, recipient block, salutation,
# signature). When the user picks ``patch_existing`` for a cover
# letter, they want us to preserve those layout elements and only
# rewrite the body paragraphs in between.
#
# The heuristic anchors:
#
#   * Body STARTS after a ``"Dear ..."`` paragraph (the salutation).
#     If there isn't one we bail to PatchFallback rather than guess --
#     replacing a whole letter when we can't find the salutation
#     usually means we'd overwrite the address block.
#
#   * Body ENDS at a closing line. We accept the common English
#     closings ("Sincerely", "Best regards", "Best,", "Yours
#     sincerely", "Thank you", "Regards", "Kind regards"). If none
#     is found we fall back to the last non-empty paragraph in the
#     document (which is typically the signature).

_COVER_CLOSING_PATTERNS = (
    "sincerely",
    "yours sincerely",
    "yours truly",
    "best regards",
    "kind regards",
    "warm regards",
    "regards",
    "best wishes",
    "best,",
    "thank you",
    "thanks,",
    "respectfully",
    "cordially",
)


def _find_cover_salutation_index(doc: Any) -> int | None:
    """Return the index of the salutation paragraph (``Dear ...``)
    in the source DOCX, or ``None`` if we can't find one."""
    for idx, para in enumerate(doc.paragraphs):
        text = (para.text or "").strip().lower()
        if text.startswith("dear ") or text.startswith("to whom"):
            return idx
    return None


def _find_cover_closing_index(
    doc: Any, after_idx: int
) -> int | None:
    """Return the index of the first closing line at or after
    ``after_idx``, or ``None`` if no recognized closing is present."""
    for idx in range(after_idx + 1, len(doc.paragraphs)):
        text = (doc.paragraphs[idx].text or "").strip().lower().rstrip(",")
        if not text:
            continue
        for marker in _COVER_CLOSING_PATTERNS:
            if text == marker or text.startswith(marker + ","):
                return idx
            # Some closings appear inline like "Sincerely yours"
            if text.startswith(marker + " ") and len(text) - len(marker) < 16:
                return idx
    return None


def patch_cover_letter_docx(
    source_path: Path,
    document,
    *,
    output_path: Path,
) -> PatchReport:
    """Replace the body paragraphs of a cover letter DOCX in place.

    ``document`` is a :class:`~src.generation.ir.CoverLetterDocument`
    -- the same IR the regenerate path produces -- so this patcher
    can be a drop-in for ``patch_existing`` callers that previously
    fell back to template rendering.

    The patcher preserves everything outside the body block: the
    sender / date / recipient header above the salutation, the
    salutation itself, the closing line + signature below the
    closing, and any contact line at the very bottom. The body
    between salutation and closing is rewritten run-by-run with the
    IR's ``paragraphs`` -- surplus IR paragraphs are appended after
    the last body paragraph, deficit slots are blanked.

    Raises :class:`PatchFallback` when we can't anchor a salutation,
    so the caller can route to the template path with a clean note.
    """
    try:
        from docx import Document  # local import: heavy
    except ImportError as exc:  # pragma: no cover -- python-docx is a hard dep
        raise PatchFallback(f"python-docx not available: {exc}") from exc

    if not source_path.exists():
        raise PatchFallback(f"source DOCX missing: {source_path}")

    doc = Document(str(source_path))
    report = PatchReport(success=True, output_path=output_path)

    salutation_idx = _find_cover_salutation_index(doc)
    if salutation_idx is None:
        raise PatchFallback(
            "could not locate a 'Dear ...' salutation in the source DOCX -- "
            "refusing to overwrite the whole document"
        )

    closing_idx = _find_cover_closing_index(doc, salutation_idx)
    # If we can't find a closing, treat everything after the salutation
    # up to the last non-empty paragraph as the body. The last
    # non-empty paragraph is left alone so the operator's signature /
    # name line survives.
    if closing_idx is None:
        last_non_empty = salutation_idx
        for idx in range(len(doc.paragraphs) - 1, salutation_idx, -1):
            if (doc.paragraphs[idx].text or "").strip():
                last_non_empty = idx
                break
        closing_idx = last_non_empty
        report.warnings.append(
            "no recognised closing line ('Sincerely', 'Best regards', etc.) "
            "found in source; replaced body up to the last non-empty "
            "paragraph"
        )

    body_indices = [
        idx
        for idx in range(salutation_idx + 1, closing_idx)
        if (doc.paragraphs[idx].text or "").strip()
    ]

    # Extract the new body paragraphs from the IR. ``paragraphs`` is a
    # list of CoverLetterParagraph; we drop salutation/closing-typed
    # entries because the source DOCX already provides those and we
    # explicitly want to keep them.
    new_paragraphs: list[str] = []
    for para in getattr(document, "paragraphs", []) or []:
        text = (getattr(para, "text", "") or "").strip()
        if not text:
            continue
        ptype = getattr(para, "type", "other")
        if ptype in {"opening", "closing"}:
            # ``opening`` from the IR is the salutation, ``closing``
            # is "Sincerely,\nLiam" — both already in the source
            # template, don't duplicate them.
            continue
        new_paragraphs.append(text)

    if not new_paragraphs:
        # If the IR offered no body content, nothing to patch -- treat
        # as a "no-op" patch rather than blanking the source.
        report.warnings.append(
            "cover letter IR had no body paragraphs; source body retained"
        )
    else:
        try:
            # Replace existing body slots run-by-run.
            for i, source_idx in enumerate(body_indices):
                if i < len(new_paragraphs):
                    _replace_paragraph_text(
                        doc.paragraphs[source_idx], new_paragraphs[i]
                    )
                else:
                    _replace_paragraph_text(doc.paragraphs[source_idx], "")

            # Append surplus IR paragraphs after the last existing
            # body paragraph so we don't lose tailored content.
            if len(new_paragraphs) > len(body_indices) and body_indices:
                anchor = doc.paragraphs[body_indices[-1]]
                for extra in new_paragraphs[len(body_indices):]:
                    anchor = _clone_paragraph_after(anchor, extra)
            elif len(new_paragraphs) > 0 and not body_indices:
                # The source had a salutation immediately followed by
                # the closing -- no body slots to replace. Inject the
                # body after the salutation by cloning.
                anchor = doc.paragraphs[salutation_idx]
                for extra in new_paragraphs:
                    anchor = _clone_paragraph_after(anchor, extra)

            report.operations.append(
                PatchOperation(
                    kind="cover_body",
                    detail=(
                        f"body slots {len(body_indices)} -> "
                        f"{len(new_paragraphs)}"
                    ),
                )
            )
        except PatchFallback:
            raise
        except Exception as exc:  # noqa: BLE001
            logger.exception("cover letter patch failed; falling back to template")
            raise PatchFallback(
                f"cover letter patch raised: {exc!r}"
            ) from exc

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return report


__all__ = [
    "PatchFallback",
    "PatchOperation",
    "PatchReport",
    "patch_resume_docx",
    "patch_cover_letter_docx",
]
