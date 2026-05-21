"""Resume importer.

Parses existing Word (.docx) or PDF resumes into structured YAML format
using Claude CLI for intelligent extraction.
"""

from __future__ import annotations

import logging
from pathlib import Path

import yaml

from src.utils.llm import generate_text

logger = logging.getLogger("autoapply.memory.resume_importer")

EXTRACTION_SYSTEM_PROMPT = """You are a resume parser.
Extract structured data from the resume text provided.
Return ONLY valid YAML (no markdown fences, no explanations) matching this exact schema:

identity:
  full_name: "..."
  email: "..."
  phone: "..."
  location: "..."
  linkedin_url: "..."
  github_url: "..."
  portfolio_url: "..."

education:
  - institution: "..."
    degree: "..."
    field: "..."
    location: "..."
    start_date: "YYYY-MM"
    end_date: "YYYY-MM"
    gpa: "..."
    relevant_courses:
      - name: "..."
        tags: ["..."]

work_experiences:
  - company: "..."
    title: "..."
    location: "..."
    start_date: "YYYY-MM"
    end_date: "YYYY-MM"
    bullets:
      - text: "exact bullet text from resume"
        tags: ["skill1", "skill2"]

projects:
  - name: "..."
    role: "..."
    description: "..."
    tech_stack: ["..."]
    links: ["https://github.com/me/project", "..."]
    bullets:
      - text: "exact bullet text from resume"
        tags: ["skill1", "skill2"]

skills:
  languages: ["..."]
  frameworks: ["..."]
  databases: ["..."]
  tools: ["..."]
  domains: ["..."]

custom_sections:
  # ANY section in the resume that does not fit the buckets above goes
  # here verbatim. Common examples: VOLUNTEER EXPERIENCE, AWARDS &
  # HONORS, PROFESSIONAL AFFILIATIONS, INTERESTS & ACTIVITIES,
  # CERTIFICATIONS, PUBLICATIONS, LANGUAGES (spoken languages, not
  # programming -- those go in skills.languages), CONFERENCES, PATENTS.
  - title: "VOLUNTEER EXPERIENCE"  # use the resume's exact heading
    entries:
      - title: "Role or award name"
        organization: "..."
        location: "..."
        start_date: "YYYY-MM"
        end_date: "YYYY-MM"
        details: "single-line description if the resume only has one line"
        bullets:
          - "exact bullet text from the resume"

Rules:
- Preserve original bullet text exactly — do not rephrase or embellish
- Tags should be lowercase single words or short phrases
  (e.g., "python", "distributed_systems", "api_design")
- Dates in YYYY-MM format. Use "Present" for current positions
- If information is not in the resume, omit that field entirely. Never
  invent placeholder values like "None", "N/A" or "Unknown" — leave the
  field out.
- For skills, categorize into the groups shown. If unsure of category, put in "tools"
- Hyperlinks in the input are rendered as ``"label (https://target)"``.
  Pull the URL out and assign it to the right structured field rather
  than leaving it inline:
    * GitHub link in the contact line → ``identity.github_url``
    * LinkedIn link in the contact line → ``identity.linkedin_url``
    * Personal site / portfolio in the contact line → ``identity.portfolio_url``
    * A URL attached to a project title (or anywhere in the project
      block that clearly points to the project itself) → add it to
      that project's ``links`` list. Multiple links are allowed.
  When you move a URL into a structured field, also remove the
  ``(https://…)`` suffix from the surrounding text so bullet/title
  strings stay clean.
- ``custom_sections`` MUST capture every heading you saw in the resume
  that is not Education / Experience / Work / Projects / Skills.
  Examples of headings that belong here: VOLUNTEER EXPERIENCE, AWARDS,
  HONORS, PROFESSIONAL AFFILIATIONS, CERTIFICATIONS, INTERESTS,
  ACTIVITIES, LANGUAGES (when listing spoken languages), PUBLICATIONS.
  Preserve the heading text exactly as it appears in the resume. Each
  entry's free-form supporting line goes into ``details``; longer
  multi-line entries become ``bullets``.
"""


# Word's XML namespaces. We only need ``w`` (main document grammar)
# and ``r`` (relationship references) to find hyperlinks; declaring
# them here keeps the XPath expressions readable.
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_DOCX_NSMAP = {"w": _W_NS, "r": _R_NS}


def _render_paragraph_with_links(paragraph) -> str:
    """Walk a python-docx paragraph's XML children in document order
    and produce a plain-text rendering where hyperlinks are written
    as ``text (url)``.

    Why we bother: ``paragraph.text`` from python-docx concatenates
    every ``<w:t>`` element it can find, which yields just the
    *visible* text. Hyperlinks live inside ``<w:hyperlink>`` wrappers
    whose ``r:id`` attribute points into the document's relationships
    map -- the URL itself never appears inside the paragraph's runs,
    so the default ``.text`` accessor silently drops it. Resumes
    routinely hang the candidate's LinkedIn/GitHub off the contact
    line as hyperlinks, and project titles often link to a repo;
    losing those means the parsed profile has no way to recover them.

    Internal hyperlinks (anchors / bookmarks) and links whose
    relationship can't be resolved fall back to plain text rather
    than emitting a misleading ``()`` suffix.
    """
    rels = paragraph.part.rels
    pieces: list[str] = []

    for child in paragraph._element.iterchildren():
        tag = child.tag
        if tag == f"{{{_W_NS}}}r":
            # Plain run: just collect any <w:t> text under it.
            for t in child.iter(f"{{{_W_NS}}}t"):
                if t.text:
                    pieces.append(t.text)
        elif tag == f"{{{_W_NS}}}hyperlink":
            inner_text = "".join(
                (t.text or "") for t in child.iter(f"{{{_W_NS}}}t")
            )
            r_id = child.get(f"{{{_R_NS}}}id")
            url: str | None = None
            if r_id and r_id in rels:
                rel = rels[r_id]
                # ``is_external`` is False for anchor-only links
                # (e.g. ``<w:hyperlink w:anchor="..."/>``); those have
                # no URL we can preserve, so render the text only.
                if getattr(rel, "is_external", False):
                    url = rel.target_ref or None
            if inner_text and url and url not in inner_text:
                pieces.append(f"{inner_text} ({url})")
            elif inner_text:
                pieces.append(inner_text)
        # Other children (bookmarks, proofErr, etc.) carry no
        # user-visible text; ignore them.

    return "".join(pieces).strip()


def extract_text_from_docx(path: Path) -> str:
    """Extract plain text from a .docx file, preserving hyperlinks.

    Hyperlinks are rendered as ``"link text (https://target)"`` inline
    with the surrounding text so the downstream LLM has both the
    human-readable label and the URL when it builds the structured
    profile (see ``_render_paragraph_with_links`` for the rationale).
    """
    from docx import Document

    doc = Document(str(path))
    paragraphs: list[str] = []

    for para in doc.paragraphs:
        rendered = _render_paragraph_with_links(para)
        if rendered:
            paragraphs.append(rendered)

    # Also extract from tables (common in resume templates). We
    # iterate cell.paragraphs rather than cell.text so hyperlinks
    # inside tables get the same treatment.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    rendered = _render_paragraph_with_links(para)
                    if rendered:
                        paragraphs.append(rendered)

    return "\n".join(paragraphs)


def extract_text_from_pdf(path: Path) -> str:
    """Extract plain text from a PDF file."""
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(str(path))
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())
        return "\n".join(text_parts)
    except ImportError:
        raise ImportError(
            "PyMuPDF (fitz) is required for PDF parsing. Install with: uv add pymupdf"
        )


def _strip_markdown_fences(text: str) -> str:
    """Remove ```yaml ... ``` style markdown fences if present."""
    cleaned = text.strip()
    if cleaned.startswith("```"):
        lines = cleaned.split("\n")
        lines = [line for line in lines if not line.strip().startswith("```")]
        cleaned = "\n".join(lines)
    return cleaned.strip()


def _strip_chatty_preamble(text: str) -> str:
    """Drop common LLM preambles like 'Here is the YAML:' that sit
    before the first real YAML key. We look for the first line that
    looks like a top-level mapping key (``identity:``, ``education:``,
    etc.) and discard anything above it. If no such line exists, the
    text is returned unchanged.
    """
    expected_top_keys = (
        "identity:",
        "education:",
        "work_experiences:",
        "projects:",
        "skills:",
        "custom_sections:",
    )
    lines = text.split("\n")
    for idx, line in enumerate(lines):
        stripped = line.lstrip()
        if any(stripped.startswith(key) for key in expected_top_keys):
            return "\n".join(lines[idx:])
    return text


def _parse_llm_yaml_response(response: str) -> dict:
    """Best-effort parser for the LLM's YAML reply.

    The LLM is instructed to return bare YAML, but in practice it
    sometimes wraps the output in markdown fences, prepends a chatty
    preamble (\"Here is the structured YAML:\"), or returns the whole
    payload double-encoded as a single string. We unwrap those cases
    before giving up so that a transient phrasing tic on the model's
    side doesn't surface as a hard parse failure to the user.

    Raises:
        ValueError: With a snippet of the actual response so the
            operator can see what the model returned.
    """
    cleaned = _strip_markdown_fences(response)
    cleaned = _strip_chatty_preamble(cleaned)

    try:
        parsed = yaml.safe_load(cleaned)
    except yaml.YAMLError as e:
        logger.error("Failed to parse LLM response as YAML: %s", e)
        logger.error("Raw LLM response (first 500 chars):\n%s", cleaned[:500])
        raise ValueError(
            f"LLM returned invalid YAML: {e}. Response started with: "
            f"{cleaned[:200]!r}"
        )

    # If the model wrapped the whole document in quotes, safe_load gives
    # us a string — try reparsing the inner content once.
    if isinstance(parsed, str):
        inner = _strip_markdown_fences(parsed)
        inner = _strip_chatty_preamble(inner)
        try:
            reparsed = yaml.safe_load(inner)
        except yaml.YAMLError:
            reparsed = None
        if isinstance(reparsed, dict):
            return reparsed

        logger.error(
            "LLM returned a scalar string instead of a YAML mapping. "
            "First 500 chars:\n%s",
            cleaned[:500],
        )
        raise ValueError(
            "LLM did not return a structured resume — got plain text. "
            f"First 200 chars: {cleaned[:200]!r}"
        )

    if not isinstance(parsed, dict):
        logger.error(
            "LLM response parsed to %s, expected dict. First 500 chars:\n%s",
            type(parsed).__name__,
            cleaned[:500],
        )
        raise ValueError(
            f"Expected a YAML mapping from the LLM, got {type(parsed).__name__}. "
            f"First 200 chars: {cleaned[:200]!r}"
        )

    return parsed


_CANONICAL_HEADING_PATTERNS = (
    r"^education",
    r"^academic\s+",
    r"^work\s+experience",
    r"^professional\s+experience",
    r"^employment",
    r"^experience$",
    r"^experiences?$",
    r"^internships?$",
    r"^projects?$",
    r"^selected\s+projects",
    r"^personal\s+projects",
    r"^technical\s+skills",
    r"^skills?$",
    r"^summary",
    r"^objective",
    r"^profile$",
    r"^contact",
)


def _looks_canonical_heading(text: str) -> bool:
    """True if ``text`` is one of the canonical section headings already
    captured by the main schema (Education / Experience / Projects /
    Skills). Match is anchored case-insensitive on a normalised version
    of the text so 'TECHNICAL SKILLS' / 'Skills' both hit."""
    import re

    normalised = text.strip().lower().rstrip(":")
    return any(re.match(pat, normalised) for pat in _CANONICAL_HEADING_PATTERNS)


def _scan_resume_headings(raw_text: str) -> list[dict]:
    """Find candidate section headings in the resume text.

    A heading is a short, standalone line that

    * is either ALL-CAPS (with optional ``&`` / ``-`` / ``/`` / digits), or
    * is Title Case AND its label looks like a canonical resume section
      heading (Awards / Volunteer / etc.), AND
    * is preceded by a blank line (or is the very top of the doc), AND
    * does not end with terminal punctuation.

    The "blank line before" gate is what stops normal entry titles
    ("Local Library", "Acme Corp") from being misread as section
    headings -- those sit immediately under their parent heading. The
    accompanying body is everything between this heading and the next
    one (or end of document).
    """
    import re

    lines = raw_text.splitlines()
    # All-caps headings: at least 4 chars so short acronyms inside an
    # entry body (UBC, USA, NASA, ACM) are not mistaken for a section
    # heading. Resume sections like "EDUCATION", "AWARDS", "SKILLS"
    # are all 5+ chars, so this is safe.
    allcaps_re = re.compile(r"^[A-Z][A-Z0-9 &/\-+,'.]{3,49}$")
    # Title-case fallback for resumes that don't use all caps. We
    # intentionally restrict this to headings that contain at least one
    # known resume-section keyword so we don't grab every Title-Case
    # entry title in the doc.
    title_case_re = re.compile(
        r"^[A-Z][A-Za-z]+(?:\s+(?:[A-Z&][A-Za-z]*|and|of|the|in))*$"
    )
    title_case_section_keywords = {
        "education", "experience", "experiences", "employment", "work",
        "internship", "internships", "projects", "skills", "summary",
        "objective", "profile", "volunteer", "volunteering", "awards",
        "honors", "honours", "affiliations", "memberships", "interests",
        "activities", "certifications", "certificates", "publications",
        "presentations", "patents", "languages", "leadership", "training",
        "courses", "coursework", "extracurriculars", "extracurricular",
        "achievements", "research", "papers", "conferences", "talks",
    }

    def _is_heading_line(text: str, prev_blank: bool) -> bool:
        if not text:
            return False
        if text.endswith((".", ",", ";", "?", "!")):
            return False
        clean = text.rstrip(":").strip()
        if not (3 <= len(clean) <= 60):
            return False
        if len(clean.split()) > 6:
            return False
        # Both heading shapes require a blank line before them so that
        # inline items inside an entry body (school acronyms, company
        # names, project names) cannot be mistaken for section headings.
        if not prev_blank:
            return False
        if allcaps_re.match(clean):
            return True
        if title_case_re.match(clean):
            lowered = clean.lower()
            return any(kw in lowered.split() for kw in title_case_section_keywords)
        return False

    heading_indices: list[tuple[int, str]] = []
    for idx, line in enumerate(lines):
        stripped = line.strip().rstrip(":")
        if not stripped:
            continue
        prev_blank = idx == 0 or not lines[idx - 1].strip()
        if _is_heading_line(stripped, prev_blank):
            heading_indices.append((idx, stripped))

    # Drop the very first "heading" if it is the candidate's name --
    # i.e. it sits before any canonical heading and is the first match.
    if heading_indices and not _looks_canonical_heading(heading_indices[0][1]):
        # Heuristic: a candidate name is usually 2-3 Title Case words
        # with no special characters and sits at the very top.
        first_idx, first_title = heading_indices[0]
        if first_idx < 3 and not any(
            ch in first_title for ch in "&/-,"
        ) and len(first_title.split()) <= 3:
            heading_indices = heading_indices[1:]

    headings: list[dict] = []
    for pos, (line_idx, title) in enumerate(heading_indices):
        end_idx = (
            heading_indices[pos + 1][0]
            if pos + 1 < len(heading_indices)
            else len(lines)
        )
        body = "\n".join(line for line in lines[line_idx + 1 : end_idx]).strip()
        if not body:
            continue
        headings.append({"title": title, "body": body})
    return headings


def _backfill_custom_sections(profile_data: dict, raw_text: str) -> dict:
    """Make sure every non-canonical heading in the resume lands in
    ``profile_data.custom_sections``, even when the primary LLM pass
    decided to skip them.

    Strategy:
    1. Walk the raw text and collect ALL candidate headings + their
       bodies (``_scan_resume_headings``).
    2. Drop headings that map to canonical sections (Education / etc.)
       or are already captured by the parsed ``custom_sections``.
    3. For each surviving heading, ask the LLM (focused, one-shot) to
       structure that single section into the ``CustomSection`` shape.
       Failures fall back to a single ``details`` entry holding the
       raw body so the candidate's content never disappears.
    """
    data = dict(profile_data or {})
    existing_titles = {
        str(section.get("title", "")).strip().lower()
        for section in (data.get("custom_sections") or [])
        if isinstance(section, dict)
    }

    new_sections: list[dict] = list(data.get("custom_sections") or [])
    for heading in _scan_resume_headings(raw_text):
        title = heading["title"]
        normalised = title.strip().lower()
        if _looks_canonical_heading(title):
            continue
        if normalised in existing_titles:
            continue
        existing_titles.add(normalised)
        new_sections.append(_structure_custom_section(title, heading["body"]))

    if new_sections:
        data["custom_sections"] = new_sections
    return data


def _structure_custom_section(title: str, body: str) -> dict:
    """Convert a single heading + raw body into a custom-section dict.

    Tries the LLM first for a clean structure; falls back to a single
    entry with the raw body in ``details`` if the LLM is unavailable or
    returns junk. Either way the content is preserved -- the goal is
    "never silently lose what the candidate wrote".
    """
    structuring_prompt = f"""You are extracting one resume section into JSON.

The section is titled: {title!r}

The raw section body from the resume is between <body> tags:
<body>
{body}
</body>

Return ONLY a JSON object of the form:
{{
  "entries": [
    {{
      "title": "...",
      "organization": "...",
      "location": "...",
      "start_date": "YYYY-MM or empty",
      "end_date": "YYYY-MM or empty",
      "details": "single-line supporting text or empty",
      "bullets": ["exact bullet text", ...]
    }}
  ]
}}

Rules:
- Preserve every line of the body somewhere in the output -- no
  content from the body may be dropped.
- One entry per item the candidate listed. If the section is a single
  free-form paragraph (e.g. INTERESTS), produce ONE entry whose
  ``details`` is that paragraph.
- If a field is not present in the body, leave it as an empty string
  -- never invent "None" / "N/A".
- Do NOT rephrase bullet text. Copy it verbatim.
"""

    try:
        from src.utils.llm import generate_json  # local import keeps the
        # heavyweight LLM dependency lazy when the heading scanner is
        # invoked from a unit test that does not need a live provider.

        result = generate_json(
            f"Structure this resume section into JSON as specified.\n\n"
            f"Section title: {title}",
            system=structuring_prompt,
            timeout=60,
        )
        if isinstance(result, dict) and isinstance(result.get("entries"), list):
            return {
                "title": title,
                "entries": [
                    _coerce_entry_payload(entry)
                    for entry in result["entries"]
                    if isinstance(entry, dict)
                ],
            }
    except Exception as exc:  # noqa: BLE001
        logger.info(
            "Custom-section structuring failed for %r (%s); falling back to raw body.",
            title,
            exc,
        )

    # Deterministic fallback: keep the entire body as a single entry's
    # details so the user's content is preserved verbatim.
    return {
        "title": title,
        "entries": [{"details": body}],
    }


def _coerce_entry_payload(entry: dict) -> dict:
    """Strip unexpected keys and coerce types on the LLM JSON output.

    Belt and braces against an LLM that returns ``"bullets": "..."``
    instead of a list, or wraps a single bullet in extra prose.
    """
    bullets_raw = entry.get("bullets") or []
    bullets: list[str] = []
    if isinstance(bullets_raw, list):
        for bullet in bullets_raw:
            if isinstance(bullet, dict):
                bullets.append(str(bullet.get("text") or "").strip())
            elif bullet is not None:
                bullets.append(str(bullet).strip())
    elif isinstance(bullets_raw, str) and bullets_raw.strip():
        bullets.append(bullets_raw.strip())
    return {
        "title": str(entry.get("title") or "").strip(),
        "organization": str(entry.get("organization") or "").strip(),
        "location": str(entry.get("location") or "").strip(),
        "start_date": str(entry.get("start_date") or "").strip(),
        "end_date": str(entry.get("end_date") or "").strip(),
        "details": str(entry.get("details") or "").strip(),
        "bullets": [b for b in bullets if b],
    }


def import_resume(resume_path: Path, output_path: Path | None = None) -> dict:
    """Import a resume file and convert to structured YAML.

    Args:
        resume_path: Path to .docx or .pdf resume file.
        output_path: Optional path to save the generated YAML.

    Returns:
        Parsed profile data as a dict.
    """
    suffix = resume_path.suffix.lower()
    if suffix == ".docx":
        raw_text = extract_text_from_docx(resume_path)
    elif suffix == ".pdf":
        raw_text = extract_text_from_pdf(resume_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}. Use .docx or .pdf")

    if not raw_text.strip():
        raise ValueError(f"No text extracted from {resume_path}")

    logger.info("Extracted %d chars from %s", len(raw_text), resume_path.name)

    # Resume import was previously routed via the Phase 17.9.5 small
    # tier "to save tokens on pure extraction". That choice silently
    # dropped uncommon sections (VOLUNTEER EXPERIENCE / AWARDS /
    # AFFILIATIONS / INTERESTS) because small models tend to follow
    # the example schema literally and skip any "extra" fields like
    # custom_sections. Bumping to primary so the schema is honoured;
    # the cost difference is negligible for a per-import call.
    prompt = f"Parse this resume into structured YAML:\n\n{raw_text}"
    response = generate_text(
        prompt, system=EXTRACTION_SYSTEM_PROMPT, timeout=180, tier="primary"
    )

    profile_data = _parse_llm_yaml_response(response)
    # Belt-and-braces: even with a sharper prompt the LLM still drops
    # uncommon sections on a bad day. Scan the raw text deterministically
    # for any heading we didn't see in the parsed YAML and call the LLM
    # a second time, narrowly, to convert each missed block into
    # ``custom_sections`` entries. This is the only way to guarantee
    # nothing the candidate wrote silently disappears.
    profile_data = _backfill_custom_sections(profile_data, raw_text)

    # Save to file if output path specified
    if output_path:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        with open(output_path, "w", encoding="utf-8") as f:
            yaml.dump(
                profile_data, f, default_flow_style=False, allow_unicode=True, sort_keys=False
            )
        logger.info("Saved structured profile to %s", output_path)

    return profile_data
