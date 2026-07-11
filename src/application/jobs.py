"""Job search and application use cases shared by CLI and Web."""

from __future__ import annotations

import logging
import os
import re
import uuid
from datetime import UTC, datetime
from pathlib import Path
from urllib.parse import urlparse

from src.application.profile import get_active_profile_path, get_profile_path
from src.core.config import PROJECT_ROOT, load_config
from src.core.state_machine import ApplicationState, AppStatus

logger = logging.getLogger("autoapply.application.jobs")

SEARCH_METADATA_KEY = "_search_filters"
MATERIAL_TYPES = {
    "resume_pdf",
    "resume_docx",
    "resume_tex",
    "cover_letter_pdf",
    "cover_letter_docx",
    "cover_letter_tex",
}
ATS_TYPES = {
    "greenhouse",
    "lever",
    "ashby",
    "linkedin",
    "adzuna",
    "workday",
    "hn",
    "remotive",
    "company_site",
    "unknown",
}
EMPLOYMENT_TYPES = {"internship", "fulltime", "parttime", "contract", "coop", "unknown"}
SENIORITY_LEVELS = {"internship", "entry", "mid", "senior", "staff", "unknown"}

# Free tier is 250 calls/day and five overnight automation plans run daily
# -- capping each search to ~10 Adzuna calls keeps one search from burning
# a large chunk of the daily quota.
ADZUNA_MAX_CALLS_PER_SEARCH = 10

PAY_RANGE_RE = re.compile(
    r"(?:\$|usd\s*)(\d{1,3}(?:,\d{3})+|\d+(?:\.\d+)?)\s*(k|m)?\s*"
    r"(?:-|to)\s*(?:\$|usd\s*)?(\d{1,3}(?:,\d{3})+|\d+(?:\.\d+)?)\s*(k|m)?",
    re.IGNORECASE,
)
PAY_FLOOR_RE = re.compile(
    r"(?:from|starting at|minimum|at least)\s+(?:\$|usd\s*)"
    r"(\d{1,3}(?:,\d{3})+|\d+(?:\.\d+)?)\s*(k|m)?",
    re.IGNORECASE,
)
PAY_CAP_RE = re.compile(
    r"(?:up to|maximum|at most)\s+(?:\$|usd\s*)"
    r"(\d{1,3}(?:,\d{3})+|\d+(?:\.\d+)?)\s*(k|m)?",
    re.IGNORECASE,
)


async def search_jobs(
    *,
    profile: str | None = "default",
    config_dir: Path = PROJECT_ROOT / "config",
    no_parse: bool = False,
    use_llm: bool = False,
    source: str = "ats",
    ats: str | None = None,
    company: str | None = None,
    score: bool = False,
    keyword: str | None = None,
    keywords: list[str] | None = None,
    search_location: str | None = None,
    time_filter: str = "week",
    experience_levels: list[str] | None = None,
    employment_types: list[str] | None = None,
    location_types: list[str] | None = None,
    locations: list[str] | None = None,
    pay_operator: str | None = None,
    pay_amount: int | None = None,
    experience_operator: str | None = None,
    experience_years: int | None = None,
    education_levels: list[str] | None = None,
    max_pages: int = 20,
    no_enrich: bool = False,
    headless: bool = False,
    require_keyword_for_linkedin: bool = False,
    warn_on_missing_profile: bool = False,
    allow_public_linkedin_fallback: bool = False,
    include_views: bool = False,
    force_refresh: bool = False,
    use_job_index: bool = False,
) -> dict:
    jobs = []
    ats_jobs: list = []
    linkedin_jobs: list = []
    adzuna_jobs: list = []
    hn_jobs: list = []
    remotive_jobs: list = []
    errors: list[str] = []
    error_code: str | None = None
    counts = {
        "ats": 0,
        "linkedin": 0,
        "linkedin_external_ats": 0,
        "adzuna": 0,
        "hn": 0,
        "remotive": 0,
        "total": 0,
    }
    experience_levels = _normalize_experience_levels(_normalize_list(experience_levels))
    employment_types = _normalize_list(employment_types)
    location_types = _normalize_location_types(_normalize_list(location_types))
    locations = _normalize_list(locations)
    education_levels = _normalize_list(education_levels)
    keywords = _normalize_string_list(keywords)
    linkedin_keywords = _resolve_linkedin_keywords(keyword, keywords)
    linkedin_search_locations = _resolve_linkedin_search_locations(
        source=source,
        search_location=search_location,
        candidate_locations=locations,
    )
    job_index_events: list[dict] = []
    search_cache_policy = _search_cache_policy() if use_job_index else None
    linkedin_detail_limits = _linkedin_detail_fetch_limits()

    if source in ("ats", "all"):
        try:
            from src.intake.search import search_jobs as search_ats_jobs

            ats_jobs = search_ats_jobs(
                profile=profile,
                config_dir=config_dir,
                companies=_build_companies_filter(config_dir, ats, company),
                parse_jds=not no_parse,
                use_llm=use_llm,
                force_refresh=force_refresh,
            )
            # Keywords must narrow ATS results too. The boards return the
            # ENTIRE company board (no server-side search exists), so
            # without this the web UI's keyword box silently did nothing
            # for ATS sources and users got every role worldwide.
            if linkedin_keywords:
                ats_jobs = [
                    job for job in ats_jobs if _job_matches_keywords(job, linkedin_keywords)
                ]
            # Workday list items carry NO description at all (see
            # src/intake/workday.py) -- fetch full JDs for jobs that
            # survived keyword filtering, like LinkedIn detail fetches did,
            # capped per tenant so one board can't dominate the run.
            _enrich_workday_job_details(ats_jobs)
            counts["ats"] = len(ats_jobs)
            jobs.extend(ats_jobs)
        except Exception as exc:
            errors.append(f"ATS: {exc}")

    if source in ("all",):
        try:
            adzuna_jobs = _search_adzuna(linkedin_keywords, search_location)
            # Adzuna results MUST pass the normal ATS-style location/keyword
            # filters (invariant #2) -- treat like ATS, not like LinkedIn.
            # The boards-style keyword narrowing below is extra precision on
            # top of Adzuna's own server-side `what` search.
            if linkedin_keywords:
                adzuna_jobs = [
                    job for job in adzuna_jobs if _job_matches_keywords(job, linkedin_keywords)
                ]
            counts["adzuna"] = len(adzuna_jobs)
            jobs.extend(adzuna_jobs)
        except Exception as exc:
            errors.append(f"Adzuna: {exc}")

    if source in ("all",):
        try:
            from src.intake.hn_hiring import fetch_latest_hn_hiring_jobs

            hn_jobs = fetch_latest_hn_hiring_jobs(force_refresh=force_refresh)
            # Same treatment as ATS/Adzuna: HN's own "search" is really just
            # the whole thread, so keyword narrowing happens here. Location
            # narrowing is generic (applied to every non-LinkedIn source
            # later in _apply_search_filters) and needs no special-casing.
            if linkedin_keywords:
                hn_jobs = [job for job in hn_jobs if _job_matches_keywords(job, linkedin_keywords)]
            counts["hn"] = len(hn_jobs)
            jobs.extend(hn_jobs)
        except Exception as exc:
            errors.append(f"HN: {exc}")

    # Remotive is a remote-only jobs board -- only worth a call when the
    # caller actually wants remote roles, so it's gated on location_types
    # rather than always firing under "all" like Adzuna/HN.
    if source in ("all",) and "remote" in location_types:
        try:
            remotive_jobs = _search_remotive(linkedin_keywords[0] if linkedin_keywords else "")
            if linkedin_keywords:
                remotive_jobs = [
                    job for job in remotive_jobs if _job_matches_keywords(job, linkedin_keywords)
                ]
            counts["remotive"] = len(remotive_jobs)
            jobs.extend(remotive_jobs)
        except Exception as exc:
            errors.append(f"Remotive: {exc}")

    # LinkedIn is deliberately NOT part of "all" -- automated LinkedIn
    # access is permanently off after an account restriction (see
    # docs/CHANGELOG.md "safety: stop all automated LinkedIn access").
    # It only runs on an explicit, standalone source="linkedin" request
    # (e.g. a user manually testing a saved session), never swept in by
    # "all" for either the overnight automation (src/orchestration/
    # plan_run.py) or a general "search everything" web/CLI call.
    if source == "linkedin":
        if not linkedin_keywords:
            if require_keyword_for_linkedin:
                errors.append("LinkedIn search requires --keyword.")
        else:
            try:
                from src.intake.linkedin import LinkedInAuthRequiredError
                from src.intake.search import search_linkedin

                linkedin_jobs = []
                search_targets = linkedin_search_locations or [""]
                for location_target in search_targets:
                    linkedin_max_pages = _linkedin_max_pages(
                        max_pages,
                        search_location=location_target,
                        experience_levels=experience_levels,
                        employment_types=employment_types,
                        location_types=location_types,
                        locations=locations,
                        pay_operator=pay_operator,
                        experience_operator=experience_operator,
                        education_levels=education_levels,
                    )
                    search_kwargs = {
                        "keywords": linkedin_keywords,
                        "location": location_target,
                        "time_filter": _normalize_time_filter(time_filter),
                        "experience_levels": _map_linkedin_experience_levels(experience_levels),
                        "job_types": _map_linkedin_job_types(employment_types),
                        "max_pages": linkedin_max_pages,
                        "enrich_details": not no_enrich,
                        "max_keyword_detail_fetches": linkedin_detail_limits[
                            "max_keyword_detail_fetches"
                        ],
                        "max_redirect_detail_fetches": linkedin_detail_limits[
                            "max_redirect_detail_fetches"
                        ],
                        "headless": headless,
                        "filter_profile": profile,
                        "config_dir": config_dir,
                        "allow_public_fallback": allow_public_linkedin_fallback,
                    }
                    if use_job_index:
                        location_jobs, job_index_event = await _search_linkedin_with_job_index(
                            search_kwargs=search_kwargs,
                            force_refresh=force_refresh
                            or not (search_cache_policy or {}).get("enabled", True),
                            freshness_hours=(search_cache_policy or {}).get("ttl_hours", 24),
                        )
                        job_index_events.append(job_index_event)
                    else:
                        location_jobs = await search_linkedin(**search_kwargs)
                    linkedin_jobs.extend(location_jobs)

                linkedin_jobs = _dedupe_jobs_by_signature(linkedin_jobs)
                counts["linkedin"] = len(linkedin_jobs)
                counts["linkedin_external_ats"] = sum(
                    1
                    for job in linkedin_jobs
                    if job.ats_type in ("greenhouse", "lever", "ashby", "workday", "company_site")
                )
                jobs.extend(linkedin_jobs)
            except LinkedInAuthRequiredError as exc:
                if error_code is None:
                    error_code = "linkedin_auth_required"
                errors.append(f"LinkedIn: {exc}")
            except Exception as exc:
                errors.append(f"LinkedIn: {exc}")

    # Cross-source dedupe: with source="all" the same posting often arrives
    # from both an ATS board and the LinkedIn search. ATS copies win (they
    # carry the full JD and a direct apply URL) because they were appended
    # first and _dedupe_jobs_by_signature keeps the first occurrence.
    jobs = _dedupe_jobs_by_signature(jobs)

    # 2026-07-08: self-growing board registry. LinkedIn postings that
    # redirect to Greenhouse/Lever reveal boards not yet in
    # companies.yaml; harvest the slugs so tomorrow's ATS pass covers
    # them. Best-effort — never blocks or fails the search.
    # 2026-07-10: Adzuna's redirect_url is included too, though in
    # practice Adzuna's landing page (adzuna.com/land/ad/...) blocks
    # unauthenticated fetches -- see docs/CHANGELOG.md -- so discovery
    # from Adzuna jobs is currently a no-op most of the time. Harmless
    # to leave wired in for whenever that changes.
    if linkedin_jobs or adzuna_jobs:
        try:
            from src.intake.board_discovery import register_discovered_boards

            register_discovered_boards(linkedin_jobs + adzuna_jobs, config_dir)
        except Exception:  # noqa: BLE001
            logger.debug("Board discovery skipped", exc_info=True)

    raw_total = len(jobs)
    fetched_jobs = list(jobs)

    if include_views and fetched_jobs:
        _prepare_jobs_for_search_filters(fetched_jobs, use_llm=use_llm)

    scored = False
    if score and fetched_jobs and include_views:
        scored, scoring_errors = _score_jobs(
            fetched_jobs, warn_on_missing_profile=warn_on_missing_profile
        )
        errors.extend(scoring_errors)

    jobs = _apply_search_filters(
        list(fetched_jobs),
        experience_levels=experience_levels,
        employment_types=employment_types,
        location_types=location_types,
        locations=locations,
        search_location=search_location,
        searched_linkedin_locations=linkedin_search_locations,
        pay_operator=pay_operator,
        pay_amount=pay_amount,
        experience_operator=experience_operator,
        experience_years=experience_years,
        education_levels=education_levels,
        use_llm=use_llm,
    )

    if score and jobs and not include_views:
        scored, scoring_errors = _score_jobs(jobs, warn_on_missing_profile=warn_on_missing_profile)
        errors.extend(scoring_errors)

    if score and include_views:
        jobs.sort(key=_job_sort_key)
        ats_jobs.sort(key=_job_sort_key)
        linkedin_jobs.sort(key=_job_sort_key)
        fetched_jobs.sort(key=_job_sort_key)

    counts["total"] = len(jobs)

    return {
        "search_params": {
            "profile": profile,
            "config_dir": str(config_dir),
            "no_parse": no_parse,
            "use_llm": use_llm,
            "source": source,
            "ats": ats,
            "company": company,
            "score": score,
            "keyword": keyword or "",
            "keywords": linkedin_keywords,
            "location": search_location or "",
            "search_locations": linkedin_search_locations,
            "time_filter": time_filter,
            "experience_levels": experience_levels,
            "employment_types": employment_types,
            "location_types": location_types,
            "locations": locations,
            "pay_operator": pay_operator,
            "pay_amount": pay_amount,
            "experience_operator": experience_operator,
            "experience_years": experience_years,
            "education_levels": education_levels,
            "max_pages": max_pages,
            "no_enrich": no_enrich,
            "headless": headless,
        },
        "jobs": [serialize_job(job) for job in jobs],
        "views": {
            "shown": [serialize_job(job) for job in jobs],
            "fetched": [serialize_job(job) for job in fetched_jobs],
            "linkedin": [serialize_job(job) for job in linkedin_jobs],
            "ats": [serialize_job(job) for job in ats_jobs],
        }
        if include_views
        else None,
        "errors": errors,
        "error": "; ".join(errors) or None,
        "error_code": error_code,
        "counts": {
            **counts,
            "raw_total": raw_total,
            "filtered_total": len(jobs),
            "total": len(jobs),
        },
        "scored": scored,
        "job_index": {
            "enabled": bool(use_job_index),
            "policy": search_cache_policy,
            "events": job_index_events,
        },
    }


async def _search_linkedin_with_job_index(
    *,
    search_kwargs: dict,
    force_refresh: bool,
    freshness_hours: int,
) -> tuple[list, dict]:
    """Run one LinkedIn search location through the Phase 13 Job Index.

    Fresh fetches still return the full ``RawJob`` objects from the scraper,
    then persist immutable snapshots so later cache hits can reconstruct the
    same UI shape from ``job_postings.latest_snapshot_id``.
    """
    from src.cache import get_cache
    from src.core.database import get_session_factory
    from src.intake.search import search_linkedin
    from src.jobs.enrich import enrich_posting
    from src.jobs.search import cached_search
    from src.jobs.store import JobIndexStore

    scraped_jobs: list = []
    params = _linkedin_job_index_params(search_kwargs)

    async def fetch_and_capture() -> list:
        result = await search_linkedin(**search_kwargs)
        scraped_jobs[:] = list(result)
        return scraped_jobs

    try:
        session_factory = get_session_factory(load_config())
        with session_factory() as session, session.begin():
            store = JobIndexStore(session)
            outcome = await cached_search(
                store=store,
                cache=get_cache(),
                source="linkedin",
                params=params,
                fetch_fn=fetch_and_capture,
                max_pages=search_kwargs.get("max_pages"),
                force_refresh=force_refresh,
                freshness_hours=freshness_hours,
            )
            if scraped_jobs:
                from src.core.models import JobPosting

                for job in scraped_jobs:
                    enriched = enrich_posting(
                        store=store,
                        source=job.source,
                        source_id=job.source_id,
                        company=job.company,
                        content=_raw_job_content(job),
                    )
                    # Ghost-posting signal: even on a fresh scrape the
                    # posting may have been first seen weeks ago. Stash
                    # first_seen_at so the quality multiplier can
                    # penalize long-lived (likely evergreen/ghost) jobs.
                    try:
                        posting = session.get(JobPosting, enriched.posting_id)
                        if posting is not None and posting.first_seen_at:
                            job.raw_data["first_seen_at"] = posting.first_seen_at.isoformat()
                    except Exception:  # noqa: BLE001 -- signal is best-effort
                        logger.debug("first_seen_at stash skipped", exc_info=True)
                jobs = scraped_jobs
            else:
                jobs = _raw_jobs_from_index_postings(session, outcome.postings)

            if outcome.refresh_failed and not jobs:
                raise RuntimeError(outcome.last_error or "LinkedIn refresh failed")

            return jobs, {
                "ok": True,
                "cached": outcome.cached,
                "stale": outcome.stale,
                "force_refresh": force_refresh,
                "query_id": str(outcome.query_id),
                "last_run_at": _isoformat(outcome.last_run_at),
                "last_success_at": _isoformat(outcome.last_success_at),
                "last_error": outcome.last_error,
                "counts": outcome.counts,
                "location": search_kwargs.get("location") or "",
            }
    except Exception as exc:
        # The Job Index should improve freshness, not make search unusable if
        # the operator forgot a migration. Fall back to a direct live pull.
        logger.warning("Job Index search path failed; falling back to live search: %s", exc)
        jobs = list(await search_linkedin(**search_kwargs))
        return jobs, {
            "ok": False,
            "cached": False,
            "stale": False,
            "force_refresh": True,
            "fallback_live": True,
            "error": str(exc),
            "location": search_kwargs.get("location") or "",
        }


def _enrich_workday_job_details(ats_jobs: list) -> None:
    """Fetch full JDs for Workday jobs that survived keyword filtering.

    Workday's list endpoint carries no description at all (see
    src/intake/workday.py's module docstring), so this is load-bearing
    for Workday jobs, not an optional enrichment. Capped per tenant
    (``DETAIL_FETCH_CAP``) so one large board can't dominate a run.
    Mutates ``ats_jobs`` in place; best-effort -- a fetch failure leaves
    the job's description unchanged (already ``None``).
    """
    from collections import defaultdict
    from concurrent.futures import ThreadPoolExecutor, as_completed

    from src.intake.workday import DETAIL_FETCH_CAP, WorkdayScraper

    by_tenant: dict[str, list[int]] = defaultdict(list)
    for idx, job in enumerate(ats_jobs):
        if job.source == "workday":
            tenant = (job.raw_data or {}).get("workday_tenant") or ""
            by_tenant[tenant].append(idx)
    if not by_tenant:
        return

    targets: list[int] = []
    for tenant, indices in by_tenant.items():
        capped = indices[:DETAIL_FETCH_CAP]
        if len(indices) > DETAIL_FETCH_CAP:
            logger.warning(
                "Workday detail-fetch cap (%d) reached for tenant '%s'; "
                "%d/%d keyword-surviving jobs left without a full JD",
                DETAIL_FETCH_CAP,
                tenant,
                len(indices) - DETAIL_FETCH_CAP,
                len(indices),
            )
        targets.extend(capped)
    if not targets:
        return

    with WorkdayScraper() as scraper:
        with ThreadPoolExecutor(max_workers=min(8, len(targets))) as pool:
            futures = {
                pool.submit(scraper.fetch_job_detail, ats_jobs[idx]): idx for idx in targets
            }
            for future in as_completed(futures):
                idx = futures[future]
                try:
                    ats_jobs[idx] = future.result()
                except Exception:  # noqa: BLE001 -- best-effort, never fail the search
                    logger.debug("Workday detail fetch skipped for job %d", idx, exc_info=True)


def _adzuna_settings() -> dict:
    raw = load_config().get("adzuna", {})
    if not isinstance(raw, dict):
        raw = {}
    try:
        results_per_query = max(int(raw.get("results_per_query", 50)), 1)
    except (TypeError, ValueError):
        results_per_query = 50
    return {
        "enabled": bool(raw.get("enabled", False)),
        "app_id": str(raw.get("app_id") or "").strip(),
        "app_key_env": str(raw.get("app_key_env") or "AUTOAPPLY_ADZUNA_KEY").strip(),
        "country": str(raw.get("country") or "us").strip(),
        "results_per_query": results_per_query,
    }


def _search_remotive(keyword: str) -> list:
    """Query the Remotive remote-jobs API once.

    Only the first search keyword is used (Remotive's own ``search``
    param handles one free-text query, unlike Adzuna's per-keyword
    ``what``) -- the caller re-applies keyword filtering client-side
    across every keyword afterward regardless, same as every other
    source here.
    """
    from src.intake.remotive import RemotiveScraper

    with RemotiveScraper() as scraper:
        return scraper.fetch_jobs(keyword)


def _search_adzuna(keywords: list[str], location: str | None) -> list:
    """Query the Adzuna search API, once per keyword, capped at
    ``ADZUNA_MAX_CALLS_PER_SEARCH`` calls.
    """
    settings = _adzuna_settings()
    if not settings["enabled"]:
        return []
    app_key = os.environ.get(settings["app_key_env"]) or ""
    if not settings["app_id"] or not app_key:
        logger.warning(
            "Adzuna enabled but app_id/%s is not set; skipping Adzuna search",
            settings["app_key_env"],
        )
        return []

    from src.intake.adzuna import AdzunaScraper

    search_keywords = keywords or [""]
    if len(search_keywords) > ADZUNA_MAX_CALLS_PER_SEARCH:
        logger.warning(
            "Adzuna call cap (%d) reached; searching only the first %d of %d keywords",
            ADZUNA_MAX_CALLS_PER_SEARCH,
            ADZUNA_MAX_CALLS_PER_SEARCH,
            len(search_keywords),
        )
        search_keywords = search_keywords[:ADZUNA_MAX_CALLS_PER_SEARCH]

    all_jobs: list = []
    with AdzunaScraper(
        app_id=settings["app_id"], app_key=app_key, country=settings["country"]
    ) as scraper:
        for kw in search_keywords:
            try:
                all_jobs.extend(
                    scraper.search(
                        keyword=kw,
                        location=location or "",
                        results_per_page=settings["results_per_query"],
                    )
                )
            except Exception as exc:  # noqa: BLE001 -- one bad keyword must not fail the search
                logger.warning("Adzuna search failed for keyword=%r: %s", kw, exc)

    return _dedupe_jobs_by_signature(all_jobs)


def _search_cache_policy() -> dict:
    raw = load_config().get("search_cache", {})
    try:
        ttl_hours = max(int(raw.get("ttl_hours", 24)), 1)
    except (TypeError, ValueError):
        ttl_hours = 24
    return {
        "enabled": bool(raw.get("enabled", True)),
        "ttl_hours": ttl_hours,
    }


def _linkedin_detail_fetch_limits() -> dict[str, int]:
    raw = load_config().get("linkedin", {})
    if not isinstance(raw, dict):
        raw = {}
    return {
        "max_keyword_detail_fetches": _positive_int_setting(
            raw.get("max_keyword_detail_fetches"), 5000
        ),
        "max_redirect_detail_fetches": _positive_int_setting(
            raw.get("max_redirect_detail_fetches"), 5000
        ),
    }


def _positive_int_setting(value, default: int) -> int:
    try:
        parsed = int(value)
    except (TypeError, ValueError):
        return default
    return max(parsed, 0)


def _linkedin_job_index_params(search_kwargs: dict) -> dict:
    return {
        "keywords": search_kwargs.get("keywords") or [],
        "location": search_kwargs.get("location") or "",
        "time_filter": search_kwargs.get("time_filter") or "week",
        "experience_levels": search_kwargs.get("experience_levels") or [],
        "job_types": search_kwargs.get("job_types") or [],
        "max_pages": search_kwargs.get("max_pages"),
        "enrich_details": bool(search_kwargs.get("enrich_details", True)),
        "max_keyword_detail_fetches": search_kwargs.get("max_keyword_detail_fetches"),
        "max_redirect_detail_fetches": search_kwargs.get("max_redirect_detail_fetches"),
        "filter_profile": search_kwargs.get("filter_profile") or "",
        "allow_public_fallback": bool(search_kwargs.get("allow_public_fallback", False)),
    }


def _raw_job_content(job) -> dict:
    requirements = (
        job.requirements.model_dump()
        if hasattr(job.requirements, "model_dump")
        else job.requirements
    )
    raw_data = dict(job.raw_data or {})
    raw_data.update(
        {
            "source": job.source,
            "source_id": job.source_id,
            "company": job.company,
            "title": job.title,
            "location": job.location,
            "employment_type": job.employment_type,
            "seniority": job.seniority,
            "description": job.description,
            "requirements": requirements,
            "application_url": job.application_url,
            "ats_type": job.ats_type,
        }
    )
    return {
        "title": job.title,
        "location": job.location,
        "employment_type": job.employment_type,
        "seniority": job.seniority,
        "description": job.description,
        "requirements": requirements,
        "application_url": job.application_url,
        "raw_data": raw_data,
    }


def _raw_jobs_from_index_postings(session, postings: list) -> list:
    from src.core.models import JobSnapshot

    jobs = []
    for posting in postings:
        snapshot = None
        if posting.latest_snapshot_id:
            snapshot = session.get(JobSnapshot, posting.latest_snapshot_id)
        jobs.append(_raw_job_from_index_posting(posting, snapshot))
    return jobs


def _raw_job_from_index_posting(posting, snapshot):
    from src.intake.schema import JobRequirements, RawJob

    raw = dict(getattr(snapshot, "raw_data", None) or {})
    # Ghost-posting signal for the scorer's quality multiplier.
    first_seen = getattr(posting, "first_seen_at", None)
    if first_seen is not None:
        raw["first_seen_at"] = first_seen.isoformat()
    requirements_payload = getattr(snapshot, "requirements", None) or raw.get("requirements") or {}
    try:
        requirements = JobRequirements(**requirements_payload)
    except TypeError:
        requirements = JobRequirements()

    employment_type = (
        raw.get("employment_type") or getattr(snapshot, "employment_type", None) or "unknown"
    )
    seniority = raw.get("seniority") or getattr(snapshot, "seniority", None) or "unknown"
    ats_type = raw.get("ats_type") or raw.get("source") or posting.source
    source = posting.source if posting.source in ATS_TYPES else "linkedin"

    return RawJob(
        id=uuid.UUID(raw["id"]) if raw.get("id") else uuid.uuid4(),
        source=source,
        source_id=posting.source_id,
        company=posting.company,
        title=raw.get("title") or getattr(snapshot, "title", None) or "Unknown Role",
        location=raw.get("location") or getattr(snapshot, "location", None),
        employment_type=employment_type if employment_type in EMPLOYMENT_TYPES else "unknown",
        seniority=seniority if seniority in SENIORITY_LEVELS else "unknown",
        description=raw.get("description") or getattr(snapshot, "description", None),
        requirements=requirements,
        application_url=(
            raw.get("application_url")
            or getattr(snapshot, "application_url", None)
            or posting.canonical_url
        ),
        ats_type=ats_type if ats_type in ATS_TYPES else "unknown",
        raw_data=raw,
        discovered_at=datetime.now(UTC),
    )


async def get_linkedin_session_status(force_refresh: bool = False) -> dict:
    try:
        from src.intake.linkedin import (
            get_linkedin_session_status as get_linkedin_session_status_intake,
        )

        return await get_linkedin_session_status_intake(force_refresh=force_refresh)
    except Exception as exc:
        logger.exception("Failed to inspect LinkedIn session")
        return {
            "ok": False,
            "authenticated": False,
            "has_session_data": False,
            "message": f"Failed to inspect LinkedIn session: {exc}",
            "error": str(exc),
            "error_code": "linkedin_session_status_failed",
        }


async def connect_linkedin_session() -> dict:
    try:
        from src.intake.linkedin import connect_linkedin_session as connect_linkedin_session_intake

        return await connect_linkedin_session_intake()
    except Exception as exc:
        logger.exception("Failed to connect LinkedIn session")
        return {
            "ok": False,
            "authenticated": False,
            "has_session_data": False,
            "message": f"Failed to connect LinkedIn session: {exc}",
            "error": str(exc),
            "error_code": "linkedin_session_connect_failed",
        }


async def resolve_manual_apply_url(url: str) -> dict:
    if not _is_linkedin_url(url):
        return {
            "ok": True,
            "url": url,
            "source_url": url,
            "ats_url": url if _detect_ats_from_url(url) else None,
            "error": None,
            "error_code": None,
        }

    try:
        from src.intake.linkedin import resolve_linkedin_apply_target

        return await resolve_linkedin_apply_target(url)
    except Exception as exc:
        logger.exception("Failed to resolve manual apply URL")
        return {
            "ok": False,
            "url": url,
            "source_url": url,
            "ats_url": None,
            "error": str(exc),
            "error_code": "manual_apply_resolution_failed",
        }


def clear_linkedin_session() -> dict:
    try:
        from src.intake.linkedin import clear_linkedin_session as clear_linkedin_session_intake

        return clear_linkedin_session_intake()
    except Exception as exc:
        logger.exception("Failed to clear LinkedIn session")
        return {
            "ok": False,
            "authenticated": False,
            "has_session_data": True,
            "message": f"Failed to clear LinkedIn session: {exc}",
            "error": str(exc),
            "error_code": "linkedin_session_clear_failed",
        }


def preview_batch_jobs(*, profile: str = "default", top_n: int = 5) -> dict:
    selected_jobs, errors, total_matches = _select_batch_jobs(profile, top_n)
    return {
        "profile": profile,
        "top_n": top_n,
        "available_matches": total_matches,
        "selected_jobs": [
            serialize_job(job, match_score=match_score) for job, match_score in selected_jobs
        ],
        "errors": errors,
    }


async def apply_to_url(
    *,
    url: str,
    auto_submit: bool = False,
    headless: bool = True,
    dry_run: bool = False,
) -> dict:
    payload = {
        "mode": "url",
        "input": {"url": url},
    }

    resolved_url = url
    if _is_linkedin_url(url):
        resolved = await resolve_manual_apply_url(url)
        resolved_target = resolved.get("ats_url") or resolved.get("url")
        if resolved_target and not _is_linkedin_url(resolved_target):
            resolved_url = resolved_target
            payload["resolved_url"] = resolved_url

    # If we are still pointing at a LinkedIn URL at this point it means
    # the LinkedIn resolver couldn't find a real external apply target.
    # That's the Easy-Apply-only case: refuse with a clear message
    # rather than letting the pipeline pretend a random link on the
    # page is the apply form. When LinkedIn Easy Apply is eventually
    # automated we'll branch here instead of erroring out.
    if _is_linkedin_url(resolved_url):
        return {
            **payload,
            "ok": False,
            "status": None,
            "job": None,
            "tracking_id": None,
            "result": None,
            "artifacts": _empty_artifacts(),
            "error": _linkedin_easy_apply_message(),
            "error_code": "linkedin_easy_apply_only",
            "dry_run": dry_run,
        }

    ats_type = _detect_ats_from_url(resolved_url)
    if not ats_type:
        return {
            **payload,
            "ok": False,
            "status": None,
            "job": None,
            "tracking_id": None,
            "result": None,
            "artifacts": _empty_artifacts(),
            "error": _unsupported_ats_message(resolved_url),
            "error_code": "unsupported_ats",
            "dry_run": dry_run,
        }

    resolved_url = _normalize_application_url_for_ats(resolved_url, ats_type)
    payload["resolved_url"] = resolved_url

    try:
        job, hydrated = _load_job_for_application(resolved_url, ats_type)
    except Exception as exc:
        return {
            **payload,
            "ok": False,
            "status": None,
            "job": None,
            "tracking_id": None,
            "result": None,
            "artifacts": _empty_artifacts(),
            "error": f"Failed to load job context: {exc}",
            "error_code": "job_load_failed",
            "dry_run": dry_run,
        }

    if not hydrated and ats_type in ("ashby", "workday", "company_site"):
        # Generic / company-site / Workday URLs cannot be hydrated from the
        # URL alone (we don't have a parser that recovers the JD), so we'd
        # be tailoring against title="Unknown Role" with no description.
        # Force the caller to apply via a stored job id (which carries real
        # JD context) instead of generating useless materials.
        return {
            **payload,
            "ok": False,
            "status": None,
            "job": serialize_job(job),
            "tracking_id": None,
            "result": None,
            "artifacts": _empty_artifacts(),
            "error": _no_job_context_message(ats_type),
            "error_code": "job_context_required",
            "dry_run": dry_run,
        }

    # Final EasyApply guard: the LinkedIn scraper persists
    # ``raw_data.linkedin_easy_apply_only`` whenever the only apply
    # path on the LinkedIn page was Easy Apply. Even if a previous
    # scrape pass stored a bogus external URL (Fitzrovia's marketing
    # homepage), the flag tells us not to send the form-filler there.
    if _job_is_easy_apply_only(job):
        return {
            **payload,
            "ok": False,
            "status": None,
            "job": serialize_job(job),
            "tracking_id": None,
            "result": None,
            "artifacts": _empty_artifacts(),
            "error": _linkedin_easy_apply_message(),
            "error_code": "linkedin_easy_apply_only",
            "dry_run": dry_run,
        }

    detected_ats = _detect_ats_from_url(job.application_url or "")
    if detected_ats:
        job.ats_type = detected_ats

    profile_data = _load_profile()
    if not profile_data:
        return {
            **payload,
            "ok": False,
            "status": None,
            "job": serialize_job(job),
            "tracking_id": None,
            "result": None,
            "artifacts": _empty_artifacts(),
            "error": "Profile not configured.",
            "error_code": "profile_missing",
            "dry_run": dry_run,
        }

    return await _run_application_for_job(
        job=job,
        profile_data=profile_data,
        auto_submit=auto_submit,
        headless=headless,
        dry_run=dry_run,
        mode="url",
        input_payload=payload["input"],
    )


async def apply_to_job_id(
    *,
    job_id: str,
    auto_submit: bool = False,
    headless: bool = True,
    dry_run: bool = False,
) -> dict:
    payload = {
        "mode": "job_id",
        "input": {"job_id": job_id},
    }

    try:
        job_uuid = uuid.UUID(job_id)
    except ValueError:
        return {
            **payload,
            "ok": False,
            "status": None,
            "job": None,
            "tracking_id": None,
            "result": None,
            "artifacts": _empty_artifacts(),
            "error": f"Invalid job ID format: {job_id}",
            "error_code": "invalid_job_id",
            "dry_run": dry_run,
        }

    try:
        from src.core.config import load_config
        from src.core.database import get_session_factory
        from src.core.models import Job

        session_factory = get_session_factory(load_config())
        with session_factory() as session:
            db_job = session.get(Job, job_uuid)
            if db_job is None:
                return {
                    **payload,
                    "ok": False,
                    "status": None,
                    "job": None,
                    "tracking_id": None,
                    "result": None,
                    "artifacts": _empty_artifacts(),
                    "error": f"Job {job_id} not found in database.",
                    "error_code": "job_not_found",
                    "dry_run": dry_run,
                }

            if not db_job.application_url:
                return {
                    **payload,
                    "ok": False,
                    "status": None,
                    "job": serialize_job(_job_to_raw_job(db_job)),
                    "tracking_id": None,
                    "result": None,
                    "artifacts": _empty_artifacts(),
                    "error": f"Job {db_job.title} at {db_job.company} has no application URL.",
                    "error_code": "missing_application_url",
                    "dry_run": dry_run,
                }

            job = _job_to_raw_job(db_job)

            # Same Easy-Apply-only guard as ``apply_to_url`` -- catch
            # the Fitzrovia regression when applying by stored job id
            # too, so the apply pipeline never tries to form-fill a
            # company's marketing homepage.
            if _job_is_easy_apply_only(job) or _is_linkedin_url(job.application_url or ""):
                return {
                    **payload,
                    "ok": False,
                    "status": None,
                    "job": serialize_job(job),
                    "tracking_id": None,
                    "result": None,
                    "artifacts": _empty_artifacts(),
                    "error": _linkedin_easy_apply_message(),
                    "error_code": "linkedin_easy_apply_only",
                    "dry_run": dry_run,
                }
    except Exception as exc:
        return {
            **payload,
            "ok": False,
            "status": None,
            "job": None,
            "tracking_id": None,
            "result": None,
            "artifacts": _empty_artifacts(),
            "error": f"Failed to load job from database: {exc}",
            "error_code": "job_load_failed",
            "dry_run": dry_run,
        }

    detected_ats = _detect_ats_from_url(job.application_url or "")
    if detected_ats:
        job.ats_type = detected_ats

    profile_data = _load_profile()
    if not profile_data:
        return {
            **payload,
            "ok": False,
            "status": None,
            "job": serialize_job(job),
            "tracking_id": None,
            "result": None,
            "artifacts": _empty_artifacts(),
            "error": "Profile not configured.",
            "error_code": "profile_missing",
            "dry_run": dry_run,
        }

    return await _run_application_for_job(
        job=job,
        profile_data=profile_data,
        auto_submit=auto_submit,
        headless=headless,
        dry_run=dry_run,
        mode="job_id",
        input_payload=payload["input"],
    )


async def generate_material_for_job(
    *,
    job_payload: dict,
    material_type: str,
    use_llm: bool = False,
    template_id: str | None = None,
    profile_id: str | None = None,
    strategy: str | None = None,
    source_document_id: str | None = None,
    patch_aggressiveness: str | None = None,
    patch_allow_reorder_sections: bool | None = None,
    patch_allow_add_remove_bullets: bool | None = None,
) -> dict:
    """Generate one selected application artifact for a web search result.

    Phase 17.8: ``strategy`` + ``source_document_id`` are optional
    per-call overrides. When omitted, the user's saved defaults for
    this document type win (see
    :mod:`src.application.material_defaults`).
    """
    from src.application.material_defaults import resolve_material_choice

    material_type = (material_type or "").strip()
    if material_type not in MATERIAL_TYPES:
        return {
            "ok": False,
            "job": None,
            "material_type": material_type,
            "artifact": None,
            "artifacts": _empty_material_artifacts(),
            "requirements": None,
            "error": "Unsupported material type.",
            "error_code": "invalid_material_type",
        }
    template_id = _clean_optional_web_payload_id(template_id)
    profile_id = _clean_optional_web_payload_id(profile_id)
    document_type = "cover_letter" if material_type.startswith("cover_letter") else "resume"
    try:
        choice = resolve_material_choice(
            document_type=document_type,
            override_strategy=strategy,
            override_template_id=template_id,
            override_document_id=source_document_id,
            override_patch_aggressiveness=patch_aggressiveness,
            override_patch_allow_reorder_sections=patch_allow_reorder_sections,
            override_patch_allow_add_remove_bullets=patch_allow_add_remove_bullets,
        )
    except ValueError as exc:
        return {
            "ok": False,
            "job": None,
            "material_type": material_type,
            "artifact": None,
            "artifacts": _empty_material_artifacts(),
            "requirements": None,
            "error": str(exc),
            "error_code": "invalid_strategy",
        }
    template_id = choice["template_id"] or template_id
    resolved_strategy = choice["strategy"]
    resolved_document_id = choice["document_id"]
    resolved_patch_aggressiveness = choice["patch_aggressiveness"]
    resolved_patch_allow_reorder = choice["patch_allow_reorder_sections"]
    resolved_patch_allow_add_remove = choice["patch_allow_add_remove_bullets"]

    try:
        job = _raw_job_from_web_payload(job_payload, use_llm=use_llm)
    except Exception as exc:
        return {
            "ok": False,
            "job": None,
            "material_type": material_type,
            "artifact": None,
            "artifacts": _empty_material_artifacts(),
            "requirements": None,
            "error": str(exc),
            "error_code": "invalid_job_payload",
        }

    profile_data = _load_profile(profile_id)
    if not profile_data:
        return {
            "ok": False,
            "job": serialize_job(job),
            "material_type": material_type,
            "artifact": None,
            "artifacts": _empty_material_artifacts(),
            "requirements": job.requirements.model_dump(),
            "error": "Profile not configured.",
            "error_code": "profile_missing",
        }

    try:
        material_result = _generate_selected_material(
            profile_data,
            job,
            material_type,
            template_id=template_id,
            strategy=resolved_strategy,
            source_document_id=resolved_document_id,
            patch_aggressiveness=resolved_patch_aggressiveness,
            patch_allow_reorder_sections=resolved_patch_allow_reorder,
            patch_allow_add_remove_bullets=resolved_patch_allow_add_remove,
        )
    except ValueError as exc:
        return {
            "ok": False,
            "job": serialize_job(job),
            "material_type": material_type,
            "artifact": None,
            "artifacts": _empty_material_artifacts(),
            "document": None,
            "validation": None,
            "requirements": job.requirements.model_dump(),
            "error": str(exc),
            "error_code": "invalid_template_id",
        }
    except Exception as exc:
        logger.warning(
            "Failed to generate %s for %s at %s: %s",
            material_type,
            job.title,
            job.company,
            exc,
        )
        return {
            "ok": False,
            "job": serialize_job(job),
            "material_type": material_type,
            "artifact": None,
            "artifacts": _empty_material_artifacts(),
            "document": None,
            "validation": None,
            "requirements": job.requirements.model_dump(),
            "error": f"Failed to generate material: {exc}",
            "error_code": "material_generation_failed",
        }

    artifacts = material_result["artifacts"]
    document = material_result.get("document")
    validation = material_result.get("validation")
    template = material_result.get("template")
    path = artifacts.get(material_type)
    if not path:
        logger.error(
            "artifact lookup failed: material_type=%s artifacts=%s "
            "strategy=%s resolved_strategy=%s source_document_id=%s "
            "strategy_notes=%s",
            material_type,
            {k: str(v) if v else None for k, v in artifacts.items()},
            strategy,
            resolved_strategy,
            resolved_document_id,
            material_result.get("strategy_notes"),
        )
        return {
            "ok": False,
            "job": serialize_job(job),
            "material_type": material_type,
            "artifact": None,
            "artifacts": _stringify_material_artifacts(artifacts),
            "document": _serialize_generation_model(document),
            "validation": _serialize_generation_model(validation),
            "template": template,
            "requirements": job.requirements.model_dump(),
            "error": "The selected artifact could not be generated.",
            "error_code": "material_generation_failed",
            "strategy_notes": material_result.get("strategy_notes") or [],
        }

    # Anything that pydantic-serializes here can in principle blow up
    # on a value the IR / validation pipeline didn't expect (e.g. a
    # Path leaking out of model_dump or a custom type the JSON mode
    # doesn't recognize). The original code let those bubble up
    # uncaught -- FastAPI then rendered an opaque ``Internal Server
    # Error`` with no traceback in our log. Wrap the post-generation
    # serialization in its own try/except so we get a structured
    # error + a real log line if a future regression bites.
    try:
        serialized_document = _serialize_generation_model(document)
        serialized_validation = _serialize_generation_model(validation)
        artifact = _serialize_material_artifact(material_type, path)
        serialized_artifacts = _stringify_material_artifacts(artifacts)
        serialized_job = serialize_job(job)
        requirements = job.requirements.model_dump()
        version = _save_generation_version(
            job=serialized_job,
            material_type=material_type,
            artifact=artifact,
            artifacts=serialized_artifacts,
            document=serialized_document,
            validation=serialized_validation,
            requirements=requirements,
        )
    except Exception as exc:
        logger.exception(
            "post-generation serialization failed for material_type=%s",
            material_type,
        )
        return {
            "ok": False,
            "job": serialize_job(job) if job else None,
            "material_type": material_type,
            "artifact": None,
            "artifacts": _stringify_material_artifacts(artifacts),
            "document": None,
            "validation": None,
            "template": template,
            "requirements": job.requirements.model_dump() if job else None,
            "error": (
                f"Generation succeeded but the response could not be "
                f"serialized: {type(exc).__name__}: {exc}"
            ),
            "error_code": "serialization_failed",
            "strategy_notes": material_result.get("strategy_notes") or [],
        }

    return {
        "ok": True,
        "job": serialized_job,
        "material_type": material_type,
        "artifact": artifact,
        "artifacts": serialized_artifacts,
        "document": serialized_document,
        "validation": serialized_validation,
        "template": template,
        "requirements": requirements,
        "version": version,
        "strategy": resolved_strategy,
        "strategy_source": choice["source"],
        "strategy_notes": material_result.get("strategy_notes") or [],
        "source_document_id": resolved_document_id,
        "patch_aggressiveness": resolved_patch_aggressiveness,
        "patch_allow_reorder_sections": resolved_patch_allow_reorder,
        "patch_allow_add_remove_bullets": resolved_patch_allow_add_remove,
        "error": None,
        "error_code": None,
    }


def list_material_templates() -> dict:
    """List available resume and cover-letter template packages for the UI."""
    from src.documents.templates import list_template_packages

    return {"templates": list_template_packages()}


def upload_material_template(
    *,
    document_type: str,
    filename: str,
    content: bytes,
    template_name: str | None = None,
) -> dict:
    """Save an uploaded DOCX or LaTeX material template package."""
    from src.documents.templates import save_uploaded_template_package

    if document_type not in {"resume", "cover_letter"}:
        return {
            "ok": False,
            "error": "Unsupported template document type.",
            "error_code": "invalid_document_type",
        }
    try:
        template = save_uploaded_template_package(
            document_type=document_type,
            filename=filename,
            content=content,
            template_name=template_name,
        )
    except ValueError as exc:
        return {"ok": False, "error": str(exc), "error_code": "invalid_template_upload"}
    except Exception as exc:
        return {"ok": False, "error": str(exc), "error_code": "template_upload_failed"}

    return {"ok": True, "template": template, **list_material_templates()}


def create_material_template(
    *,
    document_type: str,
    template_name: str | None = None,
    description: str | None = None,
) -> dict:
    """Create a blank editable LaTeX material template package."""
    from src.documents.templates import create_latex_template_package

    if document_type not in {"resume", "cover_letter"}:
        return {
            "ok": False,
            "error": "Unsupported template document type.",
            "error_code": "invalid_document_type",
        }
    try:
        template = create_latex_template_package(
            document_type=document_type,
            template_name=template_name,
            description=description,
        )
    except ValueError as exc:
        return {"ok": False, "error": str(exc), "error_code": "invalid_template"}
    except Exception as exc:
        return {"ok": False, "error": str(exc), "error_code": "template_create_failed"}

    return {"ok": True, "template": template, **list_material_templates()}


def get_material_template(*, document_type: str, template_id: str) -> dict:
    """Return a material template package, including editable content for LaTeX."""
    from src.documents.templates import get_template_package_detail

    if document_type not in {"resume", "cover_letter"}:
        return {
            "ok": False,
            "error": "Unsupported template document type.",
            "error_code": "invalid_document_type",
        }
    try:
        template = get_template_package_detail(document_type, template_id)
    except ValueError as exc:
        return {"ok": False, "error": str(exc), "error_code": "invalid_template_id"}
    except Exception as exc:
        return {"ok": False, "error": str(exc), "error_code": "template_load_failed"}

    return {"ok": True, "template": template}


def update_material_template(
    *,
    document_type: str,
    template_id: str,
    content: str,
    template_name: str | None = None,
    description: str | None = None,
    target_pages: int | None = None,
    filename_pattern: str | None = None,
    filename_custom_label: str | None = None,
    emphasis_font: str | None = None,
) -> dict:
    """Update an editable LaTeX material template package."""
    from src.documents.templates import update_latex_template_package

    if document_type not in {"resume", "cover_letter"}:
        return {
            "ok": False,
            "error": "Unsupported template document type.",
            "error_code": "invalid_document_type",
        }
    try:
        template = update_latex_template_package(
            document_type=document_type,
            template_id=template_id,
            content=content,
            template_name=template_name,
            description=description,
            target_pages=target_pages,
            filename_pattern=filename_pattern,
            filename_custom_label=filename_custom_label,
            emphasis_font=emphasis_font,
        )
    except ValueError as exc:
        return {"ok": False, "error": str(exc), "error_code": "invalid_template"}
    except Exception as exc:
        return {"ok": False, "error": str(exc), "error_code": "template_update_failed"}

    return {"ok": True, "template": template, **list_material_templates()}


def update_material_template_styles(
    *,
    document_type: str,
    template_id: str,
    overrides: dict,
    template_name: str | None = None,
    description: str | None = None,
    target_pages: int | None = None,
    filename_pattern: str | None = None,
    filename_custom_label: str | None = None,
    emphasis_font: str | None = None,
) -> dict:
    """Apply DOCX style overrides (font/size/bold/line-spacing) to a template."""
    from src.documents.templates import update_docx_template_styles

    if document_type not in {"resume", "cover_letter"}:
        return {
            "ok": False,
            "error": "Unsupported template document type.",
            "error_code": "invalid_document_type",
        }
    try:
        template = update_docx_template_styles(
            document_type=document_type,
            template_id=template_id,
            overrides=overrides or {},
            template_name=template_name,
            description=description,
            target_pages=target_pages,
            filename_pattern=filename_pattern,
            filename_custom_label=filename_custom_label,
            emphasis_font=emphasis_font,
        )
    except ValueError as exc:
        return {"ok": False, "error": str(exc), "error_code": "invalid_template"}
    except Exception as exc:
        return {"ok": False, "error": str(exc), "error_code": "template_update_failed"}

    return {"ok": True, "template": template, **list_material_templates()}


def validate_material_template(*, document_type: str, template_id: str) -> dict:
    """Validate a material template package."""
    from src.documents.templates import load_template_package, serialize_template_package

    if document_type not in {"resume", "cover_letter"}:
        return {
            "ok": False,
            "error": "Unsupported template document type.",
            "error_code": "invalid_document_type",
        }
    try:
        package = load_template_package(document_type, template_id)
        template = serialize_template_package(package)
    except ValueError as exc:
        return {"ok": False, "error": str(exc), "error_code": "invalid_template_id"}
    except Exception as exc:
        return {"ok": False, "error": str(exc), "error_code": "template_validate_failed"}

    return {"ok": True, "template": template, "validation": template["validation"]}


def delete_material_template(*, document_type: str, template_id: str) -> dict:
    """Delete a material template package, refusing for built-in defaults."""
    from src.documents.templates import delete_template_package

    if document_type not in {"resume", "cover_letter"}:
        return {
            "ok": False,
            "error": "Unsupported template document type.",
            "error_code": "invalid_document_type",
        }
    try:
        delete_template_package(document_type, template_id)
    except FileNotFoundError as exc:
        return {"ok": False, "error": str(exc), "error_code": "template_not_found"}
    except ValueError as exc:
        # _template_package_dir raises ValueError on bad ids, and
        # delete_template_package raises ValueError when refusing to
        # delete the built-in default. Both surface as a 400-shape error
        # so the API can decide between 400 and 403.
        message = str(exc)
        code = (
            "template_default_protected"
            if "default" in message.lower()
            else "invalid_template_id"
        )
        return {"ok": False, "error": message, "error_code": code}
    except Exception as exc:
        return {"ok": False, "error": str(exc), "error_code": "template_delete_failed"}

    return {"ok": True, **list_material_templates()}


async def apply_batch_jobs(
    *,
    selected_jobs: list[tuple],
    profile: str,
    top_n: int,
    auto_submit: bool = False,
    headless: bool = True,
    dry_run: bool = False,
) -> dict:
    profile_data = _load_profile()
    if not profile_data:
        return {
            "mode": "batch",
            "profile": profile,
            "top_n": top_n,
            "ok": False,
            "errors": ["Profile not configured."],
            "items": [],
            "summary": _batch_summary(),
            "dry_run": dry_run,
        }

    from src.utils.rate_limiter import RateLimiter, RateLimiterConfig

    limiter = RateLimiter(RateLimiterConfig(min_delay=5, max_delay=15))
    items = []
    summary = _batch_summary()

    for job, match_score in selected_jobs:
        if not job.application_url:
            items.append(
                {
                    "mode": "batch",
                    "input": {"profile": profile},
                    "ok": False,
                    "status": "SKIPPED",
                    "job": serialize_job(job, match_score=match_score),
                    "tracking_id": None,
                    "result": None,
                    "artifacts": _empty_artifacts(),
                    "error": "Job has no application URL.",
                    "error_code": "missing_application_url",
                    "dry_run": dry_run,
                }
            )
            summary["skipped"] += 1
            continue

        if not await limiter.can_apply():
            summary["stopped_early"] = True
            summary["stop_reason"] = "rate_limit_reached"
            break

        ats_type = _detect_ats_from_url(job.application_url)
        if not ats_type:
            items.append(
                {
                    "mode": "batch",
                    "input": {"profile": profile},
                    "ok": False,
                    "status": "SKIPPED",
                    "job": serialize_job(job, match_score=match_score),
                    "tracking_id": None,
                    "result": None,
                    "artifacts": _empty_artifacts(),
                    "error": _unsupported_ats_message(job.application_url),
                    "error_code": "unsupported_ats",
                    "dry_run": dry_run,
                }
            )
            summary["skipped"] += 1
            continue

        job.ats_type = ats_type

        item = await _run_application_for_job(
            job=job,
            profile_data=profile_data,
            auto_submit=auto_submit,
            headless=headless,
            dry_run=dry_run,
            match_score=match_score,
            mode="batch",
            input_payload={"profile": profile},
        )
        items.append(item)

        status = item["status"]
        if status == AppStatus.SUBMITTED:
            await limiter.record_application()
            summary["submitted"] += 1
        elif status == AppStatus.REVIEW_REQUIRED:
            summary["review"] += 1
        elif status == "DRY_RUN":
            summary["dry_run"] += 1
        elif status == "SKIPPED":
            summary["skipped"] += 1
        else:
            summary["failed"] += 1
            await limiter.error_cooldown()

    summary["attempted"] = len(items)

    return {
        "mode": "batch",
        "profile": profile,
        "top_n": top_n,
        "ok": True,
        "errors": [],
        "items": items,
        "summary": summary,
        "dry_run": dry_run,
    }


def serialize_job(job, match_score: float | None = None) -> dict:
    score = job.raw_data.get("match_score") if match_score is None else match_score
    metadata = _job_search_metadata(job)
    return {
        "id": str(job.id),
        "source": job.source,
        "source_id": job.source_id,
        "company": job.company,
        "title": job.title,
        "location": job.location,
        "employment_type": job.employment_type,
        "seniority": job.seniority,
        "description": job.description,
        "application_url": job.application_url,
        "ats_type": job.ats_type,
        "match_score": score,
        "best_profile": job.raw_data.get("best_profile"),
        "profile_scores": job.raw_data.get("profile_scores"),
        "disqualified": bool(job.raw_data.get("disqualified")),
        "experience_level": metadata.get("experience_level"),
        "employment_category": metadata.get("employment_category"),
        "location_type": metadata.get("location_type"),
        "education_level": metadata.get("education_level"),
        "experience_years_min": metadata.get("experience_years_min"),
        "experience_years_max": metadata.get("experience_years_max"),
        "pay_min": metadata.get("pay_min"),
        "pay_max": metadata.get("pay_max"),
        "raw_data": job.raw_data,
        "discovered_at": _isoformat(job.discovered_at),
    }


def _batch_summary() -> dict:
    return {
        "attempted": 0,
        "submitted": 0,
        "review": 0,
        "failed": 0,
        "skipped": 0,
        "dry_run": 0,
        "stopped_early": False,
        "stop_reason": None,
    }


def _build_companies_filter(
    config_dir: Path,
    ats: str | None,
    company: str | None,
) -> dict[str, list[str]] | None:
    if ats and company:
        # ats="workday" + company="<tenant>" degrades gracefully rather than
        # working: a bare tenant string can't determine host/site, so
        # src.intake.search's board loop logs a warning and skips it (same
        # self-pruning path as a malformed companies.yaml entry). Workday
        # can only be searched via the curated companies.yaml list.
        return {ats: [company]}
    if company:
        return {"greenhouse": [company], "lever": [company], "ashby": [company]}
    if ats:
        from src.intake.batch import load_company_list

        all_companies = load_company_list(config_dir / "companies.yaml")
        return {ats: all_companies.get(ats, [])}
    return None


def _apply_search_filters(
    jobs,
    *,
    experience_levels: list[str],
    employment_types: list[str],
    location_types: list[str],
    locations: list[str],
    search_location: str | None,
    searched_linkedin_locations: list[str],
    pay_operator: str | None,
    pay_amount: int | None,
    experience_operator: str | None,
    experience_years: int | None,
    education_levels: list[str],
    use_llm: bool,
):
    if not any(
        [
            experience_levels,
            employment_types,
            location_types,
            locations,
            pay_operator and pay_amount is not None,
            experience_operator and experience_years is not None,
            education_levels,
        ]
    ):
        return jobs

    _prepare_jobs_for_search_filters(jobs, use_llm=use_llm)
    filtered = []

    for job in jobs:
        metadata = _job_search_metadata(job)
        # "Unknown passes" (same convention as intake/filters.py): the
        # classifiers are heuristic and most JDs omit pay / education /
        # employment terms, so excluding unclassifiable jobs made stacked
        # Advanced filters delete nearly everything (a saved profile with
        # pay>=90k + full-time + remote once yielded "0 shown / 361
        # fetched"). Only a KNOWN violating value excludes a job.
        if not _passes_category(metadata.get("experience_level"), experience_levels):
            continue
        if not _passes_category(metadata.get("employment_category"), employment_types):
            continue
        if not _passes_category(metadata.get("location_type"), location_types):
            continue
        should_skip_location_filter = job.source == "linkedin" and bool(
            searched_linkedin_locations or search_location
        )
        if (
            locations
            and not should_skip_location_filter
            and not _matches_locations(job.location, locations)
        ):
            continue
        if not _passes_category(metadata.get("education_level"), education_levels):
            continue
        # Startup quality gate (HN "Who is hiring?" only): every other
        # source follows "unknown passes" for pay (see the docstring
        # above) because most JDs simply omit compensation and excluding
        # them nukes results. HN postings are the opposite case -- silence
        # on pay usually means a tiny non-paying/equity-only shop, which
        # the user explicitly does NOT want surfaced when they've set a
        # pay floor. The adapter sets raw_data["strict_pay"] = True (only
        # ever true for source="hn") to opt into this inversion without
        # touching the shared unknown-passes behavior below/elsewhere.
        if (
            pay_operator
            and pay_amount is not None
            and (job.raw_data or {}).get("strict_pay")
            and metadata.get("pay_min") is None
            and metadata.get("pay_max") is None
        ):
            continue
        if (
            pay_operator
            and pay_amount is not None
            and not _matches_numeric_filter(
                metadata.get("pay_min"), metadata.get("pay_max"), pay_operator, pay_amount
            )
        ):
            continue
        if (
            experience_operator
            and experience_years is not None
            and not _matches_numeric_filter(
                metadata.get("experience_years_min"),
                metadata.get("experience_years_max"),
                experience_operator,
                experience_years,
            )
        ):
            continue
        filtered.append(job)

    return filtered


def _prepare_jobs_for_search_filters(jobs, *, use_llm: bool) -> None:
    from src.intake.jd_parser import parse_requirements

    for job in jobs:
        if _requirements_empty(job.requirements) and job.description:
            try:
                job.requirements = parse_requirements(job.description, use_llm=use_llm)
            except Exception:
                logger.debug(
                    "JD parsing skipped for search filters on %s", job.title, exc_info=True
                )

        _set_job_search_metadata(
            job,
            {
                "experience_level": _classify_experience_level(job.title),
                "employment_category": _classify_employment_category(job),
                "location_type": _classify_location_type(job),
                "education_level": _normalize_education_level(
                    job.requirements.education_level, job.description
                ),
                "experience_years_min": job.requirements.experience_years_min,
                "experience_years_max": job.requirements.experience_years_max,
                **_extract_pay_range(job),
            },
        )


def _requirements_empty(requirements) -> bool:
    return not any(
        [
            requirements.education_level,
            requirements.experience_years_min,
            requirements.experience_years_max,
            requirements.remote_ok is not None,
            requirements.must_have_skills,
            requirements.preferred_skills,
        ]
    )


def _classify_experience_level(title: str) -> str:
    text = title.lower()
    if any(
        token in text
        for token in ("executive", "chief ", "vp ", "vice president", "cxo", "ceo", "cto", "cfo")
    ):
        return "executive"
    if any(token in text for token in ("director", "head of")):
        return "director"
    if any(token in text for token in ("manager", "management")):
        return "manager"
    if any(token in text for token in ("senior", "sr.", " sr ", "lead", "staff", "principal")):
        return "senior"
    if any(
        token in text
        for token in (
            "entry",
            "junior",
            "jr.",
            " jr ",
            "associate",
            "new grad",
            "intern",
            "co-op",
            "coop",
            "student",
        )
    ):
        return "entry"
    return "unknown"


def _classify_employment_category(job) -> str:
    explicit_text = " ".join(
        filter(
            None,
            [
                job.title,
                job.employment_type,
                str(job.raw_data.get("employment_type", "")),
                str(job.raw_data.get("categories", {}).get("commitment", "")),
            ],
        )
    ).lower()
    description_text = (job.description or "").lower()
    text = " ".join(filter(None, [explicit_text, description_text])).lower()

    if any(token in explicit_text for token in ("co-op", "coop")):
        return "coop"
    if any(
        token in explicit_text for token in ("internship", "intern ")
    ) or explicit_text.startswith("intern"):
        return "internship"

    if any(token in text for token in ("volunteer",)):
        return "volunteer"
    if any(token in text for token in ("freelance", "freelancer")):
        return "freelance"
    if any(token in text for token in ("seasonal",)):
        return "seasonal"
    if any(token in text for token in ("casual",)):
        return "casual"
    if any(token in text for token in ("temporary", "temp ", "temp-", "fixed term", "fixed-term")):
        return "temporary"
    if any(token in text for token in ("permanent",)):
        return "permanent"
    if any(token in text for token in ("part-time", "part time", "parttime")):
        return "part_time"
    if any(token in text for token in ("contract", "contractor")):
        return "contract"
    if any(token in text for token in ("full-time", "full time", "fulltime")):
        return "full_time"
    return "unknown"


def _classify_location_type(job) -> str:
    description = (job.description or "").lower()
    location = (job.location or "").lower()
    combined = f"{location} {description}"

    if "hybrid" in combined:
        return "hybrid"
    if job.requirements.remote_ok is True or any(
        token in combined for token in ("remote", "work from home", "wfh", "anywhere")
    ):
        return "remote"
    if any(
        token in combined
        for token in ("in-person", "in person", "onsite", "on-site", "in-office", "in office")
    ):
        return "in_person"
    if location:
        return "in_person"
    return "unknown"


def _normalize_education_level(level: str | None, description: str | None = None) -> str:
    try:
        from src.intake.jd_parser import infer_education_requirement

        inferred = infer_education_requirement(description)
        if inferred:
            level = inferred
    except Exception:  # noqa: BLE001 - metadata classification should not break search
        pass

    text = f"{level or ''}".lower()
    if any(token in text for token in ("juris doctor", " j.d", " jd")):
        return "jd"
    if any(token in text for token in ("doctor of medicine", " md", "m.d")):
        return "md"
    if "phd" in text or "doctorate" in text:
        return "phd"
    if "mba" in text:
        return "mba"
    if "master" in text or "m.s" in text or "ms " in text or "m.a" in text:
        return "master"
    if "bachelor" in text or "b.s" in text or "b.a" in text:
        return "bachelor"
    if "associate" in text:
        return "associate"
    if "high school" in text or "secondary school" in text:
        return "high_school"
    return "unknown"


def _extract_pay_range(job) -> dict:
    text = " ".join(filter(None, [job.description, str(job.raw_data)]))
    match = PAY_RANGE_RE.search(text)
    if match:
        return {
            "pay_min": _parse_money(match.group(1), match.group(2)),
            "pay_max": _parse_money(match.group(3), match.group(4)),
        }

    floor_match = PAY_FLOOR_RE.search(text)
    cap_match = PAY_CAP_RE.search(text)
    return {
        "pay_min": _parse_money(floor_match.group(1), floor_match.group(2))
        if floor_match
        else None,
        "pay_max": _parse_money(cap_match.group(1), cap_match.group(2)) if cap_match else None,
    }


def _parse_money(value: str, suffix: str | None) -> int | None:
    try:
        numeric = float(value.replace(",", ""))
    except ValueError:
        return None

    suffix_normalized = (suffix or "").lower()
    if suffix_normalized == "k":
        numeric *= 1000
    elif suffix_normalized == "m":
        numeric *= 1_000_000

    return int(numeric)


# Common aliases so "us", "usa", and "united states" all behave the same.
# Values are matched as whole words/phrases, never raw substrings.
_LOCATION_ALIASES: dict[str, tuple[str, ...]] = {
    "us": ("us", "usa", "u.s.", "united states", "america"),
    "usa": ("us", "usa", "u.s.", "united states", "america"),
    "united states": ("us", "usa", "u.s.", "united states", "america"),
    "uk": ("uk", "u.k.", "united kingdom", "england", "scotland", "wales"),
    "united kingdom": ("uk", "u.k.", "united kingdom", "england", "scotland", "wales"),
    "remote": ("remote", "anywhere", "work from home", "wfh"),
}

# US state names <-> postal abbreviations. Abbreviations are matched only in
# the ", XX" form ("Portland, OR") because many collide with English words
# ("or", "in", "me", "hi", "ok").
_US_STATE_BY_ABBREV: dict[str, str] = {
    "al": "alabama", "ak": "alaska", "az": "arizona", "ar": "arkansas",
    "ca": "california", "co": "colorado", "ct": "connecticut", "de": "delaware",
    "fl": "florida", "ga": "georgia", "hi": "hawaii", "id": "idaho",
    "il": "illinois", "in": "indiana", "ia": "iowa", "ks": "kansas",
    "ky": "kentucky", "la": "louisiana", "me": "maine", "md": "maryland",
    "ma": "massachusetts", "mi": "michigan", "mn": "minnesota", "ms": "mississippi",
    "mo": "missouri", "mt": "montana", "ne": "nebraska", "nv": "nevada",
    "nh": "new hampshire", "nj": "new jersey", "nm": "new mexico", "ny": "new york",
    "nc": "north carolina", "nd": "north dakota", "oh": "ohio", "ok": "oklahoma",
    "or": "oregon", "pa": "pennsylvania", "ri": "rhode island", "sc": "south carolina",
    "sd": "south dakota", "tn": "tennessee", "tx": "texas", "ut": "utah",
    "vt": "vermont", "va": "virginia", "wa": "washington", "wv": "west virginia",
    "wi": "wisconsin", "wy": "wyoming", "dc": "district of columbia",
}
_US_ABBREV_BY_STATE: dict[str, str] = {name: abbr for abbr, name in _US_STATE_BY_ABBREV.items()}
_US_COUNTRY_TERMS = {"us", "usa", "united states", "america"}


def _word_match(term: str, text: str) -> bool:
    return re.search(rf"(?<![a-z0-9]){re.escape(term)}(?![a-z0-9])", text) is not None


def _matches_locations(location: str | None, candidates: list[str]) -> bool:
    """Whole-word location matching with US-state and country aliases.

    The previous implementation was a raw substring test, which let short
    candidates match unrelated regions ("ny" in "Germany", "us" in
    "Australia") -- the main reason searches surfaced jobs from regions the
    user never asked for.
    """
    normalized_location = (location or "").lower()
    if not normalized_location:
        return False
    return any(
        _location_candidate_matches(candidate, normalized_location) for candidate in candidates
    )


def _location_candidate_matches(candidate: str, location: str) -> bool:
    normalized = candidate.strip().lower()
    if not normalized:
        return False

    # "City, ST" style chips: require every comma-part to match, so
    # "Portland, OR" matches "Portland, Oregon Metropolitan Area" and
    # "Portland, OR (Remote)" but not "Portland, Maine" or bare "US".
    if "," in normalized:
        parts = [part.strip() for part in normalized.split(",") if part.strip()]
        if len(parts) > 1:
            return all(_location_candidate_matches(part, location) for part in parts)

    terms = set(_LOCATION_ALIASES.get(normalized, (normalized,)))
    abbrevs: set[str] = set()

    if normalized in _US_STATE_BY_ABBREV:
        # Candidate is an abbreviation: also match the full state name.
        terms.discard(normalized)
        terms.add(_US_STATE_BY_ABBREV[normalized])
        abbrevs.add(normalized)
    elif normalized in _US_ABBREV_BY_STATE:
        abbrevs.add(_US_ABBREV_BY_STATE[normalized])

    for term in terms:
        # Skip whole-word matching for bare 2-letter state codes (handled
        # below); "us"/"uk" style aliases are unambiguous enough as words.
        if term in _US_STATE_BY_ABBREV:
            abbrevs.add(term)
            continue
        if _word_match(term, location):
            return True

    if any(
        re.search(rf",\s*{re.escape(abbrev)}(?![a-z0-9])", location) for abbrev in abbrevs
    ):
        return True

    # "united states" should also match US locations that only name a
    # city/state ("San Francisco, CA", "Dallas, Texas").
    return normalized in _US_COUNTRY_TERMS and _mentions_us_state(location)


def _mentions_us_state(location: str) -> bool:
    match = re.search(r",\s*([a-z]{2})(?![a-z0-9])", location)
    if match and match.group(1) in _US_STATE_BY_ABBREV:
        return True
    return any(_word_match(name, location) for name in _US_ABBREV_BY_STATE)


def _job_matches_keywords(job, keywords: list[str]) -> bool:
    """True if any keyword phrase appears in the job title or description."""
    haystack = f"{job.title or ''} {job.description or ''}".lower()
    return any(keyword.lower() in haystack for keyword in keywords if keyword.strip())


def _normalize_experience_levels(values: list[str]) -> list[str]:
    """Map UI/LinkedIn-style tokens onto the classifier vocabulary.

    2026-07-08: saved search profiles in the wild contain
    ``entry_level`` / ``associate`` / ``mid_senior`` — tokens that
    neither ``_classify_experience_level`` (emits ``entry`` /
    ``manager`` / ``senior`` / ``director`` / ``executive``) nor
    ``_map_linkedin_experience_levels`` understand. The mismatch was
    perverse: a filter FOR entry-level jobs excluded jobs explicitly
    titled entry-level (known ``entry`` not in ``{entry_level,
    associate}``) while passing every unclassifiable title.
    """
    aliases = {
        "entry_level": "entry",
        "entry-level": "entry",
        "associate": "entry",
        "junior": "entry",
        "internship": "entry",
        "mid_senior": "senior",
        "mid-senior": "senior",
    }
    normalized = [aliases.get(value, value) for value in values]
    # Preserve order, drop duplicates introduced by aliasing.
    return list(dict.fromkeys(normalized))


def _normalize_location_types(values: list[str]) -> list[str]:
    """Same dead-token repair for location types.

    ``_classify_location_type`` emits ``remote`` / ``hybrid`` /
    ``in_person``; saved profiles say ``onsite``, which excluded every
    explicitly on-site job while the user believed on-site was included.
    """
    aliases = {
        "onsite": "in_person",
        "on-site": "in_person",
        "on_site": "in_person",
        "office": "in_person",
    }
    normalized = [aliases.get(value, value) for value in values]
    return list(dict.fromkeys(normalized))


def _passes_category(value: str | None, selected: list[str]) -> bool:
    """Category filter with "unknown passes" semantics.

    Mirrors ``JobFilter._check_employment_type`` in intake/filters.py:
    an empty selection matches everything, an unclassifiable job is
    kept, and only a known non-matching value excludes.
    """
    if not selected:
        return True
    if value is None or value == "unknown":
        return True
    return value in selected


def _matches_numeric_filter(
    min_value: int | None,
    max_value: int | None,
    operator: str,
    target: int,
) -> bool:
    # Unknown passes: pay/experience are regex-extracted from JD text and
    # usually absent; excluding on missing data nuked most results.
    if min_value is None and max_value is None:
        return True

    lower = min_value if min_value is not None else max_value
    upper = max_value if max_value is not None else min_value

    if operator == "gt":
        return upper is not None and upper > target
    if operator == "gte":
        return upper is not None and upper >= target
    if operator == "lt":
        return lower is not None and lower < target
    if operator == "lte":
        return lower is not None and lower <= target
    return True


def _normalize_list(values: list[str] | None) -> list[str]:
    if not values:
        return []
    return [value.strip().lower() for value in values if isinstance(value, str) and value.strip()]


def _normalize_string_list(values: list[str] | None) -> list[str]:
    if not values:
        return []
    return [value.strip() for value in values if isinstance(value, str) and value.strip()]


def _resolve_linkedin_keywords(keyword: str | None, keywords: list[str]) -> list[str]:
    if keywords:
        return keywords
    if keyword and keyword.strip():
        return [keyword.strip()]
    return []


def _resolve_linkedin_search_locations(
    *,
    source: str,
    search_location: str | None,
    candidate_locations: list[str],
) -> list[str]:
    # "all" no longer runs LinkedIn (see the source == "linkedin" gate
    # above) -- only resolve locations for an explicit LinkedIn request.
    if source != "linkedin":
        return []
    if candidate_locations:
        return candidate_locations
    if search_location and search_location.strip():
        return [search_location.strip().lower()]
    return []


def _dedupe_jobs_by_signature(jobs) -> list:
    deduped = []
    seen: set[tuple[str, str, str]] = set()
    for job in jobs:
        signature = (
            (job.company or "").strip().lower(),
            (job.title or "").strip().lower(),
            (job.location or "").strip().lower(),
        )
        if signature in seen:
            continue
        seen.add(signature)
        deduped.append(job)
    return deduped


def _normalize_time_filter(value: str) -> str:
    return value if value in {"24h", "week", "month"} else ""


def _map_linkedin_experience_levels(values: list[str]) -> list[str] | None:
    mapping = {
        "entry": "entry",
        "senior": "mid_senior",
        "director": "director",
        "executive": "executive",
    }
    mapped = [mapping[value] for value in values if value in mapping]
    return mapped or None


def _map_linkedin_job_types(values: list[str]) -> list[str] | None:
    mapping = {
        "part_time": "parttime",
        "contract": "contract",
        "internship": "internship",
        "coop": "internship",
        "full_time": "fulltime",
        "temporary": "temporary",
    }
    mapped = [mapping[value] for value in values if value in mapping]
    return mapped or None


def _linkedin_max_pages(
    max_pages: int,
    *,
    search_location: str,
    experience_levels: list[str],
    employment_types: list[str],
    location_types: list[str],
    locations: list[str],
    pay_operator: str | None,
    experience_operator: str | None,
    education_levels: list[str],
) -> int:
    local_only_filters = any(
        [
            any(value in {"manager"} for value in experience_levels),
            any(
                value in {"permanent", "casual", "seasonal", "freelance", "volunteer"}
                for value in employment_types
            ),
            location_types,
            locations and not search_location,
            pay_operator,
            experience_operator,
            education_levels,
        ]
    )
    return max(max_pages, 10) if local_only_filters else max_pages


def _job_search_metadata(job) -> dict:
    return job.raw_data.get(SEARCH_METADATA_KEY, {})


def _set_job_search_metadata(job, metadata: dict) -> None:
    job.raw_data[SEARCH_METADATA_KEY] = metadata


def _score_jobs(jobs, *, warn_on_missing_profile: bool) -> tuple[bool, list[str]]:
    """Score every job against EVERY saved applicant profile, keep the best.

    2026-07-07: previously only the active profile was scored, so a job
    that perfectly matched the sales-engineer resume looked mediocre while
    the analyst profile happened to be active. Now each profile in
    ``data/profile/profiles/`` scores every job; the best-scoring profile
    wins ``match_score`` and is recorded as ``raw_data.best_profile`` so
    material generation can auto-select the matching resume. Per-profile
    scores land in ``raw_data.profile_scores`` for the UI.

    Ties prefer the active profile. A profile that fails to load or score
    is reported as a warning and skipped rather than failing the search.
    """
    from src.application.profile import get_active_profile_id, list_profiles
    from src.matching.scorer import build_scoring_context
    from src.matching.scorer import score_jobs as score_ranked_jobs
    from src.memory.profile import load_profile_yaml

    profiles = list_profiles()
    if not profiles:
        if warn_on_missing_profile:
            return False, ["No profile found -- run `autoapply init` first to enable scoring."]
        return False, []

    active_id = get_active_profile_id()
    errors: list[str] = []
    ranked_by_profile: dict[str, dict] = {}
    for meta in profiles:
        try:
            profile_data = load_profile_yaml(Path(meta["path"]))
            scoring_ctx = build_scoring_context(profile_data)
            ranked = score_ranked_jobs(jobs, scoring_ctx)
        except Exception as exc:  # noqa: BLE001 -- one bad profile must not kill search
            errors.append(f"Scoring against profile '{meta['id']}' failed: {exc}")
            continue
        ranked_by_profile[meta["id"]] = {score.job_id: score for score in ranked}

    if not ranked_by_profile:
        return False, errors or ["Scoring failed for all profiles."]

    for job in jobs:
        job_id = str(job.id)
        profile_scores: dict[str, float] = {}
        best = None
        best_profile_id = None
        for profile_id, score_by_id in ranked_by_profile.items():
            score = score_by_id.get(job_id)
            if score is None:
                continue
            profile_scores[profile_id] = score.final_score
            if (
                best is None
                or score.final_score > best.final_score
                or (score.final_score == best.final_score and profile_id == active_id)
            ):
                best = score
                best_profile_id = profile_id
        if best is None:
            continue
        job.raw_data["match_score"] = best.final_score
        job.raw_data["disqualified"] = best.disqualified
        job.raw_data["best_profile"] = best_profile_id
        job.raw_data["profile_scores"] = profile_scores
        # Phase 16.3: stash the structured breakdown alongside the
        # legacy scalar fields so the "Why was this filtered?"
        # popover can render rule_id / verdict / reason /
        # evidence_excerpt without re-scoring round-trip.
        job.raw_data["score_breakdown"] = best.to_dict()

    jobs.sort(key=_job_sort_key)
    return True, errors


def _job_sort_key(job) -> tuple:
    """Deterministic ranking: match score desc, then company / title asc.

    ``raw_data.get("match_score", 0.0)`` alone had two problems: a stored
    ``None`` crashed the comparison, and unscored ties kept raw scrape
    order, which made results look randomly shuffled between runs.
    """
    raw_score = job.raw_data.get("match_score")
    try:
        score = float(raw_score) if raw_score is not None else 0.0
    except (TypeError, ValueError):
        score = 0.0
    return (-score, (job.company or "").lower(), (job.title or "").lower())


def _select_batch_jobs(filter_profile: str, top_n: int) -> tuple[list[tuple], list[str], int]:
    from src.intake.search import search_jobs as search_ats_jobs
    from src.matching.scorer import build_scoring_context
    from src.matching.scorer import score_jobs as score_ranked_jobs

    jobs = search_ats_jobs(profile=filter_profile, parse_jds=True)
    if not jobs:
        return [], ["No matching jobs found."], 0

    profile_data = _load_profile()
    if not profile_data:
        return [], ["Profile not configured."], 0

    scoring_ctx = build_scoring_context(profile_data)
    ranked = score_ranked_jobs(jobs, scoring_ctx)
    job_by_id = {str(job.id): job for job in jobs}
    selected = []

    for score in ranked:
        job = job_by_id.get(score.job_id)
        if job is None:
            continue
        job.raw_data["match_score"] = score.final_score
        job.raw_data["disqualified"] = score.disqualified
        # Phase 16.3 explainability hook.
        job.raw_data["score_breakdown"] = score.to_dict()
        if score.disqualified:
            continue
        selected.append((job, score.final_score))
        if len(selected) >= top_n:
            break

    total_matches = sum(1 for score in ranked if not score.disqualified)
    return selected, [], total_matches


def _empty_artifacts() -> dict:
    return {
        "resume_path": None,
        "cover_letter_path": None,
        "qa_responses": None,
    }


def _empty_material_artifacts() -> dict:
    return {
        "resume_pdf": None,
        "resume_docx": None,
        "resume_tex": None,
        "cover_letter_pdf": None,
        "cover_letter_docx": None,
        "cover_letter_tex": None,
    }


def _stringify_material_artifacts(artifacts: dict) -> dict:
    combined = _empty_material_artifacts()
    combined.update(artifacts)
    return {key: str(value) if value else None for key, value in combined.items()}


def _serialize_material_artifact(material_type: str, path: Path | str) -> dict:
    # ``path`` can arrive as either a Path (the regenerate flow) or a
    # plain string (the patch_existing flow stringifies before stashing
    # into ``artifacts``). Coerce to Path so ``.name`` lookups don't
    # crash with ``AttributeError: 'str' object has no attribute 'name'``.
    path_obj = path if isinstance(path, Path) else Path(path)
    return {
        "type": material_type,
        "path": str(path_obj),
        "filename": path_obj.name,
    }


def _serialize_generation_model(value):
    if value is None:
        return None
    if hasattr(value, "model_dump"):
        return value.model_dump(mode="json")
    return value


def _save_generation_version(**kwargs) -> dict | None:
    try:
        from src.generation.versions import save_generation_version

        return save_generation_version(**kwargs)
    except Exception as exc:
        logger.warning("Generation version persistence skipped: %s", exc)
        return None


def _raw_job_from_web_payload(job_payload: dict, *, use_llm: bool):
    from src.intake.jd_parser import parse_requirements
    from src.intake.schema import JobRequirements, RawJob

    if not isinstance(job_payload, dict):
        raise ValueError("Job payload is required.")

    company = _clean_web_payload_string(job_payload.get("company"))
    title = _clean_web_payload_string(job_payload.get("title"))
    if not company or not title:
        raise ValueError("Job company and title are required.")

    description = _clean_web_payload_string(job_payload.get("description")) or None
    requirements_payload = job_payload.get("requirements")
    if isinstance(requirements_payload, dict):
        requirements = JobRequirements.model_validate(requirements_payload)
    elif description:
        requirements = parse_requirements(description, use_llm=use_llm)
    else:
        requirements = JobRequirements()

    raw_data = (
        job_payload.get("raw_data") if isinstance(job_payload.get("raw_data"), dict) else {}
    )
    source = _coerce_web_payload_value(job_payload.get("source"), ATS_TYPES, "unknown")
    ats_type = _coerce_web_payload_value(job_payload.get("ats_type"), ATS_TYPES, "unknown")
    if source == "unknown" and ats_type != "unknown":
        source = ats_type

    kwargs = {}
    try:
        if job_payload.get("id"):
            kwargs["id"] = uuid.UUID(str(job_payload["id"]))
    except (TypeError, ValueError):
        pass

    return RawJob(
        **kwargs,
        source=source,
        source_id=str(
            job_payload.get("source_id")
            or job_payload.get("id")
            or job_payload.get("application_url")
            or f"{company}:{title}"
        ),
        company=company,
        title=title,
        location=_clean_web_payload_string(job_payload.get("location")) or None,
        employment_type=_coerce_web_payload_value(
            job_payload.get("employment_type"), EMPLOYMENT_TYPES, "unknown"
        ),
        seniority=_coerce_web_payload_value(
            job_payload.get("seniority"), SENIORITY_LEVELS, "unknown"
        ),
        description=description,
        requirements=requirements,
        application_url=_clean_web_payload_string(job_payload.get("application_url")) or None,
        ats_type=ats_type,
        raw_data={**raw_data, "web_material_generation": True},
    )


def _clean_web_payload_string(value) -> str:
    return value.strip() if isinstance(value, str) else ""


def _clean_optional_web_payload_id(value) -> str | None:
    if not isinstance(value, str):
        return None
    cleaned = value.strip()
    if not cleaned or cleaned.lower() in {"null", "none", "undefined"}:
        return None
    return cleaned


def _coerce_web_payload_value(value, allowed: set[str], default: str) -> str:
    if isinstance(value, str) and value in allowed:
        return value
    return default


def _generate_selected_material(
    profile_data: dict,
    job,
    material_type: str,
    *,
    template_id: str | None = None,
    strategy: str = "regenerate",
    source_document_id: str | None = None,
    patch_aggressiveness: str = "balanced",
    patch_allow_reorder_sections: bool = True,
    patch_allow_add_remove_bullets: bool = True,
) -> dict:
    from src.documents.templates import ensure_template_package, serialize_template_package
    from src.generation.cover_letter import generate_cover_letter, generate_cover_letter_latex
    from src.generation.resume_builder import generate_resume, generate_resume_latex

    output_dir = PROJECT_ROOT / "data" / "output"
    output_dir.mkdir(parents=True, exist_ok=True)

    artifacts = _empty_material_artifacts()
    strategy_notes: list[str] = []

    # --- Strategy 3: use_library --------------------------------------
    # Drop in the user's chosen library document as-is. No LLM, no
    # rendering, no template. We just copy the file into the output
    # directory and pin the right artifact slot. This is the "I
    # already wrote the perfect resume, just apply with it" path.
    if strategy == "use_library":
        if not source_document_id:
            raise ValueError(
                "use_library strategy requires a source document id."
            )
        copied_path, used_doc = _copy_library_document_to_output(
            document_id=source_document_id,
            output_dir=output_dir,
            material_type=material_type,
        )
        source_suffix = _source_type_to_material_suffix(used_doc.source_type)
        artifact_type = f"{used_doc.document_type}_{source_suffix}"
        artifacts[artifact_type] = str(copied_path)
        strategy_notes.append(
            f"Used your library document {used_doc.original_filename!r} as-is."
        )
        return {
            "artifacts": artifacts,
            "document": None,
            "validation": None,
            "template": None,
            "strategy_notes": strategy_notes,
        }

    if material_type in {"resume_pdf", "resume_docx", "resume_tex"}:
        template_package = ensure_template_package("resume", template_id)
        _ensure_template_supports_material(template_package, material_type)
        if template_package.manifest.renderer == "latex":
            resume_files = generate_resume_latex(
                job=job,
                profile_data=profile_data,
                output_dir=output_dir,
                template_id=template_package.template_id,
            )
            artifacts["resume_pdf"] = resume_files.get("pdf")
            artifacts["resume_tex"] = resume_files.get("tex")
            if strategy == "patch_existing":
                strategy_notes.append(
                    "patch_existing isn't supported for LaTeX templates yet; "
                    "rendered fresh instead."
                )
            return {
                "artifacts": artifacts,
                "document": resume_files.get("ir"),
                "validation": resume_files.get("validation"),
                "template": serialize_template_package(template_package),
                "strategy_notes": strategy_notes,
            }

        resume_files = generate_resume(
            job=job,
            profile_data=profile_data,
            output_dir=output_dir,
            template_id=template_package.template_id,
            # ``patch_aggressiveness`` is a property of the
            # ``patch_existing`` strategy. For ``regenerate`` it has
            # no source document to be conservative about, but the
            # rewrite mode also feeds the bullet-keyword editor's
            # tone, so it's reasonable to honor the user's preference
            # in both branches. The default ``"balanced"`` matches
            # the historical behaviour.
            rewrite_mode=patch_aggressiveness,
            rewrite=True,
            use_llm=True,
        )
        artifacts["resume_pdf"] = resume_files.get("pdf")
        artifacts["resume_docx"] = resume_files.get("docx")

        if strategy == "patch_existing" and source_document_id:
            patched_path, patch_note = _try_patch_docx_from_library(
                document_id=source_document_id,
                ir=resume_files.get("ir"),
                output_dir=output_dir,
                allow_reorder_sections=patch_allow_reorder_sections,
                allow_add_remove_bullets=patch_allow_add_remove_bullets,
            )
            logger.info(
                "patch_existing for %s: patched_path=%s note=%s "
                "pre-patch artifacts[resume_docx]=%s",
                source_document_id,
                patched_path,
                patch_note,
                artifacts.get("resume_docx"),
            )
            if patched_path is not None:
                artifacts["resume_docx"] = str(patched_path)
                # The PDF rendered from the template is now stale
                # relative to the patched DOCX, but we can only drop
                # it if the caller asked for DOCX -- when the caller
                # asked for ``resume_pdf`` the PDF artifact is the
                # thing they're downloading and nuking it would
                # silently turn the request into a "could not be
                # generated" error downstream. In that case we keep
                # the fresh-template PDF and surface a note that the
                # patched DOCX is a side artifact only.
                if material_type == "resume_pdf":
                    strategy_notes.append(
                        "Patched the DOCX from your library, but the PDF "
                        "you downloaded is rendered from the template "
                        "(re-rendering the patched DOCX to PDF would need "
                        "LibreOffice/Word). Pick the DOCX format if you "
                        "want the patched output."
                    )
                else:
                    artifacts["resume_pdf"] = None
                    strategy_notes.append(
                        "Patched your library document with the tailored bullets."
                    )
            if patch_note:
                strategy_notes.append(patch_note)
            logger.info(
                "patch_existing post: artifacts[resume_docx]=%s strategy_notes=%s",
                artifacts.get("resume_docx"),
                strategy_notes,
            )

        return {
            "artifacts": artifacts,
            "document": resume_files.get("ir"),
            "validation": resume_files.get("validation"),
            "template": serialize_template_package(template_package),
            "strategy_notes": strategy_notes,
        }

    template_package = ensure_template_package("cover_letter", template_id)
    _ensure_template_supports_material(template_package, material_type)
    if template_package.manifest.renderer == "latex":
        if strategy == "patch_existing":
            strategy_notes.append(
                "patch_existing isn't supported for LaTeX cover letter "
                "templates yet; rendered fresh instead."
            )
        cover_files = generate_cover_letter_latex(
            job=job,
            profile_data=profile_data,
            output_dir=output_dir,
            template_id=template_package.template_id,
        )
        artifacts["cover_letter_pdf"] = cover_files.get("pdf")
        artifacts["cover_letter_tex"] = cover_files.get("tex")
        return {
            "artifacts": artifacts,
            "document": cover_files.get("ir"),
            "validation": cover_files.get("validation"),
            "template": serialize_template_package(template_package),
            "strategy_notes": strategy_notes,
        }

    cover_files = generate_cover_letter(
        job=job,
        profile_data=profile_data,
        output_dir=output_dir,
        template_id=template_package.template_id,
        use_llm=True,
    )
    artifacts["cover_letter_pdf"] = cover_files.get("pdf")
    artifacts["cover_letter_docx"] = cover_files.get("docx")

    if strategy == "patch_existing" and source_document_id:
        patched_path, patch_note = _try_patch_cover_letter_from_library(
            document_id=source_document_id,
            ir=cover_files.get("ir"),
            output_dir=output_dir,
        )
        if patched_path is not None:
            artifacts["cover_letter_docx"] = str(patched_path)
            # See the resume patch branch for the rationale: the
            # rendered PDF is now stale relative to the patched DOCX.
            # If the user asked for the DOCX, drop the PDF artifact so
            # the UI doesn't surface a disagreeing download. If they
            # asked for the PDF, keep the template-rendered PDF and
            # surface a note explaining the patched DOCX is a side
            # artifact only.
            if material_type == "cover_letter_pdf":
                strategy_notes.append(
                    "Patched the DOCX from your library, but the PDF "
                    "you downloaded is rendered from the template "
                    "(re-rendering the patched DOCX to PDF would need "
                    "LibreOffice/Word). Pick the DOCX format if you "
                    "want the patched output."
                )
            else:
                artifacts["cover_letter_pdf"] = None
                strategy_notes.append(
                    "Patched your library cover letter with the "
                    "tailored body paragraphs."
                )
        if patch_note:
            strategy_notes.append(patch_note)

    return {
        "artifacts": artifacts,
        "document": cover_files.get("ir"),
        "validation": cover_files.get("validation"),
        "template": serialize_template_package(template_package),
        "strategy_notes": strategy_notes,
    }


def _try_patch_docx_from_library(
    *,
    document_id: str,
    ir,
    output_dir,
    allow_reorder_sections: bool = True,
    allow_add_remove_bullets: bool = True,
):
    """Apply :func:`patch_resume_docx` to a UserDocument from the library.

    The two ``allow_*`` flags come from the user's saved patch settings
    (see ``material_defaults``). When False they constrain the patcher
    to leave structural decisions to the source DOCX rather than the
    IR, which matches the conservative / "preserve my layout" use case.

    Returns ``(output_path, note)`` where ``output_path`` is the path
    on disk of the patched docx (``None`` on failure) and ``note`` is
    an optional human-readable warning to surface back to the UI.
    """
    if ir is None:
        return None, "Could not patch your library document: no tailored content was produced."
    try:
        from uuid import UUID

        from src.core.database import get_session_factory
        from src.documents.user_documents import get_document, resolve_storage_path
        from src.generation.docx_patch import PatchFallback, patch_resume_docx

        doc_uuid = UUID(document_id)
        session_factory = get_session_factory(load_config())
        with session_factory() as session:
            row = get_document(session, doc_uuid)
            if row is None:
                return None, "Your selected library document is missing — generated fresh instead."
            if row.document_type != "resume":
                return None, (
                    "Selected library document isn't a resume — "
                    "generated fresh instead."
                )
            if row.source_type != "docx":
                return None, (
                    f"Library document is a {row.source_type.upper()} file "
                    "and can't be patched in place — generated fresh instead."
                )
            source_path = resolve_storage_path(row)

        source_text = _visible_docx_text(source_path)
        output_path = output_dir / f"patched_resume_{uuid.uuid4().hex}.docx"
        from src.maintenance.atomic import atomic_write  # noqa: PLC0415

        try:
            # Phase 18.4: write to a .tmp sibling and atomically rename
            # on success so a crash mid-patch can't leave a half-written
            # ``patched_resume_<uuid>.docx`` on disk.
            with atomic_write(output_path) as tmp_path:
                patch_resume_docx(
                    source_path,
                    ir,
                    output_path=tmp_path,
                    allow_reorder_sections=allow_reorder_sections,
                    allow_add_remove_bullets=allow_add_remove_bullets,
                )
        except PatchFallback as exc:
            return None, f"Couldn't patch your library document ({exc}); generated fresh instead."
        patched_text = _visible_docx_text(output_path)
        if source_text and patched_text == source_text:
            output_path.unlink(missing_ok=True)
            return None, (
                "Patch produced no visible resume changes; generated a fresh "
                "tailored document instead."
            )
        return output_path, None
    except Exception as exc:  # noqa: BLE001
        logger.exception("patch_existing fell over for document %s", document_id)
        return None, f"Patch failed ({exc}); generated fresh instead."


def _try_patch_cover_letter_from_library(
    *,
    document_id: str,
    ir,
    output_dir,
):
    """Apply :func:`patch_cover_letter_docx` to a UserDocument from the
    library. Mirrors ``_try_patch_docx_from_library`` for resumes but
    targets the cover-letter patcher and the cover-letter document
    type.

    Returns ``(output_path, note)`` where ``output_path`` is the path
    on disk of the patched docx (``None`` on failure) and ``note`` is
    a human-readable warning to surface back to the UI.
    """
    if ir is None:
        return None, (
            "Could not patch your library cover letter: no tailored "
            "content was produced."
        )
    try:
        from uuid import UUID

        from src.core.database import get_session_factory
        from src.documents.user_documents import get_document, resolve_storage_path
        from src.generation.docx_patch import PatchFallback, patch_cover_letter_docx

        doc_uuid = UUID(document_id)
        session_factory = get_session_factory(load_config())
        with session_factory() as session:
            row = get_document(session, doc_uuid)
            if row is None:
                return None, (
                    "Your selected library cover letter is missing — "
                    "generated fresh instead."
                )
            if row.document_type != "cover_letter":
                return None, (
                    "Selected library document isn't a cover letter — "
                    "generated fresh instead."
                )
            if row.source_type != "docx":
                return None, (
                    f"Library document is a {row.source_type.upper()} file "
                    "and can't be patched in place — generated fresh instead."
                )
            source_path = resolve_storage_path(row)

        source_text = _visible_docx_text(source_path)
        output_path = (
            output_dir / f"patched_cover_letter_{uuid.uuid4().hex}.docx"
        )
        from src.maintenance.atomic import atomic_write  # noqa: PLC0415

        try:
            # Phase 18.4: atomic_write rename so we never leave a
            # half-written ``patched_cover_letter_<uuid>.docx`` on
            # disk if the underlying patcher raises mid-write.
            with atomic_write(output_path) as tmp_path:
                patch_cover_letter_docx(source_path, ir, output_path=tmp_path)
        except PatchFallback as exc:
            return None, (
                f"Couldn't patch your library cover letter ({exc}); "
                "generated fresh instead."
            )
        patched_text = _visible_docx_text(output_path)
        if source_text and patched_text == source_text:
            output_path.unlink(missing_ok=True)
            return None, (
                "Patch produced no visible cover letter changes; generated "
                "a fresh tailored document instead."
            )
        return output_path, None
    except Exception as exc:  # noqa: BLE001
        logger.exception(
            "patch_existing for cover letter fell over for document %s",
            document_id,
        )
        return None, f"Cover letter patch failed ({exc}); generated fresh instead."


def _visible_docx_text(path) -> str:
    try:
        from docx import Document  # noqa: PLC0415
    except ImportError:
        return ""
    try:
        doc = Document(str(path))
    except Exception:
        return ""
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = " ".join((paragraph.text or "").split())
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = " ".join((paragraph.text or "").split())
                    if text:
                        parts.append(text)
    return "\n".join(parts)


def _copy_library_document_to_output(
    *,
    document_id: str,
    output_dir,
    material_type: str,
):
    """Resolve a UserDocument by id and copy it into ``output_dir``.

    This is the implementation behind the ``use_library`` strategy:
    the user opted out of any LLM rewriting and just wants their
    library document used as the application artifact. We copy the
    file (rather than symlinking or returning the library path
    directly) so the application's audit trail mirrors the
    regenerate / patch flows -- every artifact a job ever applied
    with is a snapshot in ``data/output/``.

    Returns ``(copied_path, row)``. Raises ``ValueError`` with a
    user-readable message for missing / wrong-type documents so the
    caller can surface a clean error.
    """
    import shutil
    from uuid import UUID

    from src.core.database import get_session_factory
    from src.documents.user_documents import get_document, resolve_storage_path

    expected_doc_type = (
        "resume" if material_type.startswith("resume") else "cover_letter"
    )
    expected_source = material_type.rsplit("_", 1)[-1]  # docx / pdf / tex / txt

    try:
        doc_uuid = UUID(document_id)
    except (TypeError, ValueError) as exc:
        raise ValueError(f"Library document id is not a UUID: {document_id!r}") from exc

    session_factory = get_session_factory(load_config())
    with session_factory() as session:
        row = get_document(session, doc_uuid)
        if row is None:
            raise ValueError("Selected library document is no longer in your library.")
        if row.document_type != expected_doc_type:
            raise ValueError(
                f"Library document is a {row.document_type.replace('_', ' ')}, "
                f"but a {expected_doc_type.replace('_', ' ')} was requested."
            )
        actual_source = _source_type_to_material_suffix(row.source_type)
        if actual_source != expected_source:
            raise ValueError(
                f"Library document is a {row.source_type.upper()} file, "
                f"but the requested output format is {expected_source.upper()}. "
                "Pick a library document whose extension matches the output "
                "format, or switch strategy to Regenerate / Patch."
            )
        source_path = resolve_storage_path(row)
        original_filename = row.original_filename

    if not source_path.exists():
        raise ValueError(
            f"Library file is missing on disk: {source_path}"
        )

    target_name = (
        f"{expected_doc_type}_uselib_{uuid.uuid4().hex}{source_path.suffix.lower()}"
    )
    target_path = output_dir / target_name
    from src.maintenance.atomic import atomic_write  # noqa: PLC0415

    # Phase 18.4: copy via atomic_write so a crash mid-copy can't leave
    # a half-written ``*_uselib_<uuid>.docx`` orphan in data/output.
    with atomic_write(target_path) as tmp_path:
        shutil.copy2(source_path, tmp_path)

    # Re-fetch a fresh ``row``-like object so the caller can still
    # read original_filename / source_type without re-opening the
    # session. SimpleNamespace keeps this lightweight.
    from types import SimpleNamespace

    return target_path, SimpleNamespace(
        original_filename=original_filename,
        source_type=row.source_type,
        document_type=row.document_type,
    )


def _ensure_template_supports_material(template_package, material_type: str) -> None:
    output_format = material_type.rsplit("_", 1)[-1]
    if output_format == "txt" and template_package.manifest.renderer == "docx":
        return
    if output_format not in set(template_package.manifest.supported_outputs):
        raise ValueError(
            f"Template '{template_package.template_id}' does not support "
            f"{output_format.upper()} output."
        )


def _unsupported_ats_message(url: str) -> str:
    if "linkedin.com" in url.lower():
        return (
            "This LinkedIn job did not expose an external apply target. "
            "Open the job manually if it uses LinkedIn Easy Apply."
        )
    return "Could not detect an application URL from this input."


def _no_job_context_message(ats_type: str) -> str:
    label_map = {
        "ashby": "Ashby",
        "workday": "Workday",
        "company_site": "company-site",
    }
    label = label_map.get(ats_type, ats_type)
    return (
        f"Cannot apply to a {label} URL without a stored job. "
        "Run `autoapply search` first and apply via the resulting job id, "
        "or paste the JD into the Materials page to download tailored files."
    )


def _linkedin_easy_apply_message() -> str:
    """Standard error message when a job only exposes LinkedIn Easy Apply.

    AutoApply does not yet automate LinkedIn Easy Apply: the form lives
    inside a LinkedIn modal that needs its own automation path. Until
    that ships, refusing the apply request is the honest answer --
    far better than silently submitting some random homepage form.
    """
    return (
        "This job only offers LinkedIn Easy Apply, which AutoApply does not "
        "yet automate. Open the LinkedIn posting and submit manually -- or "
        "wait for the Easy Apply automation to land."
    )


def _job_is_easy_apply_only(job) -> bool:
    """True when the scraper marked this job as LinkedIn-Easy-Apply-only.

    The scraper writes ``raw_data["linkedin_easy_apply_only"]`` during
    job-detail extraction. Reading via ``getattr`` + ``dict.get``
    keeps the helper tolerant of older scrape passes that predate the
    flag (in which case we fall through to the URL-domain check the
    caller already performs).
    """
    raw = getattr(job, "raw_data", None) or {}
    if not isinstance(raw, dict):
        return False
    return bool(raw.get("linkedin_easy_apply_only"))


def _serialize_execution_result(result) -> dict:
    if result is None:
        return {
            "status": None,
            "error": None,
            "fields_filled": 0,
            "fields_total": 0,
            "fill_details": [],
            "files_uploaded": [],
            "qa_answered": 0,
            "screenshots": [],
        }

    return {
        "status": str(result.status),
        "error": result.error or None,
        "fields_filled": result.fields_filled,
        "fields_total": result.fields_total,
        # ``getattr`` guards against older adapter shapes that don't
        # publish per-field details yet -- their payload comes back
        # with ``[]`` which the UI renders as "no per-field details".
        "fill_details": getattr(result, "fill_details", []) or [],
        "files_uploaded": result.files_uploaded,
        "qa_answered": result.qa_answered,
        "screenshots": [str(path) for path in result.screenshots],
    }


async def _run_application_for_job(
    *,
    job,
    profile_data: dict,
    auto_submit: bool,
    headless: bool,
    dry_run: bool,
    mode: str,
    input_payload: dict,
    match_score: float | None = None,
) -> dict:
    resume_path, cover_letter_path, qa_responses = await _generate_materials(profile_data, job)

    payload = {
        "mode": mode,
        "input": input_payload,
        "job": serialize_job(job, match_score=match_score),
        "tracking_id": None,
        "artifacts": {
            "resume_path": str(resume_path) if resume_path else None,
            "cover_letter_path": str(cover_letter_path) if cover_letter_path else None,
            "qa_responses": qa_responses,
        },
        "dry_run": dry_run,
        "result": None,
        "status": None,
        "ok": False,
        "error": None,
        "error_code": None,
    }

    if not resume_path:
        payload["error"] = "Cannot continue without a generated resume."
        payload["error_code"] = "resume_generation_failed"
        payload["result"] = _serialize_execution_result(None)
        return payload

    if dry_run:
        payload["status"] = "DRY_RUN"
        payload["ok"] = True
        payload["result"] = {
            "status": "DRY_RUN",
            "error": None,
            "fields_filled": 0,
            "fields_total": 0,
            "files_uploaded": [],
            "qa_answered": 0,
            "screenshots": [],
        }
        return payload

    app_id = _create_tracking_application(
        job=job,
        match_score=match_score,
        resume_path=resume_path,
        cover_letter_path=cover_letter_path,
    )
    if app_id is not None:
        payload["tracking_id"] = str(app_id)

    state = ApplicationState(str(app_id or job.id))
    result = await _execute_application(
        url=job.application_url or "",
        ats_type=job.ats_type,
        job=job,
        profile_data=profile_data,
        resume_path=resume_path,
        cover_letter_path=cover_letter_path,
        qa_responses=qa_responses,
        auto_submit=auto_submit,
        headless=headless,
        state=state,
    )

    if app_id is not None:
        _sync_tracking_application(app_id, state, result, qa_responses)

    payload["result"] = _serialize_execution_result(result)
    payload["status"] = payload["result"]["status"]
    payload["ok"] = payload["status"] in {AppStatus.SUBMITTED, AppStatus.REVIEW_REQUIRED}
    if not payload["ok"]:
        payload["error"] = payload["result"]["error"] or "Application failed."
        payload["error_code"] = "application_failed"

    return payload


async def _generate_materials(
    profile_data: dict, job
) -> tuple[Path | None, Path | None, dict | None]:
    from src.application.material_defaults import resolve_material_choice
    from src.generation.qa_responder import answer_questions

    output_dir = PROJECT_ROOT / "data" / "output"
    output_dir.mkdir(parents=True, exist_ok=True)

    resume_choice = resolve_material_choice(document_type="resume")
    resume_material_type = _application_material_type("resume", resume_choice)
    resume_files = _generate_selected_material(
        profile_data,
        job,
        resume_material_type,
        template_id=resume_choice["template_id"],
        strategy=resume_choice["strategy"],
        source_document_id=resume_choice["document_id"],
        patch_aggressiveness=resume_choice["patch_aggressiveness"],
        patch_allow_reorder_sections=resume_choice["patch_allow_reorder_sections"],
        patch_allow_add_remove_bullets=resume_choice["patch_allow_add_remove_bullets"],
    )
    resume_path = _pick_application_artifact(
        resume_files.get("artifacts") or {}, "resume", resume_choice["strategy"]
    )

    cover_choice = resolve_material_choice(document_type="cover_letter")
    cover_material_type = _application_material_type("cover_letter", cover_choice)
    cover_files = _generate_selected_material(
        profile_data,
        job,
        cover_material_type,
        template_id=cover_choice["template_id"],
        strategy=cover_choice["strategy"],
        source_document_id=cover_choice["document_id"],
        patch_aggressiveness=cover_choice["patch_aggressiveness"],
        patch_allow_reorder_sections=cover_choice["patch_allow_reorder_sections"],
        patch_allow_add_remove_bullets=cover_choice["patch_allow_add_remove_bullets"],
    )
    cover_letter_path = _pick_application_artifact(
        cover_files.get("artifacts") or {}, "cover_letter", cover_choice["strategy"]
    )

    qa_entries = [
        entry
        for entry in profile_data.get("qa_bank", [])
        if isinstance(entry, dict) and entry.get("question_pattern")
    ]
    qa_responses = None
    if qa_entries:
        answers = answer_questions(
            questions=[entry["question_pattern"] for entry in qa_entries],
            job=job,
            profile_data=profile_data,
            qa_entries=qa_entries,
            use_llm=False,
        )
        qa_responses = {
            response.question: response.answer for response in answers if response.answer
        }
        if not qa_responses:
            qa_responses = None

    return resume_path, cover_letter_path, qa_responses


def _application_material_type(document_type: str, choice: dict) -> str:
    # Patch mode produces a tailored DOCX from the user's library. Ask
    # for DOCX so the apply pipeline uploads that patched file instead
    # of a stale template PDF side-effect.
    if choice.get("strategy") == "use_library":
        suffix = _source_type_to_material_suffix(
            _library_document_source_type(choice.get("document_id"))
        ) or "pdf"
    elif choice.get("strategy") == "patch_existing":
        suffix = "docx"
    else:
        suffix = "pdf"
    prefix = "resume" if document_type == "resume" else "cover_letter"
    return f"{prefix}_{suffix}"


def _source_type_to_material_suffix(source_type: str | None) -> str | None:
    if not source_type:
        return None
    normalized = str(source_type).strip().lower()
    return "tex" if normalized == "latex" else normalized


def _library_document_source_type(document_id: str | None) -> str | None:
    if not document_id:
        return None
    try:
        from uuid import UUID

        from src.core.database import get_session_factory
        from src.documents.user_documents import get_document

        doc_uuid = UUID(document_id)
        session_factory = get_session_factory(load_config())
        with session_factory() as session:
            row = get_document(session, doc_uuid)
            return getattr(row, "source_type", None) if row is not None else None
    except Exception as exc:  # noqa: BLE001
        logger.warning("Could not resolve library document source type: %s", exc)
        return None


def _pick_application_artifact(artifacts: dict, document_type: str, strategy: str) -> Path | None:
    prefix = "resume" if document_type == "resume" else "cover_letter"
    order = (
        [f"{prefix}_docx", f"{prefix}_pdf", f"{prefix}_tex", f"{prefix}_txt"]
        if strategy in {"patch_existing", "use_library"}
        else [f"{prefix}_pdf", f"{prefix}_docx", f"{prefix}_tex", f"{prefix}_txt"]
    )
    for key in order:
        value = artifacts.get(key)
        if value:
            return Path(value)
    return None


async def _execute_application(
    *,
    url: str,
    ats_type: str,
    job=None,
    profile_data: dict,
    resume_path: Path | None,
    cover_letter_path: Path | None,
    qa_responses: dict[str, str] | None,
    auto_submit: bool,
    headless: bool,
    state=None,
):
    from src.execution.ats.ashby import AshbyAdapter
    from src.execution.ats.generic import GenericAdapter
    from src.execution.ats.greenhouse import GreenhouseAdapter
    from src.execution.ats.lever import LeverAdapter
    from src.execution.browser import BrowserManager

    adapter_map = {
        "ashby": AshbyAdapter,
        "company_site": GenericAdapter,
        "greenhouse": GreenhouseAdapter,
        "lever": LeverAdapter,
        "workday": GenericAdapter,
    }
    adapter_cls = adapter_map.get(ats_type)
    if not adapter_cls:
        raise ValueError(f"No adapter for ATS type: {ats_type}")

    if state is None:
        state = ApplicationState(str(uuid.uuid4()))

    if state.status == AppStatus.DISCOVERED:
        state.transition(AppStatus.QUALIFIED)
    if state.status == AppStatus.QUALIFIED:
        state.transition(AppStatus.MATERIALS_READY)

    application_url = _normalize_application_url_for_ats(url, ats_type)

    async with BrowserManager(headless=headless) as browser:
        adapter = adapter_cls(browser=browser)
        page = await browser.new_page()
        result = await adapter.apply(
            page=page,
            application_url=application_url,
            state=state,
            profile_data=profile_data,
            resume_path=resume_path,
            cover_letter_path=cover_letter_path,
            qa_responses=qa_responses,
            auto_submit=auto_submit,
            job_context=job,
        )
        return result


def _detect_ats_from_url(url: str) -> str | None:
    url_lower = url.lower()
    if not url_lower.startswith(("http://", "https://")):
        return None
    if "linkedin.com" in url_lower:
        return None
    if "ashbyhq.com" in url_lower:
        return "ashby"
    if "greenhouse.io" in url_lower:
        return "greenhouse"
    if "lever.co" in url_lower:
        return "lever"
    if "workday" in url_lower or "myworkdayjobs.com" in url_lower:
        return "workday"
    return "company_site"


def _normalize_application_url_for_ats(url: str, ats_type: str) -> str:
    if ats_type != "ashby":
        return url

    parsed = urlparse(url)
    path = parsed.path.rstrip("/")
    if not path or path.endswith("/application"):
        return url

    return parsed._replace(path=f"{path}/application").geturl()


def _is_linkedin_url(url: str) -> bool:
    return "linkedin.com" in (url or "").lower()


def _load_profile(profile_id: str | None = None) -> dict | None:
    profile_path = get_profile_path(profile_id) if profile_id else get_active_profile_path()
    if profile_path is None or not profile_path.exists():
        return None

    from src.memory.profile import load_profile_yaml

    return load_profile_yaml(profile_path)


def _load_job_for_application(url: str, ats_type: str) -> tuple:
    """Return ``(raw_job, hydrated)`` for the given apply URL.

    ``hydrated`` is True when we have real job context (a stored job from
    the database or a freshly scraped one from a supported ATS). It is
    False when we had to fall back to ``_synthesize_job_from_url``, which
    produces a placeholder Job with ``title="Unknown Role"`` and no
    description. Callers must refuse to generate tailored materials in
    that case so we never run the LLM/template pipeline on an empty JD.
    """

    db_job = _find_db_job_by_url(url)
    if db_job is not None:
        return _job_to_raw_job(db_job), True

    index_job = _find_index_job_by_url(url)
    if index_job is not None:
        return index_job, True

    fetched_job = _fetch_job_from_ats(url, ats_type)
    if fetched_job is not None:
        return fetched_job, True

    return _synthesize_job_from_url(url, ats_type), False


def _find_db_job_by_url(url: str):
    try:
        from src.core.config import load_config
        from src.core.database import get_session_factory
        from src.core.models import Job

        session_factory = get_session_factory(load_config())
        with session_factory() as session:
            return session.query(Job).filter(Job.application_url == url).first()
    except Exception as exc:
        logger.debug("DB lookup skipped for %s: %s", url, exc)
        return None


def _find_index_job_by_url(url: str):
    try:
        from sqlalchemy import or_, select

        from src.core.config import load_config
        from src.core.database import get_session_factory
        from src.core.models import JobPosting, JobSnapshot

        session_factory = get_session_factory(load_config())
        with session_factory() as session:
            stmt = (
                select(JobPosting, JobSnapshot)
                .join(JobSnapshot, JobSnapshot.id == JobPosting.latest_snapshot_id)
                .where(
                    or_(
                        JobPosting.canonical_url == url,
                        JobSnapshot.application_url == url,
                    )
                )
                .order_by(JobSnapshot.scraped_at.desc())
                .limit(1)
            )
            row = session.execute(stmt).first()
            if row is None:
                return None
            posting, snapshot = row
            return _raw_job_from_index_posting(posting, snapshot)
    except Exception as exc:
        logger.debug("Job Index lookup skipped for %s: %s", url, exc)
        return None


def _fetch_job_from_ats(url: str, ats_type: str):
    locator = _parse_ats_job_locator(url, ats_type)
    if locator is None:
        return None

    company_slug, job_id = locator

    try:
        if ats_type == "greenhouse":
            from src.intake.greenhouse import GreenhouseScraper

            with GreenhouseScraper() as scraper:
                return scraper.fetch_job(company_slug, job_id)
        if ats_type == "lever":
            from src.intake.lever import LeverScraper

            with LeverScraper() as scraper:
                return scraper.fetch_job(company_slug, job_id)
    except Exception as exc:
        logger.warning("Failed to fetch %s job details for %s: %s", ats_type, url, exc)

    return None


def _parse_ats_job_locator(url: str, ats_type: str) -> tuple[str, str] | None:
    parsed = urlparse(url)
    parts = [part for part in parsed.path.split("/") if part]

    if ats_type == "greenhouse":
        if len(parts) >= 3 and parts[1] == "jobs":
            return parts[0], parts[2]
        if "jobs" in parts:
            idx = parts.index("jobs")
            if idx > 0 and idx + 1 < len(parts):
                return parts[idx - 1], parts[idx + 1]

    if ats_type == "lever" and len(parts) >= 2:
        return parts[0], parts[1]

    return None


def _synthesize_job_from_url(url: str, ats_type: str):
    from src.intake.schema import RawJob

    parsed = urlparse(url)
    parts = [part for part in parsed.path.split("/") if part]
    company = parts[0].replace("-", " ").replace("_", " ").title() if parts else parsed.netloc
    source_id = parts[-1] if parts else parsed.netloc

    return RawJob(
        source=ats_type,
        source_id=source_id,
        company=company or "Unknown Company",
        title="Unknown Role",
        application_url=url,
        ats_type=ats_type,
        description=None,
    )


def _job_to_raw_job(job):
    from src.intake.schema import JobRequirements, RawJob

    return RawJob(
        id=job.id,
        source=job.source or "unknown",
        source_id=job.source_id or str(job.id),
        company=job.company,
        title=job.title,
        location=job.location,
        employment_type=job.employment_type or "unknown",
        seniority=job.seniority or "unknown",
        description=job.description,
        requirements=JobRequirements.model_validate(job.requirements or {}),
        application_url=job.application_url,
        ats_type=job.ats_type or job.source or "unknown",
        raw_data=job.raw_data or {},
        discovered_at=job.discovered_at,
        expires_at=job.expires_at,
    )


def _create_tracking_application(
    *,
    job,
    match_score: float | None,
    resume_path: Path,
    cover_letter_path: Path | None,
) -> uuid.UUID | None:
    try:
        from src.core.config import load_config
        from src.core.database import get_session_factory
        from src.tracker.database import create_application

        session_factory = get_session_factory(load_config())
        with session_factory() as session:
            db_job = _get_or_create_job_record(session, job)
            application = create_application(
                session,
                db_job.id,
                match_score=match_score,
                resume_version=str(resume_path),
                cover_letter_version=str(cover_letter_path) if cover_letter_path else None,
            )
            session.commit()
            return application.id
    except Exception as exc:
        logger.warning("Tracking create skipped for %s at %s: %s", job.title, job.company, exc)
        return None


def _sync_tracking_application(app_id: uuid.UUID, state, result, qa_responses: dict | None) -> None:
    try:
        from src.core.config import load_config
        from src.core.database import get_session_factory
        from src.tracker.database import sync_state_to_db

        session_factory = get_session_factory(load_config())
        with session_factory() as session:
            sync_state_to_db(
                session,
                app_id,
                state,
                {
                    "fields_filled": result.fields_filled,
                    "fields_total": result.fields_total,
                    "files_uploaded": result.files_uploaded,
                    "qa_responses": qa_responses,
                    "screenshot_paths": [str(path) for path in result.screenshots],
                    # Phase 18.5: forward the per-field record published
                    # by the ATS adapter. ``getattr`` keeps older custom
                    # adapters that haven't migrated to the new attr
                    # from blowing up here.
                    "fill_details": getattr(result, "fill_details", []) or [],
                },
            )
            session.commit()
    except Exception as exc:
        logger.warning("Tracking sync skipped for application %s: %s", app_id, exc)


def _get_or_create_job_record(session, job):
    from src.core.models import Job

    existing = (
        session.query(Job)
        .filter(
            Job.source == job.source,
            Job.company == job.company,
            Job.source_id == job.source_id,
        )
        .first()
    )
    if existing is not None:
        return existing

    if job.application_url:
        existing = session.query(Job).filter(Job.application_url == job.application_url).first()
        if existing is not None:
            return existing

    db_job = Job(
        id=job.id,
        source=job.source,
        source_id=job.source_id,
        company=job.company,
        title=job.title,
        location=job.location,
        employment_type=job.employment_type,
        seniority=job.seniority,
        description=job.description,
        requirements=job.requirements.model_dump(),
        visa_sponsorship=job.requirements.visa_sponsorship,
        ats_type=job.ats_type,
        application_url=job.application_url,
        raw_data=job.raw_data,
        discovered_at=job.discovered_at,
        expires_at=job.expires_at,
    )
    session.add(db_job)
    session.flush()
    return db_job


def _isoformat(value) -> str | None:
    return value.isoformat() if value is not None else None
