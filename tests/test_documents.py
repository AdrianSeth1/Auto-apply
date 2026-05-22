"""Tests for the document processing layer."""

from datetime import UTC, datetime
from pathlib import Path
from subprocess import CompletedProcess
from unittest.mock import patch

import pytest
from docx import Document

from src.documents.docx_engine import (
    _resolved_section_order as _resolved_section_order_docx,
)
from src.documents.docx_engine import (
    build_cover_letter_from_ir,
    build_resume,
    build_resume_from_ir,
    create_default_template,
    substitute_placeholders,
)
from src.documents.file_manager import (
    get_output_paths,
    make_filename,
    make_filename_from_pattern,
    next_template_sequence,
)
from src.documents.latex_engine import (
    _resolved_section_order as _resolved_section_order_latex,
)
from src.documents.latex_engine import (
    build_resume_tex_from_ir,
    compile_latex_to_pdf,
    latex_escape,
)
from src.documents.templates import (
    create_latex_template_package,
    discover_templates,
    editable_style_options,
    ensure_template_package,
    get_template_package_detail,
    get_template_path,
    list_template_packages,
    register_template,
    save_uploaded_template_package,
    update_docx_template_styles,
    update_latex_template_package,
)
from src.generation.ir import (
    CoverLetterDocument,
    CoverLetterParagraph,
    ResumeBullet,
    ResumeDocument,
    ResumeItem,
)

TMP_DIR = Path("data/output/_test")

SAMPLE_IDENTITY = {
    "full_name": "Jane Doe",
    "email": "jane@example.com",
    "phone": "+1 604-000-0000",
    "location": "Vancouver, BC",
    "linkedin_url": "linkedin.com/in/janedoe",
    "github_url": "github.com/janedoe",
}

SAMPLE_EDUCATION = [
    {
        "institution": "UBC",
        "degree": "Master of Science",
        "field": "Computer Science",
        "start_date": "2024-09",
        "end_date": "2026-04",
        "gpa": "4.0/4.0",
        "relevant_courses": [
            {"name": "Distributed Systems", "tags": ["distributed"]},
        ],
    }
]

SAMPLE_EXPERIENCES = [
    {
        "company": "Stripe",
        "title": "Software Engineer Intern",
        "location": "San Francisco, CA",
        "start_date": "2025-05",
        "end_date": "2025-08",
        "bullets": [
            {
                "text": "Built payment retry logic handling 1M+ transactions/day",
                "tags": ["backend"],
            },
        ],
    }
]

SAMPLE_PROJECTS = [
    {
        "name": "AutoApply",
        "role": "Lead Developer",
        "tech_stack": ["Python", "Playwright"],
        "bullets": [
            {
                "text": "Developed an AI agent for automated job applications",
                "tags": ["python", "ai"],
            },
        ],
    }
]

SAMPLE_SKILLS = {
    "languages": ["Python", "Java"],
    "frameworks": ["FastAPI", "React"],
    "databases": ["PostgreSQL"],
    "tools": ["Docker"],
}


@pytest.fixture(autouse=True)
def cleanup():
    yield
    import shutil

    if TMP_DIR.exists():
        shutil.rmtree(TMP_DIR)


class TestDocxEngine:
    def test_create_default_template(self):
        template_path = TMP_DIR / "default_template.docx"
        result = create_default_template(template_path)
        assert result.exists()
        doc = Document(str(result))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "{{FULL_NAME}}" in full_text
        assert "{{EDUCATION_BLOCK}}" in full_text
        assert "{{EXPERIENCE_BLOCK}}" in full_text

    def test_substitute_placeholders(self):
        template_path = TMP_DIR / "template.docx"
        create_default_template(template_path)
        doc = Document(str(template_path))
        substitute_placeholders(doc, {"FULL_NAME": "Jane Doe", "EMAIL": "jane@example.com"})
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "Jane Doe" in full_text
        assert "jane@example.com" in full_text

    def test_build_resume(self):
        template_path = TMP_DIR / "template.docx"
        create_default_template(template_path)
        output_path = TMP_DIR / "resume_test.docx"

        result = build_resume(
            template_path=template_path,
            identity=SAMPLE_IDENTITY,
            education=SAMPLE_EDUCATION,
            experiences=SAMPLE_EXPERIENCES,
            projects=SAMPLE_PROJECTS,
            skills=SAMPLE_SKILLS,
            selected_bullets={},
            output_path=output_path,
        )

        assert result.exists()
        doc = Document(str(result))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "Jane Doe" in full_text
        assert "Stripe" in full_text
        assert "AutoApply" in full_text
        assert "Python" in full_text
        assert "UBC" in full_text

    def test_build_resume_from_ir(self):
        template_path = TMP_DIR / "template_ir.docx"
        create_default_template(template_path)
        output_path = TMP_DIR / "resume_ir.docx"
        document = ResumeDocument(
            target_role="Backend Intern",
            company="Stripe",
            header=SAMPLE_IDENTITY,
            education=SAMPLE_EDUCATION,
            skills=SAMPLE_SKILLS,
            section_order=["header", "projects", "skills", "experience", "education"],
            experiences=[
                ResumeItem(
                    source_id="experience:stripe",
                    source_type="experience",
                    name="Stripe",
                    organization="Stripe",
                    title="Software Engineer Intern",
                    location="San Francisco, CA",
                    start_date="2025-05",
                    end_date="2025-08",
                    bullets=[
                        ResumeBullet(
                            text="Built payment retry logic handling 1M+ transactions/day",
                            source_id="experience:stripe:bullet:0",
                            source_type="experience",
                            source_entity="Stripe - Software Engineer Intern",
                        )
                    ],
                )
            ],
            projects=[],
        )

        result = build_resume_from_ir(template_path, document, output_path)

        assert result.exists()
        doc = Document(str(result))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "Jane Doe" in full_text
        assert "Stripe" in full_text
        assert "payment retry" in full_text
        paragraph_text = [p.text for p in doc.paragraphs]
        assert paragraph_text.index("Skills") < paragraph_text.index("Experience")
        assert paragraph_text.index("Experience") < paragraph_text.index("Education")

    def test_build_resume_skips_empty_gpa(self, tmp_path):
        """Profile fields the user did not fill out must NOT appear as
        literal "None" / "N/A" on the rendered resume. This is the
        regression for the GPA: None bug -- f-string interpolation of
        edu.get('gpa') used to slip through even when the field was
        missing."""
        template_path = tmp_path / "template_no_gpa.docx"
        create_default_template(template_path)
        output_path = tmp_path / "resume_no_gpa.docx"
        education = [
            {
                "institution": "Test U",
                "degree": "BSc",
                "field": "CS",
                # gpa is deliberately absent
                "start_date": "2024-09",
                "end_date": "2028-05",
            },
            {
                "institution": "Test U2",
                "degree": "BSc",
                "field": "CS",
                "gpa": "None",  # also catch the literal string "None"
                "start_date": "2024-09",
                "end_date": "2028-05",
            },
        ]
        document = ResumeDocument(
            target_role="SWE",
            company="Acme",
            header=SAMPLE_IDENTITY,
            education=education,
            skills={},
            section_order=["header", "education"],
            experiences=[],
            projects=[],
        )

        result = build_resume_from_ir(template_path, document, output_path)
        doc = Document(str(result))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "GPA: None" not in full_text
        assert "None" not in full_text  # not anywhere as a stray field

    def test_build_resume_renders_inline_bold_and_italic_runs(self, tmp_path):
        """LLM-emitted inline ``**bold**`` and ``*italic*`` markers must
        survive into the rendered DOCX as run-level formatting (not
        leak through as literal asterisks)."""
        from src.generation.ir import ResumeBullet, ResumeDocument, ResumeItem

        template_path = tmp_path / "template_runs.docx"
        create_default_template(template_path)
        output_path = tmp_path / "resume_runs.docx"
        document = ResumeDocument(
            target_role="SWE",
            company="Acme",
            header=SAMPLE_IDENTITY,
            education=[],
            skills={},
            section_order=["header", "experience"],
            experiences=[
                ResumeItem(
                    source_id="exp:0",
                    source_type="experience",
                    name="Acme",
                    title="SWE",
                    organization="Acme",
                    bullets=[
                        ResumeBullet(
                            text="Built **REST APIs** with *FastAPI* serving 1M+ req/day",
                            source_id="b:0",
                            source_type="experience",
                            source_entity="Acme",
                        ),
                    ],
                )
            ],
            projects=[],
        )

        result = build_resume_from_ir(template_path, document, output_path)

        doc = Document(str(result))
        target_para = None
        for paragraph in doc.paragraphs:
            if "REST APIs" in paragraph.text:
                target_para = paragraph
                break
        assert target_para is not None, "bullet missing from rendered DOCX"
        # Asterisks should NOT appear in the rendered text -- they
        # were structural markup, not content.
        assert "**" not in target_para.text
        assert "*F" not in target_para.text
        # At least one run is bold ("REST APIs") and one is italic ("FastAPI").
        runs = list(target_para.runs)
        assert any(run.bold and "REST APIs" in run.text for run in runs)
        assert any(run.italic and "FastAPI" in run.text for run in runs)

    def test_build_resume_emphasis_font_swaps_bold_run_font(self, tmp_path):
        """When the manifest declares an ``emphasis_font`` the bold runs
        emitted by inline markup pick up that font, while plain runs
        keep the body font. Regression for the second-font feature."""
        from src.documents.templates import default_manifest
        from src.generation.ir import ResumeBullet, ResumeDocument, ResumeItem

        manifest = default_manifest("resume")
        manifest = manifest.model_copy(update={"emphasis_font": "Georgia"})

        template_path = tmp_path / "template_emphasis.docx"
        create_default_template(template_path)
        output_path = tmp_path / "resume_emphasis.docx"
        document = ResumeDocument(
            target_role="SWE",
            company="Acme",
            header=SAMPLE_IDENTITY,
            education=[],
            skills={},
            section_order=["header", "experience"],
            experiences=[
                ResumeItem(
                    source_id="exp:0",
                    source_type="experience",
                    name="Acme",
                    title="SWE",
                    organization="Acme",
                    bullets=[
                        ResumeBullet(
                            text="Owned **payment retries** on a high-traffic service",
                            source_id="b:0",
                            source_type="experience",
                            source_entity="Acme",
                        ),
                    ],
                )
            ],
            projects=[],
        )

        result = build_resume_from_ir(
            template_path, document, output_path, manifest=manifest
        )

        doc = Document(str(result))
        target_para = next(
            p for p in doc.paragraphs if "payment retries" in p.text
        )
        bold_runs = [run for run in target_para.runs if run.bold]
        plain_runs = [run for run in target_para.runs if not run.bold]
        assert bold_runs, "no bold run rendered"
        assert any(run.font.name == "Georgia" for run in bold_runs), (
            "emphasis_font='Georgia' should swap onto bold runs"
        )
        # Plain runs must NOT take Georgia just because the bold ones did.
        for run in plain_runs:
            assert run.font.name != "Georgia"

    def test_build_resume_inserts_divider_after_section(self, tmp_path):
        """Fit Planner output may include ``dividers_after`` to request
        a horizontal rule below a section. Verify the renderer emits
        the rule as an empty paragraph with a bottom border."""
        from src.generation.ir import ResumeDocument

        template_path = tmp_path / "template_hr.docx"
        create_default_template(template_path)
        output_path = tmp_path / "resume_hr.docx"
        document = ResumeDocument(
            target_role="SWE",
            company="Acme",
            header=SAMPLE_IDENTITY,
            education=SAMPLE_EDUCATION,
            skills=SAMPLE_SKILLS,
            section_order=["header", "education", "skills"],
            experiences=[],
            projects=[],
            dividers_after=["education"],
        )

        result = build_resume_from_ir(template_path, document, output_path)
        doc = Document(str(result))

        # Locate the paragraph right after the last education paragraph.
        # The HR is rendered as an empty paragraph carrying a w:pBdr
        # element with a bottom border. python-docx exposes the XML
        # element on paragraph._p so we can introspect.
        paragraphs = list(doc.paragraphs)
        hr_count = 0
        for paragraph in paragraphs:
            if paragraph.text.strip():
                continue
            p_pr = paragraph._p.find(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr"
            )
            if p_pr is None:
                continue
            p_bdr = p_pr.find(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pBdr"
            )
            if p_bdr is not None and p_bdr.find(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bottom"
            ) is not None:
                hr_count += 1
        assert hr_count >= 1, "expected at least one HR paragraph"

    def test_build_resume_renders_custom_sections(self, tmp_path):
        """Profile sections outside the canonical buckets (VOLUNTEER,
        AWARDS, etc.) must show up on the rendered resume. Otherwise
        the user's content silently disappears between import and
        generation."""
        from src.generation.ir import CustomSection, CustomSectionEntry

        template_path = tmp_path / "template_custom.docx"
        create_default_template(template_path)
        output_path = tmp_path / "resume_custom.docx"
        document = ResumeDocument(
            target_role="SWE",
            company="Acme",
            header=SAMPLE_IDENTITY,
            education=SAMPLE_EDUCATION,
            skills=SAMPLE_SKILLS,
            section_order=["header", "education", "skills"],
            experiences=[],
            projects=[],
            custom_sections=[
                CustomSection(
                    title="VOLUNTEER EXPERIENCE",
                    entries=[
                        CustomSectionEntry(
                            title="Tutor",
                            organization="Local Library",
                            start_date="2023-01",
                            end_date="2023-08",
                            bullets=["Mentored 12 students in introductory CS"],
                        ),
                    ],
                ),
                CustomSection(
                    title="AWARDS",
                    entries=[
                        CustomSectionEntry(
                            title="Dean's List",
                            details="2023, 2024",
                        ),
                    ],
                ),
            ],
        )

        result = build_resume_from_ir(template_path, document, output_path)
        doc = Document(str(result))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "VOLUNTEER EXPERIENCE" in full_text
        assert "Tutor" in full_text
        assert "Mentored 12 students" in full_text
        assert "AWARDS" in full_text
        assert "Dean's List" in full_text

    def test_template_package_renderer_uses_named_styles(self, tmp_path):
        package = ensure_template_package("resume", template_root=tmp_path)
        output_path = tmp_path / "resume_named_styles.docx"
        document = ResumeDocument(
            target_role="Backend Intern",
            company="Stripe",
            header=SAMPLE_IDENTITY,
            education=SAMPLE_EDUCATION,
            skills=SAMPLE_SKILLS,
            section_order=["header", "skills"],
            experiences=[],
            projects=[],
        )

        result = build_resume_from_ir(
            package.template_path,
            document,
            output_path,
            manifest=package.manifest,
        )

        doc = Document(str(result))
        styles_by_text = {paragraph.text: paragraph.style.name for paragraph in doc.paragraphs}
        assert styles_by_text["Jane Doe"] == "Resume.Name"
        assert styles_by_text["Skills"] == "Resume.SectionHeading"
        assert "{{resume.sections}}" not in " ".join(p.text for p in doc.paragraphs)


class TestSectionOrderResolution:
    """Regression guard: `_resolved_section_order` must respect the
    caller's explicit ordering. Earlier behavior appended any default
    section missing from the order, which silently tacked Summary onto
    the end of student/intern resumes whose order deliberately
    excluded it."""

    def _doc(self, order):
        return ResumeDocument(
            target_role="Backend Intern",
            company="Stripe",
            header=SAMPLE_IDENTITY,
            section_order=order,
        )

    def test_docx_omitted_section_stays_omitted(self):
        # Student-style order with no "summary".
        order = ["header", "education", "skills", "projects", "experience"]
        resolved = _resolved_section_order_docx(self._doc(order))
        assert "summary" not in resolved
        assert resolved == order

    def test_latex_omitted_section_stays_omitted(self):
        order = ["header", "education", "skills", "projects", "experience"]
        resolved = _resolved_section_order_latex(self._doc(order))
        assert "summary" not in resolved
        assert resolved == order

    def test_empty_order_falls_back_to_default(self):
        # An empty list is the documented "use defaults" signal.
        # The default order never includes "summary".
        resolved = _resolved_section_order_docx(self._doc([]))
        assert "summary" not in resolved
        assert resolved[0] == "header"

    def test_summary_in_explicit_order_is_filtered(self):
        # Even if a legacy caller / manifest asks for summary, it must
        # be filtered out -- the system never renders a Summary section.
        order = ["header", "summary", "skills", "experience"]
        resolved_docx = _resolved_section_order_docx(self._doc(order))
        resolved_tex = _resolved_section_order_latex(self._doc(order))
        assert "summary" not in resolved_docx
        assert "summary" not in resolved_tex


class TestFileManager:
    def test_make_filename_resume(self):
        date = datetime(2026, 4, 2, tzinfo=UTC)
        name = make_filename("resume", "Stripe", "Backend Engineer", date)
        assert name == "resume_stripe_backend_engineer_2026-04-02.docx"

    def test_make_filename_cover(self):
        date = datetime(2026, 4, 2, tzinfo=UTC)
        name = make_filename("cover", "Google LLC", "SWE Intern", date, ext="pdf")
        assert name == "cover_google_llc_swe_intern_2026-04-02.pdf"

    def test_make_filename_special_chars(self):
        date = datetime(2026, 4, 2, tzinfo=UTC)
        name = make_filename("resume", "A/B & Co.", "C++ Dev", date)
        # Should not contain special chars
        assert "/" not in name
        assert "&" not in name

    def test_get_output_paths(self):
        date = datetime(2026, 4, 2, tzinfo=UTC)
        paths = get_output_paths(TMP_DIR, "Stripe", "Backend Intern", date)
        assert "resume_docx" in paths
        assert "resume_pdf" in paths
        assert "resume_tex" in paths
        assert "cover_docx" in paths
        assert "cover_pdf" in paths
        assert "cover_tex" in paths
        assert paths["resume_docx"].suffix == ".docx"

    def test_make_filename_from_pattern_profile_seq(self):
        name = make_filename_from_pattern(
            doc_type="resume",
            ext="docx",
            pattern="type_profile_seq",
            company="Stripe",
            role="Backend Intern",
            profile_name="Jane Doe",
            seq=3,
        )
        assert name == "resume_jane_doe_003.docx"

    def test_make_filename_from_pattern_custom_seq_fallback(self):
        # Empty custom label falls back to a slug of the company.
        name = make_filename_from_pattern(
            doc_type="cover",
            ext="pdf",
            pattern="type_custom_seq",
            company="ACME Robotics",
            role="ignored",
            custom_label="",
            seq=12,
        )
        assert name == "cover_acme_robotics_012.pdf"

    def test_next_template_sequence_persists_monotonically(self, tmp_path):
        first = next_template_sequence(tmp_path, "resume:ats_v1")
        second = next_template_sequence(tmp_path, "resume:ats_v1")
        other = next_template_sequence(tmp_path, "resume:other")
        third = next_template_sequence(tmp_path, "resume:ats_v1")
        assert (first, second, third) == (1, 2, 3)
        assert other == 1

    def test_get_output_paths_respects_pattern(self, tmp_path):
        paths = get_output_paths(
            tmp_path,
            "Stripe",
            "Backend Intern",
            pattern="type_profile_seq",
            profile_name="Jane Doe",
            template_id="ats_v1",
        )
        assert paths["resume_docx"].name == "resume_jane_doe_001.docx"
        assert paths["cover_docx"].name == "cover_jane_doe_001.docx"
        assert paths["resume_pdf"].suffix == ".pdf"
        assert paths["resume_tex"].suffix == ".tex"


class TestTemplateRegistry:
    def test_register_and_get(self):
        template_path = TMP_DIR / "my_template.docx"
        create_default_template(template_path)
        register_template("test_template", template_path)
        assert get_template_path("test_template") == template_path

    def test_get_missing_template_raises(self):
        with pytest.raises(KeyError):
            get_template_path("nonexistent_template_xyz")

    def test_discover_templates(self):
        # Create two templates
        create_default_template(TMP_DIR / "modern.docx")
        create_default_template(TMP_DIR / "classic.docx")
        discover_templates(TMP_DIR)
        # Both should now be registered
        assert get_template_path("modern") is not None
        assert get_template_path("classic") is not None

    def test_list_template_packages(self, tmp_path):
        resume_package = ensure_template_package("resume", template_root=tmp_path)
        ensure_template_package("cover_letter", template_root=tmp_path)
        (resume_package.directory / "preview.pdf").write_bytes(b"pdf")

        templates = list_template_packages(template_root=tmp_path)

        assert templates["resume"][0]["template_id"] == "ats_single_column_v1"
        assert templates["resume"][0]["preview_pdf"] == "preview.pdf"
        assert not Path(templates["resume"][0]["preview_pdf"]).is_absolute()
        assert templates["resume"][0]["validation"]["ok"] is True
        assert templates["cover_letter"][0]["template_id"] == "classic_v1"

        cover_doc = Document(str(tmp_path / "cover_letter" / "classic_v1" / "template.docx"))
        cover_texts = [paragraph.text for paragraph in cover_doc.paragraphs]
        body_style = cover_doc.styles["CoverLetter.Body"]
        assert "Dear Hiring Manager," in cover_texts
        assert "Sincerely," in cover_texts
        assert "{{date}}" in cover_texts
        assert "{{recipient.block}}" in cover_texts
        assert "Enclosure" in cover_texts
        assert body_style.font.name == "Times New Roman"
        assert body_style.font.size.pt == 11
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        header_style = cover_doc.styles["CoverLetter.Header"]
        assert header_style.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.RIGHT
        assert round(cover_doc.sections[0].left_margin.inches, 2) == 0.85

    def test_build_cover_letter_from_ir_renders_classic_frame(self, tmp_path):
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        package = ensure_template_package("cover_letter", template_root=tmp_path)
        output_path = tmp_path / "cover.docx"
        document = CoverLetterDocument(
            recipient={"company": "TestCo", "location": "Calgary, AB (Hybrid)"},
            applicant={
                "name": "Test User",
                "location": "Vancouver, BC (Remote)",
                "phone": "+1 604-000-0000",
                "email": "test@example.com",
                "linkedin_url": "linkedin.com/in/test",
                "portfolio_url": "test.example.com",
            },
            paragraphs=[
                CoverLetterParagraph(type="opening", text="Opening body paragraph."),
                CoverLetterParagraph(type="closing", text="Closing body paragraph."),
            ],
        )

        result = build_cover_letter_from_ir(
            document,
            output_path,
            template_path=package.template_path,
            manifest=package.manifest,
        )

        doc = Document(str(result))
        texts = [paragraph.text for paragraph in doc.paragraphs]
        assert doc.styles["CoverLetter.Header"].paragraph_format.alignment == (
            WD_ALIGN_PARAGRAPH.RIGHT
        )
        assert "Test User" in texts
        assert any("Vancouver, BC" in text and "test@example.com" in text for text in texts)
        assert not any("Remote" in text or "Hybrid" in text for text in texts)
        assert not any("linkedin" in text.lower() or "test.example.com" in text for text in texts)
        assert any(" 202" in text and "," in text for text in texts)
        assert any(
            "Hiring Team" in text and "TestCo" in text and "Calgary, AB" in text
            for text in texts
        )
        assert texts.index("Dear Hiring Manager,") < texts.index("Opening body paragraph.")
        assert texts.index("Closing body paragraph.") < texts.index("Sincerely,")
        assert "Enclosure" in texts

    def test_save_uploaded_template_package(self, tmp_path):
        upload_docx = tmp_path / "upload.docx"
        doc = Document()
        doc.add_paragraph("Uploaded resume template")
        doc.save(str(upload_docx))

        template = save_uploaded_template_package(
            document_type="resume",
            filename="upload.docx",
            content=upload_docx.read_bytes(),
            template_name="My Resume Template",
            template_root=tmp_path / "templates",
        )

        assert template["template_id"] == "my_resume_template"
        assert template["validation"]["ok"] is True

    def test_template_package_rejects_path_traversal_id(self, tmp_path):
        with pytest.raises(ValueError, match="Invalid template id"):
            ensure_template_package(
                "resume",
                "..\\escape",
                template_root=tmp_path / "templates",
            )

        assert not (tmp_path / "escape").exists()


class TestLatexTemplates:
    def test_create_latex_template_package(self, tmp_path):
        template = create_latex_template_package(
            document_type="resume",
            template_name="Technical LaTeX Resume",
            template_root=tmp_path / "templates",
        )

        assert template["template_id"] == "technical_latex_resume"
        assert template["renderer"] == "latex"
        assert template["supported_outputs"] == ["tex", "pdf"]
        assert template["validation"]["ok"] is True

        detail = get_template_package_detail(
            "resume",
            template["template_id"],
            template_root=tmp_path / "templates",
        )
        assert "{{resume.sections}}" in detail["content"]

    def test_uploaded_latex_template_reports_missing_marker(self, tmp_path):
        template = save_uploaded_template_package(
            document_type="resume",
            filename="plain.tex",
            content=b"\\documentclass{article}\n\\begin{document}No marker\\end{document}\n",
            template_name="Plain LaTeX",
            template_root=tmp_path / "templates",
        )

        assert template["renderer"] == "latex"
        assert template["validation"]["ok"] is False
        assert template["validation"]["issues"][0]["type"] == "missing_block"
        assert template["validation"]["issues"][0]["severity"] == "error"
        assert "Add this marker exactly" in template["validation"]["issues"][0]["message"]

    def test_update_latex_template_package(self, tmp_path):
        template = create_latex_template_package(
            document_type="cover_letter",
            template_name="Editable Cover",
            template_root=tmp_path / "templates",
        )
        updated = update_latex_template_package(
            document_type="cover_letter",
            template_id=template["template_id"],
            template_name="Updated Cover",
            description="Edited in tests.",
            content="\\documentclass{article}\n\\begin{document}\n{{cover_letter.body}}\n\\end{document}\n",
            template_root=tmp_path / "templates",
        )

        assert updated["name"] == "Updated Cover"
        assert updated["description"] == "Edited in tests."
        assert updated["validation"]["ok"] is True

    def test_editable_style_options_lists_resume_and_cover_letter(self):
        resume_styles = editable_style_options("resume")
        cover_styles = editable_style_options("cover_letter")

        resume_keys = {entry["key"] for entry in resume_styles}
        cover_keys = {entry["key"] for entry in cover_styles}
        assert {"name", "normal", "section_heading", "bullet"} <= resume_keys
        assert {"header", "body", "signature"} <= cover_keys
        body_entry = next(entry for entry in cover_styles if entry["key"] == "body")
        assert body_entry["defaults"]["font"] == "Times New Roman"
        assert body_entry["supports_line_spacing"] is True
        normal_entry = next(entry for entry in resume_styles if entry["key"] == "normal")
        assert normal_entry["defaults"]["font"] == "Arial"

    def test_update_docx_template_styles_applies_overrides(self, tmp_path):
        ensure_template_package("cover_letter", template_root=tmp_path)
        updated = update_docx_template_styles(
            document_type="cover_letter",
            template_id="classic_v1",
            overrides={
                "body": {"font": "Garamond", "size": 12, "line_spacing": 1.2},
                "header": {"bold": False},
            },
            template_root=tmp_path,
        )

        assert updated["style_overrides"]["body"]["font"] == "Garamond"
        assert updated["style_overrides"]["body"]["size"] == 12
        assert updated["style_overrides"]["body"]["line_spacing"] == 1.2
        assert updated["style_overrides"]["header"]["bold"] is False

        cover_doc = Document(
            str(tmp_path / "cover_letter" / "classic_v1" / "template.docx")
        )
        body_style = cover_doc.styles["CoverLetter.Body"]
        assert body_style.font.name == "Garamond"
        assert body_style.font.size.pt == 12
        assert body_style.paragraph_format.line_spacing == pytest.approx(1.2)
        header_style = cover_doc.styles["CoverLetter.Header"]
        assert header_style.font.bold is False

    def test_update_template_settings_target_pages_and_filename_pattern(self, tmp_path):
        ensure_template_package("resume", template_root=tmp_path)
        updated = update_docx_template_styles(
            document_type="resume",
            template_id="ats_single_column_v1",
            overrides={},
            target_pages=2,
            filename_pattern="type_custom_seq",
            filename_custom_label="ML Engineering",
            template_root=tmp_path,
        )

        manifest = updated["manifest"]
        assert manifest["target_pages"] == 2
        assert manifest["capacity"]["max_pages"] == 2
        assert manifest["filename_pattern"] == "type_custom_seq"
        assert manifest["filename_custom_label"] == "ML Engineering"

    def test_update_template_settings_rejects_unknown_filename_pattern(self, tmp_path):
        ensure_template_package("resume", template_root=tmp_path)
        with pytest.raises(ValueError, match="filename_pattern"):
            update_docx_template_styles(
                document_type="resume",
                template_id="ats_single_column_v1",
                overrides={},
                filename_pattern="does_not_exist",
                template_root=tmp_path,
            )

    def test_update_template_settings_rejects_out_of_range_target_pages(self, tmp_path):
        ensure_template_package("resume", template_root=tmp_path)
        with pytest.raises(ValueError, match="target_pages"):
            update_docx_template_styles(
                document_type="resume",
                template_id="ats_single_column_v1",
                overrides={},
                target_pages=42,
                template_root=tmp_path,
            )

    def test_style_overrides_survive_ensure_template_package(self, tmp_path):
        """Regenerate path calls ensure_template_package -> _ensure_required_markers,
        which previously re-applied default styles on the default cover letter
        template and silently undid the user's Template Library edits."""
        ensure_template_package("cover_letter", template_root=tmp_path)
        update_docx_template_styles(
            document_type="cover_letter",
            template_id="classic_v1",
            overrides={"body": {"font": "Garamond", "size": 13}},
            template_root=tmp_path,
        )

        # Simulate a regenerate kick: the materials pipeline calls
        # ensure_template_package which used to clobber the override.
        package = ensure_template_package(
            "cover_letter", "classic_v1", template_root=tmp_path
        )
        cover_doc = Document(str(package.template_path))
        body_style = cover_doc.styles["CoverLetter.Body"]
        assert body_style.font.name == "Garamond"
        assert body_style.font.size.pt == 13

    def test_update_docx_template_styles_rejects_latex(self, tmp_path):
        template = create_latex_template_package(
            document_type="resume",
            template_name="LaTeX",
            template_root=tmp_path / "templates",
        )
        with pytest.raises(ValueError, match="DOCX"):
            update_docx_template_styles(
                document_type="resume",
                template_id=template["template_id"],
                overrides={},
                template_root=tmp_path / "templates",
            )

    def test_update_docx_template_styles_rejects_unknown_key(self, tmp_path):
        ensure_template_package("resume", template_root=tmp_path)
        with pytest.raises(ValueError, match="Unknown style key"):
            update_docx_template_styles(
                document_type="resume",
                template_id="ats_single_column_v1",
                overrides={"not_a_real_style": {"size": 12}},
                template_root=tmp_path,
            )

    def test_get_template_package_detail_includes_editable_styles(self, tmp_path):
        ensure_template_package("resume", template_root=tmp_path)
        detail = get_template_package_detail(
            "resume", "ats_single_column_v1", template_root=tmp_path
        )

        assert detail["editable_styles"], "expected DOCX templates to expose editable_styles"
        assert detail["style_overrides"] == {}
        assert detail["content"] is None

    def test_latex_escape(self):
        escaped = latex_escape(r"R&D_50% C# {x} \ path")

        assert r"R\&D\_50\% C\# \{x\}" in escaped
        assert r"\textbackslash{} path" in escaped

    def test_latex_inline_renders_bold_and_italic_markers(self):
        from src.documents.latex_engine import latex_inline

        rendered = latex_inline("Built **REST APIs** with *FastAPI*")
        # Bold + italic markers must turn into LaTeX commands rather
        # than leaking as literal asterisks.
        assert r"\textbf{REST APIs}" in rendered
        assert r"\textit{FastAPI}" in rendered
        assert "*" not in rendered

    def test_latex_inline_still_escapes_specials(self):
        from src.documents.latex_engine import latex_inline

        rendered = latex_inline("Tuned **A&B_50%** retries")
        assert r"\textbf{A\&B\_50\%}" in rendered
        # Plain ``&`` outside markers also gets escaped via latex_escape.
        rendered2 = latex_inline("R&D")
        assert rendered2 == r"R\&D"

    def test_build_resume_tex_from_ir(self, tmp_path):
        template = create_latex_template_package(
            document_type="resume",
            template_name="Render LaTeX Resume",
            template_root=tmp_path / "templates",
        )
        package = ensure_template_package(
            "resume",
            template["template_id"],
            template_root=tmp_path / "templates",
        )
        output_path = tmp_path / "resume.tex"
        document = ResumeDocument(
            target_role="Backend Intern",
            company="Stripe",
            header={**SAMPLE_IDENTITY, "full_name": "Jane & Doe"},
            education=SAMPLE_EDUCATION,
            skills=SAMPLE_SKILLS,
            section_order=["header", "skills", "experience", "education"],
            experiences=[
                ResumeItem(
                    source_id="experience:stripe",
                    source_type="experience",
                    name="Stripe",
                    organization="Stripe",
                    title="Software Engineer Intern",
                    bullets=[
                        ResumeBullet(
                            text="Built R&D tooling with C# and 50% less toil",
                            source_id="experience:stripe:bullet:0",
                            source_type="experience",
                            source_entity="Stripe",
                        )
                    ],
                )
            ],
            projects=[],
        )

        result = build_resume_tex_from_ir(
            package.template_path,
            document,
            output_path,
            manifest=package.manifest,
        )

        text = result.read_text(encoding="utf-8")
        assert r"Jane \& Doe" in text
        assert r"\section*{Skills}" in text
        assert r"R\&D tooling with C\# and 50\% less toil" in text
        assert "{{resume.sections}}" not in text

    def test_latex_renderer_rejects_missing_marker(self, tmp_path):
        template = create_latex_template_package(
            document_type="resume",
            template_name="Broken LaTeX Resume",
            template_root=tmp_path / "templates",
        )
        package = ensure_template_package(
            "resume",
            template["template_id"],
            template_root=tmp_path / "templates",
        )
        package.template_path.write_text(
            "\\documentclass{article}\n\\begin{document}No marker\\end{document}\n",
            encoding="utf-8",
        )

        with pytest.raises(ValueError, match="missing block marker"):
            build_resume_tex_from_ir(
                package.template_path,
                ResumeDocument(
                    target_role="Backend Intern",
                    company="Stripe",
                    header=SAMPLE_IDENTITY,
                ),
                tmp_path / "out.tex",
                manifest=package.manifest,
            )

    def test_compile_latex_to_pdf_disables_shell_escape(self, tmp_path):
        tex_path = tmp_path / "resume.tex"
        tex_path.write_text("\\documentclass{article}\n\\begin{document}Hi\\end{document}\n")
        pdf_path = tmp_path / "resume.pdf"

        def fake_run(command, *, cwd, **kwargs):
            Path(cwd, "main.pdf").write_bytes(b"%PDF")
            Path(cwd, "main.log").write_text("compile log", encoding="utf-8")
            return CompletedProcess(command, 0, stdout="", stderr="")

        def fake_which(name):
            return "pdflatex" if name == "pdflatex" else None

        with (
            patch("src.documents.latex_engine.shutil.which", side_effect=fake_which),
            patch("src.documents.latex_engine.subprocess.run", side_effect=fake_run) as run,
        ):
            result = compile_latex_to_pdf(tex_path, pdf_path)

        assert result == pdf_path
        assert pdf_path.exists()
        assert pdf_path.with_suffix(".log").read_text(encoding="utf-8") == "compile log"
        first_command = run.call_args_list[0].args[0]
        assert "-no-shell-escape" in first_command
