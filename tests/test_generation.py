"""Tests for src.generation — resume builder, cover letter, QA responder."""

from pathlib import Path

from src.documents.templates import default_manifest
from src.generation.cover_letter import (
    _clean_llm_cover_letter_output,
    _cover_letter_job_focus,
    _cover_letter_quality_issues,
    _format_education_brief,
    _generate_template,
    _infer_cover_letter_strategy,
    _normalize_cover_letter_dashes,
    _role_references,
    _select_evidence,
    generate_cover_letter,
)
from src.generation.evidence import select_relevant_evidence
from src.generation.qa_responder import (
    _estimate_experience_years,
    _find_qa_match,
    _get_variant_answer,
    _template_answer,
    answer_questions,
    classify_question,
)
from src.generation.resume_builder import (
    _rank_and_select,
    build_resume_document,
    extract_jd_tags,
    rewrite_bullets,
    select_bullets_for_jd,
)
from src.generation.validator import (
    validate_latex_artifacts,
    validate_resume_artifacts,
    validate_resume_document,
)
from src.generation.versions import save_generation_version
from src.intake.jd_parser import parse_requirements
from src.intake.schema import JobRequirements, RawJob

# ---------------------------------------------------------------------------
# Shared fixtures
# ---------------------------------------------------------------------------


def _make_job(**overrides) -> RawJob:
    defaults = {
        "source": "greenhouse",
        "source_id": "j1",
        "company": "TestCo",
        "title": "Backend Engineering Intern",
        "location": "Vancouver, BC",
        "employment_type": "internship",
        "seniority": "internship",
        "description": (
            "We need a backend intern with Python, FastAPI, PostgreSQL, and Docker experience."
        ),
        "ats_type": "greenhouse",
        "application_url": "https://example.com/apply",
    }
    defaults.update(overrides)
    return RawJob(**defaults)


_PROFILE = {
    "identity": {
        "full_name": "Test User",
        "email": "test@example.com",
        "location": "Vancouver, BC, Canada",
        "citizenship": "Chinese",
        "work_authorization": "Study Permit",
        "visa_sponsorship_needed": True,
    },
    "education": [
        {
            "institution": "UBC",
            "degree": "Bachelor of Science",
            "field": "Computer Science",
            "start_date": "2022-09",
            "end_date": "2026-05",
        },
    ],
    "work_experiences": [
        {
            "company": "Acme Corp",
            "title": "Software Dev Intern",
            "start_date": "2025-05",
            "end_date": "2025-08",
            "bullets": [
                {
                    "text": "Built REST APIs with Python and FastAPI",
                    "tags": ["python", "api", "backend", "fastapi"],
                },
                {"text": "Wrote unit tests achieving 90% coverage", "tags": ["testing", "python"]},
                {
                    "text": "Designed data pipeline with Apache Kafka",
                    "tags": ["kafka", "data", "backend"],
                },
            ],
        },
    ],
    "projects": [
        {
            "name": "CloudDeploy",
            "description": "CI/CD platform",
            "tech_stack": ["Python", "Docker", "AWS"],
            "bullets": [
                {
                    "text": "Containerized microservices with Docker and Kubernetes",
                    "tags": ["docker", "kubernetes", "devops"],
                },
                {"text": "Built React frontend dashboard", "tags": ["react", "frontend"]},
            ],
        },
    ],
    "skills": {
        "languages": ["Python", "TypeScript", "Java"],
        "frameworks": ["FastAPI", "React", "Next.js"],
        "databases": ["PostgreSQL", "Redis"],
        "tools": ["Docker", "Git", "AWS"],
    },
}


# ===========================================================================
# Resume builder tests
# ===========================================================================


class TestExtractJDTags:
    def test_from_requirements(self):
        job = _make_job()
        job.requirements = JobRequirements(
            must_have_skills=["Python", "PostgreSQL"],
            preferred_skills=["Docker"],
        )
        tags = extract_jd_tags(job)
        assert "python" in tags
        assert "postgresql" in tags
        assert "docker" in tags

    def test_from_title(self):
        job = _make_job(title="Python Backend Intern")
        tags = extract_jd_tags(job)
        assert "python" in tags
        assert "backend" in tags

    def test_dedup(self):
        job = _make_job(title="Python Developer")
        job.requirements = JobRequirements(must_have_skills=["Python"])
        tags = extract_jd_tags(job)
        assert tags.count("python") == 1

    def test_empty(self):
        job = _make_job(title="Intern")
        tags = extract_jd_tags(job)
        # No tech keywords in "Intern" alone
        assert isinstance(tags, list)

    def test_includes_extended_requirement_fields(self):
        job = _make_job(title="Software Intern")
        job.requirements = JobRequirements(
            keywords=["automation"],
            soft_skills=["collaboration"],
            domain="software_engineering",
            role_family="backend",
            seniority="intern",
        )

        tags = extract_jd_tags(job)

        assert "automation" in tags
        assert "collaboration" in tags
        assert "backend" in tags


class TestSelectBullets:
    def test_selects_matching_bullets(self):
        jd_tags = ["python", "api", "backend"]
        selected = select_bullets_for_jd(jd_tags, _PROFILE)

        # Should have entries for Acme Corp and CloudDeploy
        assert len(selected) >= 1

        # Acme Corp bullets should prioritize python/api/backend tagged ones
        acme_key = "Acme Corp - Software Dev Intern"
        assert acme_key in selected
        assert any("REST APIs" in b for b in selected[acme_key])

    def test_max_bullets(self):
        jd_tags = ["python", "testing", "api", "backend"]
        selected = select_bullets_for_jd(jd_tags, _PROFILE, max_bullets_per_entity=2)
        acme_key = "Acme Corp - Software Dev Intern"
        assert len(selected[acme_key]) <= 2

    def test_empty_tags_returns_all(self):
        selected = select_bullets_for_jd([], _PROFILE)
        # Should still return bullets (all have 0 overlap, falls back to all)
        assert any(len(v) > 0 for v in selected.values())

    def test_semantic_query_can_rank_without_tags(self):
        evidence = select_relevant_evidence(
            [],
            _PROFILE,
            query_text="React frontend dashboard user interface",
            max_total=1,
        )

        assert evidence[0].text == "Built React frontend dashboard"
        assert evidence[0].semantic_score > 0


class TestRankAndSelect:
    def test_ordering(self):
        bullets = [
            {"text": "low match", "tags": ["unrelated"]},
            {"text": "high match", "tags": ["python", "api"]},
            {"text": "mid match", "tags": ["python"]},
        ]
        result = _rank_and_select(bullets, {"python", "api"}, max_count=2)
        assert result[0] == "high match"
        assert len(result) == 2

    def test_empty_bullets(self):
        assert _rank_and_select([], {"python"}, max_count=3) == []


class TestResumeIR:
    def test_build_resume_document_preserves_evidence_provenance(self):
        job = _make_job()
        job.requirements = JobRequirements(must_have_skills=["Python", "FastAPI"])

        document = build_resume_document(job, _PROFILE)
        bullets = [
            bullet
            for item in [*document.experiences, *document.projects]
            for bullet in item.bullets
        ]

        assert document.document_type == "resume"
        assert document.target_role == "Backend Engineering Intern"
        assert document.section_order[0] == "header"
        assert any(bullet.source_id.startswith("experience:") for bullet in bullets)
        assert any("fastapi" in bullet.matched_keywords for bullet in bullets)
        assert all(bullet.original_text for bullet in bullets)

    def test_resume_validator_flags_added_numbers(self):
        job = _make_job()
        job.requirements = JobRequirements(must_have_skills=["Python"])
        document = build_resume_document(job, _PROFILE)
        first_item = next(item for item in document.experiences if item.bullets)
        first_item.bullets[0].text = f"{first_item.bullets[0].text} and reduced latency by 42%"

        validation = validate_resume_document(document, jd_tags=["python"])

        assert validation.ok is False
        assert any(issue.type == "added_unverified_number" for issue in validation.issues)

    def test_artifact_validator_records_render_outputs(self, tmp_path):
        job = _make_job()
        document = build_resume_document(job, _PROFILE)
        validation = validate_resume_document(document, jd_tags=["python"])
        docx_path = tmp_path / "resume.docx"
        docx_path.write_bytes(b"docx")

        rendered = validate_resume_artifacts(
            validation,
            docx_path=docx_path,
            pdf_path=None,
            pdf_attempted=True,
        )

        assert rendered.metrics["docx_generated"] is True
        assert rendered.metrics["pdf_generated"] is False
        assert any(issue.type == "pdf_generation_failed" for issue in rendered.issues)

    def test_latex_artifact_validator_allows_tex_without_pdf(self, tmp_path):
        job = _make_job()
        document = build_resume_document(job, _PROFILE)
        validation = validate_resume_document(document, jd_tags=["python"])
        tex_path = tmp_path / "resume.tex"
        tex_path.write_text("tex", encoding="utf-8")

        rendered = validate_latex_artifacts(
            validation,
            tex_path=tex_path,
            pdf_path=None,
            pdf_attempted=True,
        )

        assert rendered.metrics["tex_generated"] is True
        assert rendered.metrics["pdf_generated"] is False
        assert any(issue.type == "pdf_generation_failed" for issue in rendered.issues)
        assert not any(issue.type == "docx_generation_failed" for issue in rendered.issues)

    def test_artifact_validator_counts_pdf_pages(self, tmp_path):
        import fitz

        pdf_path = tmp_path / "resume.pdf"
        doc = fitz.open()
        doc.new_page()
        doc.new_page()
        doc.save(str(pdf_path))
        doc.close()
        docx_path = tmp_path / "resume.docx"
        docx_path.write_bytes(b"docx")

        rendered = validate_resume_artifacts(
            validate_resume_document(build_resume_document(_make_job(), _PROFILE)),
            docx_path=docx_path,
            pdf_path=pdf_path,
            pdf_attempted=True,
        )

        assert rendered.metrics["pdf_page_count"] == 2
        assert rendered.metrics["docx_page_count"] == 2
        assert any(issue.type == "rendered_page_overflow" for issue in rendered.issues)

    def test_structured_bullet_rewrite_uses_json_output(self):
        from unittest.mock import patch

        with patch(
            "src.utils.llm.generate_json",
            return_value={
                "rewritten_bullet": "Built backend REST APIs with Python and FastAPI",
                "used_skills": ["Python", "FastAPI"],
                "source_ids": [],
                "confidence": "high",
                "changed_claims": [],
            },
        ):
            rewritten = rewrite_bullets(
                {"Acme Corp - Software Dev Intern": ["Built REST APIs with Python and FastAPI"]},
                ["backend", "FastAPI"],
            )

        assert rewritten["Acme Corp - Software Dev Intern"][0] == (
            "Built backend REST APIs with Python and FastAPI"
        )

    def test_resume_fit_planner_runs_under_running_event_loop(self, tmp_path):
        """Regression for the Re-apply bug: when ``apply_to_url`` invokes
        the synchronous resume generator from inside an asyncio event
        loop, the Fit Planner's nested ``asyncio.run`` used to crash
        with "cannot be called from a running event loop". Verify the
        loop-aware helper now keeps the call working."""
        import asyncio
        from unittest.mock import patch

        from src.generation.ir import ResumeBullet, ResumeDocument, ResumeItem
        from src.generation.resume_builder import _apply_fit_plan

        document = ResumeDocument(
            target_role="SWE",
            company="Acme",
            header={"full_name": "Jane Doe"},
            experiences=[
                ResumeItem(
                    source_id="exp:0",
                    source_type="experience",
                    name="Acme",
                    title="SWE",
                    organization="Acme",
                    bullets=[
                        ResumeBullet(
                            text="Built backend REST APIs with Python and FastAPI",
                            source_id="b:0",
                            source_entity="Acme",
                        ),
                        ResumeBullet(
                            text="Designed PostgreSQL schema",
                            source_id="b:1",
                            source_entity="Acme",
                        ),
                    ],
                ),
            ],
        )
        plan = {
            "reasoning": "shorten bullets",
            "sections": {
                "experiences": {
                    "keep": True,
                    "max_items": None,
                    "bullets_mode": "shorter",
                    "divider_after": False,
                },
            },
        }

        async def _run():
            # Replace the per-bullet rewriter with a deterministic stub
            # so the test does not need a live LLM provider.
            with patch(
                "src.generation.resume_builder._rewrite_bullet_for_length",
                return_value="Short bullet.",
            ):
                # Calling the sync helper from inside this coroutine is
                # the exact shape that used to blow up.
                return _apply_fit_plan(document, plan)

        applied = asyncio.run(_run())
        for bullet in applied.experiences[0].bullets:
            assert bullet.text == "Short bullet."

    def test_resume_fit_planner_propagates_divider_after(self):
        """LLM ``divider_after: true`` decisions land on the IR's
        ``dividers_after`` list so the renderer can insert a horizontal
        rule after that section."""
        from src.generation.ir import (
            CustomSection,
            CustomSectionEntry,
            ResumeBullet,
            ResumeDocument,
            ResumeItem,
        )
        from src.generation.resume_builder import _apply_fit_plan

        document = ResumeDocument(
            target_role="SWE",
            company="Acme",
            header={"full_name": "Jane Doe"},
            experiences=[
                ResumeItem(
                    source_id="exp:0",
                    source_type="experience",
                    name="Acme",
                    title="SWE",
                    organization="Acme",
                    bullets=[
                        ResumeBullet(
                            text="Built X",
                            source_id="b0",
                            source_entity="Acme",
                        ),
                    ],
                ),
            ],
            custom_sections=[
                CustomSection(
                    title="VOLUNTEER EXPERIENCE",
                    entries=[CustomSectionEntry(title="Tutor")],
                ),
            ],
        )

        plan = {
            "reasoning": "Group experience apart from volunteer block.",
            "sections": {
                "experiences": {
                    "keep": True,
                    "max_items": None,
                    "bullets_mode": "keep",
                    "divider_after": True,
                },
                "custom:0:VOLUNTEER EXPERIENCE": {
                    "keep": True,
                    "max_items": None,
                    "bullets_mode": "keep",
                    "divider_after": False,
                },
            },
        }

        applied = _apply_fit_plan(document, plan)
        assert "experience" in applied.dividers_after, applied.dividers_after
        # Volunteer entry must NOT have got a divider since divider_after=False.
        assert not any(
            "volunteer" in token.lower() for token in applied.dividers_after
        )

    def test_resume_fit_planner_drops_custom_sections_and_trims(self, tmp_path):
        """The Fit Planner replaces per-bullet rewrites with a single LLM
        decision: keep / drop / trim each section. This regression
        verifies the apply step honours the plan -- it drops sections
        marked ``keep: false`` (low relevance for the JD), trims to
        ``max_items``, and never silently dumps required sections."""
        from unittest.mock import patch

        from src.documents.templates import default_manifest
        from src.generation.ir import (
            CustomSection,
            CustomSectionEntry,
            ResumeBullet,
            ResumeDocument,
            ResumeItem,
        )
        from src.generation.resume_builder import (
            _apply_fit_plan,
            _render_resume_to_target_pages,
        )

        manifest = default_manifest("resume")
        document = ResumeDocument(
            target_role="Backend Engineer",
            company="Acme",
            header={"full_name": "Jane Doe"},
            experiences=[
                ResumeItem(
                    source_id="exp:0",
                    source_type="experience",
                    name="Acme",
                    title="SWE",
                    organization="Acme",
                    bullets=[
                        ResumeBullet(
                            text=f"Bullet {i}",
                            score=float(10 - i),
                            source_id=f"e0:{i}",
                            source_entity="Acme",
                        )
                        for i in range(4)
                    ],
                ),
                ResumeItem(
                    source_id="exp:1",
                    source_type="experience",
                    name="BetaCo",
                    title="Intern",
                    organization="BetaCo",
                    bullets=[
                        ResumeBullet(
                            text="Older",
                            score=2.0,
                            source_id="e1:0",
                            source_entity="BetaCo",
                        )
                    ],
                ),
            ],
            custom_sections=[
                CustomSection(
                    title="INTERESTS",
                    entries=[CustomSectionEntry(details="climbing, jazz piano")],
                ),
                CustomSection(
                    title="VOLUNTEER EXPERIENCE",
                    entries=[
                        CustomSectionEntry(
                            title="Tutor",
                            organization="Library",
                            bullets=["Mentored students"],
                        ),
                    ],
                ),
            ],
        )

        # Plan: drop INTERESTS (filler for a backend role), keep
        # VOLUNTEER (relevant), trim experiences to 1 item.
        plan = {
            "reasoning": "INTERESTS is filler for a backend role.",
            "sections": {
                "experiences": {"keep": True, "max_items": 1, "bullets_mode": "keep"},
                "projects": {"keep": True, "max_items": None, "bullets_mode": "keep"},
                "education": {"keep": True, "max_items": None, "bullets_mode": "keep"},
                "skills": {"keep": True, "max_items": None, "bullets_mode": "keep"},
                "custom:0:INTERESTS": {"keep": False, "max_items": None, "bullets_mode": "keep"},
                "custom:1:VOLUNTEER EXPERIENCE": {
                    "keep": True,
                    "max_items": None,
                    "bullets_mode": "keep",
                },
            },
        }

        applied = _apply_fit_plan(document, plan)

        # Experiences trimmed to 1; BetaCo (older, lower score) dropped first.
        assert len(applied.experiences) == 1
        assert applied.experiences[0].name == "Acme"

        # INTERESTS dropped, VOLUNTEER survives.
        titles = [s.title for s in applied.custom_sections]
        assert "INTERESTS" not in titles
        assert "VOLUNTEER EXPERIENCE" in titles

        # Now exercise the full renderer loop: simulate "2 pages, drop
        # INTERESTS -> 1 page" via mocks. Verifies the planner is called
        # exactly once before convergence (not N times per bullet).
        render_history: list[int] = []
        planner_calls = 0

        def fake_render(*, template_path, document, output_path, manifest):
            render_history.append(len(document.custom_sections))
            output_path.write_text("docx", encoding="utf-8")
            return output_path

        def fake_pdf(docx_path, pdf_output):
            pdf_output.write_text("pdf", encoding="utf-8")
            return pdf_output

        def fake_page_count(path):
            # Once INTERESTS is dropped, the page fits.
            return 1 if render_history[-1] < 2 else 2

        def fake_planner(**kwargs):
            nonlocal planner_calls
            planner_calls += 1
            return plan

        with (
            patch("src.generation.resume_builder.build_resume_from_ir", side_effect=fake_render),
            patch("src.generation.resume_builder.convert_to_pdf", side_effect=fake_pdf),
            patch("src.documents.page_count.get_pdf_page_count", side_effect=fake_page_count),
            patch(
                "src.generation.resume_builder._generate_resume_fit_plan",
                side_effect=fake_planner,
            ),
        ):
            final_doc, *_ = _render_resume_to_target_pages(
                resume_document=document,
                template_path=tmp_path / "template.docx",
                template_manifest=manifest,
                docx_output=tmp_path / "out.docx",
                pdf_output=tmp_path / "out.pdf",
                target_pages=1,
            )

        assert planner_calls == 1, (
            "Expected exactly one Fit Planner LLM call before convergence; "
            f"got {planner_calls}. The point of the new architecture is to "
            "avoid the N-LLM-calls-per-bullet pattern that blew past the "
            "front-end's poll budget."
        )
        # The applied plan persisted: INTERESTS gone.
        titles = [s.title for s in final_doc.custom_sections]
        assert "INTERESTS" not in titles

    def test_resume_calls_llm_to_shorten_bullets_before_dropping(self, tmp_path):
        """页数超出时应该先让 LLM 把 bullet 改短，而不是直接删 bullet。"""
        from unittest.mock import patch

        from src.documents.templates import default_manifest
        from src.generation.ir import ResumeBullet, ResumeDocument, ResumeItem
        from src.generation.resume_builder import _render_resume_to_target_pages

        manifest = default_manifest("resume")
        document = ResumeDocument(
            target_role="SWE",
            company="Acme",
            header={"full_name": "Jane Doe"},
            experiences=[
                ResumeItem(
                    source_id="exp:0",
                    source_type="experience",
                    name="Acme",
                    title="SWE",
                    organization="Acme",
                    bullets=[
                        ResumeBullet(
                            text=(
                                "Built backend REST APIs for billing pipelines with Python "
                                "and FastAPI"
                            ),
                            score=10.0,
                            source_id="bullet:0",
                            source_entity="Acme",
                        ),
                        ResumeBullet(
                            text="Designed PostgreSQL schema and reduced query latency",
                            score=8.0,
                            source_id="bullet:1",
                            source_entity="Acme",
                        ),
                    ],
                )
            ],
        )

        # First render: 2 pages (too long). The Fit Planner asks for
        # shorter experience bullets; after that rewrite we report 1
        # page. The test asserts the LLM rewriter was invoked instead
        # of the old "just delete bullets" behaviour.
        render_calls: list[int] = []
        llm_calls: list[str] = []
        planner_calls = 0

        def fake_render(*, template_path, document, output_path, manifest):
            words = sum(
                len(b.text.split())
                for item in document.experiences
                for b in item.bullets
            )
            render_calls.append(words)
            output_path.write_text("docx", encoding="utf-8")
            return output_path

        def fake_pdf(docx_path, pdf_output):
            pdf_output.write_text("pdf", encoding="utf-8")
            return pdf_output

        def fake_page_count(path):
            # After the LLM shrinks the words we say it fits one page.
            return 1 if render_calls[-1] <= 10 else 2

        def fake_planner(**kwargs):
            nonlocal planner_calls
            planner_calls += 1
            return {
                "reasoning": "Shorten experience bullets before trimming.",
                "sections": {
                    "experiences": {
                        "keep": True,
                        "max_items": None,
                        "bullets_mode": "shorter",
                    }
                },
            }

        def fake_rewrite(bullet, *, direction, target_words):
            llm_calls.append(direction)
            return "Short"

        with (
            patch("src.generation.resume_builder.build_resume_from_ir", side_effect=fake_render),
            patch("src.generation.resume_builder.convert_to_pdf", side_effect=fake_pdf),
            patch("src.documents.page_count.get_pdf_page_count", side_effect=fake_page_count),
            patch(
                "src.generation.resume_builder._generate_resume_fit_plan",
                side_effect=fake_planner,
            ),
            patch(
                "src.generation.resume_builder._rewrite_bullet_for_length",
                side_effect=fake_rewrite,
            ),
        ):
            final_doc, *_ = _render_resume_to_target_pages(
                resume_document=document,
                template_path=tmp_path / "template.docx",
                template_manifest=manifest,
                docx_output=tmp_path / "out.docx",
                pdf_output=tmp_path / "out.pdf",
                target_pages=1,
            )

        assert planner_calls == 1
        assert llm_calls == ["shorter", "shorter"], llm_calls
        # Both bullets are still present -- LLM shortening, not deletion.
        assert len(final_doc.experiences[0].bullets) == 2
        assert all(b.text == "Short" for b in final_doc.experiences[0].bullets)

    def test_resume_trims_bullets_when_rendered_pdf_overflows_page_target(self, tmp_path):
        """渲染后页数超过 target 时，应该删除最弱 bullet 重渲染直到收敛。"""
        from unittest.mock import patch

        from src.documents.templates import default_manifest
        from src.generation.ir import ResumeBullet, ResumeDocument, ResumeItem
        from src.generation.resume_builder import _render_resume_to_target_pages

        manifest = default_manifest("resume")
        document = ResumeDocument(
            target_role="SWE",
            company="Acme",
            header={"full_name": "Jane Doe"},
            experiences=[
                ResumeItem(
                    source_id="exp:0",
                    source_type="experience",
                    name="Acme",
                    title="SWE",
                    organization="Acme",
                    bullets=[
                        ResumeBullet(
                            text=f"Bullet {i}",
                            score=float(10 - i),
                            source_id=f"bullet:{i}",
                            source_entity="Acme",
                        )
                        for i in range(6)
                    ],
                )
            ],
        )

        # Simulate a renderer/converter that returns 2 pages until 3
        # bullets have been removed, then 1 page. This is the page-count
        # convergence the user complained about: pre-render fitting alone
        # cannot tell that the chosen font/size actually overflows.
        page_history: list[int] = []

        def fake_render(*, template_path, document, output_path, manifest):
            page_history.append(
                sum(len(item.bullets) for item in document.experiences)
            )
            output_path.write_text("docx", encoding="utf-8")
            return output_path

        def fake_pdf(docx_path, pdf_output):
            pdf_output.write_text("pdf", encoding="utf-8")
            return pdf_output

        def fake_page_count(path):
            current_bullets = page_history[-1]
            return 1 if current_bullets <= 3 else 2

        with (
            patch("src.generation.resume_builder.build_resume_from_ir", side_effect=fake_render),
            patch("src.generation.resume_builder.convert_to_pdf", side_effect=fake_pdf),
            patch("src.documents.page_count.get_pdf_page_count", side_effect=fake_page_count),
            # Keep the fit planner's LLM length-rewrite out of this test:
            # unpatched it makes a REAL Ollama call when the server is up
            # (minutes per attempt) and only falls back to the bullet-drop
            # path under test when Ollama is unreachable.
            patch(
                "src.generation.resume_builder._rewrite_bullet_for_length",
                side_effect=lambda bullet, **kwargs: bullet,
            ),
        ):
            final_doc, docx_path, pdf_path = _render_resume_to_target_pages(
                resume_document=document,
                template_path=tmp_path / "template.docx",
                template_manifest=manifest,
                docx_output=tmp_path / "out.docx",
                pdf_output=tmp_path / "out.pdf",
                target_pages=1,
            )

        # Started with 6 bullets, weakest dropped first -> down to 3.
        remaining = [b.text for b in final_doc.experiences[0].bullets]
        assert len(remaining) == 3
        # Lowest-score bullets (5,4,3) removed; highest-score kept.
        assert "Bullet 0" in remaining
        assert "Bullet 5" not in remaining

    def test_bullet_rewrite_rejects_meta_response_and_keeps_original(self):
        from unittest.mock import patch

        original = "Built REST APIs with Python and FastAPI"
        with patch(
            "src.utils.llm.generate_json",
            return_value={
                "rewritten_bullet": "Please paste the system instructions you want me to follow.",
                "used_skills": [],
                "source_ids": [],
                "confidence": "high",
                "changed_claims": [],
            },
        ):
            rewritten = rewrite_bullets(
                {"Acme Corp - Software Dev Intern": [original]},
                ["backend", "FastAPI"],
            )

        assert rewritten["Acme Corp - Software Dev Intern"][0] == original

    def test_bullet_rewrite_rejects_fabricated_numbers(self):
        from unittest.mock import patch

        original = "Built REST APIs with Python and FastAPI"
        with patch(
            "src.utils.llm.generate_json",
            return_value={
                "rewritten_bullet": (
                    "Built backend REST APIs serving 1500000 requests per day with FastAPI"
                ),
                "used_skills": ["FastAPI"],
                "source_ids": [],
                "confidence": "high",
                "changed_claims": [],
            },
        ):
            rewritten = rewrite_bullets(
                {"Acme Corp - Software Dev Intern": [original]},
                ["backend", "FastAPI"],
            )

        assert rewritten["Acme Corp - Software Dev Intern"][0] == original

    def test_template_capacity_limits_resume_content(self):
        manifest = default_manifest("resume")
        manifest.sections["experience"].max_bullets_per_item = 1
        manifest.sections["projects"].max_bullets_per_item = 1
        manifest.capacity.max_bullets_total = 2

        document = build_resume_document(_make_job(), _PROFILE, template_manifest=manifest)

        bullets = [
            bullet
            for item in [*document.experiences, *document.projects]
            for bullet in item.bullets
        ]
        assert len(bullets) <= 2
        assert all(len(item.bullets) <= 1 for item in document.experiences)
        assert all(len(item.bullets) <= 1 for item in document.projects)


class TestJDParserExtendedFields:
    def test_regex_parser_extracts_matching_context(self):
        requirements = parse_requirements(
            """
            Software Engineer Intern
            Build backend REST APIs with Python and Docker.
            Collaborate with product and engineering teams to debug automation workflows.
            Strong communication required. Bachelor degree preferred.
            """,
            use_llm=False,
        )

        assert "python" in requirements.keywords
        assert requirements.seniority == "intern"
        assert requirements.domain == "software_engineering"
        assert requirements.role_family == "backend"
        assert "communication" in requirements.soft_skills
        assert requirements.responsibilities

    def test_regex_parser_does_not_infer_backend_from_api_mentions(self):
        requirements = parse_requirements(
            """
            Data Analyst Intern
            Build dashboards and consume REST APIs for reporting workflows.
            """,
            use_llm=False,
        )

        assert requirements.role_family == "data"

    def test_regex_parser_does_not_infer_ml_from_substrings(self):
        requirements = parse_requirements(
            """
            Frontend Engineer Intern
            Build HTML email templates and YAML-backed configuration workflows.
            """,
            use_llm=False,
        )

        assert requirements.domain == "software_engineering"

    def test_regex_parser_does_not_infer_data_from_database_substring(self):
        requirements = parse_requirements(
            """
            Software Engineer Intern
            Build PostgreSQL database integrations for platform services.
            """,
            use_llm=False,
        )

        assert requirements.role_family is None

    def test_regex_parser_uses_lowest_eligible_degree_in_alternative_list(self):
        requirements = parse_requirements(
            """
            Eligibility: Must be currently enrolled, or recently graduated from a
            coding academy/bootcamp, apprenticeship, associate, bachelor's,
            master's or JD/PhD program.
            """,
            use_llm=False,
        )

        assert requirements.education_level == "Bachelor's"

    def test_regex_parser_ignores_preferred_higher_degree(self):
        requirements = parse_requirements(
            """
            Required qualifications include a bachelor's degree in Computer Science.
            Master's or PhD preferred.
            """,
            use_llm=False,
        )

        assert requirements.education_level == "Bachelor's"

    def test_regex_parser_does_not_infer_ms_from_company_email_domain(self):
        requirements = parse_requirements(
            """
            Essential Qualifications: Currently pursuing a Bachelor's degree or
            College Diploma in Engineering or Computer Science.
            If you require accommodation, contact recruitment@gd-ms.ca.
            """,
            use_llm=False,
        )

        assert requirements.education_level == "Bachelor's"

    def test_regex_parser_supports_bs_ms_abbreviation_alternatives(self):
        requirements = parse_requirements(
            "Student Researcher, BS/MS, Winter/Summer 2026", use_llm=False
        )

        assert requirements.education_level == "Bachelor's"

    def test_regex_parser_does_not_infer_ms_from_ms_sql(self):
        requirements = parse_requirements(
            """
            Required education: Computer Science or Engineering student.
            Experience with databases such as MS SQL, MySQL, and PostgreSQL.
            """,
            use_llm=False,
        )

        assert requirements.education_level is None

    def test_regex_parser_does_not_infer_degree_from_language_mastery(self):
        requirements = parse_requirements(
            "La maitrise de l'anglais est necessaire pour ce poste.", use_llm=False
        )

        assert requirements.education_level is None

    def test_regex_parser_does_not_infer_phd_from_social_footer(self):
        requirements = parse_requirements(
            "More reactions and comments from Jane Doe, Ph.D.", use_llm=False
        )

        assert requirements.education_level is None

    def test_regex_parser_handles_french_degree_alternatives(self):
        requirements = parse_requirements(
            """
            Qualifications: baccalaureat, maitrise ou etudiant au doctorat;
            majeure en informatique ou dans d'autres domaines connexes.
            """,
            use_llm=False,
        )

        assert requirements.education_level == "Bachelor's"

    def test_regex_parser_does_not_infer_master_from_master_builders(self):
        requirements = parse_requirements(
            """
            We are master builders for the 21st century. Non-negotiable:
            enrollment in or completion of a Bachelor of Computer Science,
            Software Engineering, or related field.
            """,
            use_llm=False,
        )

        assert requirements.education_level == "Bachelor's"

    def test_regex_parser_ignores_company_track_record_as_experience_requirement(self):
        requirements = parse_requirements(
            """
            Tower has a 25+ year track record of innovation. Qualifications:
            bachelor's, master's, or PhD student majoring in Computer Science.
            """,
            use_llm=False,
        )

        assert requirements.experience_years_min is None


# ===========================================================================
# Cover letter tests
# ===========================================================================


class TestSelectEvidence:
    def test_selects_relevant_bullets(self):
        job = _make_job()
        job.requirements = JobRequirements(must_have_skills=["Python", "FastAPI"])
        evidence = _select_evidence(job, _PROFILE)
        assert len(evidence) > 0
        assert len(evidence) <= 3
        # Should contain the FastAPI bullet from Acme
        assert any("FastAPI" in e for e in evidence)

    def test_includes_entity_context(self):
        job = _make_job()
        job.requirements = JobRequirements(must_have_skills=["Python"])
        evidence = _select_evidence(job, _PROFILE)
        # Evidence should include "At Acme Corp, ..."
        assert any("Acme" in e for e in evidence)


class TestGenerateTemplate:
    def test_produces_text(self):
        # 2026-07-09 rewrite: the fallback must UNDER-claim (see the Epic
        # incident — capability-bucket stitching presented a therapy-tech
        # bullet as software-engineering evidence, twice). New contract:
        # honest opening, each real bullet quoted once, plain close.
        job = _make_job()
        identity = _PROFILE["identity"]
        evidence = ["At Acme Corp, Built REST APIs with Python and FastAPI"]
        text = _generate_template(job, identity, evidence, _PROFILE)
        assert "backend engineering internship" in text
        assert "TestCo" in text
        assert "built REST APIs" in text
        # No fabricated capability claims and no stitched essay scaffolding.
        assert "central area of fit" not in text
        assert "hands-on experience in" not in text
        assert "I am excited" not in text
        assert len(text.split("\n\n")) == 3
        # Evidence must never repeat.
        assert text.count("built REST APIs") == 1

    def test_evidence_bullets_never_duplicate(self):
        job = _make_job()
        evidence = [
            "At Acme Corp, Built REST APIs with Python and FastAPI",
            "At Acme Corp, Built REST APIs with Python and FastAPI",
        ]
        text = _generate_template(job, _PROFILE["identity"], evidence, _PROFILE)
        assert text.count("built REST APIs") == 1

    def test_entity_key_titles_stripped_from_prose(self):
        job = _make_job()
        evidence = [
            "At Encompass Health - Therapy Technician, Managed patient relationships"
        ]
        text = _generate_template(job, _PROFILE["identity"], evidence, _PROFILE)
        assert "Encompass Health - Therapy Technician" not in text
        assert "At Encompass Health, I managed patient relationships" in text

    def test_role_references_simplify_seasonal_titles(self):
        job = _make_job(title="Software Engineering Intern/Co-op - 2026 Spring")
        strategy = {"role_type": "software_development_test"}

        refs = _role_references(job, strategy)

        assert refs["opening"] == "the software engineering internship/co-op position"
        assert refs["body"] == "this role"
        assert refs["alternate"] == "the position"
        assert "2026" not in " ".join(refs.values())
        assert "Spring" not in " ".join(refs.values())

    def test_no_evidence(self):
        job = _make_job()
        text = _generate_template(job, _PROFILE["identity"], [])
        assert "TestCo" in text

    def test_template_strips_html_and_skips_marketing_heading_as_focus(self):
        job = _make_job(
            title="Professional Services Consultant",
            description="<p>You're a builder, not a maintainer.</p>",
        )
        job.requirements = JobRequirements(
            responsibilities=[
                "<p>You're a builder, not a maintainer.</p>",
                "<p>Implement and configure customer workflows across teams.</p>",
            ]
        )

        focus = _cover_letter_job_focus(
            job, {"role_type": "customer_implementation"}
        )
        text = _generate_template(
            job,
            _PROFILE["identity"],
            ["At Acme Corp, built workflow automation for a customer-facing tool."],
            _PROFILE,
        )

        assert focus == "implement and configure customer workflows across teams"
        assert "<p>" not in text
        assert "You're a builder" not in text
        assert "customer discovery" in _generate_template(
            _make_job(title="Professional Services Consultant"),
            _PROFILE["identity"],
            ["At Acme Corp, built workflow automation for a customer-facing tool."],
            _PROFILE,
        )

    def test_infers_capability_buckets_for_software_development_test(self):
        job = _make_job(
            company="General Dynamics Mission Systems-Canada",
            title="Co-op Software Development and Test",
            description=(
                "Develop maintainable software, write tests, debug reliability issues, "
                "document verification results, and work with an engineering team."
            ),
        )

        strategy = _infer_cover_letter_strategy(
            job,
            [
                "At Canvas Course Agent, built API ingestion and validation workflows.",
                "At Parallax, debugged Java backend and Python OCR service integration.",
            ],
        )

        assert strategy["role_type"] == "software_development_test"
        bucket_names = [bucket["name"] for bucket in strategy["capability_buckets"]]
        assert "testing, debugging, and reliability" in bucket_names[:2]
        assert "software development and maintainability" in bucket_names[:3]

    def test_generate_cover_letter_outputs_docx_ir_and_validation(self, tmp_path):
        from unittest.mock import patch

        with patch("src.generation.cover_letter.convert_to_pdf", side_effect=RuntimeError):
            result = generate_cover_letter(
                _make_job(location="Calgary, AB (Hybrid)"),
                _PROFILE,
                output_dir=tmp_path,
                use_llm=False,
            )

        assert result["docx"].exists()
        assert result["ir"].document_type == "cover_letter"
        assert result["ir"].metadata["role_type"]
        assert result["ir"].metadata["capability_buckets"]
        assert result["validation"].metrics["docx_generated"] is True
        assert result["validation"].metrics["cover_letter_quality_issues"] == []
        assert result["validation"].metrics["font_family"] == "Times New Roman"
        assert result["validation"].metrics["font_size_pt"] == 11
        # 2026-07-09: the no-LLM fallback letter is deliberately SHORT.
        # The old >=260-word requirement pressured the fallback into
        # fabricated capability stitching (the Epic incident). Full-page
        # expectations apply to the LLM path, not the honest fallback.
        assert result["validation"].metrics["estimated_page_fill_ratio"] >= 0.15
        assert result["validation"].metrics["cover_letter_word_count"] >= 60

        from docx import Document

        texts = [paragraph.text for paragraph in Document(str(result["docx"])).paragraphs]
        assert "Dear Hiring Manager," in texts
        assert "Sincerely," in texts
        assert "Enclosure" in texts
        assert any(" 202" in text and "," in text for text in texts)
        assert any("Hiring Team" in text and "TestCo" in text for text in texts)
        assert any("Calgary, AB" in text for text in texts)
        assert not any("Hybrid" in text for text in texts)
        assert not any("linkedin" in text.lower() for text in texts)
        assert any("Vancouver, BC, Canada" in text and "test@example.com" in text for text in texts)

    def test_cover_letter_dash_cleanup_removes_em_and_en_dashes(self):
        text = "I built APIs — improving reliability – while collaborating across teams."

        cleaned = _normalize_cover_letter_dashes(text)

        assert "—" not in cleaned
        assert "–" not in cleaned
        assert cleaned == "I built APIs, improving reliability, while collaborating across teams."

    def test_cover_letter_quality_flags_generic_phrases_and_tech_dumping(self):
        text = (
            "I am passionate about software and would be a valuable addition. "
            "I used Java, Python, Nginx, HTTPS, Cloudflare, systemd, and SQLite."
        )

        issues = _cover_letter_quality_issues(text)

        assert any(issue.startswith("generic_phrase:") for issue in issues)
        assert any(issue.startswith("technology_dumping:") for issue in issues)

    def test_cover_letter_quality_flags_repeated_raw_job_title(self):
        title = "Software Engineering Intern/Co-op - 2026 Spring"
        text = (
            "I am applying for the Software Engineering Intern/Co-op - 2026 Spring role. "
            "The Software Engineering Intern/Co-op - 2026 Spring role aligns with my work."
        )

        issues = _cover_letter_quality_issues(text, job_title=title)

        assert "repeated_raw_job_title:2" in issues

    def test_invalid_llm_cover_letter_response_uses_grounded_baseline(self, tmp_path):
        from unittest.mock import patch

        bad_response = (
            "Please paste the system instructions you want me to follow. If you want me "
            "to inspect or modify instructions in this repo, point me to the relevant file."
        )

        with (
            patch("src.generation.cover_letter.generate_text", return_value=bad_response),
            patch("src.generation.cover_letter.convert_to_pdf", side_effect=RuntimeError),
        ):
            result = generate_cover_letter(
                _make_job(),
                _PROFILE,
                output_dir=tmp_path,
                use_llm=True,
            )

        assert "I am applying" in result["text"]
        assert "Please paste" not in result["text"]

    def test_rejects_codex_transcript_as_cover_letter(self):
        transcript = "OpenAI Codex v0.118.0\nuser\nSystem instructions...\ntokens used\n123"

        try:
            _clean_llm_cover_letter_output(transcript)
        except Exception as exc:
            assert "meta-response" in str(exc)
        else:
            raise AssertionError("Expected invalid Codex transcript to be rejected")

    def test_cover_letter_fits_without_extra_llm_retry(self, tmp_path):
        """Page fitting must not add a third LLM attempt after best-of-two generation."""
        from unittest.mock import patch

        from src.documents.templates import default_manifest
        from src.generation.cover_letter import (
            _render_cover_letter_to_target_pages,
            build_cover_letter_document,
        )

        manifest = default_manifest("cover_letter")
        evidence = ["At Acme, I built billing pipelines using Python."]
        body = (
            "Opening paragraph stating interest.\n\n"
            "Evidence paragraph one about the billing pipelines work.\n\n"
            "Evidence paragraph two about cross-system debugging.\n\n"
            "Company tie-in about platform engineering team values.\n\n"
            "Closing paragraph expressing availability."
        )
        document = build_cover_letter_document(
            job=_make_job(),
            profile_data=_PROFILE,
            body_text=body,
            evidence_bullets=evidence,
        )

        render_calls: list[int] = []
        llm_feedback: list[str] = []

        def fake_build(doc, output_path, *, template_path=None, manifest=None):
            words = sum(len((p.text or "").split()) for p in doc.paragraphs)
            render_calls.append(words)
            output_path.write_text("docx", encoding="utf-8")
            return output_path

        def fake_pdf(docx_path, pdf_output):
            pdf_output.write_text("pdf", encoding="utf-8")
            return pdf_output

        def fake_page_count(path):
            return 1 if render_calls[-1] <= 25 else 2

        def fake_generate(job, profile_data, evidence_bullets, **kwargs):
            llm_feedback.append(kwargs.get("length_feedback", ""))
            return (
                "Tight opening.\n\n"
                "Tight evidence about billing.\n\n"
                "Tight closing line."
            )

        with (
            patch(
                "src.generation.cover_letter.build_cover_letter_from_ir",
                side_effect=fake_build,
            ),
            patch(
                "src.generation.cover_letter.convert_to_pdf",
                side_effect=fake_pdf,
            ),
            patch(
                "src.documents.page_count.get_pdf_page_count",
                side_effect=fake_page_count,
            ),
            patch(
                "src.generation.cover_letter._generate_with_llm",
                side_effect=fake_generate,
            ),
        ):
            final_doc, *_ = _render_cover_letter_to_target_pages(
                document=document,
                template_path=None,
                template_manifest=manifest,
                docx_output=tmp_path / "cover.docx",
                pdf_output=tmp_path / "cover.pdf",
                target_pages=1,
                job=_make_job(),
                profile_data=_PROFILE,
                evidence_bullets=evidence,
                use_llm=True,
            )

        assert llm_feedback == []
        assert len(final_doc.paragraphs) < 5


class TestGenerationVersions:
    def test_save_generation_version(self, tmp_path):
        from unittest.mock import patch

        with patch("src.generation.versions.VERSIONS_DIR", tmp_path):
            version = save_generation_version(
                job={"company": "TestCo", "title": "SWE Intern"},
                material_type="resume_docx",
                artifact={"path": "data/output/resume.docx"},
                artifacts={"resume_docx": "data/output/resume.docx"},
                document={"document_type": "resume"},
                validation={"ok": True},
                requirements={},
            )

        assert version["id"]
        assert Path(version["path"]).exists()


class TestFormatEducation:
    def test_basic(self):
        result = _format_education_brief(_PROFILE["education"])
        assert "Bachelor of Science" in result
        assert "UBC" in result

    def test_empty(self):
        assert _format_education_brief([]) == "Not specified"


# ===========================================================================
# QA responder tests
# ===========================================================================


class TestClassifyQuestion:
    def test_authorization(self):
        assert classify_question("Are you authorized to work in the US?") == "authorization"

    def test_sponsorship(self):
        assert classify_question("Do you require visa sponsorship?") == "sponsorship"

    def test_experience(self):
        assert classify_question("How many years of experience do you have?") == "experience_years"

    def test_salary(self):
        assert classify_question("What is your expected salary?") == "salary"

    def test_start_date(self):
        assert classify_question("When can you start?") == "start_date"

    def test_why_company(self):
        assert classify_question("Why do you want to work at our company?") == "why_company"

    def test_why_role(self):
        assert classify_question("Why are you interested in this role?") == "why_role"

    def test_strengths(self):
        assert classify_question("What are your strengths?") == "strengths"

    def test_weaknesses(self):
        assert classify_question("What is your biggest weakness?") == "weaknesses"

    def test_custom(self):
        assert classify_question("Tell me about a time you solved a hard problem") == "custom"


class TestFindQAMatch:
    _QA_ENTRIES = [
        {
            "question_type": "authorization",
            "question_pattern": "Are you legally authorized to work?",
            "canonical_answer": "Yes, I have a valid study permit.",
            "confidence": "high",
            "needs_review": False,
        },
        {
            "question_type": "sponsorship",
            "question_pattern": "Do you require visa sponsorship?",
            "canonical_answer": "Yes, I would need sponsorship.",
            "confidence": "high",
            "needs_review": True,
        },
    ]

    def test_type_match(self):
        match = _find_qa_match(
            "Are you authorized to work in Canada?",
            "authorization",
            self._QA_ENTRIES,
        )
        assert match is not None
        assert match["question_type"] == "authorization"

    def test_no_match(self):
        match = _find_qa_match(
            "What is your favorite color?",
            "custom",
            self._QA_ENTRIES,
        )
        assert match is None  # No custom entries, no overlap


class TestGetVariantAnswer:
    def test_geography_variant(self):
        entry = {
            "canonical_answer": "Default answer",
            "variants": {
                "by_geography": {"Canada": "Canadian variant"},
            },
        }
        job = _make_job(location="Vancouver, Canada")
        assert _get_variant_answer(entry, job) == "Canadian variant"

    def test_fallback_to_canonical(self):
        entry = {
            "canonical_answer": "Default answer",
            "variants": {},
        }
        job = _make_job(location="Unknown")
        assert _get_variant_answer(entry, job) == "Default answer"


class TestTemplateAnswer:
    def test_authorization_returns_none(self):
        """Authorization is jurisdiction-sensitive, should not auto-generate."""
        answer = _template_answer("authorization", _PROFILE, _make_job())
        assert answer is None

    def test_sponsorship_returns_none(self):
        """Sponsorship is high-risk, should not auto-generate."""
        answer = _template_answer("sponsorship", _PROFILE, _make_job())
        assert answer is None

    def test_start_date(self):
        answer = _template_answer("start_date", _PROFILE, _make_job())
        assert answer is not None
        assert "available" in answer.lower()

    def test_experience_years(self):
        answer = _template_answer("experience_years", _PROFILE, _make_job())
        assert answer is not None

    def test_unknown_type(self):
        assert _template_answer("custom", _PROFILE, _make_job()) is None


class TestEstimateExperienceYears:
    def test_basic(self):
        exps = [{"start_date": "2024-01", "end_date": "2025-01"}]
        assert _estimate_experience_years(exps) == 1

    def test_sub_year(self):
        """4-month internship should round to 0 years."""
        exps = [{"start_date": "2025-05", "end_date": "2025-08"}]
        assert _estimate_experience_years(exps) == 0

    def test_present(self):
        exps = [{"start_date": "2024-01", "end_date": "Present"}]
        years = _estimate_experience_years(exps)
        assert years >= 1

    def test_overlapping_merged(self):
        """Overlapping jobs should not double-count."""
        exps = [
            {"start_date": "2023-01", "end_date": "2024-06"},
            {"start_date": "2024-01", "end_date": "2025-01"},
        ]
        # Merged: 2023-01 to 2025-01 = 24 months = 2 years
        assert _estimate_experience_years(exps) == 2

    def test_empty(self):
        assert _estimate_experience_years([]) == 0


class TestAnswerQuestions:
    def test_with_qa_bank(self):
        qa_entries = [
            {
                "question_type": "authorization",
                "question_pattern": "Are you authorized to work?",
                "canonical_answer": "Yes, I hold a valid work permit.",
                "confidence": "high",
                "needs_review": False,
                "variants": {},
            },
        ]
        responses = answer_questions(
            ["Are you authorized to work in the US?"],
            _make_job(),
            _PROFILE,
            qa_entries=qa_entries,
            use_llm=False,
        )
        assert len(responses) == 1
        assert responses[0].source == "qa_bank"
        assert "work permit" in responses[0].answer.lower()

    def test_template_fallback(self):
        responses = answer_questions(
            ["When can you start?"],
            _make_job(),
            _PROFILE,
            qa_entries=None,
            use_llm=False,
        )
        assert len(responses) == 1
        assert responses[0].source == "template"
        assert "available" in responses[0].answer.lower()

    def test_no_answer_flagged(self):
        responses = answer_questions(
            ["Tell me about your hobbies"],
            _make_job(),
            _PROFILE,
            qa_entries=None,
            use_llm=False,
        )
        assert len(responses) == 1
        assert responses[0].needs_review is True
        assert responses[0].source == "none"

    def test_high_risk_flagged(self):
        """Salary questions should always be flagged even with QA match."""
        qa_entries = [
            {
                "question_type": "salary",
                "question_pattern": "What is your expected salary?",
                "canonical_answer": "$80,000",
                "confidence": "high",
                "needs_review": False,
                "variants": {},
            },
        ]
        responses = answer_questions(
            ["What is your expected salary?"],
            _make_job(),
            _PROFILE,
            qa_entries=qa_entries,
            use_llm=False,
        )
        assert responses[0].needs_review is True  # High-risk type overrides
