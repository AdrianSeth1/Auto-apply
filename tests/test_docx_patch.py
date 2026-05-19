"""Phase 15.2: tests for the DOCX patch mode.

We build small in-memory DOCX fixtures with python-docx, run the
patcher against a representative ResumeDocument IR, and read the
output back to verify:

* The Summary section is stripped from the source DOCX entirely --
  generated resumes never include one regardless of source content.
* Skills section is rewritten in-place under its heading.
* Bullets are swapped run-by-run; surplus bullets are appended; deficit
  bullets are blanked (not physically deleted).
* Sections not in ``section_order`` get their body blanked.
* :class:`PatchFallback` is raised for missing source files so the
  Phase 15.5 router can route to the template path.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from src.generation.docx_patch import (
    PatchFallback,
    patch_resume_docx,
)
from src.generation.ir import ResumeBullet, ResumeDocument, ResumeItem


def _build_source(path: Path) -> None:
    doc = Document()
    doc.add_heading("Summary", level=1)
    doc.add_paragraph("Old summary about the candidate.")
    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Must Have: legacy-python, legacy-sql")
    doc.add_paragraph("Preferred: react, docker")
    doc.add_heading("Experience", level=1)
    doc.add_paragraph("Original Co — Software Engineer", style="Heading 2")
    doc.add_paragraph("Built the old thing", style="List Bullet")
    doc.add_paragraph("Shipped some feature", style="List Bullet")
    doc.add_heading("Projects", level=1)
    doc.add_paragraph("Side Project", style="Heading 2")
    doc.add_paragraph("Old bullet for side project", style="List Bullet")
    doc.add_heading("Awards", level=1)
    doc.add_paragraph("Some award text.")
    doc.save(str(path))


def _bullet(text: str) -> ResumeBullet:
    return ResumeBullet(
        text=text,
        source_id="exp-1",
        source_type="experience",
        source_entity="Original Co",
    )


def _ir(experience_bullets: list[str], project_bullets: list[str] | None = None) -> ResumeDocument:
    return ResumeDocument(
        target_role="Software Engineer Intern",
        company="Test Co",
        header={"name": "Test"},
        skills={
            "must_have": ["python", "fastapi"],
            "preferred": ["postgres", "redis"],
        },
        experiences=[
            ResumeItem(
                source_id="exp-1",
                source_type="experience",
                name="Original Co — Software Engineer",
                bullets=[_bullet(t) for t in experience_bullets],
            )
        ],
        projects=[
            ResumeItem(
                source_id="proj-1",
                source_type="project",
                name="Side Project",
                bullets=[_bullet(t) for t in (project_bullets or [])],
            )
        ],
    )


# ---- Summary stripping -----------------------------------------------


def test_summary_section_is_stripped(tmp_path: Path) -> None:
    """Generated resumes never include a Summary, regardless of what
    the user-uploaded source DOCX contained."""
    src = tmp_path / "source.docx"
    out = tmp_path / "out.docx"
    _build_source(src)
    document = _ir(experience_bullets=["Designed scalable API layer."])
    report = patch_resume_docx(src, document, output_path=out)
    assert report.success is True
    doc = Document(str(out))
    paragraph_texts = [p.text for p in doc.paragraphs]
    # No paragraph should contain the original summary copy or a
    # "Summary" heading anywhere in the output.
    assert all("Old summary about the candidate" not in t for t in paragraph_texts)
    assert all(t.strip().lower() != "summary" for t in paragraph_texts)


# ---- Skills ---------------------------------------------------------


def test_skills_section_rewritten(tmp_path: Path) -> None:
    src = tmp_path / "source.docx"
    out = tmp_path / "out.docx"
    _build_source(src)
    document = _ir(experience_bullets=["x"])
    patch_resume_docx(src, document, output_path=out)
    doc = Document(str(out))
    skills_idx = next(
        i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Skills"
    )
    body = doc.paragraphs[skills_idx + 1 : skills_idx + 3]
    bodies = [p.text for p in body if p.text.strip()]
    assert any("Must Have" in line and "python" in line for line in bodies)
    assert any("Preferred" in line and "postgres" in line for line in bodies)


# ---- Bullets --------------------------------------------------------


def test_experience_bullets_are_swapped_in_place(tmp_path: Path) -> None:
    src = tmp_path / "source.docx"
    out = tmp_path / "out.docx"
    _build_source(src)
    document = _ir(experience_bullets=["NEW: built and shipped", "NEW: scaled service"])
    patch_resume_docx(src, document, output_path=out)
    doc = Document(str(out))
    bullets = [p.text for p in doc.paragraphs if p.style.name.startswith("List")]
    assert "NEW: built and shipped" in bullets
    assert "NEW: scaled service" in bullets
    assert "Built the old thing" not in bullets


def test_surplus_bullets_are_appended(tmp_path: Path) -> None:
    src = tmp_path / "source.docx"
    out = tmp_path / "out.docx"
    _build_source(src)
    document = _ir(
        experience_bullets=[
            "bullet 1",
            "bullet 2",
            "bullet 3 (overflow)",
        ]
    )
    patch_resume_docx(src, document, output_path=out)
    doc = Document(str(out))
    bullets = [
        p.text for p in doc.paragraphs if p.style.name.startswith("List") and p.text.strip()
    ]
    assert "bullet 3 (overflow)" in bullets


def test_deficit_bullets_are_blanked_not_removed(tmp_path: Path) -> None:
    src = tmp_path / "source.docx"
    out = tmp_path / "out.docx"
    _build_source(src)
    document = _ir(experience_bullets=["only one"])
    patch_resume_docx(src, document, output_path=out)
    doc = Document(str(out))
    # The original DOCX had two List Bullet paragraphs in Experience;
    # we should still see two paragraphs styled as List Bullet under
    # Experience, just the second one is empty.
    bullets_in_exp = [
        p
        for p in doc.paragraphs
        if p.style.name.startswith("List")
    ]
    assert len(bullets_in_exp) >= 2
    assert any(p.text == "only one" for p in bullets_in_exp)
    assert any(p.text == "" for p in bullets_in_exp)


# ---- Section visibility ---------------------------------------------


def test_sections_not_in_section_order_are_blanked(tmp_path: Path) -> None:
    src = tmp_path / "source.docx"
    out = tmp_path / "out.docx"
    _build_source(src)
    document = _ir(experience_bullets=["x"])
    document.section_order = ["summary", "skills", "experience", "projects"]
    patch_resume_docx(src, document, output_path=out)
    doc = Document(str(out))
    # The Awards section should still have its heading but no body.
    awards_idx = next(
        i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Awards"
    )
    # The single body paragraph that followed should now be empty.
    assert doc.paragraphs[awards_idx + 1].text == ""


# ---- Fallback signal ------------------------------------------------


def test_missing_source_raises_patchfallback(tmp_path: Path) -> None:
    document = _ir(experience_bullets=["x"])
    with pytest.raises(PatchFallback):
        patch_resume_docx(
            tmp_path / "does-not-exist.docx",
            document,
            output_path=tmp_path / "out.docx",
        )


# ---- Report ---------------------------------------------------------


def test_report_records_what_was_changed(tmp_path: Path) -> None:
    src = tmp_path / "source.docx"
    out = tmp_path / "out.docx"
    _build_source(src)
    document = _ir(experience_bullets=["new b1", "new b2"])
    report = patch_resume_docx(src, document, output_path=out)
    kinds = {op.kind for op in report.operations}
    # ``section_drop`` covers the summary strip; ``skills`` and
    # ``bullet`` cover the in-place rewrites.
    assert {"section_drop", "skills", "bullet"} <= kinds
    assert report.output_path == out


# ---- Patch policy knobs ---------------------------------------------


def test_allow_add_remove_bullets_false_preserves_count(tmp_path: Path) -> None:
    """When ``allow_add_remove_bullets=False`` the source DOCX's bullet
    count wins: surplus IR bullets are dropped (with a warning) and
    deficit slots keep their original text. This is the user's
    "preserve my exact structure" mode."""
    src = tmp_path / "source.docx"
    out = tmp_path / "out.docx"
    _build_source(src)
    # The source has 2 experience bullets. Feed IR with 3 (one surplus).
    document = _ir(
        experience_bullets=[
            "tailored bullet 1",
            "tailored bullet 2",
            "SHOULD BE DROPPED — surplus",
        ]
    )
    report = patch_resume_docx(
        src,
        document,
        output_path=out,
        allow_add_remove_bullets=False,
    )
    doc = Document(str(out))
    bullets = [
        p.text
        for p in doc.paragraphs
        if p.style.name.startswith("List") and p.text.strip()
    ]
    # Exactly the two source slots, no third one appended.
    assert "tailored bullet 1" in bullets
    assert "tailored bullet 2" in bullets
    assert "SHOULD BE DROPPED — surplus" not in bullets
    # The drop is reported.
    assert any("dropped" in w.lower() for w in report.warnings)


def test_allow_add_remove_bullets_false_keeps_original_when_ir_short(
    tmp_path: Path,
) -> None:
    """If the IR is *shorter* than the source, the conservative
    setting leaves the unmatched source bullets in place rather than
    blanking them."""
    src = tmp_path / "source.docx"
    out = tmp_path / "out.docx"
    _build_source(src)
    document = _ir(experience_bullets=["only one tailored bullet"])
    patch_resume_docx(
        src,
        document,
        output_path=out,
        allow_add_remove_bullets=False,
    )
    doc = Document(str(out))
    bullets = [
        p.text
        for p in doc.paragraphs
        if p.style.name.startswith("List")
    ]
    # Slot 0: new bullet. Slot 1: original source text retained.
    assert "only one tailored bullet" in bullets
    assert any("Shipped some feature" in b for b in bullets)
    # And nothing got blanked.
    assert not any(b == "" for b in bullets)


def test_allow_reorder_sections_false_skips_visibility_pass(
    tmp_path: Path,
) -> None:
    """When ``allow_reorder_sections=False`` the patcher should
    refrain from blanking sections that the IR's ``section_order``
    omits. The source DOCX's Awards section must survive even though
    the IR doesn't list it."""
    src = tmp_path / "source.docx"
    out = tmp_path / "out.docx"
    _build_source(src)
    document = _ir(experience_bullets=["x"])
    document.section_order = ["summary", "skills", "experience", "projects"]
    patch_resume_docx(
        src,
        document,
        output_path=out,
        allow_reorder_sections=False,
    )
    doc = Document(str(out))
    # Awards heading is still there with its body intact -- compare
    # to ``test_sections_not_in_section_order_are_blanked`` above,
    # which under the default policy blanks the body.
    awards_idx = next(
        i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Awards"
    )
    assert doc.paragraphs[awards_idx + 1].text != ""


# ---- Cover letter patcher -------------------------------------------


def _build_cover_source(path: Path) -> None:
    """Build a fake cover letter DOCX with header / salutation /
    body / closing / signature layout."""
    from docx import Document as _Document

    doc = _Document()
    doc.add_paragraph("Liam Liu")
    doc.add_paragraph("liam@example.com · (555) 123-4567")
    doc.add_paragraph("")
    doc.add_paragraph("January 15, 2026")
    doc.add_paragraph("")
    doc.add_paragraph("Acme Hiring Team")
    doc.add_paragraph("123 Market Street")
    doc.add_paragraph("")
    doc.add_paragraph("Dear Hiring Manager,")
    doc.add_paragraph("Original first body paragraph mentioning Python and React.")
    doc.add_paragraph("Original second body paragraph about teamwork.")
    doc.add_paragraph("Original third body paragraph wrapping up.")
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph("Liam Liu")
    doc.save(str(path))


def _cover_ir(body_texts: list[str]):
    """Build a minimal CoverLetterDocument IR with the given body
    paragraphs (all typed ``experience_evidence`` for simplicity)."""
    from src.generation.ir import CoverLetterDocument, CoverLetterParagraph

    return CoverLetterDocument(
        recipient={"company": "Acme"},
        applicant={"name": "Liam Liu"},
        paragraphs=[
            CoverLetterParagraph(type="experience_evidence", text=t)
            for t in body_texts
        ],
    )


def test_patch_cover_letter_replaces_body_only(tmp_path: Path) -> None:
    """The salutation, address block, closing line, and signature
    must all survive the patch. Only the paragraphs between
    salutation and closing are replaced."""
    from src.generation.docx_patch import patch_cover_letter_docx

    src = tmp_path / "cover.docx"
    out = tmp_path / "out.docx"
    _build_cover_source(src)
    ir = _cover_ir(
        [
            "Tailored body about distributed systems.",
            "Tailored body about scaling APIs.",
            "Tailored body about Acme's mission.",
        ]
    )

    patch_cover_letter_docx(src, ir, output_path=out)

    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]

    # Header survives.
    assert "Liam Liu" in texts
    assert "liam@example.com · (555) 123-4567" in texts
    assert "January 15, 2026" in texts
    assert "Acme Hiring Team" in texts
    # Salutation survives.
    assert "Dear Hiring Manager," in texts
    # Original body is gone, tailored body is in.
    assert "Original first body paragraph mentioning Python and React." not in texts
    assert "Tailored body about distributed systems." in texts
    assert "Tailored body about scaling APIs." in texts
    assert "Tailored body about Acme's mission." in texts
    # Closing + signature survive.
    assert "Sincerely," in texts
    # Signature line at the end (the second "Liam Liu" -- header copy
    # is the first occurrence).
    assert texts.count("Liam Liu") >= 2


def test_patch_cover_letter_raises_fallback_without_salutation(
    tmp_path: Path,
) -> None:
    """Without a salutation we have no anchor for where the body
    starts -- bail rather than risk overwriting the address block."""
    from docx import Document as _Document

    from src.generation.docx_patch import PatchFallback, patch_cover_letter_docx

    src = tmp_path / "no_salutation.docx"
    doc = _Document()
    doc.add_paragraph("Liam Liu")
    doc.add_paragraph("Some opening line that isn't a salutation.")
    doc.add_paragraph("Some body content.")
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph("Liam Liu")
    doc.save(str(src))

    ir = _cover_ir(["new body"])

    with pytest.raises(PatchFallback):
        patch_cover_letter_docx(src, ir, output_path=tmp_path / "out.docx")


def test_patch_cover_letter_appends_when_ir_has_more_paragraphs(
    tmp_path: Path,
) -> None:
    """The source has 3 body paragraphs; the IR has 5. The patcher
    must keep all 5 (replace 3, append 2 after the last)."""
    from src.generation.docx_patch import patch_cover_letter_docx

    src = tmp_path / "cover.docx"
    out = tmp_path / "out.docx"
    _build_cover_source(src)
    ir = _cover_ir([f"tailored body paragraph #{i}" for i in range(1, 6)])

    patch_cover_letter_docx(src, ir, output_path=out)

    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]

    for i in range(1, 6):
        assert f"tailored body paragraph #{i}" in texts
    # And the closing still comes AFTER the last tailored paragraph.
    closing_idx = texts.index("Sincerely,")
    fifth_idx = texts.index("tailored body paragraph #5")
    assert fifth_idx < closing_idx


def test_patch_cover_letter_skips_ir_paragraphs_typed_as_closing(
    tmp_path: Path,
) -> None:
    """IR paragraphs typed ``opening``/``closing`` would duplicate
    the source DOCX's salutation/closing; the patcher must filter
    them out."""
    from src.generation.docx_patch import patch_cover_letter_docx
    from src.generation.ir import CoverLetterDocument, CoverLetterParagraph

    src = tmp_path / "cover.docx"
    out = tmp_path / "out.docx"
    _build_cover_source(src)
    ir = CoverLetterDocument(
        recipient={"company": "Acme"},
        applicant={"name": "Liam Liu"},
        paragraphs=[
            CoverLetterParagraph(type="opening", text="Dear Hiring Manager,"),
            CoverLetterParagraph(
                type="experience_evidence",
                text="Real tailored body about scaling.",
            ),
            CoverLetterParagraph(
                type="closing", text="Sincerely,\nLiam Liu"
            ),
        ],
    )

    patch_cover_letter_docx(src, ir, output_path=out)

    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]

    # The synthetic salutation/closing from the IR did NOT get
    # duplicated -- the original ones from the source are preserved
    # and the tailored body is in between.
    assert texts.count("Dear Hiring Manager,") == 1
    assert texts.count("Sincerely,") == 1
    assert "Real tailored body about scaling." in texts
