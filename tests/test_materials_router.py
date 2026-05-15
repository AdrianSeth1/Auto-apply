"""Phase 15.5: tests for the materials router.

Pure-Python: the router stays decoupled from SQLAlchemy (we use a
``SourceResumeView`` value object) so these tests run without a DB.
The LaTeX compile path is monkey-patched so we do not need a real
compiler installed.
"""

from __future__ import annotations

import uuid
from pathlib import Path
from typing import Any

import pytest
from docx import Document

from src.documents.templates import (
    LatexConfig,
    LatexFieldMapping,
    TemplateManifest,
    TemplatePackage,
)
from src.generation.docx_patch import PatchFallback
from src.generation.ir import ResumeBullet, ResumeDocument, ResumeItem
from src.generation.materials_router import (
    MaterialsBindings,
    MaterialsRouterError,
    SourceResumeView,
    generate_materials,
)

# ---- Fixtures --------------------------------------------------------


def _ir() -> ResumeDocument:
    return ResumeDocument(
        target_role="Software Engineer Intern",
        company="Acme",
        header={"name": "Alice"},
        summary=["new summary"],
        skills={"must_have": ["python"]},
        experiences=[
            ResumeItem(
                source_id="exp-1",
                source_type="experience",
                name="Initech — SWE",
                bullets=[
                    ResumeBullet(
                        text="new bullet",
                        source_id="exp-1",
                        source_type="experience",
                        source_entity="Initech",
                    )
                ],
            )
        ],
    )


def _build_docx_source(path: Path) -> None:
    doc = Document()
    doc.add_heading("Summary", level=1)
    doc.add_paragraph("Old summary placeholder.")
    doc.add_heading("Experience", level=1)
    doc.add_paragraph("Initech — SWE", style="Heading 2")
    doc.add_paragraph("Old bullet", style="List Bullet")
    doc.save(str(path))


def _latex_package(tmp_path: Path) -> TemplatePackage:
    pkg_dir = tmp_path / "latex_pkg"
    pkg_dir.mkdir()
    template_tex = pkg_dir / "template.tex"
    template_tex.write_text(
        r"\documentclass{article}\begin{document}"
        + "\n{{resume.commands}}\n"
        + r"\end{document}",
        encoding="utf-8",
    )
    manifest = TemplateManifest(
        template_id="latex-router-test",
        document_type="resume",
        template_format="latex",
        renderer="latex",
        latex=LatexConfig(
            compile_engine="pdflatex",
            field_mappings=[
                LatexFieldMapping(ir_field="header.name", command="header", arity=1)
            ],
        ),
    )
    return TemplatePackage(
        template_id="latex-router-test",
        document_type="resume",
        directory=pkg_dir,
        template_path=template_tex,
        manifest_path=pkg_dir / "manifest.json",
        manifest=manifest,
    )


# ---- patch_existing dispatch ----------------------------------------


def test_patch_existing_requires_source(tmp_path: Path) -> None:
    with pytest.raises(MaterialsRouterError, match="requires source"):
        generate_materials(
            document=_ir(),
            mode="patch_existing",
            output_dir=tmp_path,
        )


def test_patch_existing_docx_happy_path(tmp_path: Path) -> None:
    source_path = tmp_path / "src.docx"
    _build_docx_source(source_path)
    source = SourceResumeView(
        id=uuid.uuid4(),
        source_type="docx",
        editable=True,
        absolute_path=source_path,
    )
    out = tmp_path / "out"
    outcome = generate_materials(
        document=_ir(),
        mode="patch_existing",
        output_dir=out,
        source=source,
    )
    assert outcome.decision == "patch_docx_ok"
    assert outcome.output_paths == [out / "patched_resume.docx"]
    assert outcome.bindings.source_resume_id == source.id


def test_patch_existing_docx_fallback_uses_template_when_provided(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """A PatchFallback (e.g. missing source file) routes to the template
    if one is provided. We force PatchFallback by deleting the source
    after the View is built."""
    source_path = tmp_path / "ghost.docx"
    source = SourceResumeView(
        id=uuid.uuid4(),
        source_type="docx",
        editable=True,
        absolute_path=source_path,
    )
    pkg = _latex_package(tmp_path)

    # Stub the LaTeX compile so the fallback path completes without a
    # real compiler.
    def _stub_render_and_compile(
        template_path: Path, document: Any, *, manifest: Any, package_dir: Any,
        tex_output: Path, pdf_output: Path | None = None, **_: Any,
    ) -> tuple[Path, Path]:
        tex_output.parent.mkdir(parents=True, exist_ok=True)
        tex_output.write_text("ok", encoding="utf-8")
        pdf = pdf_output or tex_output.with_suffix(".pdf")
        pdf.write_bytes(b"%PDF-1.4\n")
        return tex_output, pdf

    monkeypatch.setattr(
        "src.generation.materials_router.render_and_compile",
        _stub_render_and_compile,
    )

    outcome = generate_materials(
        document=_ir(),
        mode="patch_existing",
        output_dir=tmp_path / "out",
        source=source,
        template_package=pkg,
    )
    assert outcome.decision == "patch_docx_fallback"
    assert outcome.mode == "patch_existing"  # mode is preserved
    assert outcome.output_paths  # template path produced
    assert any("docx patch failed" in w for w in outcome.warnings)


def test_patch_existing_pdf_routes_to_template(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    source = SourceResumeView(
        id=uuid.uuid4(),
        source_type="pdf",
        editable=False,
        absolute_path=tmp_path / "irrelevant.pdf",
    )
    pkg = _latex_package(tmp_path)

    def _stub(
        template_path: Path, document: Any, *, manifest: Any, package_dir: Any,
        tex_output: Path, pdf_output: Path | None = None, **_: Any,
    ) -> tuple[Path, Path]:
        tex_output.parent.mkdir(parents=True, exist_ok=True)
        tex_output.write_text("ok", encoding="utf-8")
        pdf = pdf_output or tex_output.with_suffix(".pdf")
        pdf.write_bytes(b"%PDF-1.4\n")
        return tex_output, pdf

    monkeypatch.setattr("src.generation.materials_router.render_and_compile", _stub)

    outcome = generate_materials(
        document=_ir(),
        mode="patch_existing",
        output_dir=tmp_path / "out",
        source=source,
        template_package=pkg,
    )
    assert outcome.decision == "generate_latex"
    assert any("not editable" not in w for w in outcome.warnings)
    assert any("D024" in w for w in outcome.warnings)


def test_patch_existing_pdf_without_fallback_reports_unsupported(tmp_path: Path) -> None:
    source = SourceResumeView(
        id=uuid.uuid4(),
        source_type="pdf",
        editable=False,
        absolute_path=tmp_path / "x.pdf",
    )
    outcome = generate_materials(
        document=_ir(),
        mode="patch_existing",
        output_dir=tmp_path,
        source=source,
    )
    assert outcome.decision == "unsupported"
    assert outcome.failure is not None and "not editable" in outcome.failure


def test_patch_existing_latex_source_routes_to_template(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    source = SourceResumeView(
        id=uuid.uuid4(),
        source_type="latex",
        editable=True,
        absolute_path=tmp_path / "src.tex",
    )
    pkg = _latex_package(tmp_path)
    monkeypatch.setattr(
        "src.generation.materials_router.render_and_compile",
        lambda *a, **kw: (
            kw.get("tex_output").write_text("ok") or kw.get("tex_output"),
            kw.get("pdf_output") or kw.get("tex_output").with_suffix(".pdf"),
        )[0:0]  # discard, see below
        or (
            kw["tex_output"],
            (kw.get("pdf_output") or kw["tex_output"].with_suffix(".pdf")).write_bytes(
                b"%PDF-1.4\n"
            )
            or (kw.get("pdf_output") or kw["tex_output"].with_suffix(".pdf")),
        ),
    )
    outcome = generate_materials(
        document=_ir(),
        mode="patch_existing",
        output_dir=tmp_path / "out",
        source=source,
        template_package=pkg,
    )
    assert outcome.decision == "generate_latex"
    assert any("rendered fresh" in w for w in outcome.warnings)


# ---- generate_from_template -----------------------------------------


def test_generate_from_template_requires_template(tmp_path: Path) -> None:
    with pytest.raises(MaterialsRouterError, match="requires template_package"):
        generate_materials(
            document=_ir(),
            mode="generate_from_template",
            output_dir=tmp_path,
        )


def test_generate_from_template_latex_path(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    pkg = _latex_package(tmp_path)
    seen_outputs: list[tuple[Path, Path]] = []

    def _stub(
        template_path: Path, document: Any, *, manifest: Any, package_dir: Any,
        tex_output: Path, pdf_output: Path | None = None, **_: Any,
    ) -> tuple[Path, Path]:
        tex_output.parent.mkdir(parents=True, exist_ok=True)
        tex_output.write_text("ok", encoding="utf-8")
        pdf = pdf_output or tex_output.with_suffix(".pdf")
        pdf.write_bytes(b"%PDF-1.4\n")
        seen_outputs.append((tex_output, pdf))
        return tex_output, pdf

    monkeypatch.setattr("src.generation.materials_router.render_and_compile", _stub)

    outcome = generate_materials(
        document=_ir(),
        mode="generate_from_template",
        output_dir=tmp_path / "out",
        template_package=pkg,
    )
    assert outcome.decision == "generate_latex"
    assert len(outcome.output_paths) == 2  # .tex + .pdf
    assert outcome.bindings.template_package_id == "latex-router-test"


def test_generate_from_template_latex_propagates_render_error(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    pkg = _latex_package(tmp_path)
    from src.documents.latex_renderer import ManifestRenderError

    def _stub_fail(*_args: Any, **_kwargs: Any) -> tuple[Path, Path]:
        raise ManifestRenderError("simulated failure")

    monkeypatch.setattr("src.generation.materials_router.render_and_compile", _stub_fail)

    outcome = generate_materials(
        document=_ir(),
        mode="generate_from_template",
        output_dir=tmp_path / "out",
        template_package=pkg,
    )
    assert outcome.decision == "generate_latex"
    assert outcome.failure is not None and "simulated failure" in outcome.failure
    assert outcome.output_paths == []


def test_generate_from_template_latex_without_manifest_latex_block_falls_back_to_placeholder(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """Codex P2 fix: a LaTeX package with no Phase 15.3 ``latex``
    block falls back to the placeholder-based engine instead of
    returning 'unsupported'. We monkey-patch the placeholder helper
    so the test stays decoupled from a real compiler."""
    pkg_dir = tmp_path / "legacy_pkg"
    pkg_dir.mkdir()
    template_tex = pkg_dir / "template.tex"
    template_tex.write_text(
        "\\documentclass{article}\\begin{document}{{resume.sections}}\\end{document}",
        encoding="utf-8",
    )
    manifest = TemplateManifest(
        template_id="legacy",
        document_type="resume",
        template_format="latex",
        renderer="latex",
        # no latex= block -- legacy shape
    )
    pkg = TemplatePackage(
        template_id="legacy",
        document_type="resume",
        directory=pkg_dir,
        template_path=template_tex,
        manifest_path=pkg_dir / "manifest.json",
        manifest=manifest,
    )

    monkeypatch.setattr(
        "src.generation.materials_router._render_placeholder_latex",
        lambda tp, doc, tex_out, pdf_out: (
            (tex_out.write_text("ok", encoding="utf-8") or tex_out),
            (pdf_out.write_bytes(b"%PDF-1.4\n") or pdf_out),
        ),
    )

    outcome = generate_materials(
        document=_ir(),
        mode="generate_from_template",
        output_dir=tmp_path / "out",
        template_package=pkg,
    )
    assert outcome.decision == "generate_latex"
    assert outcome.failure is None
    assert len(outcome.output_paths) == 2


# ---- Bindings carry through ------------------------------------------


def test_bindings_inherit_explicit_values(tmp_path: Path) -> None:
    source_path = tmp_path / "src.docx"
    _build_docx_source(source_path)
    source = SourceResumeView(
        id=uuid.uuid4(),
        source_type="docx",
        editable=True,
        absolute_path=source_path,
    )
    explicit_snapshot = uuid.uuid4()
    outcome = generate_materials(
        document=_ir(),
        mode="patch_existing",
        output_dir=tmp_path / "out",
        source=source,
        bindings=MaterialsBindings(
            job_snapshot_id=explicit_snapshot,
            profile_version="v42",
            trace_id="20260515T0000Z-aabb",
            tenant_id="t-1",
        ),
    )
    assert outcome.bindings.job_snapshot_id == explicit_snapshot
    assert outcome.bindings.profile_version == "v42"
    assert outcome.bindings.trace_id == "20260515T0000Z-aabb"
    assert outcome.bindings.tenant_id == "t-1"
    assert outcome.bindings.source_resume_id == source.id


def test_bindings_to_dict_renders_strings() -> None:
    bindings = MaterialsBindings(
        job_snapshot_id=uuid.uuid4(),
        source_resume_id=uuid.uuid4(),
        template_package_id="ats_single_column_v1",
        profile_version="v1",
        trace_id="abc",
        tenant_id="t",
    )
    payload = bindings.to_dict()
    assert isinstance(payload["job_snapshot_id"], str)
    assert isinstance(payload["source_resume_id"], str)
    assert payload["template_package_id"] == "ats_single_column_v1"


# ---- Mode validation -------------------------------------------------


def test_unknown_mode_errors(tmp_path: Path) -> None:
    # ``mode`` is a Literal, but Python doesn't enforce at runtime.
    with pytest.raises(MaterialsRouterError, match="unknown mode"):
        generate_materials(
            document=_ir(),
            mode="invalid_mode",  # type: ignore[arg-type]
            output_dir=tmp_path,
        )


# ---- PatchFallback contract (docs) -----------------------------------


def test_patchfallback_is_exposed_from_docx_patch_module() -> None:
    """The router catches PatchFallback specifically, so this guards
    against renaming the exception class elsewhere without updating
    the router."""
    assert PatchFallback is not None
