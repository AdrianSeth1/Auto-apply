"""Tests for the defensive YAML-response parser in resume_importer.

The LLM is *supposed* to return a bare YAML mapping, but in practice it
sometimes wraps the document in markdown fences, prepends a chatty
preamble, or double-encodes the whole thing as a string scalar. These
tests pin the parser's tolerance for those modes and verify that genuine
garbage produces an actionable error message (not a confusing one like
"got <class 'str'>").
"""

from __future__ import annotations

from pathlib import Path

import pytest

from src.memory.resume_importer import (
    _parse_llm_yaml_response,
    extract_text_from_docx,
)

VALID_YAML = """\
identity:
  full_name: "Liam Liu"
  email: "liam@example.com"
education: []
work_experiences: []
projects: []
skills:
  languages: ["python"]
"""


def test_parses_bare_yaml() -> None:
    parsed = _parse_llm_yaml_response(VALID_YAML)
    assert parsed["identity"]["full_name"] == "Liam Liu"


def test_strips_markdown_fences() -> None:
    wrapped = f"```yaml\n{VALID_YAML}\n```"
    parsed = _parse_llm_yaml_response(wrapped)
    assert parsed["identity"]["full_name"] == "Liam Liu"


def test_strips_chatty_preamble() -> None:
    chatty = "Here is the structured YAML for the resume:\n\n" + VALID_YAML
    parsed = _parse_llm_yaml_response(chatty)
    assert parsed["identity"]["full_name"] == "Liam Liu"


def test_strips_fences_then_preamble() -> None:
    """Fence-stripping runs first, so a fenced block whose first
    content line is a chatty preamble still resolves cleanly."""
    response = (
        f"```yaml\nHere is the structured YAML for the resume:\n\n{VALID_YAML}\n```"
    )
    parsed = _parse_llm_yaml_response(response)
    assert parsed["identity"]["full_name"] == "Liam Liu"


def test_plain_text_response_raises_with_snippet() -> None:
    response = "I'm sorry, but I cannot parse this resume."
    with pytest.raises(ValueError) as exc:
        _parse_llm_yaml_response(response)
    # The actual response content should be in the error so the user
    # can tell what went wrong, instead of the cryptic "got <class 'str'>".
    assert "cannot parse this resume" in str(exc.value)


def test_yaml_list_response_raises_with_type_info() -> None:
    response = "- python\n- javascript\n"
    with pytest.raises(ValueError) as exc:
        _parse_llm_yaml_response(response)
    assert "list" in str(exc.value).lower()


def test_malformed_yaml_raises_with_snippet() -> None:
    response = "identity:\n  full_name: \"unclosed\n"
    with pytest.raises(ValueError) as exc:
        _parse_llm_yaml_response(response)
    assert "invalid YAML" in str(exc.value)


# ---- DOCX hyperlink extraction ----------------------------------------


def _build_docx_with_hyperlink(
    path: Path,
    *,
    leading_text: str,
    link_text: str,
    link_url: str,
) -> None:
    """Build a one-paragraph DOCX with a hyperlink inside it.

    python-docx doesn't ship a high-level hyperlink helper, so we
    poke at the underlying XML the same way the upstream issue
    threads recommend (Relationship + raw OOXML element). The
    result is a real .docx that mirrors what Word produces when a
    user types a URL and presses space.
    """
    from docx import Document
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()
    paragraph = doc.add_paragraph(leading_text)

    rel_id = paragraph.part.relate_to(link_url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = link_text
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)

    doc.save(str(path))


def test_extract_inlines_hyperlink_url(tmp_path: Path) -> None:
    src = tmp_path / "linked.docx"
    _build_docx_with_hyperlink(
        src,
        leading_text="Contact: ",
        link_text="GitHub",
        link_url="https://github.com/liam",
    )

    rendered = extract_text_from_docx(src)

    assert "GitHub" in rendered
    assert "https://github.com/liam" in rendered
    assert "GitHub (https://github.com/liam)" in rendered


def test_extract_does_not_double_emit_url_when_label_is_the_url(
    tmp_path: Path,
) -> None:
    """If the link's display text is already the URL itself (e.g.
    a raw https:// pasted into Word), don't repeat it as ``url
    (url)``."""
    src = tmp_path / "linked.docx"
    _build_docx_with_hyperlink(
        src,
        leading_text="See ",
        link_text="https://example.com/me",
        link_url="https://example.com/me",
    )

    rendered = extract_text_from_docx(src)

    assert rendered.count("https://example.com/me") == 1
    assert "(https://example.com/me)" not in rendered


def test_extract_preserves_plain_paragraphs(tmp_path: Path) -> None:
    """Regression: the hyperlink-aware walker must not regress
    plain paragraphs without links."""
    from docx import Document

    src = tmp_path / "plain.docx"
    doc = Document()
    doc.add_paragraph("Liam Liu")
    doc.add_paragraph("Software Engineer Intern")
    doc.save(str(src))

    rendered = extract_text_from_docx(src)

    assert "Liam Liu" in rendered
    assert "Software Engineer Intern" in rendered


# ----------------------------- Custom sections ------------------------------


def test_scan_resume_headings_finds_volunteer_and_awards() -> None:
    """Deterministic scanner must pull non-canonical headings even when
    the LLM-extracted YAML doesn't include them, so the candidate's
    Volunteer / Awards / Affiliations content cannot silently disappear."""
    from src.memory.resume_importer import _looks_canonical_heading, _scan_resume_headings

    text = """LIAM FROST
liam@example.com

EDUCATION
University X
BSc CS

VOLUNTEER EXPERIENCE
Local Library
- Tutored 12 students

AWARDS & HONORS
Dean's List 2023
"""

    headings = _scan_resume_headings(text)
    titles = [h["title"] for h in headings]
    assert "EDUCATION" in titles
    assert "VOLUNTEER EXPERIENCE" in titles
    assert "AWARDS & HONORS" in titles
    # The candidate's name at the very top is NOT a section heading.
    assert "LIAM FROST" not in titles
    # _looks_canonical_heading must classify these correctly.
    assert _looks_canonical_heading("EDUCATION")
    assert not _looks_canonical_heading("VOLUNTEER EXPERIENCE")
    assert not _looks_canonical_heading("AWARDS & HONORS")


def test_backfill_custom_sections_adds_missed_headings(monkeypatch) -> None:
    """When the LLM forgets to populate ``custom_sections``, the
    deterministic backfill must rescue every non-canonical heading."""
    from src.memory import resume_importer

    raw_text = """JANE DOE
jane@example.com

EDUCATION
University X
BSc CS

VOLUNTEER EXPERIENCE
Local Library
- Tutored 12 students

AWARDS & HONORS
Dean's List 2023
"""

    profile_data = {
        "identity": {"full_name": "Jane Doe"},
        "education": [{"institution": "University X"}],
        # LLM forgot custom_sections entirely
    }

    # Force the structuring LLM call to fail so we exercise the
    # deterministic fallback path.
    def _explode(*args, **kwargs):
        raise RuntimeError("simulated LLM outage")

    monkeypatch.setattr(
        "src.utils.llm.generate_json", _explode, raising=False
    )

    enriched = resume_importer._backfill_custom_sections(profile_data, raw_text)

    titles = [s["title"] for s in enriched["custom_sections"]]
    assert "VOLUNTEER EXPERIENCE" in titles
    assert "AWARDS & HONORS" in titles
    for section in enriched["custom_sections"]:
        # Fallback path leaves the raw body in details so content
        # is preserved verbatim, never lost.
        assert section["entries"]
        first = section["entries"][0]
        assert first.get("details")
